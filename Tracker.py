# Standard library imports
import sys
import os
import json
import getpass
import logging
import threading
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple, Union, Set, Callable, TypeVar, Generic, TypedDict
from dataclasses import dataclass, field
from pathlib import Path
from enum import Enum
from contextlib import contextmanager

# Third-party imports
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logging.warning("pandas not available. Some features may be limited.")

# Markdown rendering support
try:
    import markdown2
    MARKDOWN2_AVAILABLE = True
except ImportError:
    MARKDOWN2_AVAILABLE = False

try:
    import mistune
    MISTUNE_AVAILABLE = True
except ImportError:
    MISTUNE_AVAILABLE = False

import sqlite3
import tempfile
from PyQt6.QtSql import QSqlDatabase, QSqlQuery

# PyQt6 imports - organized by module
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QTabWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QTableWidget, QTableWidgetItem, QMessageBox, QFileDialog,
    QLabel, QLineEdit, QComboBox, QCalendarWidget, QDialog, QInputDialog, 
    QMenu, QFrame, QGraphicsDropShadowEffect, QListWidget, QHeaderView, 
    QListWidgetItem, QTreeWidget, QTreeWidgetItem, QCheckBox, QGridLayout, 
    QProgressBar, QGraphicsView, QGraphicsScene, QToolBar, QStyledItemDelegate, 
    QAbstractItemView, QToolTip, QTableView, QToolButton, QGroupBox, 
    QSpinBox, QTextEdit, QScrollArea, QSplitter, QStackedWidget, QTreeView,
    QButtonGroup
)
from PyQt6.QtCore import (
    Qt, QTimer, QDate, QRectF, QRect, QPropertyAnimation, QEasingCurve, 
    QAbstractAnimation, QObject, pyqtSignal, QModelIndex, QFileSystemWatcher, QSize
)
from PyQt6.QtGui import (
    QAction, QColor, QPalette, QFont, QIcon, QPixmap, QPainter, QPen, 
    QBrush, QPainterPath, QKeySequence, QShortcut, QIntValidator, QStandardItemModel, QStandardItem, QTextCursor
)
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog


# SVG export support has been removed

# Matplotlib imports
from matplotlib.backends.backend_qtagg import FigureCanvasQTAgg as FigureCanvas
import matplotlib.pyplot as plt

# Silence Qt stylesheet 'unknown property' warnings (e.g., 'transform') without changing logic
try:
    from PyQt6.QtCore import qInstallMessageHandler, QtMsgType
    def _qt_msg_handler(mode, context, message):
        try:
            text = str(message)
            lower = text.lower()
            # Filter typical stylesheet unknown property warnings
            if (
                "unknown property" in lower or
                "is an unknown property" in lower or
                "unknown shorthand" in lower
            ):
                return
            # Forward other Qt messages to stderr
            try:
                sys.stderr.write(text + "\n")
            except Exception as e:
                # Logger not yet defined, use print for now
                print(f"Error in Qt message handler: {e}")
        except Exception as e:
            # Logger not yet defined, use print for now
            print(f"Error in Qt message handler setup: {e}")
    qInstallMessageHandler(_qt_msg_handler)
except Exception as e:
    # Logger not yet defined, use print for now
    print(f"Failed to install Qt message handler: {e}")

# Type variables for generic types
T = TypeVar('T')
K = TypeVar('K')
V = TypeVar('V')

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('tracker.log'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Error Handling Systems
    
# ASCII-safe icons for log levels (avoid Windows console encoding issues)
LEVEL_EMOJI = {
    "INFO": "[i]",
    "SUCCESS": "[+]",
    "WARNING": "[!]",
    "ERROR": "[x]",
}

# ErrorHandler class moved to line 2506 - see consolidated version there

class MemoryManager:
    """Memory management and cleanup system"""
    
    def __init__(self):
        self._object_pool = {}
        self._cleanup_callbacks = []
    
    def register_cleanup(self, obj: Any, cleanup_func: Callable):
        """Register an object for cleanup when it's no longer needed"""
        obj_id = id(obj)
        self._cleanup_callbacks.append((obj_id, cleanup_func))
    
    def cleanup_object(self, obj: Any):
        """Clean up a specific object"""
        obj_id = id(obj)
        for i, (registered_id, cleanup_func) in enumerate(self._cleanup_callbacks):
            if registered_id == obj_id:
                try:
                    cleanup_func()
                except Exception as e:
                    # Use the consolidated ErrorHandler from line 2506
                    pass  # Will be fixed when we update all references
                self._cleanup_callbacks.pop(i)
                break
    
    def cleanup_all(self):
        """Clean up all registered objects"""
        for obj_id, cleanup_func in self._cleanup_callbacks:
            try:
                cleanup_func()
            except Exception as e:
                ErrorHandler.handle_ui_error("memory_cleanup_all", e)
        self._cleanup_callbacks.clear()

# Global instances
memory_manager = MemoryManager()

# Input Validation and Data Validation Systems
class InputValidator:
    """Comprehensive input validation system"""
    
    # Constants for validation
    MAX_TEXT_LENGTH = 255
    MAX_HOURS_PER_DAY = 999
    MAX_DESCRIPTION_LENGTH = 1000
    MIN_HOURS = 0
    MAX_DECIMAL_PLACES = 2
    
    @staticmethod
    def validate_text_length(value: str, field_name: str, max_length: int = MAX_TEXT_LENGTH) -> Tuple[bool, str]:
        """Validate text length"""
        if not value or not str(value).strip():
            return True, ""  # Empty values are allowed
        
        if len(str(value)) > max_length:
            return False, f"{field_name} must be {max_length} characters or less"
        return True, ""
    
    @staticmethod
    def validate_hours(value: str, field_name: str = "Hours") -> Tuple[bool, str]:
        """Validate hours values"""
        if not value or not str(value).strip():
            return True, ""  # Empty hours are allowed
        
        try:
            hours = float(value)
            if hours < InputValidator.MIN_HOURS or hours > InputValidator.MAX_HOURS_PER_DAY:
                return False, f"{field_name} must be between {InputValidator.MIN_HOURS} and {InputValidator.MAX_HOURS_PER_DAY}"
            
            # Check decimal places
            decimal_places = len(str(hours).split('.')[-1]) if '.' in str(hours) else 0
            if decimal_places > InputValidator.MAX_DECIMAL_PLACES:
                return False, f"{field_name} can have at most {InputValidator.MAX_DECIMAL_PLACES} decimal places"
            
            return True, ""
        except ValueError:
            return False, f"{field_name} must be a valid number"
    
    @staticmethod
    def validate_date_range(start_date: datetime, end_date: datetime) -> Tuple[bool, str]:
        """Validate date range"""
        if start_date > end_date:
            return False, "Start date cannot be after end date"
        
        # Check if date range is too large (more than 1 year)
        if (end_date - start_date).days > 365:
            return False, "Date range cannot exceed 1 year"
        
        return True, ""
    
    @staticmethod
    def validate_email(email: str) -> Tuple[bool, str]:
        """Validate email format"""
        if not email:
            return True, ""  # Empty email is allowed
        
        import re
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(pattern, email):
            return False, "Invalid email format"
        return True, ""
    
    @staticmethod
    def validate_required_field(value: Any, field_name: str) -> Tuple[bool, str]:
        """Validate required fields"""
        if value is None or (isinstance(value, str) and not value.strip()):
            return False, f"{field_name} is required"
        return True, ""

class DataValidator:
    """Data validation and integrity checking system"""
    
    def __init__(self):
        self.validation_rules = {}
        self.custom_validators = {}
    
    def add_validation_rule(self, field_name: str, rule: Callable[[Any], Tuple[bool, str]]):
        """Add a custom validation rule for a field"""
        self.validation_rules[field_name] = rule
    
    def validate_data(self, data: Dict[str, Any]) -> Tuple[bool, List[str]]:
        """Validate data against all rules"""
        errors = []
        
        for field_name, value in data.items():
            # Check if there's a custom rule for this field
            if field_name in self.validation_rules:
                is_valid, error_msg = self.validation_rules[field_name](value)
                if not is_valid:
                    errors.append(f"{field_name}: {error_msg}")
            
            # Apply standard validations based on field name
            if "hours" in field_name.lower():
                is_valid, error_msg = InputValidator.validate_hours(str(value), field_name)
                if not is_valid:
                    errors.append(error_msg)
            
            if "email" in field_name.lower():
                is_valid, error_msg = InputValidator.validate_email(str(value))
                if not is_valid:
                    errors.append(error_msg)
            
            if "description" in field_name.lower():
                is_valid, error_msg = InputValidator.validate_text_length(
                    str(value), field_name, InputValidator.MAX_DESCRIPTION_LENGTH
                )
                if not is_valid:
                    errors.append(error_msg)
        
        return len(errors) == 0, errors
    
    def validate_data_integrity(self, data_list: List[Dict[str, Any]]) -> Tuple[bool, List[str]]:
        """Validate data integrity across multiple records"""
        errors = []
        
        # Check for duplicate IDs
        ids = [record.get('id') for record in data_list if record.get('id')]
        if len(ids) != len(set(ids)):
            errors.append("Duplicate IDs found in data")
        
        # Check for orphaned references
        # This would be implemented based on specific business rules
        
        return len(errors) == 0, errors

# Global validator instance
data_validator = DataValidator()

# Responsive Layout and UI Performance Systems
class ResponsiveLayoutManager:
    """Responsive layout management for different screen sizes"""
    
    def __init__(self):
        self.breakpoints = {
            'mobile': 480,
            'tablet': 768,
            'desktop': 1024,
            'large': 1440
        }
        self.current_breakpoint = self.get_current_breakpoint()
        self.layout_callbacks = {}
    
    def get_screen_size(self) -> Tuple[int, int]:
        """Get current screen size"""
        app = QApplication.instance()
        if app:
            screen = app.primaryScreen()
            if screen:
                geometry = screen.availableGeometry()
                return geometry.width(), geometry.height()
        return 1024, 768  # Default fallback
    
    def get_current_breakpoint(self) -> str:
        """Get current breakpoint based on screen size"""
        width, height = self.get_screen_size()
        
        if width <= self.breakpoints['mobile']:
            return 'mobile'
        elif width <= self.breakpoints['tablet']:
            return 'tablet'
        elif width <= self.breakpoints['desktop']:
            return 'desktop'
        else:
            return 'large'
    
    def register_layout_callback(self, breakpoint: str, callback: Callable):
        """Register a callback for layout changes at specific breakpoints"""
        if breakpoint not in self.layout_callbacks:
            self.layout_callbacks[breakpoint] = []
        self.layout_callbacks[breakpoint].append(callback)
    
    def apply_responsive_layout(self, widget: QWidget):
        """Apply responsive layout to a widget"""
        current_breakpoint = self.get_current_breakpoint()
        
        if current_breakpoint != self.current_breakpoint:
            self.current_breakpoint = current_breakpoint
            
            # Execute callbacks for the new breakpoint
            if current_breakpoint in self.layout_callbacks:
                for callback in self.layout_callbacks[current_breakpoint]:
                    try:
                        callback(widget)
                    except Exception as e:
                        ErrorHandler.handle_ui_error("responsive_layout", e)

class VirtualScrollingTable(QTableWidget):
    """Table widget with virtual scrolling for performance"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self._data = []
        self._visible_rows = 50  # Number of rows to render at once
        self._scroll_offset = 0
        self._total_rows = 0
        self._item_pool = []  # Object pool for table items
        
    def set_data(self, data: List[Dict[str, Any]]):
        """Set data for virtual scrolling"""
        self._data = data
        self._total_rows = len(data)
        self.setRowCount(min(self._visible_rows, self._total_rows))
        self._update_visible_items()
    
    def _update_visible_items(self):
        """Update visible items based on scroll position"""
        start_row = self._scroll_offset
        end_row = min(start_row + self._visible_rows, self._total_rows)
        
        for row in range(self.rowCount()):
            data_row = start_row + row
            if data_row < self._total_rows:
                self._populate_row(row, self._data[data_row])
            else:
                self._clear_row(row)
    
    def _populate_row(self, row: int, data: Dict[str, Any]):
        """Populate a row with data"""
        for col, (key, value) in enumerate(data.items()):
            if col < self.columnCount():
                item = self._get_pooled_item()
                item.setText(str(value))
                self.setItem(row, col, item)
    
    def _clear_row(self, row: int):
        """Clear a row and return items to pool"""
        for col in range(self.columnCount()):
            item = self.takeItem(row, col)
            if item:
                self._return_to_pool(item)
    
    def _get_pooled_item(self) -> QTableWidgetItem:
        """Get an item from the pool or create a new one"""
        if self._item_pool:
            return self._item_pool.pop()
        return QTableWidgetItem()
    
    def _return_to_pool(self, item: QTableWidgetItem):
        """Return an item to the pool"""
        item.setText("")
        self._item_pool.append(item)


class ProgressiveLoader:
    """Progressive loading system for large datasets"""
    
    def __init__(self, chunk_size: int = 100):
        self.chunk_size = chunk_size
        self.loaded_chunks = 0
        self.total_chunks = 0
        self.data = []
        self.load_callbacks = []
    
    def load_data_progressively(self, data_source: List[Any], callback: Callable):
        """Load data in chunks with progress callbacks"""
        self.data = data_source
        self.total_chunks = (len(data_source) + self.chunk_size - 1) // self.chunk_size
        self.loaded_chunks = 0
        self.load_callbacks.append(callback)
        
        self._load_next_chunk()
    
    def _load_next_chunk(self):
        """Load the next chunk of data"""
        start_idx = self.loaded_chunks * self.chunk_size
        end_idx = min(start_idx + self.chunk_size, len(self.data))
        
        chunk = self.data[start_idx:end_idx]
        
        # Process chunk
        for callback in self.load_callbacks:
            try:
                # Ensure UI-bound callbacks run on the main thread
                if QApplication.instance() is not None:
                    QTimer.singleShot(0, lambda cb=callback, ch=chunk, lc=self.loaded_chunks, tc=self.total_chunks: cb(ch, lc, tc))
                else:
                    callback(chunk, self.loaded_chunks, self.total_chunks)
            except Exception as e:
                ErrorHandler.handle_ui_error("progressive_loader", e)
        
        self.loaded_chunks += 1
        
        # Continue loading if there are more chunks
        if self.loaded_chunks < self.total_chunks:
            # Use singleShot only if QApplication is available
            if QApplication.instance() is not None:
                QTimer.singleShot(10, self._load_next_chunk)  # Small delay to keep UI responsive
            else:
                # Fallback: use threading for delay
                import threading
                def delayed_load():
                    threading.Event().wait(0.01)  # 10ms delay
                    self._load_next_chunk()
                threading.Thread(target=delayed_load, daemon=True).start()

# Global instances
responsive_manager = ResponsiveLayoutManager()

# Feature flags for optional subsystems
FEATURE_FLAGS = {
    'search_engine': True,
    'collaboration': False,
    'report_builder': False,
    'scheduled_reports': False,
    'data_manager': True,
    'help_system': False,
    'documentation_generator': False,
}

class FeatureTogglesDialog(QDialog):
    """Simple dialog to toggle optional subsystems at runtime (requires restart to fully apply)."""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Feature Toggles")
        self.setModal(True)
        self.resize(420, 300)
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("Enable or disable optional subsystems. Changes take effect after restart."))
        self._checks: dict[str, QCheckBox] = {}
        grid = QGridLayout()
        row = 0
        for key, value in FEATURE_FLAGS.items():
            label = key.replace('_', ' ').title()
            cb = QCheckBox(label)
            cb.setChecked(bool(value))
            self._checks[key] = cb
            grid.addWidget(cb, row, 0)
            row += 1
        layout.addLayout(grid)
        buttons = QHBoxLayout()
        buttons.addStretch(1)
        btn_ok = QPushButton("OK"); btn_cancel = QPushButton("Cancel")
        btn_ok.setObjectName("primary")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)
        buttons.addWidget(btn_ok); buttons.addWidget(btn_cancel)
        layout.addLayout(buttons)
    
    def result_flags(self) -> dict:
        return {k: self._checks[k].isChecked() for k in self._checks}

def install_feature_toggle_ui(window: QMainWindow):
    """Attach a Settings->Feature Toggles menu item and shortcut to the main window."""
    try:
        def open_dialog():
            dlg = FeatureTogglesDialog(window)
            if dlg.exec() == QDialog.DialogCode.Accepted:
                updated = dlg.result_flags()
                FEATURE_FLAGS.update(updated)
                try:
                    # Persist to SQLite settings
                    if hasattr(window, '_save_backend_sqlite'):
                        window._log_change("Settings", "System", f"Updated feature flags: {updated}")
                        window._save_backend_sqlite()
                except Exception as e:
                    ErrorHandler.handle_ui_error("persist feature flags", e)
                QMessageBox.information(window, "Feature Toggles", "Changes saved. Please restart the application to apply.")
        # Menu
        try:
            mb = window.menuBar() if hasattr(window, 'menuBar') else None
            if mb is not None:
                settings_menu = None
                for i in range(mb.actions().__len__()):
                    act = mb.actions()[i]
                    if act.text() and 'Settings' in act.text():
                        settings_menu = act.menu()
                        break
                if settings_menu is None:
                    settings_menu = QMenu("Settings", window)
                    mb.addMenu(settings_menu)
                act_toggle = QAction("Feature Toggles...", window)
                act_toggle.setShortcut(QKeySequence("Ctrl+,"))
                act_toggle.triggered.connect(open_dialog)
                settings_menu.addAction(act_toggle)
        except Exception as e:
            ErrorHandler.handle_ui_error("install menu feature toggles", e)
        # Shortcut (fallback)
        try:
            sc = QShortcut(QKeySequence("Ctrl+Alt+F"), window)
            sc.activated.connect(open_dialog)
        except Exception as e:
            ErrorHandler.handle_ui_error("install shortcut feature toggles", e)
    except Exception as e:
        ErrorHandler.handle_ui_error("install feature toggle ui", e)

def initialize_global_timers():
    """Initialize all global timers when QApplication is available"""
    try:
        if QApplication.instance() is not None:
            # Re-initialize timers for managers that were created before QApplication
            if hasattr(scheduled_report_manager, '_initialize_timer'):
                scheduled_report_manager._initialize_timer()
            if hasattr(data_manager, '_initialize_timer'):
                data_manager._initialize_timer()
    except Exception as e:
        logger.warning(f"Failed to initialize global timers: {e}")

# Comprehensive Search System
class SearchEngine:
    """Advanced search functionality with multiple search modes"""
    
    def __init__(self):
        self.search_index = {}
        self.search_history = []
        self.max_history = 100
    
    def index_data(self, data: List[Dict[str, Any]], id_field: str = 'id'):
        """Index data for fast searching"""
        self.search_index = {}
        
        for item in data:
            item_id = item.get(id_field, str(len(self.search_index)))
            searchable_text = self._extract_searchable_text(item)
            self.search_index[item_id] = {
                'data': item,
                'searchable_text': searchable_text.lower(),
                'keywords': self._extract_keywords(searchable_text)
            }
    
    def _extract_searchable_text(self, item: Dict[str, Any]) -> str:
        """Extract all searchable text from an item"""
        searchable_fields = []
        for key, value in item.items():
            if isinstance(value, (str, int, float)):
                searchable_fields.append(str(value))
        return ' '.join(searchable_fields)
    
    def _extract_keywords(self, text: str) -> Set[str]:
        """Extract keywords from text"""
        import re
        # Remove special characters and split into words
        words = re.findall(r'\b\w+\b', text.lower())
        return set(words)
    
    def search(self, query: str, search_type: str = 'contains') -> List[Dict[str, Any]]:
        """Search for items matching the query"""
        if not query.strip():
            return []
        
        # Add to search history
        self._add_to_history(query)
        
        query_lower = query.lower()
        results = []
        
        for item_id, indexed_item in self.search_index.items():
            if self._matches_query(indexed_item, query_lower, search_type):
                results.append(indexed_item['data'])
        
        return results
    
    def _matches_query(self, indexed_item: Dict, query: str, search_type: str) -> bool:
        """Check if an indexed item matches the query"""
        searchable_text = indexed_item['searchable_text']
        keywords = indexed_item['keywords']
        
        if search_type == 'contains':
            return query in searchable_text
        elif search_type == 'exact':
            return query == searchable_text
        elif search_type == 'starts_with':
            return searchable_text.startswith(query)
        elif search_type == 'ends_with':
            return searchable_text.endswith(query)
        elif search_type == 'keywords':
            query_keywords = set(query.split())
            return query_keywords.issubset(keywords)
        elif search_type == 'fuzzy':
            return self._fuzzy_match(searchable_text, query)
        
        return False
    
    def _fuzzy_match(self, text: str, query: str) -> bool:
        """Simple fuzzy matching implementation"""
        if len(query) < 3:
            return query in text
        
        # Simple Levenshtein distance check
        query_words = query.split()
        text_words = text.split()
        
        for query_word in query_words:
            for text_word in text_words:
                if self._levenshtein_distance(query_word, text_word) <= 2:
                    return True
        return False
    
    def _levenshtein_distance(self, s1: str, s2: str) -> int:
        """Calculate Levenshtein distance between two strings"""
        if len(s1) < len(s2):
            return self._levenshtein_distance(s2, s1)
        
        if len(s2) == 0:
            return len(s1)
        
        previous_row = list(range(len(s2) + 1))
        for i, c1 in enumerate(s1):
            current_row = [i + 1]
            for j, c2 in enumerate(s2):
                insertions = previous_row[j + 1] + 1
                deletions = current_row[j] + 1
                substitutions = previous_row[j] + (c1 != c2)
                current_row.append(min(insertions, deletions, substitutions))
            previous_row = current_row
        
        return previous_row[-1]
    
    def _add_to_history(self, query: str):
        """Add query to search history"""
        if query not in self.search_history:
            self.search_history.insert(0, query)
            if len(self.search_history) > self.max_history:
                self.search_history = self.search_history[:self.max_history]
    
    def get_search_suggestions(self, partial_query: str) -> List[str]:
        """Get search suggestions based on partial query"""
        suggestions = []
        partial_lower = partial_query.lower()
        
        # Add matching history items
        for history_item in self.search_history:
            if partial_lower in history_item.lower():
                suggestions.append(history_item)
        
        # Add matching keywords from index
        all_keywords = set()
        for indexed_item in self.search_index.values():
            all_keywords.update(indexed_item['keywords'])
        
        for keyword in all_keywords:
            if partial_lower in keyword and keyword not in suggestions:
                suggestions.append(keyword)
        
        return suggestions[:10]  # Limit to 10 suggestions

# Collaboration and Multi-User Support System
class CollaborationManager:
    """Multi-user collaboration and change tracking system"""
    
    def __init__(self):
        self.active_users = {}
        self.change_log = []
        self.comments = {}
        self.annotations = {}
        self.lock_manager = LockManager()
    
    def register_user(self, user_id: str, user_info: Dict[str, Any]):
        """Register a new user"""
        self.active_users[user_id] = {
            'info': user_info,
            'last_seen': datetime.now(),
            'current_activity': None
        }
    
    def log_change(self, user_id: str, change_type: str, item_id: str, 
                   old_value: Any, new_value: Any, timestamp: datetime = None):
        """Log a change made by a user"""
        if timestamp is None:
            timestamp = datetime.now()
        
        change_entry = {
            'id': len(self.change_log),
            'user_id': user_id,
            'change_type': change_type,
            'item_id': item_id,
            'old_value': old_value,
            'new_value': new_value,
            'timestamp': timestamp
        }
        
        self.change_log.append(change_entry)
        
        # Update user activity
        if user_id in self.active_users:
            self.active_users[user_id]['last_seen'] = timestamp
            self.active_users[user_id]['current_activity'] = f"{change_type} on {item_id}"
    
    def add_comment(self, user_id: str, item_id: str, comment: str, 
                   parent_comment_id: str = None):
        """Add a comment to an item"""
        comment_id = f"{item_id}_{len(self.comments)}"
        
        comment_entry = {
            'id': comment_id,
            'user_id': user_id,
            'item_id': item_id,
            'comment': comment,
            'parent_id': parent_comment_id,
            'timestamp': datetime.now(),
            'replies': []
        }
        
        self.comments[comment_id] = comment_entry
        
        # Add to parent comment's replies if applicable
        if parent_comment_id and parent_comment_id in self.comments:
            self.comments[parent_comment_id]['replies'].append(comment_id)
        
        return comment_id
    
    def get_item_comments(self, item_id: str) -> List[Dict[str, Any]]:
        """Get all comments for an item"""
        item_comments = []
        for comment in self.comments.values():
            if comment['item_id'] == item_id and comment['parent_id'] is None:
                item_comments.append(comment)
        
        # Sort by timestamp
        item_comments.sort(key=lambda x: x['timestamp'])
        return item_comments
    
    def get_change_history(self, item_id: str) -> List[Dict[str, Any]]:
        """Get change history for an item"""
        item_changes = []
        for change in self.change_log:
            if change['item_id'] == item_id:
                item_changes.append(change)
        
        # Sort by timestamp
        item_changes.sort(key=lambda x: x['timestamp'])
        return item_changes

class LockManager:
    """Resource locking system for collaboration"""
    
    def __init__(self):
        self.locks = {}
        self.lock_timeout = 300  # 5 minutes
    
    def acquire_lock(self, user_id: str, resource_id: str) -> bool:
        """Acquire a lock on a resource"""
        current_time = datetime.now()
        
        # Check if resource is already locked
        if resource_id in self.locks:
            lock_info = self.locks[resource_id]
            
            # Check if lock has expired
            if (current_time - lock_info['timestamp']).seconds > self.lock_timeout:
                # Lock expired, remove it
                del self.locks[resource_id]
            else:
                # Resource is locked by another user
                return False
        
        # Acquire the lock
        self.locks[resource_id] = {
            'user_id': user_id,
            'timestamp': current_time
        }
        return True
    
    def release_lock(self, user_id: str, resource_id: str) -> bool:
        """Release a lock on a resource"""
        if resource_id in self.locks:
            lock_info = self.locks[resource_id]
            if lock_info['user_id'] == user_id:
                del self.locks[resource_id]
                return True
        return False
    
    def is_locked(self, resource_id: str) -> bool:
        """Check if a resource is locked"""
        if resource_id not in self.locks:
            return False
        
        current_time = datetime.now()
        lock_info = self.locks[resource_id]
        
        # Check if lock has expired
        if (current_time - lock_info['timestamp']).seconds > self.lock_timeout:
            del self.locks[resource_id]
            return False
        
        return True
    
    def get_lock_owner(self, resource_id: str) -> Optional[str]:
        """Get the user ID of the lock owner"""
        if resource_id in self.locks:
            lock_info = self.locks[resource_id]
            current_time = datetime.now()
            
            # Check if lock has expired
            if (current_time - lock_info['timestamp']).seconds > self.lock_timeout:
                del self.locks[resource_id]
                return None
            
            return lock_info['user_id']
        return None
# Global instances will be created after all classes are defined

# Reporting and Analytics System
class ReportBuilder:
    """Custom report builder with multiple output formats"""
    
    def __init__(self):
        self.report_templates = {}
        self.scheduled_reports = {}
        self.export_formats = ['pdf', 'csv', 'html', 'json']
    
    def create_report_template(self, template_name: str, fields: List[str], 
                             filters: Dict[str, Any] = None, 
                             grouping: List[str] = None,
                             sorting: List[Tuple[str, str]] = None):
        """Create a custom report template"""
        self.report_templates[template_name] = {
            'fields': fields,
            'filters': filters or {},
            'grouping': grouping or [],
            'sorting': sorting or [],
            'created_at': datetime.now()
        }
    
    def generate_report(self, template_name: str, data: List[Dict[str, Any]], 
                       output_format: str = 'pdf') -> str:
        """Generate a report from template and data"""
        if template_name not in self.report_templates:
            raise ValueError(f"Report template '{template_name}' not found")
        
        template = self.report_templates[template_name]
        
        # Apply filters
        filtered_data = self._apply_filters(data, template['filters'])
        
        # Apply grouping
        if template['grouping']:
            filtered_data = self._apply_grouping(filtered_data, template['grouping'])
        
        # Apply sorting
        if template['sorting']:
            filtered_data = self._apply_sorting(filtered_data, template['sorting'])
        
        # Generate output
        if output_format == 'pdf':
            return self._generate_pdf_report(filtered_data, template)
        elif output_format == 'csv':
            return self._generate_csv_report(filtered_data, template)
        elif output_format == 'html':
            return self._generate_html_report(filtered_data, template)
        elif output_format == 'json':
            return self._generate_json_report(filtered_data, template)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")
    
    def _apply_filters(self, data: List[Dict[str, Any]], filters: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Apply filters to data"""
        filtered_data = data.copy()
        
        for field, filter_value in filters.items():
            if isinstance(filter_value, dict):
                # Range filter
                if 'min' in filter_value and 'max' in filter_value:
                    filtered_data = [item for item in filtered_data 
                                   if filter_value['min'] <= item.get(field, 0) <= filter_value['max']]
                # Date range filter
                elif 'start_date' in filter_value and 'end_date' in filter_value:
                    start_date = filter_value['start_date']
                    end_date = filter_value['end_date']
                    filtered_data = [item for item in filtered_data 
                                   if start_date <= item.get(field, datetime.min) <= end_date]
            else:
                # Exact match filter
                filtered_data = [item for item in filtered_data 
                               if item.get(field) == filter_value]
        
        return filtered_data
    
    def _apply_grouping(self, data: List[Dict[str, Any]], grouping_fields: List[str]) -> List[Dict[str, Any]]:
        """Apply grouping to data"""
        grouped_data = {}
        
        for item in data:
            group_key = tuple(item.get(field, '') for field in grouping_fields)
            if group_key not in grouped_data:
                grouped_data[group_key] = []
            grouped_data[group_key].append(item)
        
        # Convert back to list with group headers
        result = []
        for group_key, group_items in grouped_data.items():
            # Add group header
            group_header = {}
            for i, field in enumerate(grouping_fields):
                group_header[field] = group_key[i]
            group_header['_is_group_header'] = True
            result.append(group_header)
            
            # Add group items
            result.extend(group_items)
        
        return result
    
    def _apply_sorting(self, data: List[Dict[str, Any]], sorting_rules: List[Tuple[str, str]]) -> List[Dict[str, Any]]:
        """Apply sorting to data"""
        def sort_key(item):
            key_values = []
            for field, direction in sorting_rules:
                value = item.get(field, '')
                if direction.lower() == 'desc':
                    # For descending order, we'll negate numeric values
                    if isinstance(value, (int, float)):
                        value = -value
                key_values.append(value)
            return key_values
        
        return sorted(data, key=sort_key)
    
    def _generate_pdf_report(self, data: List[Dict[str, Any]], template: Dict) -> str:
        """Generate PDF report"""
        try:
            from reportlab.lib.pagesizes import letter  # type: ignore[import]
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer  # type: ignore[import]
            from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import]
            from reportlab.lib import colors  # type: ignore[import]
            
            # Create PDF file
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            doc = SimpleDocTemplate(filename, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title = Paragraph("Custom Report", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Create table data
            if data:
                table_data = [template['fields']]  # Header row
                for item in data:
                    if not item.get('_is_group_header', False):
                        row = [str(item.get(field, '')) for field in template['fields']]
                        table_data.append(row)
                
                # Create table
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
            
            # Build PDF
            doc.build(story)
            return filename
            
        except ImportError:
            logger.error("ReportLab not available. PDF generation disabled.")
            return "Error: ReportLab package not installed. Please install it with: pip install reportlab"
    
    
    def _generate_csv_report(self, data: List[Dict[str, Any]], template: Dict) -> str:
        """Generate CSV report"""
        import csv
        
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        
        with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=template['fields'])
            writer.writeheader()
            
            for item in data:
                if not item.get('_is_group_header', False):
                    row = {field: item.get(field, '') for field in template['fields']}
                    writer.writerow(row)
        
        return filename
    
    def _generate_html_report(self, data: List[Dict[str, Any]], template: Dict) -> str:
        """Generate HTML report"""
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Custom Report</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                table { border-collapse: collapse; width: 100%; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .group-header { background-color: #e6f3ff; font-weight: bold; }
            </style>
        </head>
        <body>
            <h1>Custom Report</h1>
            <p>Generated on: {}</p>
            <table>
                <thead>
                    <tr>
        """.format(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        
        # Add header row
        for field in template['fields']:
            html_content += f"<th>{field}</th>"
        html_content += "</tr></thead><tbody>"
        
        # Add data rows
        for item in data:
            if item.get('_is_group_header', False):
                html_content += '<tr class="group-header">'
            else:
                html_content += '<tr>'
            
            for field in template['fields']:
                html_content += f"<td>{item.get(field, '')}</td>"
            html_content += '</tr>'
        
        html_content += "</tbody></table></body></html>"
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return filename
    
    def _generate_json_report(self, data: List[Dict[str, Any]], template: Dict) -> str:
        """Generate JSON report"""
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        report_data = {
            'template': template,
            'generated_at': datetime.now().isoformat(),
            'data': [item for item in data if not item.get('_is_group_header', False)]
        }
        
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=2, ensure_ascii=False, default=str)
        
        return filename
    
    def _generate_text_report(self, data: List[Dict[str, Any]], template: Dict) -> str:
        """Generate simple text report as fallback"""
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        with open(filename, 'w', encoding='utf-8') as f:
            f.write("Custom Report\n")
            f.write("=" * 50 + "\n")
            f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            
            if data:
                # Write header
                f.write("\t".join(template['fields']) + "\n")
                f.write("-" * (len("\t".join(template['fields'])) * 2) + "\n")
                
                # Write data
                for item in data:
                    if not item.get('_is_group_header', False):
                        row = [str(item.get(field, '')) for field in template['fields']]
                        f.write("\t".join(row) + "\n")
        
        return filename

class ScheduledReportManager:
    """Manager for scheduled reports"""
    
    def __init__(self):
        self.scheduled_reports = {}
        self.report_timer = None
        self._initialize_timer()
    
    def _initialize_timer(self):
        """Initialize timer only when QApplication is available"""
        if QApplication.instance() is not None:
            self.report_timer = QTimer()
            self.report_timer.timeout.connect(self._check_scheduled_reports)
            self.report_timer.start(60000)  # Check every minute
    
    def schedule_report(self, report_id: str, template_name: str, 
                       schedule_type: str, schedule_value: str,
                       output_format: str = 'pdf', email_recipients: List[str] = None):
        """Schedule a report to be generated automatically"""
        self.scheduled_reports[report_id] = {
            'template_name': template_name,
            'schedule_type': schedule_type,  # 'daily', 'weekly', 'monthly'
            'schedule_value': schedule_value,  # time of day, day of week, day of month
            'output_format': output_format,
            'email_recipients': email_recipients or [],
            'last_run': None,
            'next_run': self._calculate_next_run(schedule_type, schedule_value)
        }
    
    def _calculate_next_run(self, schedule_type: str, schedule_value: str) -> datetime:
        """Calculate the next run time for a scheduled report"""
        now = datetime.now()
        
        if schedule_type == 'daily':
            # schedule_value should be in format "HH:MM"
            hour, minute = map(int, schedule_value.split(':'))
            next_run = now.replace(hour=hour, minute=minute, second=0, microsecond=0)
            if next_run <= now:
                next_run += timedelta(days=1)
        
        elif schedule_type == 'weekly':
            # schedule_value should be day of week (0=Monday, 6=Sunday) and time
            day_of_week, time_str = schedule_value.split(' ')
            hour, minute = map(int, time_str.split(':'))
            days_ahead = int(day_of_week) - now.weekday()
            if days_ahead <= 0:  # Target day already happened this week
                days_ahead += 7
            next_run = now + timedelta(days=days_ahead)
            next_run = next_run.replace(hour=hour, minute=minute, second=0, microsecond=0)
        
        elif schedule_type == 'monthly':
            # schedule_value should be day of month and time
            day_of_month, time_str = schedule_value.split(' ')
            hour, minute = map(int, time_str.split(':'))
            next_run = now.replace(day=int(day_of_month), hour=hour, minute=minute, second=0, microsecond=0)
            if next_run <= now:
                # Move to next month
                if now.month == 12:
                    next_run = next_run.replace(year=now.year + 1, month=1)
                else:
                    next_run = next_run.replace(month=now.month + 1)
        
        return next_run
    
    def _check_scheduled_reports(self):
        """Check if any scheduled reports need to be run"""
        now = datetime.now()
        
        for report_id, report_config in self.scheduled_reports.items():
            if report_config['next_run'] <= now:
                try:
                    self._run_scheduled_report(report_id, report_config)
                    
                    # Update next run time
                    report_config['last_run'] = now
                    report_config['next_run'] = self._calculate_next_run(
                        report_config['schedule_type'], 
                        report_config['schedule_value']
                    )
                except Exception as e:
                    ErrorHandler.handle_ui_error("scheduled_report", e)
    
    def _run_scheduled_report(self, report_id: str, report_config: Dict):
        """Run a scheduled report"""
        # This would integrate with the main application's data
        # For now, we'll just log that the report should be run
        logger.info(f"Running scheduled report: {report_id}")
        
        # In a real implementation, this would:
        # 1. Get the current data from the application
        # 2. Generate the report using ReportBuilder
        # 3. Send email if recipients are specified
        # 4. Store the report file

# Global instances will be created after all classes are defined

# Data Management and Backup System
class DataManager:
    """Comprehensive data management with backup, restore, and versioning"""
    
    def __init__(self):
        self.data_versions = {}
        self.backup_schedule = {}
        self.export_formats = ['json', 'csv']
        self.backup_timer = None
        self._initialize_timer()
    
    def _initialize_timer(self):
        """Initialize timer only when QApplication is available"""
        if QApplication.instance() is not None:
            self.backup_timer = QTimer()
            self.backup_timer.timeout.connect(self._perform_scheduled_backup)
            self.backup_timer.start(3600000)  # Check every hour
    
    def create_data_version(self, data: List[Dict[str, Any]], version_name: str = None) -> str:
        """Create a versioned snapshot of data"""
        if version_name is None:
            version_name = f"version_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        version_data = {
            'id': version_name,
            'data': data.copy(),
            'created_at': datetime.now(),
            'size': len(data),
            'checksum': self._calculate_checksum(data)
        }
        
        self.data_versions[version_name] = version_data
        return version_name
    
    def restore_data_version(self, version_name: str) -> List[Dict[str, Any]]:
        """Restore data from a specific version"""
        if version_name not in self.data_versions:
            raise ValueError(f"Version '{version_name}' not found")
        
        version_data = self.data_versions[version_name]
        
        # Verify data integrity
        if not self._verify_checksum(version_data['data'], version_data['checksum']):
            raise ValueError("Data integrity check failed for version")
        
        return version_data['data'].copy()
    
    def get_version_history(self) -> List[Dict[str, Any]]:
        """Get list of all data versions"""
        versions = []
        for version_name, version_data in self.data_versions.items():
            versions.append({
                'name': version_name,
                'created_at': version_data['created_at'],
                'size': version_data['size'],
                'checksum': version_data['checksum']
            })
        
        # Sort by creation date (newest first)
        versions.sort(key=lambda x: x['created_at'], reverse=True)
        return versions
    
    def delete_version(self, version_name: str) -> bool:
        """Delete a specific data version"""
        if version_name in self.data_versions:
            del self.data_versions[version_name]
            return True
        return False
    
    def export_data(self, data: List[Dict[str, Any]], format_type: str, 
                   filename: str = None) -> str:
        """Export data in various formats"""
        if format_type not in self.export_formats:
            raise ValueError(f"Unsupported export format: {format_type}")
        
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"export_{timestamp}.{format_type}"
        
        if format_type == 'json':
            return self._export_json(data, filename)
        elif format_type == 'csv':
            return self._export_csv(data, filename)
        
        return filename
    
    def import_data(self, filename: str, format_type: str = None) -> List[Dict[str, Any]]:
        """Import data from various formats"""
        if format_type is None:
            format_type = filename.split('.')[-1].lower()
        
        if format_type not in self.export_formats:
            raise ValueError(f"Unsupported import format: {format_type}")
        
        if format_type == 'json':
            return self._import_json(filename)
        elif format_type == 'csv':
            return self._import_csv(filename)
        
        return []
    
    def schedule_backup(self, schedule_type: str, schedule_value: str, 
                       backup_format: str = 'json', max_backups: int = 10):
        """Schedule automatic data backups"""
        self.backup_schedule = {
            'type': schedule_type,  # 'daily', 'weekly', 'monthly'
            'value': schedule_value,  # time or day
            'format': backup_format,
            'max_backups': max_backups,
            'last_backup': None,
            'next_backup': self._calculate_next_backup(schedule_type, schedule_value)
        }
    
    def _calculate_checksum(self, data: List[Dict[str, Any]]) -> str:
        """Calculate checksum for data integrity verification"""
        import hashlib
        
        # Convert data to string for hashing
        data_str = json.dumps(data, sort_keys=True, default=str)
        return hashlib.md5(data_str.encode()).hexdigest()
    
    def _verify_checksum(self, data: List[Dict[str, Any]], expected_checksum: str) -> bool:
        """Verify data integrity using checksum"""
        actual_checksum = self._calculate_checksum(data)
        return actual_checksum == expected_checksum
    
    def _calculate_next_backup(self, schedule_type: str, schedule_value: str) -> datetime:
        """Calculate next backup time"""
        now = datetime.now()
        
        if schedule_type == 'daily':
            hour, minute = map(int, schedule_value.split(':'))
            next_backup = now.replace(hour=hour, minute=minute, second=0, microsecond=0)
            if next_backup <= now:
                next_backup += timedelta(days=1)
        
        elif schedule_type == 'weekly':
            day_of_week, time_str = schedule_value.split(' ')
            hour, minute = map(int, time_str.split(':'))
            days_ahead = int(day_of_week) - now.weekday()
            if days_ahead <= 0:
                days_ahead += 7
            next_backup = now + timedelta(days=days_ahead)
            next_backup = next_backup.replace(hour=hour, minute=minute, second=0, microsecond=0)
        
        elif schedule_type == 'monthly':
            day_of_month, time_str = schedule_value.split(' ')
            hour, minute = map(int, time_str.split(':'))
            next_backup = now.replace(day=int(day_of_month), hour=hour, minute=minute, second=0, microsecond=0)
            if next_backup <= now:
                if now.month == 12:
                    next_backup = next_backup.replace(year=now.year + 1, month=1)
                else:
                    next_backup = next_backup.replace(month=now.month + 1)
        
        return next_backup
    
    def _perform_scheduled_backup(self):
        """Perform scheduled backup if needed"""
        if not self.backup_schedule:
            return
        
        now = datetime.now()
        if self.backup_schedule['next_backup'] <= now:
            try:
                # This would integrate with the main application's data
                # For now, we'll just log that backup should be performed
                logger.info("Performing scheduled backup")
                
                # Update next backup time
                self.backup_schedule['last_backup'] = now
                self.backup_schedule['next_backup'] = self._calculate_next_backup(
                    self.backup_schedule['type'],
                    self.backup_schedule['value']
                )
                
            except Exception as e:
                ErrorHandler.handle_ui_error("scheduled_backup", e)
    
    def _export_json(self, data: List[Dict[str, Any]], filename: str) -> str:
        """Export data to JSON format"""
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False, default=str)
        return filename
    
    def _export_csv(self, data: List[Dict[str, Any]], filename: str) -> str:
        """Export data to CSV format"""
        import csv
        
        if not data:
            return filename
        
        with open(filename, 'w', newline='', encoding='utf-8') as csvfile:
            fieldnames = data[0].keys()
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(data)
        
        return filename
    
    
    # XML and SQL export have been removed to reduce maintenance surface
    
    def _import_json(self, filename: str) -> List[Dict[str, Any]]:
        """Import data from JSON format"""
        with open(filename, 'r', encoding='utf-8') as f:
            return json.load(f)
    
    def _import_csv(self, filename: str) -> List[Dict[str, Any]]:
        """Import data from CSV format"""
        import csv
        
        data = []
        with open(filename, 'r', encoding='utf-8') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                data.append(row)
        return data
    
    
    # XML and SQL import have been removed
# Global instances will be created after all classes are defined

# Help System and Documentation
class HelpSystem:
    """Comprehensive help system with user manual and context-sensitive help"""
    
    def __init__(self):
        self.help_topics = {}
        self.tutorials = {}
        self.faq = {}
        self._initialize_help_content()
    
    def _initialize_help_content(self):
        """Initialize help content and documentation"""
        self.help_topics = {
            'getting_started': {
                'title': 'Getting Started',
                'content': '''
                Welcome to the Enhanced Tracker Application!
                
                This application provides comprehensive time tracking, project management, 
                and reporting capabilities with advanced features including:
                
                • Thread-safe operations and error handling
                • Responsive layouts for different screen sizes
                • Comprehensive search functionality
                • Multi-user collaboration features
                • Advanced reporting and analytics
                • Data backup and versioning
                • Input validation and data integrity
                
                To get started:
                1. Configure your preferences in the Settings menu
                2. Set up your projects and tasks
                3. Start tracking your time
                4. Generate reports to analyze your productivity
                ''',
                'keywords': ['start', 'begin', 'setup', 'configuration']
            },
            'time_tracking': {
                'title': 'Time Tracking',
                'content': '''
                Time Tracking Features:
                
                • Start/Stop timers for tasks
                • Manual time entry with validation
                • Automatic time calculation
                • Time categorization and tagging
                • Export time data in multiple formats
                
                Tips:
                • Use keyboard shortcuts for quick timer control
                • Set up automatic reminders for breaks
                • Categorize time entries for better reporting
                • Use the search function to find specific entries
                ''',
                'keywords': ['timer', 'time', 'tracking', 'hours', 'duration']
            },
            'reports': {
                'title': 'Reports and Analytics',
                'content': '''
                Reporting Features:
                
                • Custom report builder with multiple templates
                • Scheduled reports (daily, weekly, monthly)
                • Export to PDF, CSV, HTML, and JSON
                • Data visualization with charts and graphs
                • Advanced filtering and grouping options
                
                Creating Reports:
                1. Go to Reports menu
                2. Select "Create New Report"
                3. Choose fields and filters
                4. Set grouping and sorting options
                5. Generate and export your report
                ''',
                'keywords': ['report', 'analytics', 'export', 'pdf', 'csv']
            },
            'collaboration': {
                'title': 'Collaboration Features',
                'content': '''
                Multi-User Collaboration:
                
                • Real-time user activity tracking
                • Comments and annotations on items
                • Change tracking and audit logs
                • Resource locking to prevent conflicts
                • User presence indicators
                
                Working with Teams:
                • Register users in the system
                • Assign tasks and projects
                • Monitor team activity
                • Review change history
                • Add comments for communication
                ''',
                'keywords': ['collaboration', 'team', 'user', 'comment', 'lock']
            },
            'data_management': {
                'title': 'Data Management',
                'content': '''
                Data Management Features:
                
                • Automatic data backup and restore
                • Data versioning and history
                • Export/import in multiple formats
                • Data integrity validation
                • Scheduled backups
                
                Backup and Restore:
                1. Go to Data menu
                2. Select "Backup Data" for manual backup
                3. Use "Restore Data" to recover from backup
                4. Set up scheduled backups in preferences
                
                Data Export/Import:
                • Supported formats: JSON, CSV, SQL
                • Full data export with all relationships
                • Import validation with error reporting
                ''',
                'keywords': ['backup', 'restore', 'export', 'import', 'data']
            },
            'search': {
                'title': 'Search Functionality',
                'content': '''
                Advanced Search Features:
                
                • Full-text search across all data
                • Multiple search modes: contains, exact, fuzzy
                • Search suggestions and history
                • Keyword-based searching
                • Real-time search results
                
                Search Modes:
                • Contains: Find items containing the search term
                • Exact: Find exact matches
                • Starts With: Find items starting with the term
                • Ends With: Find items ending with the term
                • Keywords: Find items containing all keywords
                • Fuzzy: Find similar matches with typos
                ''',
                'keywords': ['search', 'find', 'query', 'filter']
            },
            'troubleshooting': {
                'title': 'Troubleshooting',
                'content': '''
                Common Issues and Solutions:
                
                Application Crashes:
                • Check the error log in the application directory
                • Restart the application
                • Restore from the most recent backup
                
                Performance Issues:
                • Use virtual scrolling for large datasets
                • Enable background processing
                • Clear old data and backups
                
                Data Issues:
                • Run data validation from the Data menu
                • Check for duplicate entries
                • Verify data integrity
                
                UI Issues:
                • Adjust screen resolution settings
                • Use responsive layout options
                • Reset window geometry in preferences
                ''',
                'keywords': ['error', 'crash', 'performance', 'issue', 'problem']
            }
        }
        
        self.faq = {
            'q1': {
                'question': 'How do I backup my data?',
                'answer': 'Go to Data menu > Backup Data, or set up automatic backups in Preferences > Data Management.'
            },
            'q2': {
                'question': 'Can I work with multiple users?',
                'answer': 'Yes! The application supports multi-user collaboration with real-time activity tracking and change management.'
            },
            'q3': {
                'question': 'What file formats can I export to?',
                'answer': 'You can export to PDF, CSV, HTML, JSON, and SQL formats.'
            },
            'q4': {
                'question': 'How do I create custom reports?',
                'answer': 'Go to Reports menu > Create New Report, then select fields, filters, and formatting options.'
            },
            'q5': {
                'question': 'Is my data secure?',
                'answer': 'Yes, the application includes data validation, integrity checks, and secure backup mechanisms.'
            }
        }
    
    def get_help_topic(self, topic_id: str) -> Dict[str, Any]:
        """Get help content for a specific topic"""
        return self.help_topics.get(topic_id, {})
    
    def search_help(self, query: str) -> List[Dict[str, Any]]:
        """Search help topics by query"""
        results = []
        query_lower = query.lower()
        
        for topic_id, topic_data in self.help_topics.items():
            # Search in title, content, and keywords
            searchable_text = f"{topic_data['title']} {topic_data['content']} {' '.join(topic_data['keywords'])}"
            
            if query_lower in searchable_text.lower():
                results.append({
                    'id': topic_id,
                    'title': topic_data['title'],
                    'relevance': self._calculate_relevance(query_lower, searchable_text.lower())
                })
        
        # Sort by relevance
        results.sort(key=lambda x: x['relevance'], reverse=True)
        return results
    
    def _calculate_relevance(self, query: str, text: str) -> int:
        """Calculate relevance score for search results"""
        score = 0
        
        # Title matches get higher score
        if query in text:
            score += text.count(query) * 2
        
        # Keyword matches
        query_words = query.split()
        for word in query_words:
            if word in text:
                score += 1
        
        return score
    
    def get_faq(self) -> List[Dict[str, str]]:
        """Get frequently asked questions"""
        return list(self.faq.values())
    
    def get_context_help(self, context: str) -> str:
        """Get context-sensitive help based on current UI context"""
        context_help_map = {
            'main_window': 'Use the main menu to access all features. The toolbar provides quick access to common functions.',
            'time_entry': 'Enter time in hours:minutes format (e.g., 8:30 for 8 hours 30 minutes). Use decimal format for partial hours.',
            'project_management': 'Create projects to organize your work. Assign tasks to projects for better tracking and reporting.',
            'report_generation': 'Select the data fields you want to include in your report. Use filters to narrow down the data.',
            'data_export': 'Choose the appropriate format for your needs. CSV is good for analysis, PDF for presentation.',
            'user_management': 'Register users to enable collaboration features. Each user can have different permissions and roles.',
            'search_results': 'Use the search filters to refine your results. Click on any result to view details.',
            'settings': 'Configure application preferences here. Changes are saved automatically.'
        }
        
        return context_help_map.get(context, 'No specific help available for this context.')

class DocumentationGenerator:
    """Generate documentation for the application"""
    
    def __init__(self):
        self.documentation = {}
    
    def generate_user_manual(self) -> str:
        """Generate a comprehensive user manual"""
        manual = """
# Enhanced Tracker Application - User Manual

## Table of Contents
1. Introduction
2. Getting Started
3. Features Overview
4. User Interface Guide
5. Time Tracking
6. Project Management
7. Reporting and Analytics
8. Collaboration Features
9. Data Management
10. Troubleshooting
11. Keyboard Shortcuts
12. FAQ

## 1. Introduction

The Enhanced Tracker Application is a comprehensive time tracking and project management solution designed for individuals and teams. It provides advanced features for productivity tracking, data analysis, and collaboration.

### Key Features
- Thread-safe operations with comprehensive error handling
- Responsive layouts for different screen sizes
- Advanced search and filtering capabilities
- Multi-user collaboration with real-time updates
- Custom report generation with multiple export formats
- Data backup, versioning, and integrity validation
- Input validation and data corruption prevention

## 2. Getting Started

### Installation
1. Ensure Python 3.8+ is installed
2. Install required dependencies: PyQt6, pandas, matplotlib
3. Run the application: python Tracker.py

### Initial Setup
1. Configure user preferences in Settings
2. Set up your first project
3. Create initial tasks
4. Start tracking time

## 3. Features Overview

### Core Features
- **Time Tracking**: Start/stop timers, manual entry, automatic calculation
- **Project Management**: Organize work into projects and tasks
- **Reporting**: Generate custom reports with multiple export options
- **Search**: Advanced search across all data with multiple modes
- **Collaboration**: Multi-user support with real-time updates
- **Data Management**: Backup, restore, versioning, and validation

### Advanced Features
- **Thread Safety**: All operations are thread-safe for stability
- **Error Handling**: Comprehensive error handling prevents crashes
- **Memory Management**: Efficient memory usage with object pooling
- **Responsive Design**: Adapts to different screen sizes
- **Input Validation**: Prevents data corruption with validation rules
- **Performance Optimization**: Virtual scrolling and background processing

## 4. User Interface Guide

### Main Window
The main window contains:
- Menu bar with all application functions
- Toolbar with quick access buttons
- Tabbed interface for different views
- Status bar with progress indicators
- Search bar for quick data access

### Navigation
- Use the menu bar to access all features
- Toolbar provides quick access to common functions
- Tab navigation for switching between views
- Keyboard shortcuts for power users

## 5. Time Tracking

### Starting a Timer
1. Select a task from the task list
2. Click the "Start Timer" button or press Ctrl+T
3. The timer will run in the background
4. Click "Stop Timer" or press Ctrl+T again to stop

### Manual Time Entry
1. Go to Time Entry tab
2. Select date and task
3. Enter hours worked
4. Add description if needed
5. Click "Save Entry"

### Time Validation
- Hours must be between 0 and 999
- Decimal places limited to 2
- Date ranges cannot exceed 1 year
- Required fields must be filled

## 6. Project Management

### Creating Projects
1. Go to Projects tab
2. Click "New Project"
3. Enter project details
4. Set start and end dates
5. Save the project

### Managing Tasks
1. Select a project
2. Click "Add Task"
3. Enter task details
4. Set priority and status
5. Assign to team members (if applicable)

## 7. Reporting and Analytics

### Creating Reports
1. Go to Reports menu
2. Select "Create New Report"
3. Choose data fields to include
4. Set filters and criteria
5. Configure grouping and sorting
6. Generate the report

### Report Templates
- Pre-built templates for common reports
- Custom templates for specific needs
- Scheduled reports for automation
- Multiple export formats

### Export Options
- PDF for presentation
- CSV for analysis and data processing
- HTML for web viewing
- JSON for integration

## 8. Collaboration Features

### User Management
1. Go to Users menu
2. Click "Add User"
3. Enter user details
4. Set permissions and roles
5. Save user information

### Real-time Collaboration
- See who's currently working
- View user activity in real-time
- Lock resources to prevent conflicts
- Track all changes with audit logs

### Comments and Annotations
- Add comments to any item
- Reply to existing comments
- View comment history
- Get notifications for new comments

## 9. Data Management

### Backup and Restore
1. Go to Data menu
2. Select "Backup Data" for manual backup
3. Use "Restore Data" to recover from backup
4. Set up scheduled backups in preferences

### Data Versioning
- Automatic version creation
- Manual version snapshots
- Version comparison
- Rollback to previous versions

### Export and Import
- Export data in multiple formats
- Import data with validation
- Data integrity checking
- Error reporting for imports

## 10. Troubleshooting

### Common Issues
- Application crashes: Check error logs, restart application
- Performance issues: Use virtual scrolling, clear old data
- Data problems: Run validation, check for duplicates
- UI issues: Adjust resolution, reset window geometry

### Error Handling
- Comprehensive error logging
- User-friendly error messages
- Automatic error recovery
- Data corruption prevention

## 11. Keyboard Shortcuts

- Ctrl+N: New item
- Ctrl+S: Save
- Ctrl+O: Open
- Ctrl+F: Search
- Ctrl+T: Start/Stop timer
- F1: Help
- Ctrl+Q: Quit

## 12. FAQ

**Q: How do I backup my data?**
A: Go to Data menu > Backup Data, or set up automatic backups in Preferences.

**Q: Can I work with multiple users?**
A: Yes! The application supports multi-user collaboration with real-time updates.

**Q: What file formats can I export to?**
A: You can export to PDF, CSV, HTML, JSON, and SQL formats.

**Q: How do I create custom reports?**
A: Go to Reports menu > Create New Report, then select fields and options.

**Q: Is my data secure?**
A: Yes, the application includes data validation, integrity checks, and secure backups.

---

For additional support, please refer to the built-in help system or contact your system administrator.
        """
        return manual
    
    def generate_api_documentation(self) -> str:
        """Generate API documentation for developers"""
        api_doc = """
# Enhanced Tracker Application - API Documentation

## Overview
This document describes the internal API and architecture of the Enhanced Tracker Application.

## Core Systems

### Thread Safety System
- `ThreadSafeTimer`: Thread-safe timer implementation
- `ErrorHandler`: Centralized error handling
- `MemoryManager`: Memory management and cleanup

### Input Validation System
- `InputValidator`: Comprehensive input validation
- `DataValidator`: Data integrity checking

### UI Performance System
- `ResponsiveLayoutManager`: Responsive layout management
- `VirtualScrollingTable`: Virtual scrolling for performance
- `BackgroundProcessor`: Background task processing
- `ProgressiveLoader`: Progressive data loading

### Search System
- `SearchEngine`: Advanced search functionality
- Multiple search modes and fuzzy matching

### Collaboration System
- `CollaborationManager`: Multi-user collaboration
- `LockManager`: Resource locking system

### Reporting System
- `ReportBuilder`: Custom report generation
- `ScheduledReportManager`: Automated report scheduling

### Data Management System
- `DataManager`: Data backup, restore, and versioning
- Multiple export/import formats

## Architecture Patterns

### Singleton Pattern
Used for global managers (ErrorHandler, MemoryManager)

### Observer Pattern
Used for UI updates and notifications

### Factory Pattern
Used for creating different types of reports and exports

### Strategy Pattern
Used for different search modes and validation rules

## Error Handling Strategy
- Comprehensive error logging
- User-friendly error messages
- Automatic error recovery
- Data integrity validation

## Performance Optimizations
- Virtual scrolling for large datasets
- Background processing for heavy operations
- Object pooling for memory efficiency
- Progressive loading for better responsiveness

## Security Considerations
- Input validation and sanitization
- Data integrity checking
- Secure backup mechanisms
- Access control for multi-user scenarios
        """
        return api_doc

# Global instances will be created after all classes are defined

# User Preferences System
class UserPreferences:
    """User preferences management system backed by SQLite settings"""
    
    def __init__(self, backend_sqlite_path: str | None = None):
        self.preferences = {
            "auto_save": True,
            "auto_save_interval": 30,  # seconds
            "show_tooltips": True,
            "tooltip_delay": 1000,  # milliseconds
            "theme": "default",
            "font_size": 10,
            "window_geometry": None,
            "last_export_path": "",
            "last_import_path": "",
            "confirm_deletions": True,
            "show_progress_bars": True,
            "enable_sounds": False,
            "backup_count": 5,
            "backup_directory": "",  # Directory for SQLite backups
            "log_level": "INFO"
        }
        self.backend_sqlite_path = backend_sqlite_path
        self.load_preferences()
    
    def _sqlite_connect(self, db_path: str) -> sqlite3.Connection:
        """Create SQLite connection with WAL mode enabled"""
        conn = sqlite3.connect(db_path, timeout=30, isolation_level=None)
        # Enable WAL for better concurrency
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA synchronous=NORMAL")
        conn.execute("PRAGMA cache_size=10000")
        conn.execute("PRAGMA temp_store=MEMORY")
        return conn
    
    
    def _ensure_backend_and_sheet(self):
        """Ensure backend file and Settings sheet exist (Excel only)."""
        try:
            if not self.backend_sqlite_path:
                return None, None
            # Skip Excel handling if using SQLite backend
            if str(self.backend_sqlite_path).lower().endswith((".sqlite", ".db")):
                return None, None
            from openpyxl import load_workbook, Workbook
            if not os.path.exists(self.backend_sqlite_path):
                wb = Workbook()
                try:
                    wb.remove(wb.active)
                except Exception as e:
                    logger.debug(f"Error removing active sheet: {e}")
                ws = wb.create_sheet("Settings")
                ws.cell(row=1, column=1, value="Setting")
                ws.cell(row=1, column=2, value="Value")
                ws.cell(row=1, column=3, value="Description")
                wb.save(self.backend_sqlite_path)
                return wb, ws
            wb = load_workbook(self.backend_sqlite_path)
            if 'Settings' not in wb.sheetnames:
                ws = wb.create_sheet('Settings')
                ws.cell(row=1, column=1, value="Setting")
                ws.cell(row=1, column=2, value="Value")
                ws.cell(row=1, column=3, value="Description")
            else:
                ws = wb['Settings']
            return wb, ws
        except Exception as e:
            logger.warning(f"Failed to ensure backend Settings sheet: {e}")
            return None, None
    
    def load_preferences(self) -> None:
        """Load preferences from SQLite settings."""
        try:
            # SQLite-backed preferences
            if self.backend_sqlite_path and str(self.backend_sqlite_path).lower().endswith((".sqlite", ".db")):
                try:
                    with self._sqlite_connect(str(self.backend_sqlite_path)) as conn:
                        cur = conn.cursor()
                        cur.execute("CREATE TABLE IF NOT EXISTS settings (setting TEXT PRIMARY KEY, value TEXT, description TEXT)")
                        cur.execute("SELECT setting, value FROM settings")
                        rows = cur.fetchall()
                    loaded = {}
                    for key, raw_val in rows:
                        if key is None:
                            continue
                        if raw_val is None:
                            val = None
                        else:
                            try:
                                val = json.loads(raw_val)
                            except Exception:
                                sval = str(raw_val)
                                if sval.lower() in ("true", "false"):
                                    val = sval.lower() == "true"
                                else:
                                    try:
                                        val = int(sval)
                                    except Exception:
                                        try:
                                            val = float(sval)
                                        except Exception:
                                            val = sval
                        loaded[str(key)] = val
                    self.preferences.update({k: loaded.get(k, v) for k, v in self.preferences.items()})
                except Exception as e:
                    logger.warning(f"Failed to load preferences from SQLite settings: {e}")
                return
            # No backend path specified, use defaults
        except Exception as e:
            logger.warning(f"Failed to load preferences: {e}")
    
    def save_preferences(self) -> None:
        """Save preferences to SQLite settings."""
        try:
            # SQLite-backed preferences
            if self.backend_sqlite_path and str(self.backend_sqlite_path).lower().endswith((".sqlite", ".db")):
                try:
                    with self._sqlite_connect(str(self.backend_sqlite_path)) as conn:
                        cur = conn.cursor()
                        cur.execute("CREATE TABLE IF NOT EXISTS settings (setting TEXT PRIMARY KEY, value TEXT, description TEXT)")
                        for key, value in self.preferences.items():
                            try:
                                sval = json.dumps(value)
                            except Exception:
                                sval = str(value)
                            cur.execute("INSERT OR REPLACE INTO settings(setting, value, description) VALUES(?,?,?)", (key, sval, None))
                except Exception as e:
                    logger.error(f"Failed to save preferences to SQLite settings: {e}")
                return
            # SQLite-backed preferences only
                return
        except Exception as e:
            logger.error(f"Failed to save preferences to Settings: {e}")
    
    def get(self, key: str, default=None):
        return self.preferences.get(key, default)
    
    def set(self, key: str, value) -> None:
        self.preferences[key] = value
        self.save_preferences()
    
    def reset_to_defaults(self) -> None:
        self.preferences = {
            "auto_save": True,
            "auto_save_interval": 30,
            "show_tooltips": True,
            "tooltip_delay": 1000,
            "theme": "default",
            "font_size": 10,
            "window_geometry": None,
            "last_export_path": "",
            "last_import_path": "",
            "confirm_deletions": True,
            "show_progress_bars": True,
            "enable_sounds": False,
            "backup_count": 5,
            "backup_directory": "",  # Directory for SQLite backups
            "log_level": "INFO"
        }
        self.save_preferences()

# Progress Indicator System
class ProgressManager:
    """Centralized progress indicator management"""
    
    def __init__(self, main_window):
        self.main_window = main_window
        self.progress_bar = None
        self.status_label = None
        self.operation_timer = None
    
    def show_progress(self, message: str = "Processing...", maximum: int = 100) -> None:
        """Show progress bar with message"""
        if not self.main_window:
            return
        
        # Create progress bar if it doesn't exist
        if not self.progress_bar:
            self.progress_bar = QProgressBar()
            self.progress_bar.setVisible(False)
            self.main_window.statusBar().addPermanentWidget(self.progress_bar)
        
        # Create status label if it doesn't exist
        if not self.status_label:
            self.status_label = QLabel()
            self.main_window.statusBar().addPermanentWidget(self.status_label)
        
        self.progress_bar.setMaximum(maximum)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(True)
        self.status_label.setText(message)
        self.status_label.setVisible(True)
        
        # Auto-hide after 30 seconds if not updated
        if self.operation_timer:
            self.operation_timer.stop()
        self.operation_timer = QTimer()
        self.operation_timer.timeout.connect(self.hide_progress)
        self.operation_timer.setSingleShot(True)
        self.operation_timer.start(30000)
    
    def update_progress(self, value: int, message: str = None) -> None:
        """Update progress bar value and message"""
        if self.progress_bar and self.progress_bar.isVisible():
            self.progress_bar.setValue(value)
            if message and self.status_label:
                self.status_label.setText(message)
    
    def hide_progress(self) -> None:
        """Hide progress bar and status label"""
        if self.progress_bar:
            self.progress_bar.setVisible(False)
        if self.status_label:
            self.status_label.setVisible(False)
        if self.operation_timer:
            self.operation_timer.stop()

# Undo/Redo System
class UndoRedoManager:
    """Centralized undo/redo management system"""
    
    def __init__(self, max_history: int = 50):
        self.undo_stack = []
        self.redo_stack = []
        self.max_history = max_history
        self.current_operation = None
    
    def add_operation(self, operation_type: str, pane_name: str, row: int, col: int, 
                     old_value: str, new_value: str, description: str = "") -> None:
        """Add operation to undo stack"""
        operation = {
            "type": operation_type,
            "pane_name": pane_name,
            "row": row,
            "col": col,
            "old_value": old_value,
            "new_value": new_value,
            "description": description,
            "timestamp": datetime.now()
        }
        
        # Clear redo stack when new operation is added
        self.redo_stack.clear()
        
        # Add to undo stack
        self.undo_stack.append(operation)
        
        # Limit stack size
        if len(self.undo_stack) > self.max_history:
            self.undo_stack.pop(0)
    
    def can_undo(self) -> bool:
        """Check if undo is possible"""
        return len(self.undo_stack) > 0
    
    def can_redo(self) -> bool:
        """Check if redo is possible"""
        return len(self.redo_stack) > 0
    
    def get_undo_description(self) -> str:
        """Get description of next undo operation"""
        if self.undo_stack:
            return self.undo_stack[-1].get("description", "Undo")
        return ""
    
    def get_redo_description(self) -> str:
        """Get description of next redo operation"""
        if self.redo_stack:
            return self.redo_stack[-1].get("description", "Redo")
        return ""
    
    def clear_history(self) -> None:
        """Clear all undo/redo history"""
        self.undo_stack.clear()
        self.redo_stack.clear()

# Centralized Tooltip System
class TooltipManager:
    """Centralized tooltip management system"""
    
    @staticmethod
    def set_tooltip(widget, text: str, delay: int = 1000) -> None:
        """Set tooltip with consistent styling"""
        if widget and text:
            widget.setToolTip(text)
            # Only set duration for widgets that support it
            if hasattr(widget, 'setToolTipDuration'):
                widget.setToolTipDuration(delay)
    
    @staticmethod
    def get_common_tooltips() -> dict:
        """Get common tooltip texts"""
        return {
            "save": "Save current data (Ctrl+S)",
            "load": "Load data from file (Ctrl+O)",
            "export": "Export data to Excel (Ctrl+E)",
            "add_row": "Add new row (Ctrl+N)",
            "delete_row": "Delete selected row (Delete)",
            "edit_cell": "Double-click to edit cell",
            "sort_column": "Click header to sort column",
            "filter": "Use filter to narrow down results",
            "refresh": "Refresh data (F5)",
            "undo": "Undo last action (Ctrl+Z)",
            "redo": "Redo last action (Ctrl+Y)",
            "search": "Search across all data (Ctrl+K)",
            "help": "Get help and shortcuts (F1)",
            "close": "Close dialog (Esc)",
            "ok": "Confirm action (Enter)",
            "cancel": "Cancel action (Esc)"
        }
# Reusable controls bar for pane tables
class ControlsBar(QWidget):
    def __init__(self, pane_name: str, parent: QWidget | None = None) -> None:
        super().__init__(parent)
        lay = QHBoxLayout(self)
        lay.setContentsMargins(0, 0, 0, 0)
        lay.setSpacing(8)

        self.add_btn_primary = QPushButton("+ Add"); self.add_btn_primary.setObjectName("primary")
        self.btn_del = QPushButton("Delete Row"); self.btn_del.setObjectName("secondary")
        self.btn_imp = QPushButton("Import"); self.btn_imp.setObjectName("secondary")
        self.btn_exp = QPushButton("Export"); self.btn_exp.setObjectName("secondary")
        self.btn_export_view = QPushButton("Export Current View"); self.btn_export_view.setObjectName("secondary")
        self.btn_sample = QPushButton("📋 Sample"); self.btn_sample.setObjectName("secondary")

        for w in (self.add_btn_primary, self.btn_del, self.btn_imp, self.btn_exp, self.btn_export_view, self.btn_sample):
            lay.addWidget(w)

        self.search = QLineEdit(); self.search.setPlaceholderText(f"Search {pane_name}...")
        self.btn_search = QPushButton("Search"); self.btn_search.setObjectName("secondary")
        self.btn_clear = QPushButton("Clear"); self.btn_clear.setObjectName("secondary")

        lay.addStretch(1)
        lay.addWidget(self.search)
        lay.addWidget(self.btn_search)
        lay.addWidget(self.btn_clear)

        # Status & RAG filters per pane
        self.status_combo = QComboBox()
        try:
            self.status_combo.addItems(["All"] + STATUS_OPTIONS)
        except Exception as e:
            logger.debug(f"Error adding status options: {e}")
            self.status_combo.addItems(["All"])  # fallback
        self.rag_combo = QComboBox()
        try:
            self.rag_combo.addItems(["All"] + RAG_OPTIONS)
        except Exception as e:
            logger.debug(f"Error adding RAG options: {e}")
            self.rag_combo.addItems(["All"])  # fallback
        lay.addWidget(QLabel("Status:"))
        lay.addWidget(self.status_combo)
        lay.addWidget(QLabel("RAG:"))
        lay.addWidget(self.rag_combo)
# Centralized Error Handling System
class ErrorHandler:
    """Centralized error handling and logging system"""
    
    _main_window: Optional[QMainWindow] = None  # Reference to main window for UI updates
    
    @classmethod
    def set_main_window(cls, main_window: QMainWindow) -> None:
        """Set reference to main window for UI updates"""
        cls._main_window = main_window
    
    @classmethod
    def _show_toast(cls, message: str, level: str = "ERROR") -> None:
        """Show toast notification if main window is available"""
        if cls._main_window and hasattr(cls._main_window, '_show_toast'):
            try:
                cls._main_window._show_toast(message, level=level)
            except Exception as e:
                logger.error(f"Failed to show toast: {e}")
    
    @classmethod
    def _update_status_bar(cls, message: str) -> None:
        """Update status bar if available"""
        if cls._main_window and hasattr(cls._main_window, 'statusBar'):
            try:
                cls._main_window.statusBar().showMessage(message, 5000)  # 5 second timeout
            except Exception as e:
                logger.error(f"Failed to update status bar: {e}")
    
        
    
    @classmethod
    def handle_ui_error(cls, operation: str, error: Exception, context: Optional[Dict[str, Any]] = None) -> None:
        """Handle UI-related errors"""
        error_msg = f"UI operation failed: {operation}"
        logger.error(f"{error_msg} - {type(error).__name__}: {str(error)}", exc_info=True)
        
        if context:
            logger.debug(f"Error context: {context}")
        
        # Show toast notification for UI errors
        cls._show_toast(f"UI Error: {operation}", "ERROR")
        cls._update_status_bar(f"UI operation failed: {operation}")
    
    @classmethod
    def handle_data_error(cls, operation: str, error: Exception, context: Optional[Dict[str, Any]] = None) -> None:
        """Handle data processing errors"""
        error_msg = f"Data operation failed: {operation}"
        logger.error(f"{error_msg} - {type(error).__name__}: {str(error)}", exc_info=True)
        
        if context:
            logger.debug(f"Error context: {context}")
        
        # Show toast notification for data errors
        cls._show_toast(f"Data Error: {operation}", "ERROR")
        cls._update_status_bar(f"Data operation failed: {operation}")
    
    @classmethod
    def handle_validation_error(cls, field: str, value: str, error: str) -> None:
        """Handle validation errors"""
        error_msg = f"Validation failed for field '{field}': {error}"
        logger.warning(f"{error_msg} - Value: '{value}'")
        
        # Show user-friendly validation error
        try:
            from PyQt6.QtWidgets import QMessageBox
            QMessageBox.warning(None, "Validation Error", 
                              f"Invalid input for {field}:\n{error}\n\nPlease correct and try again.")
        except Exception as e:
            logger.error(f"Failed to show validation dialog: {e}")
        
        # Show toast notification
        cls._show_toast(f"Validation Error: {field}", "WARNING")
        cls._update_status_bar(f"Invalid input for {field}")
    
    @classmethod
    def handle_success(cls, operation: str, details: Optional[str] = None) -> None:
        """Handle successful operations with user feedback"""
        success_msg = f"Operation completed: {operation}"
        logger.info(f"{success_msg} - {details}" if details else success_msg)
        
        # Show success toast
        cls._show_toast(f"Success: {operation}", "SUCCESS")
        cls._update_status_bar(f"Successfully completed: {operation}")
    
    @classmethod
    def handle_warning(cls, operation: str, warning: str) -> None:
        """Handle warnings with user feedback"""
        warning_msg = f"Warning in {operation}: {warning}"
        logger.warning(warning_msg)
        
        # Show warning toast
        cls._show_toast(f"Warning: {operation}", "WARNING")
        cls._update_status_bar(f"Warning: {operation}")
    
    @classmethod
    def handle_info(cls, operation: str, info: str) -> None:
        """Handle informational messages"""
        info_msg = f"Info: {operation} - {info}"
        logger.info(info_msg)
        
        # Show info toast
        cls._show_toast(f"Info: {operation}", "INFO")
        cls._update_status_bar(f"Info: {operation}")
    
    @classmethod
    def log_operation(cls, operation: str, success: bool, details: Optional[str] = None) -> None:
        """Log operation results"""
        level = logging.INFO if success else logging.ERROR
        message = f"Operation {'completed' if success else 'failed'}: {operation}"
        if details:
            message += f" - {details}"
        logger.log(level, message)
        
        # Update status bar
        if success:
            cls._update_status_bar(f"Completed: {operation}")
        else:
            cls._update_status_bar(f"Failed: {operation}")

# Modern configuration constants
DEFAULT_DATE_FORMAT = "%Y-%m-%d"
DEFAULT_DATETIME_FORMAT = "%Y-%m-%d %H:%M:%S"
MAX_TEXT_LENGTH = 255
MAX_HOURS_PER_DAY = 999
MIN_YEAR = 1900
MAX_FUTURE_YEARS = 10

# Enhanced Input Validation System
class InputValidator:
    """Comprehensive input validation system"""
    
    @staticmethod
    def validate_date(date_str: str, field_name: str = "Date") -> tuple[bool, str]:
        """Validate date format and range"""
        if not date_str or not str(date_str).strip():
            return True, ""  # Empty dates are allowed
        
        try:
            date_obj = datetime.strptime(str(date_str), "%Y-%m-%d")
            # Check if date is reasonable (not too far in past or future)
            today = datetime.now()
            if date_obj.year < 1900 or date_obj.year > today.year + 10:
                return False, f"Date must be between 1900 and {today.year + 10}"
            return True, ""
        except ValueError:
            return False, f"Date must be in YYYY-MM-DD format"
    
    @staticmethod
    def validate_due_date(date_str: str, field_name: str = "Due Date") -> tuple[bool, str]:
        """Validate due date format and ensure it's not in the past"""
        if not date_str or not str(date_str).strip():
            return True, ""  # Empty dates are allowed
        
        try:
            date_obj = datetime.strptime(str(date_str), "%Y-%m-%d")
            today = datetime.now().date()
            
            # Check if date is reasonable (not too far in past or future)
            if date_obj.year < 1900 or date_obj.year > today.year + 10:
                return False, f"Date must be between 1900 and {today.year + 10}"
            
            # Check if due date is not in the past
            if date_obj.date() < today:
                return False, f"{field_name} cannot be in the past. Please select today's date or a future date."
            
            return True, ""
        except ValueError:
            return False, f"Date must be in YYYY-MM-DD format"
    
    @staticmethod
    def validate_rag_status(value: str) -> Tuple[bool, str]:
        """Validate RAG status values using modern enums."""
        valid_values = [rag.value for rag in RAGStatus] + ["NA", ""]
        if str(value).strip() in valid_values:
            return True, ""
        return False, f"RAG status must be one of: {', '.join(valid_values)}"
    
    @staticmethod
    def validate_status(value: str) -> Tuple[bool, str]:
        """Validate status values using modern enums."""
        valid_values = [status.value for status in Status] + [""]
        if str(value).strip() in valid_values:
            return True, ""
        return False, f"Status must be one of: {', '.join(valid_values)}"
    
    @staticmethod
    def validate_priority(value: str) -> Tuple[bool, str]:
        """Validate priority values using modern enums."""
        valid_values = [priority.value for priority in Priority] + ["NA", ""]
        if str(value).strip() in valid_values:
            return True, ""
        return False, f"Priority must be one of: {', '.join(valid_values)}"
    
    @staticmethod
    def validate_hours(value: str) -> Tuple[bool, str]:
        """Validate hours values using modern constants."""
        if not value or not str(value).strip():
            return True, ""  # Empty hours are allowed
        
        try:
            hours = float(value)
            if hours < 0 or hours > MAX_HOURS_PER_DAY:
                return False, f"Hours must be between 0 and {MAX_HOURS_PER_DAY}"
            return True, ""
        except ValueError:
            return False, "Hours must be a valid number"
    
    @staticmethod
    def validate_email(email: str) -> Tuple[bool, str]:
        """Validate email format using modern regex patterns."""
        if not email or not str(email).strip():
            return True, ""  # Empty emails are allowed
        
        import re
        # Modern email regex pattern
        pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if re.match(pattern, str(email).strip()):
            return True, ""
        return False, "Invalid email format"
    
    @staticmethod
    def validate_required_field(value: str, field_name: str) -> Tuple[bool, str]:
        """Validate required fields with modern error messages."""
        if not value or not str(value).strip():
            return False, f"{field_name} is required"
        return True, ""
    
    @staticmethod
    def validate_text_length(value: str, field_name: str, max_length: int = MAX_TEXT_LENGTH) -> Tuple[bool, str]:
        """Validate text length using modern constants."""
        if len(str(value)) > max_length:
            return False, f"{field_name} must be {max_length} characters or less"
        return True, ""
    
    @staticmethod
    def validate_project_id(project_id: str) -> Tuple[bool, str]:
        """Validate project ID format using modern regex patterns."""
        if not project_id or not str(project_id).strip():
            return True, ""  # Empty project IDs are allowed
        
        # Project ID should be alphanumeric with possible hyphens/underscores
        import re
        if re.match(r'^[A-Za-z0-9_-]+$', str(project_id).strip()):
            return True, ""
        return False, "Project ID must contain only letters, numbers, hyphens, and underscores"
    
    @staticmethod
    def validate_enum_value(value: str, enum_class: type, field_name: str) -> Tuple[bool, str]:
        """Generic validator for enum values."""
        if not value or not str(value).strip():
            return True, ""  # Empty values are allowed
        
        valid_values = [item.value for item in enum_class]
        if str(value).strip() in valid_values:
            return True, ""
        return False, f"{field_name} must be one of: {', '.join(valid_values)}"

# Modern Enums for better type safety and maintainability
class ComplexityLevel(Enum):
    """Project complexity levels"""
    LOW = "Low"
    MEDIUM = "Medium"
    HIGH = "High"

class ConnectivityType(Enum):
    """Network connectivity options"""
    FACILITY_INTERNET = "Facility Internet"
    FIRST_MILE_MPLS_LAST_MILE_INTERNET = "First Mile MPLS + Last Mile Internet"
    FIRST_MILE_MPLS_LAST_MILE_S2S_TUNNEL = "First Mile MPLS + Last Mile S2S Tunnel Over Internet"
    FIRST_MILE_MPLS_LAST_MILE_CLIENT_MANAGED = "First Mile MPLS + Last Mile Client Managed"
    S2S_TUNNEL_OVER_INTERNET = "S2S Tunnel Over Internet"
    CLIENT_PROVIDED_P2P = "Client-Provided Point-to-Point Link"
    ZTNA_OVER_INTERNET = "Zero Trust Network Access (ZTNA) over Internet"
    C2S_OVER_INTERNET = "C2S over Internet (VDI Solutions)"
    CLIENT_PROVIDED_C2S_VPN = "Client Provided C2S VPN"
    END_TO_END_MPLS = "End-to-End MPLS"
    CLIENT_BASED_VPN = "Client-Based VPN (Remote Access VPN)"
    SITE_TO_SITE_VPN = "Site-to-Site VPN Tunnel (S2S)"

class VoiceSolution(Enum):
    """Voice solution options"""
    YES = "Yes"
    NO = "No"

class ContactCenter(Enum):
    """Contact center options"""
    AMAZON_CONNECT = "Amazon Connect"
    CISCO_WEBEX = "Cisco WebEx"
    NICEX = "NiceX"
    GENESYS_CLOUD = "Genesys Cloud"
    ZOOM_CONTACT_CENTER = "Zoom Contact Center"

# Legacy constants for backward compatibility
COMPLEXITY_OPTIONS = [level.value for level in ComplexityLevel]
CONNECTIVITY_OPTIONS = [conn.value for conn in ConnectivityType]
VOICE_SOLUTION_OPTIONS = [voice.value for voice in VoiceSolution]
CONTACT_CENTER_OPTIONS = [center.value for center in ContactCenter]

# Modern data structures using TypedDict for better type safety
class PaneColumns(TypedDict):
    """Type definition for pane column configurations"""
    Initiatives: List[str]
    Potential_Issues: List[str]  # Note: This is the TypedDict key, but the actual data uses "Potential Issues"
    Activities: List[str]
    Client_Visits_Audits: List[str]  # Note: This is the TypedDict key, but the actual data uses "Client Visits / Audits"
    Accolades: List[str]
    Leave_Tracker: List[str]
    Project_Details: List[str]  # Note: This is the TypedDict key, but the actual data uses "Project Details"

class ProjectData(TypedDict):
    """Type definition for project data structure"""
    project_name: str
    project_id: str
    location: str
    complexity: str
    primary_ism: str
    secondary_ism: str
    ism_hours: str
    cxl_name: str
    sdl_name: str
    connectivity_type: str
    audits_in_deal: str
    voice_solution: str
    contact_center: str

class TaskData(TypedDict):
    """Type definition for task data structure"""
    action: str
    description: str
    tracker: str
    ownership: str
    remarks: str
    start_date: str
    end_date: str
    status: str
    priority: str
    efforts: str
    additional_remarks: str

class ValidationResult(TypedDict):
    """Type definition for validation results"""
    is_valid: bool
    error_message: str
    field_name: str
    value: str

class ChangeLogEntry(TypedDict):
    """Type definition for change log entries"""
    timestamp: str
    user: str
    action: str
    details: str
    pane: str
    row: int
    column: int
    old_value: str
    new_value: str

# Modern enums for all options
class Status(Enum):
    """Task status options"""
    NA = "NA"
    YET_TO_START = "Yet to Start"
    IN_PROGRESS = "In Progress"
    COMPLETED = "Completed"

class Priority(Enum):
    """Priority levels"""
    CRITICAL = "Critical"
    HIGH = "High"
    MEDIUM = "Medium"
    LOW = "Low"

class RAGStatus(Enum):
    """RAG (Red-Amber-Green) status"""
    RED = "Red"
    AMBER = "Amber"
    GREEN = "Green"

class TrackerType(Enum):
    """Tracker types"""
    ADHOC = "Adhoc"
    GLOBAL_INITIATIVE = "Global Initiative"
    OTHERS = "Others"

class AuditStatus(Enum):
    """Audit status options"""
    YET_TO_START = "Yet to Start"
    IN_PROGRESS = "In Progress"
    COMPLETED_WITH_FINDINGS = "Completed - With Findings"
    COMPLETED_WITHOUT_FINDINGS = "Completed - Without Findings"

class LeaveType(Enum):
    """Leave type options"""
    ANNUAL = "Annual"
    SICK = "Sick"
    PERSONAL = "Personal"
    EMERGENCY = "Emergency"
    MATERNITY = "Maternity"
    PATERNITY = "Paternity"

class DurationType(Enum):
    """Duration type options"""
    FULL_DAY = "Full Day"
    HALF_DAY = "Half Day"
    HOURLY = "Hourly"

class TaskType(Enum):
    """Task type options"""
    TECHNICAL = "Technical"
    ADMINISTRATIVE = "Administrative"
    MEETING = "Meeting"
    TRAINING = "Training"
    DOCUMENTATION = "Documentation"
    TESTING = "Testing"
    DEPLOYMENT = "Deployment"

class AuditType(Enum):
    """Audit type options"""
    SECURITY = "Security"
    COMPLIANCE = "Compliance"
    PERFORMANCE = "Performance"
    QUALITY = "Quality"
    FINANCIAL = "Financial"
    OPERATIONAL = "Operational"

# Modern pane columns configuration
PANE_COLUMNS: PaneColumns = {
    "Initiatives": [
        "Action", "Description", "Tracker", "Ownership", "Remarks",
        "Start Date", "End Date", "Status", "Priority", "Efforts", "Additional Remarks"
    ],
    "Potential Issues": [
        "Task Type", "Status", "Created Date", "Due Date", "Project Name",
        "Project ID", "RAG Status", "ISM Name", "Action Owner", "Ageing", "Ageing RAG",
        "Added in PI Tool", "Leads Attention Required", "Description", "Updates"
    ],
    "Activities": [
        "ISM Name", "Project Name", "Activity/Issue", "Start Date",
        "Target Date", "Support Required", "Status", "RAG", "Brief Update"
    ],
    "Client Visits / Audits": [
        "Project Name", "Audit Scope", "Audit Type", "ISM Name",
        "Audit Start Date", "Audit End Date", "Status", "RAG Status",
        "Detailed Action Plan/Status", "Remarks"
    ],
    "Accolades": [
        "Month", "ISMT Enterprise ID", "Appreciator Enterprise ID",
        "Appreciator Designation", "Project Name", "Accolades Description", "Fortnight"
    ],
    "Leave Tracker": [
        "Date", "Type", "Duration", "Description", "ISM Name",
        "Approval Status", "Approver Enterprise ID", "Approver Name",
        "Requested By Enterprise ID", "Requested By Name", "Decision Date", "Approval Comments"
    ],
    "Project Details": [
        "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM",
        "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type",
        "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
    ],
}

# Legacy constants for backward compatibility - using enum values
STATUS_OPTIONS = [status.value for status in Status]
PRIORITY_OPTIONS = [priority.value for priority in Priority]
RAG_OPTIONS = [rag.value for rag in RAGStatus]
PRIORITY_FULL = [priority.value for priority in Priority]
TRACKER_OPTIONS = [tracker.value for tracker in TrackerType]
PI_TOOL_OPTIONS = ["Yes", "No", "NA"]
AUDIT_STATUS_OPTIONS = [status.value for status in AuditStatus]
AUDITS_IN_DEAL_OPTIONS = ["Yes", "No"]

# Additional modern constants using enums
LEAVE_TYPE_OPTIONS = [leave_type.value for leave_type in LeaveType]
DURATION_TYPE_OPTIONS = [duration.value for duration in DurationType]
TASK_TYPE_OPTIONS = [task_type.value for task_type in TaskType]
AUDIT_TYPE_OPTIONS = [audit_type.value for audit_type in AuditType]

# Remove duplicate constants - already defined above


def default_row_for_columns(columns: List[str]) -> List[str]:
    """Generate default values for a row based on column types.
    
    Args:
        columns: List of column names
        
    Returns:
        List of default values corresponding to each column
    """
    values: List[str] = []
    for col in columns:
        if "Date" in col:
            values.append(datetime.today().strftime("%Y-%m-%d"))
        elif "RAG" in col:
            values.append(RAGStatus.GREEN.value)
        elif col == "Status":
            values.append(Status.NA.value)
        else:
            values.append("")
    return values


class AnimatedLoadButton(QPushButton):
    """Modern animated button with pulse effect and proper type hints."""
    
    def __init__(self, text: str = "Load Data", parent: Optional[QWidget] = None) -> None:
        super().__init__(text, parent)
        self.setObjectName("animatedLoadButton")
        self._setup_styling()
        self._setup_animations()
        self.start_pulse()
    
    def _setup_styling(self) -> None:
        """Configure button styling with modern CSS."""
        self.setStyleSheet("""
            QPushButton#animatedLoadButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #4CAF50, stop:0.5 #45a049, stop:1 #4CAF50);
                border: none;
                border-radius: 20px;
                color: white;
                font-weight: bold;
                font-size: 12px;
                padding: 8px 16px;
                min-width: 100px;
            }
            QPushButton#animatedLoadButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #45a049, stop:0.5 #4CAF50, stop:1 #45a049);
            }
            QPushButton#animatedLoadButton:pressed {
                background: #3d8b40;
            }
        """)
    
    def _setup_animations(self) -> None:
        """Initialize animation properties."""
        # Main animation
        self.animation = QPropertyAnimation(self, b"geometry")
        self.animation.setDuration(1000)
        self.animation.setEasingCurve(QEasingCurve.Type.InOutQuad)
        
        # Pulse animation
        self.pulse_animation = QPropertyAnimation(self, b"geometry")
        self.pulse_animation.setDuration(1500)
        self.pulse_animation.setEasingCurve(QEasingCurve.Type.InOutSine)
        self.pulse_animation.setLoopCount(-1)  # Infinite loop
    
    def start_pulse(self) -> None:
        """Start the pulsing animation."""
        if self.pulse_animation.state() == QAbstractAnimation.State.Running:
            return
            
        current_rect = self.geometry()
        # Create a subtle pulse effect
        pulse_rect = current_rect.adjusted(-2, -1, 2, 1)
        
        self.pulse_animation.setStartValue(current_rect)
        self.pulse_animation.setEndValue(pulse_rect)
        self.pulse_animation.start()
    
    def stop_pulse(self):
        """Stop the pulsing animation"""
        self.pulse_animation.stop()
    
    def animate_click(self):
        """Animate button click"""
        current_rect = self.geometry()
        click_rect = current_rect.adjusted(1, 1, -1, -1)
        
        self.animation.setStartValue(current_rect)
        self.animation.setEndValue(click_rect)
        self.animation.finished.connect(lambda: self.animation.setEndValue(current_rect))
        self.animation.start()

class DatePickerDialog(QDialog):
    def __init__(self, parent=None, initial_date: QDate | None = None):
        super().__init__(parent)
        self.setWindowTitle("Select Date")
        layout = QVBoxLayout(self)
        self.calendar = QCalendarWidget()
        if isinstance(initial_date, QDate):
            self.calendar.setSelectedDate(initial_date)
        layout.addWidget(self.calendar)
        btn = QPushButton("Select")
        btn.clicked.connect(self.accept)
        layout.addWidget(btn)

    def selected_date(self) -> str:
        return self.calendar.selectedDate().toString("yyyy-MM-dd")
class LeaveCalendar(QCalendarWidget):
    def __init__(self, main_window):
        super().__init__()
        self.main = main_window
        self.setMouseTracking(True)
        self._cell_rects: dict[str, QRect] = {}
        # Set today's date as selected by default
        today = QDate.currentDate()
        self.setSelectedDate(today)
        # Track the page (month/year) currently displayed, not just the selected date
        try:
            self._page_year = today.year()
            self._page_month = today.month()
        except Exception as e:
            logger.debug(f"Error setting calendar page: {e}")
            self._page_year, self._page_month = QDate.currentDate().year(), QDate.currentDate().month()
        try:
            self.currentPageChanged.connect(self._on_page_changed)
        except Exception as e:
            logger.warning(f"Failed to connect calendar page change signal: {e}")

    def _on_page_changed(self, year: int, month: int) -> None:
        self._page_year = int(year)
        self._page_month = int(month)
        self._cell_rects.clear()
        try:
            # Some Qt builds expose no viewport() on QCalendarWidget; use standard repaint
            self.update()
        except Exception as e:
            logger.warning(f"Failed to update calendar viewport: {e}")

    def _entries_for_date(self, qdate: QDate) -> list[tuple[str, str, str, str]]:
        rows = self.main.data.get("Leave Tracker", [])
        cols = PANE_COLUMNS["Leave Tracker"]
        didx = cols.index("Date"); tidx = cols.index("Type"); duridx = cols.index("Duration"); descidx = cols.index("Description"); ismidx = cols.index("ISM Name")
        ds = qdate.toString("yyyy-MM-dd")
        out = []
        sel_ism = None
        try:
            sel_ism = self.main.leave_ism_combo.currentText()
        except Exception as e:
            logger.warning(f"Failed to get selected ISM from combo: {e}")
            sel_ism = "All ISMs"
        for r in rows:
            if didx < len(r) and r[didx] == ds:
                name = r[ismidx] if ismidx < len(r) else ""
                # Apply ISM filter unless All ISMs is selected
                if sel_ism and sel_ism != "All ISMs" and name != sel_ism:
                    continue
                typ = r[tidx] if tidx < len(r) else ""
                dur = r[duridx] if duridx < len(r) else "Full Day"
                desc = r[descidx] if descidx < len(r) else ""
                out.append((str(name), str(typ), str(dur), str(desc)))
        return out

    def paintCell(self, painter, rect, date):
        # Hide only next-month trailing days based on the page month/year
        try:
            page_year = getattr(self, "_page_year", self.selectedDate().year())
            page_month = getattr(self, "_page_month", self.selectedDate().month())
            first_day = QDate(page_year, page_month, 1)
            last_day = first_day.addMonths(1).addDays(-1)
            # Clear hover map when we hit the first day of the page month
            if date == first_day:
                self._cell_rects.clear()
            if date > last_day:
                from PyQt6.QtGui import QColor, QBrush
                painter.save()
                painter.fillRect(rect.adjusted(1, 1, -1, -1), QBrush(QColor("#ffffff")))
                painter.restore()
                return
        except Exception as e:
            ErrorHandler.handle_ui_error("paint cell", e)
        super().paintCell(painter, rect, date)
        # Gray out weekends
        try:
            if date.dayOfWeek() in (6, 7):
                from PyQt6.QtGui import QColor, QBrush
                painter.save()
                shade = QColor(243, 244, 246, 120)  # translucent gray so date remains visible
                painter.fillRect(rect.adjusted(1, 1, -1, -1), QBrush(shade))
                painter.restore()
        except Exception as e:
            ErrorHandler.handle_ui_error("paint weekend", e)

        # Draw ISM avatars at the bottom-right of the cell
        try:
            # Skip trailing next-month cells
            page_year = getattr(self, "_page_year", self.selectedDate().year())
            page_month = getattr(self, "_page_month", self.selectedDate().month())
            first_day = QDate(page_year, page_month, 1)
            last_day = first_day.addMonths(1).addDays(-1)
            if date > last_day:
                return
            entries = self._entries_for_date(date)
            # unique names to avoid duplicates in same day
            names = []
            for nm, _, _, _ in entries:
                if nm and nm not in names:
                    names.append(nm)
            count = len(names)
            if count > 0:
                # Compute dynamic size; wrap to second row if needed
                available_w = max(0, rect.width() - 6)
                spacing = 2
                min_size = 6
                # Try single row first
                size = max(min_size, min(18, int((available_w - spacing * (count - 1)) / max(1, count))))
                total_w = count * size + (count - 1) * spacing
                rows_needed = 1
                if total_w > available_w and count > 1:
                    # Wrap into 2 rows as needed
                    rows_needed = 2
                    per_row = (available_w + spacing) // (min_size + spacing)
                    per_row = max(1, int(per_row))
                    rows_needed = min(2, ((count - 1) // per_row) + 1)
                    # Recompute per-row size for better fit
                    per_row_count = min(count, int(per_row))
                    size = max(min_size, min(16, int((available_w - spacing * (per_row_count - 1)) / max(1, per_row_count))))
                from PyQt6.QtGui import QPainter
                y_bottom = rect.bottom() - size - 2
                x_right = rect.right() - 2
                cur_row = 0
                idx = 0
                while idx < count and cur_row < rows_needed:
                    remaining = count - idx
                    # Items in this row
                    per_row_count = remaining if rows_needed == 1 else min(remaining, max(1, (available_w + spacing) // (size + spacing)))
                    row_w = per_row_count * size + (per_row_count - 1) * spacing
                    start_x = x_right - row_w
                    y = y_bottom - cur_row * (size + 2)
                    for j in range(per_row_count):
                        nm = names[idx]
                        pix = self.main._default_avatar(nm, size)
                        painter.drawPixmap(start_x + j * (size + spacing), y, pix)
                        idx += 1
                    cur_row += 1
            # Cache rect by date string for hover lookup
            self._cell_rects[date.toString("yyyy-MM-dd")] = QRect(rect)
        except Exception as e:
            ErrorHandler.handle_ui_error("paint avatars", e)

    def mouseMoveEvent(self, event):
        try:
            pos = event.position().toPoint()
        except Exception as e:
            logger.debug(f"Error getting event position: {e}")
            pos = event.pos()
        # Find which date cell contains the mouse
        hovered_date = None
        for dstr, r in self._cell_rects.items():
            if r.contains(pos):
                hovered_date = dstr
                break
        if hovered_date:
            try:
                qd = QDate.fromString(hovered_date, "yyyy-MM-dd")
                entries = self._entries_for_date(qd)
                if entries:
                    lines = []
                    for nm, typ, dur, desc in entries:
                        lines.append(f"{nm}: {typ} ({dur}) — {desc}")
                    QToolTip.showText(self.mapToGlobal(pos), "\n".join(lines), self)
                else:
                    QToolTip.hideText()
            except Exception as e:
                logger.warning(f"Failed to show tooltip for date {hovered_date}: {e}")
        else:
            QToolTip.hideText()
        super().mouseMoveEvent(event)

    def contextMenuEvent(self, event):
        try:
            pos = event.pos()
            target_date = None
            for dstr, r in self._cell_rects.items():
                if r.contains(pos):
                    target_date = dstr
                    break
            if not target_date:
                return
            menu = QMenu(self)
            add_act = QAction("Add Leave", self)
            add_range_act = QAction("Add Leave Range", self)
            edit_act = QAction("Edit Leave", self)
            del_act = QAction("Delete Leave", self)
            add_act.triggered.connect(lambda: self.main._open_leave_add_dialog_for_date(target_date))
            add_range_act.triggered.connect(lambda: self.main._open_leave_add_range_for_date(target_date))
            edit_act.triggered.connect(lambda: self.main._open_leave_edit_dialog_for_date(target_date))
            del_act.triggered.connect(lambda: self.main._open_leave_delete_for_date(target_date))
            menu.addAction(add_act)
            menu.addAction(add_range_act)
            menu.addAction(edit_act)
            menu.addAction(del_act)
            menu.exec(self.mapToGlobal(pos))
        except Exception as e:
            ErrorHandler.handle_ui_error("context menu", e)

class ISMHoursDelegate(QStyledItemDelegate):
    """Thread-safe delegate for ISM Hours column validation"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
    
    def createEditor(self, parent, option, index):
        """Create editor widget for ISM Hours"""
        editor = QLineEdit(parent)
        editor.setPlaceholderText("Enter hours (e.g., 40, 120)")
        
        # Add input validation
        from PyQt6.QtGui import QIntValidator
        validator = QIntValidator(0, 9999, editor)  # Allow 0-9999 hours
        editor.setValidator(validator)
        
        return editor
    
    def setEditorData(self, editor, index):
        """Set data in editor"""
        try:
            value = index.model().data(index, Qt.ItemDataRole.EditRole)
            if isinstance(editor, QLineEdit):
                editor.setText(str(value) if value else "")
        except Exception as e:
            ErrorHandler.handle_ui_error("set editor data", e)
    
    def setModelData(self, editor, model, index):
        """Set data from editor to model"""
        try:
            if isinstance(editor, QLineEdit):
                text = editor.text().strip()
                if text:
                    # Validate the input
                    try:
                        hours = int(text)
                        if 0 <= hours <= 9999:
                            model.setData(index, str(hours), Qt.ItemDataRole.EditRole)
                        else:
                            # Show error and revert
                            model.setData(index, "0", Qt.ItemDataRole.EditRole)
                    except ValueError:
                        model.setData(index, "0", Qt.ItemDataRole.EditRole)
                else:
                    model.setData(index, "0", Qt.ItemDataRole.EditRole)
        except Exception as e:
            ErrorHandler.handle_ui_error("set model data", e)

class InlineComboDelegate(QStyledItemDelegate):
    def __init__(self, main_window, pane_name: str | None = None):
        super().__init__(main_window)
        self.main = main_window
        self.pane_name = pane_name
        # Ensure we don't install any event filters that could cause threading issues
        self._event_filter_installed = False

    def _options_for(self, pane_name: str, column_name: str) -> list[str] | None:
        if column_name == "Status":
            return STATUS_OPTIONS
        if "RAG" in column_name:
            return RAG_OPTIONS
        if pane_name == "Initiatives" and column_name == "Priority":
            return PRIORITY_FULL
        if pane_name == "Initiatives" and column_name == "Tracker":
            return TRACKER_OPTIONS
        if pane_name == "Potential Issues" and column_name == "Added in PI Tool":
            return PI_TOOL_OPTIONS
        if pane_name == "Potential Issues" and column_name == "Leads Attention Required":
            return ["Yes", "No"]
        if pane_name == "Activities" and column_name == "Support Required":
            return ["Yes", "No"]
        if pane_name == "Client Visits & Audits" and column_name == "Status":
            return AUDIT_STATUS_OPTIONS
        if column_name in ("Project Name", "Project ID"):
            # Skip dropdown for Project Details pane - allow free text input
            if pane_name == "Project Details":
                return None
            # Use repository for other panes
            if column_name == "Project Name":
                return [name for name, _ in self.main.projects]
            return [pid for _, pid in self.main.projects]
        # Org directory-backed fields
        if column_name in ("ISM Name", "Ownership"):
            rows = self.main._collect_org_directory_rows()
            names = sorted({n for (n, _, _, _, _, _) in rows if n})
            # Fallback to manual directory if org empty
            if not names:
                names = sorted(set((self.main.ism_directory or []) + list(self.main._collect_all_isms()) + [self.main.logged_in_user]))
            return names
        if column_name in ("ISMT Enterprise ID", "Enterprise ID", "ISM Enterprise ID"):
            rows = self.main._collect_org_directory_rows()
            ids = sorted({e for (_, _, e, _, _, _) in rows if e})
            return ids
        if column_name == "Audit Type":
            return [
                "Internal - Technology Scope", "Internal - No Scope for Technology",
                "External - Technology Scope", "External - No Scope for Technology",
                "Client Visit - Technology Scope", "Client Visit - No Scope for Technology",
            ]
        if pane_name == "Accolades" and column_name == "Fortnight":
            return ["1st Fortnight", "2nd Fortnight"]
        if pane_name == "Accolades" and column_name == "Month":
            # Provide a list of recent months in MMM-YY
            try:
                from datetime import date
                today = date.today()
                months = []
                y = today.year; m = today.month
                for _ in range(24):  # last 24 months
                    months.append(date(y, m, 1).strftime("%b-%y"))
                    m -= 1
                    if m == 0:
                        m = 12; y -= 1
                return months
            except Exception as e:
                ErrorHandler.handle_data_error("get months for date", e)
                return []
        # Project Details specific dropdowns
        if pane_name == "Project Details":
            if column_name == "Complexity":
                # Use hardcoded options to avoid threading issues
                return ["Low", "Medium", "High"]
            if column_name == "Connectivity Type":
                # Updated options per latest requirements (hardcoded to avoid threading issues)
                return [
                    "Facility Internet",
                    "Facility Internet with C2S",
                    "Facility Internet with S2S",
                    "Dedicated MPLS (Accenture Managed)",
                    "Dedicated MPLS (Client Managed)",
                    "Dedicated Point-to-Point (P2P) MPLS",
                    "End-to-End MPLS",
                    "MPLS (Client Managed)",
                    "Dedicated VRF",
                    "Dedicated VRF + Internet Last Mile",
                    "Dedicated VRF + Deal Dedicated Internet",
                    "Dedicated VRF + S2S Internet Last Mile",
                    "GDN Shared VRF + Client Managed MPLS",
                    "Shared GDN VRF + Internet Last Mile",
                    "Shared GDN VRF + C2S Internet Last Mile",
                    "Shared GDN VRF + Last Mile S2S Internet",
                    "Client-Provided Point-to-Point Link",
                    "S2S Tunnel Over Internet",
                    "First Mile MPLS + Last Mile Internet",
                    "First Mile MPLS + Last Mile S2S Tunnel Over Internet",
                    "First Mile MPLS + Last Mile C2S Over Internet",
                    "First Mile MPLS + Last Mile Client Managed",
                    "C2S Over Internet (VDI Solutions)",
                    "Client Provided C2S VPN",
                    "Client-Based VPN (Remote Access VPN)",
                    "Zero Trust Network Access (ZTNA) over Internet",
                    "MAN Link",
                ]
            if column_name == "Audits in Deal":
                return AUDITS_IN_DEAL_OPTIONS
            if column_name == "Voice Solution":
                # Use hardcoded options to avoid threading issues
                return ["Yes", "No"]
            if column_name == "Contact Center":
                # Use hardcoded options to avoid threading issues
                return [
                    "Amazon Connect",
                    "Cisco WebEx",
                    "NiceX",
                    "Genesys Cloud",
                    "Zoom Contact Center",
                    "Local PSTN",
                    "NA"
                ]
            if column_name in ("Primary ISM", "Secondary ISM"):
                # Get ISMs from Org Chart directory - keep dynamic for ISM columns
                org_rows = self.main._collect_org_directory_rows()
                isms = sorted({name for (name, _, _, _, _, _) in org_rows if name})
                # Fallback to collected ISMs if no org data
                if not isms:
                    isms = sorted(self.main._collect_all_isms())
                return isms
        return None

    def createEditor(self, parent, option, index):
        try:
            # Create editor directly
            
            pane_name = self.main.tabs.tabText(self.main.tabs.currentIndex()) if self.pane_name is None else self.pane_name
            if pane_name not in PANE_COLUMNS:
                return super().createEditor(parent, option, index)
            column_name = PANE_COLUMNS[pane_name][index.column()]
            opts = self._options_for(pane_name, column_name)
            if opts:
                # Custom rich combo for better ISM selection UX
                combo = QComboBox(parent)
                combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
                combo.setEditable(True)
                # Ensure no event filters are installed to prevent threading issues
                if not self._event_filter_installed:
                    self._event_filter_installed = True
                # Widen popup and add nicer spacing
                try:
                    combo.view().setMinimumWidth(320)
                except Exception as e:
                    logger.warning(f"Failed to set combo view minimum width: {e}")
                if column_name in ("ISM Name", "Ownership", "Primary ISM", "Secondary ISM"):
                    try:
                        # For Project Details, use simpler approach to avoid threading issues
                        if pane_name == "Project Details":
                            # Simple dropdown without completer to avoid threading issues
                            combo.addItems(opts)
                            for i in range(combo.count()):
                                name = combo.itemText(i)
                                try:
                                    icon = QIcon(self.main._default_avatar(name, 18))
                                    combo.setItemIcon(i, icon)
                                except Exception as e:
                                    ErrorHandler.handle_ui_error("set item icon", e)
                            return combo
                        else:
                            # For other panes, use simplified approach to avoid threading issues
                            # Skip completer setup to prevent threading issues
                            combo.addItems(opts)
                            for i in range(combo.count()):
                                name = combo.itemText(i)
                                try:
                                    icon = QIcon(self.main._default_avatar(name, 18))
                                    combo.setItemIcon(i, icon)
                                except Exception as e:
                                    ErrorHandler.handle_ui_error("set item icon", e)
                            combo.setProperty("ism-rich", False)  # Disable rich mode to avoid threading issues
                            return combo
                    except Exception as e:
                        logger.debug(f"Error creating combo with icons: {e}")
                        # Fallback to simple items with icons
                        combo.addItems(opts)
                        for i in range(combo.count()):
                            name = combo.itemText(i)
                            try:
                                icon = QIcon(self.main._default_avatar(name, 18))
                                combo.setItemIcon(i, icon)
                            except Exception as e:
                                ErrorHandler.handle_ui_error("set item icon", e)
                        return combo
                # Non-ISM combos - handle Project Details specially to avoid threading issues
                if pane_name == "Project Details":
                    # For Project Details, use simple approach without completer
                    combo.addItems(opts)
                    return combo
                else:
                    # For other panes, use normal approach
                    combo.addItems(opts)
                    return combo
            return super().createEditor(parent, option, index)
        except Exception as e:
            logger.debug(f"Error creating editor: {e}")
            return super().createEditor(parent, option, index)

    def setEditorData(self, editor, index):
        try:
            if isinstance(editor, QComboBox):
                val = index.data()
                # For Project Details, use simple text matching
                if self.pane_name == "Project Details":
                    idx = editor.findText(str(val) if val else "")
                    if idx >= 0:
                        editor.setCurrentIndex(idx)
                    else:
                        editor.setCurrentText(str(val) if val else "")
                else:
                    # For other panes, use simplified approach
                        idx = editor.findText(str(val) if val else "")
                        if idx >= 0:
                            editor.setCurrentIndex(idx)
                        else:
                        # Add new item if not found
                            editor.addItem(str(val) if val else "")
                            try:
                                icon = QIcon(self.main._default_avatar(str(val) if val else "", 18))
                                editor.setItemIcon(editor.count() - 1, icon)
                            except Exception as e:
                                ErrorHandler.handle_ui_error("set item icon", e)
                            editor.setCurrentIndex(editor.count() - 1)
            else:
                super().setEditorData(editor, index)
        except Exception as e:
            ErrorHandler.handle_ui_error("set editor data", e)
            try:
                super().setEditorData(editor, index)
            except Exception as fallback_e:
                ErrorHandler.handle_ui_error("fallback set editor data", fallback_e)

    def setModelData(self, editor, model, index):
        try:
            pane_name = self.main.tabs.tabText(self.main.tabs.currentIndex()) if self.pane_name is None else self.pane_name
            if isinstance(editor, QComboBox):
                # For Project Details, use simple text value
                if self.pane_name == "Project Details":
                    value = editor.currentText()
                else:
                    # For other panes, use simplified approach
                        value = editor.currentText()
                try:
                    row = index.row(); col = index.column()
                    self.main._set_cell_value(pane_name, row, col, value)
                    # Optional cross-fill between Name and Enterprise ID using org directory
                    cols = PANE_COLUMNS[pane_name]
                    rows = self.main._collect_org_directory_rows()
                    name_to_id = {n: e for (n, _, e, _, _, _) in rows if n and e}
                    id_to_name = {e: n for (n, _, e, _, _, _) in rows if n and e}
                    # Project repo cross-fill
                    proj_name_to_id = {n: pid for (n, pid) in self.main.projects}
                    proj_id_to_name = {pid: n for (n, pid) in self.main.projects}
                    # If editing a name/ownership, try fill enterprise id in same row
                    if cols[col] in ("ISM Name", "Ownership") and name_to_id.get(value):
                        target_cols = [c for c in ("ISMT Enterprise ID", "ISM Enterprise ID", "Enterprise ID") if c in cols]
                        if target_cols:
                            tcol = cols.index(target_cols[0])
                            self.main._set_cell_value_internal(pane_name, row, tcol, name_to_id[value], raw_override=True)
                    # If editing an enterprise id, try fill name if empty
                    if cols[col] in ("ISMT Enterprise ID", "ISM Enterprise ID", "Enterprise ID") and id_to_name.get(value):
                        target_cols = [c for c in ("ISM Name", "Ownership") if c in cols]
                        if target_cols:
                            tcol = cols.index(target_cols[0])
                            # only set if currently empty
                            try:
                                cur = self.main.data[pane_name][row][tcol]
                            except Exception:
                                cur = ""
                            if not cur:
                                self.main._set_cell_value_internal(pane_name, row, tcol, id_to_name[value], raw_override=True)
                    # If editing Project Name, set Project ID
                    if cols[col] == "Project Name" and value in proj_name_to_id:
                        if "Project ID" in cols:
                            tcol = cols.index("Project ID")
                            self.main._set_cell_value_internal(pane_name, row, tcol, proj_name_to_id[value], raw_override=True)
                    # If editing Project ID, set Project Name if empty
                    if cols[col] == "Project ID" and value in proj_id_to_name:
                        if "Project Name" in cols:
                            tcol = cols.index("Project Name")
                            try:
                                cur = self.main.data[pane_name][row][tcol]
                            except Exception:
                                cur = ""
                            if not cur:
                                self.main._set_cell_value_internal(pane_name, row, tcol, proj_id_to_name[value], raw_override=True)
                except Exception as e:
                    ErrorHandler.handle_ui_error("set cell value", e)
                    # Fallback: try to set the value directly
                    try:
                        if hasattr(index, 'model'):
                            index.model().setData(index, value)
                    except Exception as fallback_e:
                        ErrorHandler.handle_ui_error("fallback set data", fallback_e)
            else:
                super().setModelData(editor, model, index)
        except Exception as e:
            logger.warning(f"Failed to set model data in delegate: {e}")
            try:
                super().setModelData(editor, model, index)
            except Exception as fallback_error:
                logger.error(f"Fallback setModelData also failed: {fallback_error}")

class MultiLineTextDelegate(QStyledItemDelegate):
    """Delegate that provides a QTextEdit for multi-line text editing."""
    def __init__(self, parent=None):
        super().__init__(parent)

    def createEditor(self, parent, option, index):
        try:
            editor = QTextEdit(parent)
            editor.setAcceptRichText(False)
            editor.setMinimumHeight(80)
            editor.setTabChangesFocus(True)
            return editor
        except Exception as e:
            ErrorHandler.handle_ui_error("create multiline editor", e)
            return super().createEditor(parent, option, index)

    def setEditorData(self, editor, index):
        try:
            value = index.model().data(index, Qt.ItemDataRole.EditRole)
            editor.setPlainText(str(value) if value else "")
        except Exception as e:
            ErrorHandler.handle_ui_error("set multiline editor data", e)
            try:
                super().setEditorData(editor, index)
            except Exception:
                pass

    def setModelData(self, editor, model, index):
        try:
            text = editor.toPlainText()
            model.setData(index, text, Qt.ItemDataRole.EditRole)
        except Exception as e:
            ErrorHandler.handle_ui_error("set multiline model data", e)
            try:
                super().setModelData(editor, model, index)
            except Exception:
                pass

    def updateEditorGeometry(self, editor, option, index):
        editor.setGeometry(option.rect)

class LoadingOverlay(QWidget):
    """Loading overlay widget for long operations"""
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.setStyleSheet("""
            QWidget {
                background-color: rgba(0, 0, 0, 150);
                border-radius: 10px;
            }
        """)
        
        layout = QVBoxLayout(self)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # Loading spinner
        self.spinner = QLabel("⏳")
        self.spinner.setStyleSheet("""
            QLabel {
                font-size: 48px;
                color: white;
                background: transparent;
            }
        """)
        self.spinner.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.spinner)
        
        # Loading text
        self.text_label = QLabel("Loading...")
        self.text_label.setStyleSheet("""
            QLabel {
                font-size: 16px;
                color: white;
                background: transparent;
                font-weight: bold;
            }
        """)
        self.text_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.text_label)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid white;
                border-radius: 5px;
                text-align: center;
                color: white;
                font-weight: bold;
            }
            QProgressBar::chunk {
                background-color: #4CAF50;
                border-radius: 3px;
            }
        """)
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)
        
        self.hide()
    
    def show_loading(self, message="Loading...", show_progress=False):
        """Show loading overlay with message"""
        self.text_label.setText(message)
        self.progress_bar.setVisible(show_progress)
        if show_progress:
            self.progress_bar.setValue(0)
        self.show()
        QApplication.processEvents()
    
    def update_progress(self, value, message=None):
        """Update progress bar value"""
        self.progress_bar.setValue(value)
        if message:
            self.text_label.setText(message)
        QApplication.processEvents()
    
    def hide_loading(self):
        """Hide loading overlay"""
        self.hide()

class NotificationManager:
    """Modern notification manager with proper type hints and enum-based notification types."""
    
    def __init__(self, parent: QWidget) -> None:
        """Initialize notification manager.
        
        Args:
            parent: Parent widget for notifications
        """
        self.parent = parent
        self._toast: Optional[QLabel] = None
    
    def show_success(self, message: str, duration: int = 3000) -> None:
        """Show success notification with modern styling.
        
        Args:
            message: Success message to display
            duration: Display duration in milliseconds
        """
        self._show_toast(message, "#4CAF50", duration)
    
    def show_warning(self, message: str, duration: int = 4000) -> None:
        """Show warning notification with modern styling.
        
        Args:
            message: Warning message to display
            duration: Display duration in milliseconds
        """
        self._show_toast(message, "#FF9800", duration)
    
    def show_error(self, message: str, duration: int = 5000) -> None:
        """Show error notification with modern styling.
        
        Args:
            message: Error message to display
            duration: Display duration in milliseconds
        """
        self._show_toast(message, "#F44336", duration)
    
    def show_info(self, message: str, duration: int = 3000) -> None:
        """Show info notification with modern styling.
        
        Args:
            message: Info message to display
            duration: Display duration in milliseconds
        """
        self._show_toast(message, "#2196F3", duration)
    
    def _show_toast(self, text: str, color: str, duration: int) -> None:
        """Show toast notification - DISABLED to prevent duplicate notifications.
        
        Args:
            text: Notification text
            color: Background color
            duration: Display duration
        """
        # Toast notifications disabled to avoid duplicate "Data Saved" messages
        # Only the save status label near Load Data button is used now
        pass

class ValidationRule:
    """Modern base class for validation rules with proper type hints."""
    
    def validate(self, value: str, context: Optional[Dict[str, Any]] = None) -> Tuple[bool, str]:
        """Validate a value and return (is_valid, error_message).
        
        Args:
            value: Value to validate
            context: Optional context for validation
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        return True, ""
    
    def get_error_message(self, field_name: str) -> str:
        """Get formatted error message for field.
        
        Args:
            field_name: Name of the field
            
        Returns:
            Formatted error message
        """
        return f"Invalid {field_name}"

class ProjectNameValidator(ValidationRule):
    """Modern validator for Project Name field with comprehensive validation."""
    
    def validate(self, value: str, context: Optional[Dict[str, Any]] = None) -> Tuple[bool, str]:
        """Validate project name with modern error messages.
        
        Args:
            value: Project name to validate
            context: Optional validation context
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not value or not value.strip():
            return False, "Project name is required"
        
        trimmed_value = value.strip()
        if len(trimmed_value) > MAX_TEXT_LENGTH:
            return False, f"Project name must be {MAX_TEXT_LENGTH} characters or less"
        if len(trimmed_value) < 2:
            return False, "Project name must be at least 2 characters"
        
        return True, ""

class ProjectIDValidator(ValidationRule):
    """Modern validator for Project ID field with comprehensive validation."""
    
    def validate(self, value: str, context: Optional[Dict[str, Any]] = None) -> Tuple[bool, str]:
        """Validate project ID with modern error messages.
        
        Args:
            value: Project ID to validate
            context: Optional validation context
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not value or not value.strip():
            return False, "Project ID is required"
        
        trimmed_value = value.strip()
        if len(trimmed_value) > 50:
            return False, "Project ID must be 50 characters or less"
        if len(trimmed_value) < 2:
            return False, "Project ID must be at least 2 characters"
        
        return True, ""

class ISMHoursValidator(ValidationRule):
    """Modern validator for ISM Hours field with comprehensive validation."""
    
    def validate(self, value: str, context: Optional[Dict[str, Any]] = None) -> Tuple[bool, str]:
        """Validate ISM hours with modern error messages.
        
        Args:
            value: ISM hours to validate
            context: Optional validation context
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not value or not value.strip():
            return True, ""  # Empty is allowed (will default to 0)
        
        try:
            hours = int(value.strip())
            if hours < 0:
                return False, "ISM Hours cannot be negative"
            if hours > 999999:
                return False, "ISM Hours cannot exceed 999,999"
            return True, ""
        except ValueError:
            return False, "ISM Hours must be a valid number"
class ValidationEngine:
    """Modern centralized validation engine with comprehensive type hints."""
    
    def __init__(self) -> None:
        """Initialize validation engine with modern validators."""
        self.rules: Dict[str, ValidationRule] = {
            "Project Name": ProjectNameValidator(),
            "Project ID": ProjectIDValidator(),
            "ISM Hours": ISMHoursValidator(),
        }
    
    def validate_field(self, field_name: str, value: str, context: Optional[Dict[str, Any]] = None) -> Tuple[bool, str]:
        """Validate a specific field with comprehensive rules.
        
        Args:
            field_name: Name of the field to validate
            value: Value to validate
            context: Optional validation context
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if field_name in self.rules:
            return self.rules[field_name].validate(value, context)
        
        # Add validation for common field types using modern enums
        if "Due Date" in field_name or "Target Date" in field_name or "End Date" in field_name:
            return self._validate_due_date(value, field_name)
        elif "Date" in field_name:
            return self._validate_date(value)
        elif "RAG" in field_name:
            return self._validate_rag(value)
        elif "Status" in field_name:
            return self._validate_status(value)
        elif "Priority" in field_name:
            return self._validate_priority(value)
        elif "Hours" in field_name:
            return self._validate_hours(value)
        
        return True, ""  # No validation rule, assume valid
    
    def _validate_date(self, value: str) -> Tuple[bool, str]:
        """Validate date format using modern constants."""
        if not value or not str(value).strip():
            return True, ""  # Empty dates are allowed
        
        try:
            datetime.strptime(str(value), DEFAULT_DATE_FORMAT)
            return True, ""
        except ValueError:
            return False, f"Date must be in {DEFAULT_DATE_FORMAT} format"
    
    def _validate_due_date(self, value: str, field_name: str) -> Tuple[bool, str]:
        """Validate due date format and ensure it's not in the past."""
        return InputValidator.validate_due_date(value, field_name)
    
    def _validate_rag(self, value: str) -> Tuple[bool, str]:
        """Validate RAG status using modern enums."""
        valid_rag = [rag.value for rag in RAGStatus] + ["NA", ""]
        if str(value).strip() in valid_rag:
            return True, ""
        return False, f"RAG must be one of: {', '.join(valid_rag)}"
    
    def _validate_status(self, value: str) -> Tuple[bool, str]:
        """Validate status values using modern enums."""
        valid_status = [status.value for status in Status] + [""]
        if str(value).strip() in valid_status:
            return True, ""
        return False, f"Status must be one of: {', '.join(valid_status)}"
    
    def _validate_priority(self, value: str) -> Tuple[bool, str]:
        """Validate priority values using modern enums."""
        valid_priority = [priority.value for priority in Priority] + ["NA", ""]
        if str(value).strip() in valid_priority:
            return True, ""
        return False, f"Priority must be one of: {', '.join(valid_priority)}"
    
    def _validate_hours(self, value: str) -> Tuple[bool, str]:
        """Validate hours values using modern constants."""
        if not value or not str(value).strip():
            return True, ""  # Empty hours are allowed
        
        try:
            hours = float(str(value))
            if hours < 0:
                return False, "Hours cannot be negative"
            if hours > MAX_HOURS_PER_DAY * 30:  # Reasonable monthly limit
                return False, "Hours seems unreasonably high"
            return True, ""
        except ValueError:
            return False, "Hours must be a valid number"
    
    def validate_project_details_row(self, row_data: List[str]) -> Tuple[bool, List[str]]:
        """Validate an entire Project Details row with modern type hints.
        
        Args:
            row_data: List of values for the row
            
        Returns:
            Tuple of (is_valid, list_of_errors)
        """
        errors: List[str] = []
        columns = PANE_COLUMNS["Project Details"]
        
        for i, (column_name, value) in enumerate(zip(columns, row_data)):
            is_valid, error_msg = self.validate_field(column_name, str(value))
            if not is_valid:
                errors.append(f"{column_name}: {error_msg}")
        
        return len(errors) == 0, errors
    
    def validate_pane_row(self, pane_name: str, row_data: list) -> tuple[bool, list[str]]:
        """Validate a row for any pane"""
        if pane_name not in PANE_COLUMNS:
            return False, [f"Unknown pane: {pane_name}"]
        
        errors = []
        columns = PANE_COLUMNS[pane_name]
        
        # Ensure row has correct number of columns
        if len(row_data) != len(columns):
            errors.append(f"Expected {len(columns)} columns, got {len(row_data)}")
            return False, errors
        
        for i, (column_name, value) in enumerate(zip(columns, row_data)):
            is_valid, error_msg = self.validate_field(column_name, str(value))
            if not is_valid:
                errors.append(f"{column_name}: {error_msg}")
        
        return len(errors) == 0, errors

class KeyboardShortcutManager:
    """Manages keyboard shortcuts for the application"""
    def __init__(self, main_window):
        self.main_window = main_window
        self._setup_shortcuts()
    
    def _setup_shortcuts(self):
        """Setup all keyboard shortcuts"""
        # File operations
        QShortcut(QKeySequence.StandardKey.New, self.main_window, self.main_window._add_project)
        QShortcut(QKeySequence.StandardKey.Save, self.main_window, self.main_window._save_backend_sqlite)
        QShortcut(QKeySequence.StandardKey.Open, self.main_window, self.main_window.animated_load_data)
        QShortcut(QKeySequence("Ctrl+E"), self.main_window, self._export_current_pane)
        QShortcut(QKeySequence("Ctrl+I"), self.main_window, self._import_current_pane)
        
        # Edit operations
        QShortcut(QKeySequence.StandardKey.Undo, self.main_window, self.main_window._undo_last_edit)
        QShortcut(QKeySequence.StandardKey.Redo, self.main_window, self.main_window._redo_last_edit)
        QShortcut(QKeySequence.StandardKey.Delete, self.main_window, self._delete_selected)
        QShortcut(QKeySequence("Ctrl+D"), self.main_window, self._duplicate_selected)
        QShortcut(QKeySequence("Ctrl+A"), self.main_window, self._select_all)
        
        # Navigation
        QShortcut(QKeySequence("Ctrl+F"), self.main_window, self.main_window.open_global_search)
        QShortcut(QKeySequence("Ctrl+K"), self.main_window, self.main_window.open_global_search)
        QShortcut(QKeySequence("Ctrl+H"), self.main_window, self.main_window.show_help_dialog)
        QShortcut(QKeySequence("F1"), self.main_window, self.main_window.show_help_dialog)
        QShortcut(QKeySequence("F5"), self.main_window, self._refresh_current_pane)
        
        # Tab navigation
        QShortcut(QKeySequence("Ctrl+Tab"), self.main_window, self._next_tab)
        QShortcut(QKeySequence("Ctrl+Shift+Tab"), self.main_window, self._previous_tab)
        QShortcut(QKeySequence("Ctrl+1"), self.main_window, lambda: self._goto_tab(0))
        QShortcut(QKeySequence("Ctrl+2"), self.main_window, lambda: self._goto_tab(1))
        QShortcut(QKeySequence("Ctrl+3"), self.main_window, lambda: self._goto_tab(2))
        QShortcut(QKeySequence("Ctrl+4"), self.main_window, lambda: self._goto_tab(3))
        QShortcut(QKeySequence("Ctrl+5"), self.main_window, lambda: self._goto_tab(4))
        
        # Settings and preferences
        QShortcut(QKeySequence("Ctrl+,"), self.main_window, self._open_preferences)
        QShortcut(QKeySequence("Ctrl+Alt+S"), self.main_window, self._toggle_auto_save)
        
        # Project Details specific
        QShortcut(QKeySequence("Ctrl+Shift+A"), self.main_window, self.main_window._add_project)
        QShortcut(QKeySequence("Ctrl+Shift+D"), self.main_window, self._delete_selected_project)
    
    def _delete_selected(self):
        """Delete selected item in current context"""
        if hasattr(self.main_window, 'projects_table') and self.main_window.projects_table.hasFocus():
            self.main_window._delete_project()
        else:
            # Try to delete from current pane
            current_tab = self.main_window.tabs.currentWidget()
            if hasattr(current_tab, 'table'):
                self.main_window.delete_selected(current_tab.objectName())
    
    def _delete_selected_project(self):
        """Delete selected project from Project Details"""
        if hasattr(self.main_window, 'projects_table'):
            self.main_window._delete_project()
    
    def _export_current_pane(self):
        """Export current pane data"""
        try:
            current_tab = self.main_window.tabs.currentWidget()
            if hasattr(current_tab, 'objectName'):
                pane_name = current_tab.objectName()
                if pane_name in PANE_COLUMNS:
                    self.main_window.export_pane(pane_name)
        except Exception as e:
            ErrorHandler.handle_ui_error("export current pane", e)
    
    def _import_current_pane(self):
        """Import data to current pane"""
        try:
            current_tab = self.main_window.tabs.currentWidget()
            if hasattr(current_tab, 'objectName'):
                pane_name = current_tab.objectName()
                if pane_name in PANE_COLUMNS:
                    self.main_window.import_pane(pane_name)
        except Exception as e:
            ErrorHandler.handle_ui_error("import current pane", e)
    
    def _duplicate_selected(self):
        """Duplicate selected row in current pane"""
        try:
            current_tab = self.main_window.tabs.currentWidget()
            if hasattr(current_tab, 'table'):
                table = current_tab.table
                selected_rows = table.selectionModel().selectedRows()
                if selected_rows:
                    pane_name = current_tab.objectName()
                    row = selected_rows[0].row()
                    self.main_window._duplicate_row(pane_name, row)
        except Exception as e:
            ErrorHandler.handle_ui_error("duplicate selected", e)
    
    def _select_all(self):
        """Select all rows in current pane"""
        try:
            current_tab = self.main_window.tabs.currentWidget()
            if hasattr(current_tab, 'table'):
                current_tab.table.selectAll()
        except Exception as e:
            ErrorHandler.handle_ui_error("select all", e)
    
    def _refresh_current_pane(self):
        """Refresh current pane"""
        try:
            current_tab = self.main_window.tabs.currentWidget()
            if hasattr(current_tab, 'objectName'):
                pane_name = current_tab.objectName()
                if pane_name in PANE_COLUMNS:
                    self.main_window.rebuild_table(pane_name)
                    self.main_window._show_toast("Pane refreshed", level="SUCCESS")
        except Exception as e:
            ErrorHandler.handle_ui_error("refresh current pane", e)
    
    def _next_tab(self):
        """Switch to next tab"""
        try:
            current_index = self.main_window.tabs.currentIndex()
            next_index = (current_index + 1) % self.main_window.tabs.count()
            self.main_window.tabs.setCurrentIndex(next_index)
        except Exception as e:
            ErrorHandler.handle_ui_error("next tab", e)
    
    def _previous_tab(self):
        """Switch to previous tab"""
        try:
            current_index = self.main_window.tabs.currentIndex()
            prev_index = (current_index - 1) % self.main_window.tabs.count()
            self.main_window.tabs.setCurrentIndex(prev_index)
        except Exception as e:
            ErrorHandler.handle_ui_error("previous tab", e)
    
    def _goto_tab(self, index: int):
        """Go to specific tab by index"""
        try:
            if 0 <= index < self.main_window.tabs.count():
                self.main_window.tabs.setCurrentIndex(index)
        except Exception as e:
            ErrorHandler.handle_ui_error("goto tab", e)
    
    def _open_preferences(self):
        """Open preferences dialog"""
        try:
            self.main_window._show_preferences_dialog()
        except Exception as e:
            ErrorHandler.handle_ui_error("open preferences", e)
    
    def _toggle_auto_save(self):
        """Toggle auto-save functionality"""
        try:
            current_state = self.main_window.preferences.get("auto_save", True)
            new_state = not current_state
            self.main_window.preferences.set("auto_save", new_state)
            
            if new_state:
                self.main_window._setup_auto_save_timer()
                self.main_window._show_toast("Auto-save enabled", level="SUCCESS")
            else:
                if hasattr(self.main_window, 'auto_save_timer'):
                    self.main_window.auto_save_timer.stop()
                self.main_window._show_toast("Auto-save disabled", level="INFO")
        except Exception as e:
            ErrorHandler.handle_ui_error("toggle auto-save", e)


class ISMHoursDelegate(QStyledItemDelegate):
    """Thread-safe delegate for ISM Hours column to validate numeric input only"""
    def __init__(self, parent=None):
        super().__init__(parent)
        # Ensure we don't install any event filters that could cause threading issues
        self._event_filter_installed = False
    
    def createEditor(self, parent, option, index):
        try:
            # Create editor directly
            
            editor = QLineEdit(parent)
            editor.setValidator(QIntValidator(0, 999999))  # Allow only positive integers
            # Ensure no event filters are installed
            if not self._event_filter_installed:
                self._event_filter_installed = True
            return editor
        except Exception as e:
            ErrorHandler.handle_ui_error("create ISM hours editor", e)
            # Fallback to default editor
            return super().createEditor(parent, option, index)
    
    def setEditorData(self, editor, index):
        value = index.model().data(index, Qt.ItemDataRole.EditRole)
        editor.setText(str(value) if value else "0")
    
    def setModelData(self, editor, model, index):
        value = editor.text().strip()
        if value:
            try:
                # Validate that it's a positive integer
                int_value = int(value)
                if int_value >= 0:
                    model.setData(index, str(int_value), Qt.ItemDataRole.EditRole)
                else:
                    model.setData(index, "0", Qt.ItemDataRole.EditRole)
            except ValueError:
                model.setData(index, "0", Qt.ItemDataRole.EditRole)
        else:
            model.setData(index, "0", Qt.ItemDataRole.EditRole)
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Accenture ISMT Task Management Dashboard")
        
        # Set flexible window size that adapts to display constraints
        self.setMinimumSize(1200, 800)
        self.resize(1400, 900)
        # Ensure window fits on screen
        screen = QApplication.primaryScreen()
        if screen:
            screen_geometry = screen.availableGeometry()
            # Adjust size if it's too large for the screen
            if self.width() > screen_geometry.width() or self.height() > screen_geometry.height():
                self.resize(
                    min(self.width(), screen_geometry.width() - 100),
                    min(self.height(), screen_geometry.height() - 100)
                )
        self.data: dict[str, list[list[str]]] = {name: [] for name in PANE_COLUMNS}
        self.logged_in_user: str = getpass.getuser()
        self.projects: list[tuple[str, str]] = []  # list of (Project Name, Project ID)
        self.change_log: list[str] = []
        self.change_log_data: list[list[str]] = []  # For detailed change tracking
        self._edit_undo_stack: list[tuple[str,int,int,str,str]] = []  # (pane,row,col,old,new)
        self._edit_redo_stack: list[tuple[str,int,int,str,str]] = []
        # Persistent history of Action Owner team names for Potential Issues
        self.team_history: list[str] = []
        # Persistent ISM directory managed by user
        self.ism_directory: list[str] = []
        # Import log (global)
        self._imports_log: list[tuple[str,str,str,int]] = []  # (timestamp, pane, source, rows)
        # UI state persistence
        self.empty_states: dict[str, QLabel] = {}
        self._column_orders: dict[str, list[int]] = {}
        self._hidden_columns: dict[str, list[int]] = {}
        self._org_zoom: float = 1.0
        self._org_color_scheme: str = "designation"  # "level", "designation", or "combined"
        # Dismissed notifications storage
        self._dismissed_notifications: set[str] = set()  # Store notification IDs to filter them out
        # Simple in-memory queue for leave approval notifications
        self._leave_approval_queue: list[dict] = []

        # Initialize backend path - SQLite is now the primary backend
        try:
            self.backend_sqlite_path = self._get_backend_path_from_settings()
        except Exception:
            # Fallback to SQLite default
            script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
            self.backend_sqlite_path = os.path.join(script_dir, "backend_data.sqlite")

        central = QWidget()
        self.setCentralWidget(central)
        self.root_layout = QVBoxLayout(central)

        # Accent color used across UI
        self.accent = "#2a7de1"
        self._toast: QLabel | None = None
        # Initialize new managers
        self.loading_overlay = LoadingOverlay(self)
        self.notifications = NotificationManager(self)
        self.shortcuts = KeyboardShortcutManager(self)
        self.validation_engine = ValidationEngine()
        
        # Initialize new systems
        self.preferences = UserPreferences(getattr(self, '__backend_path__', None) or self.backend_sqlite_path)
        self.progress_manager = ProgressManager(self)
        self.undo_redo_manager = UndoRedoManager()
        # Project Details edit state guards
        self._pd_edit_in_progress: bool = False
        self._pd_save_timer: QTimer | None = None
        
        # Initialize error handler with main window reference
        ErrorHandler.set_main_window(self)
        
        # CRITICAL FIX: Add error recovery mechanism
        self._error_recovery_enabled = True
        self._last_error_time = None
        self._error_count = 0
        
        # Initialize status bar for user feedback
        self.statusBar().showMessage("Ready", 0)  # 0 = permanent message
        
        # Memory optimization
        self._cleanup_timers = []
        self._cleanup_widgets = []
        
        # Initialize timer cleanup tracking
        self._active_timers = []
        
        # Thread safety for UI operations
        self._ui_operation_queue = []
        self._ui_operation_timer = None

        # Initialize helper methods for UI components
        self._init_ui_helpers()

        # Header bar with title and quick action
        header = QWidget()
        header.setObjectName("headerBar")
        header_layout = QHBoxLayout(header)
        header_layout.setContentsMargins(12, 8, 12, 8)
        title = QLabel("Accenture ISMT Task Management Dashboard")
        title.setObjectName("appTitle")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setWeight(QFont.Weight.Medium)
        title.setFont(title_font)
        header_layout.addWidget(title)
        header_layout.addStretch(1)
        
        # Right-aligned button group
        button_group = QHBoxLayout()
        button_group.setSpacing(6)  # Reduced spacing for better icon alignment
        try:
            button_group.setContentsMargins(0, 0, 0, 0)
        except Exception as e:
            logger.debug(f"Failed to set button group margins: {e}")
        # Save Status Indicator
        self.save_status_label = QLabel("")
        self.save_status_label.setFixedSize(120, 36)  # Set same fixed size as buttons
        self.save_status_label.setStyleSheet("""
            QLabel {
                color: #27ae60;
                font-size: 12px;
                font-weight: bold;
                padding: 4px 8px;
                background-color: rgba(39, 174, 96, 0.1);
                border: 1px solid #27ae60;
                border-radius: 4px;
                min-width: 120px;
                text-align: center;
            }
        """)
        self.save_status_label.setVisible(False)
        # Save All button
        save_all_btn = QPushButton("💾 Save All")
        save_all_btn.setToolTip("Save all data to backend")
        try:
            from PyQt6.QtWidgets import QSizePolicy
            save_all_btn.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
            save_all_btn.setFixedSize(120, 36)
        except Exception:
            pass
        save_all_btn.setStyleSheet(
            """
            QPushButton {
                background-color: #2ecc71;
                color: white;
                border: none;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover { background-color: #27ae60; }
            QPushButton:pressed { background-color: #1e8449; }
            """
        )
        def _do_save_all():
            try:
                self._save_backend_sqlite()
            except Exception as e:
                ErrorHandler.handle_ui_error("save all", e)
        save_all_btn.clicked.connect(_do_save_all)
        button_group.addWidget(save_all_btn)

        # Save & Exit button
        save_exit_btn = QPushButton("✔ Save and Exit")
        save_exit_btn.setToolTip("Save all data and exit the application")
        try:
            from PyQt6.QtWidgets import QSizePolicy
            save_exit_btn.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
            save_exit_btn.setFixedSize(130, 36)
        except Exception:
            pass
        save_exit_btn.setStyleSheet(
            """
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover { background-color: #c0392b; }
            QPushButton:pressed { background-color: #922b21; }
            """
        )
        def _do_save_and_exit():
            try:
                self._save_backend_sqlite()
            except Exception as e:
                ErrorHandler.handle_ui_error("save and exit", e)
            try:
                self.close()
            except Exception:
                pass
        save_exit_btn.clicked.connect(_do_save_and_exit)
        button_group.addWidget(save_exit_btn)

        button_group.addWidget(self.save_status_label)
        
        quick_export = QPushButton("Export All")
        quick_export.setObjectName("primary")
        quick_export.clicked.connect(self.export_all_sheets)
        try:
            from PyQt6.QtWidgets import QSizePolicy
            quick_export.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
            quick_export.setFixedSize(120, 36)  # Set same fixed size as Load Data button
        except Exception as e:
            ErrorHandler.handle_ui_error("set size policy", e)
        button_group.addWidget(quick_export)
        
        # Refresh All button (rebuilds/refreshes all panes)
        refresh_all_btn = QPushButton("Refresh All")
        try:
            from PyQt6.QtWidgets import QSizePolicy
            refresh_all_btn.setSizePolicy(QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Fixed)
            refresh_all_btn.setFixedSize(120, 36)
        except Exception:
            pass
        refresh_all_btn.clicked.connect(self.refresh_all_panes_now)
        button_group.addWidget(refresh_all_btn)
        
        # Notification Bell Icon with Count Badge - Compact to match profile icon
        self.notification_bell_container = QWidget()
        self.notification_bell_container.setFixedSize(48, 48)  # Match compact profile icon size
        notification_layout = QVBoxLayout(self.notification_bell_container)
        notification_layout.setContentsMargins(0, 0, 0, 0)
        notification_layout.setSpacing(0)
        
        self.notification_bell = QToolButton()
        self.notification_bell.setIcon(QIcon(self._create_bell_icon()))
        self.notification_bell.setToolTip("Notifications")
        self.notification_bell.setFixedSize(48, 48)  # Match compact profile icon size
        self.notification_bell.setIconSize(QSize(36, 36))  # Compact icon size
        self.notification_bell.setStyleSheet("""
            QToolButton {
                border: none;
                background: transparent;
                padding: 1px;
                border-radius: 6px;
                margin: 1px;
            }
            QToolButton:hover {
                background-color: rgba(42, 125, 225, 0.12);
                border: 1px solid rgba(42, 125, 225, 0.4);
                transform: scale(1.05);
            }
            QToolButton:pressed {
                background-color: rgba(42, 125, 225, 0.25);
                transform: scale(0.98);
            }
        """)
        self.notification_bell.clicked.connect(self.show_notifications_dialog)
        # Ensure the bell icon has proper mouse tracking and click handling
        self.notification_bell.setMouseTracking(True)
        self.notification_bell.setAttribute(Qt.WidgetAttribute.WA_Hover, True)
        
        # Notification count badge
        self.notification_count_badge = QLabel()
        self.notification_count_badge.setFixedSize(20, 20)
        self.notification_count_badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.notification_count_badge.setStyleSheet("""
            QLabel {
                background-color: #e74c3c;
                color: white;
                border-radius: 9px;
                font-size: 10px;
                font-weight: bold;
                border: 2px solid white;
            }
        """)
        self.notification_count_badge.hide()  # Initially hidden
        
        # Position the badge in the top-right corner of the bell
        notification_layout.addWidget(self.notification_bell)
        self.notification_count_badge.setParent(self.notification_bell_container)
        self.notification_count_badge.move(32, 1)  # Adjusted position for compact bell (48x48)
        
        button_group.addWidget(self.notification_bell_container)
        
        # User Profile Icon - Compact and subtle
        self.profile_icon = QToolButton()
        self.profile_icon.setIcon(QIcon(self._default_avatar(self.logged_in_user, 36)))  # Compact avatar
        # Set tooltip to show ISM name for better clarity
        display_name = self._get_user_display_name(self.logged_in_user)
        self.profile_icon.setToolTip(f"Profile: {display_name}")
        self.profile_icon.setFixedSize(48, 48)  # Compact button
        self.profile_icon.setIconSize(QSize(36, 36))  # Compact icon size
        self.profile_icon.setStyleSheet("""
            QToolButton {
                border: none;
                background: transparent;
                padding: 1px;
                border-radius: 6px;
                margin: 1px;
            }
            QToolButton:hover {
                background-color: rgba(42, 125, 225, 0.12);
                border: 1px solid rgba(42, 125, 225, 0.4);
                transform: scale(1.05);
            }
            QToolButton:pressed {
                background-color: rgba(42, 125, 225, 0.25);
                transform: scale(0.98);
            }
        """)
        self.profile_icon.clicked.connect(self.show_profile_dialog)
        # Ensure the profile icon has proper mouse tracking and click handling
        self.profile_icon.setMouseTracking(True)
        self.profile_icon.setAttribute(Qt.WidgetAttribute.WA_Hover, True)
        button_group.addWidget(self.profile_icon)
        
        # Add the button group to header layout using a container widget for stable right alignment
        try:
            from PyQt6.QtWidgets import QWidget as _QWidget, QSizePolicy as _QSizePolicy
            button_group_container = _QWidget()
            button_group_container.setLayout(button_group)
            button_group_container.setSizePolicy(_QSizePolicy.Policy.Maximum, _QSizePolicy.Policy.Fixed)
            header_layout.addWidget(button_group_container, alignment=Qt.AlignmentFlag.AlignRight)
        except Exception as e:
            logger.debug(f"Failed to add button group container: {e}")
            # Fallback: add layout directly
            header_layout.addLayout(button_group)
            try:
                header_layout.setAlignment(button_group, Qt.AlignmentFlag.AlignRight)
            except Exception as alignment_error:
                logger.debug(f"Failed to set button group alignment: {alignment_error}")
        self.root_layout.addWidget(header)

        # Menu + Dark mode
        menu_bar = self.menuBar()
        file_menu = menu_bar.addMenu("File")
        load_data_action = QAction("Load Data File...", self)
        def do_load_data():
            # Use the animated load data function
            self.animated_load_data()
        load_data_action.triggered.connect(do_load_data)
        file_menu.addAction(load_data_action)
        create_backend_action = QAction("Create New Backend File...", self)
        def do_create_backend():
            path, _ = QFileDialog.getSaveFileName(self, "Create New Backend File", "backend_data.sqlite", "SQLite (*.sqlite)")
            if not path:
                return
            try:
                self.backend_sqlite_path = path
                if str(path).lower().endswith((".sqlite", ".db")):
                    self._create_new_backend_sqlite(path)
                else:
                    self._create_new_backend_file()
                self.__backend_path__ = path
                QMessageBox.information(self, "Backend File", f"Successfully created new backend file:\n{path}")
            except Exception as e:
                QMessageBox.critical(self, "Create Backend File", f"Failed to create backend file:\n{str(e)}")
        create_backend_action.triggered.connect(do_create_backend)
        file_menu.addAction(create_backend_action)
        reset_backend_action = QAction("Reset Current Backend File", self)
        def do_reset_backend():
            if not hasattr(self, '__backend_path__') or not self.__backend_path__:
                QMessageBox.warning(self, "Reset Backend", "No backend file is currently loaded.")
                return
            reply = QMessageBox.question(self, "Reset Backend File", 
                                       f"Are you sure you want to reset the backend file?\n\n{self.__backend_path__}\n\nThis will clear all data in the file.",
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                try:
                    # Clear all data
                    for pane in self.data:
                        self.data[pane] = []
                    # Rebuild tables (skip Leave Tracker and Project Details as they use custom UI)
                    for pane in self.data:
                        if pane not in ["Leave Tracker", "Project Details"]:
                            self.rebuild_table(pane)
                    # Save empty backend file
                    self._save_backend_sqlite()
                    # Update dashboard and home stats
                    self.update_dashboard()
                    self.update_home_stats()  # Update home page stats
                    QMessageBox.information(self, "Reset Backend", "Backend file has been reset successfully.")
                except Exception as e:
                    QMessageBox.critical(self, "Reset Backend", f"Failed to reset backend file:\n{str(e)}")
        reset_backend_action.triggered.connect(do_reset_backend)
        file_menu.addAction(reset_backend_action)
        file_menu.addSeparator()
        export_all_action = QAction("Export All to Excel...", self)
        export_all_action.triggered.connect(self.export_all_sheets)
        file_menu.addAction(export_all_action)

        # Ensure Settings menu exists
        try:
            settings_menu = None
            for act in menu_bar.actions():
                if act.text() and 'Settings' in act.text():
                    settings_menu = act.menu()
                    break
            if settings_menu is None:
                settings_menu = menu_bar.addMenu("Settings")
            # Add Backend File settings action
            act_backend = QAction("Backend File...", self)
            def open_backend_settings():
                self._open_backend_path_settings_dialog()
            act_backend.triggered.connect(open_backend_settings)
            settings_menu.addAction(act_backend)
            # Add Preferences dialog (includes auto-save timer settings)
            act_prefs = QAction("Preferences...", self)
            act_prefs.setShortcut(QKeySequence("Ctrl+,"))
            act_prefs.triggered.connect(self._show_preferences_dialog)
            settings_menu.addAction(act_prefs)
        except Exception as e:
            ErrorHandler.handle_ui_error("install backend settings menu", e)

        view_menu = menu_bar.addMenu("View")
        directory_menu = menu_bar.addMenu("Directory")
        manage_isms_action = QAction("Manage ISM Directory...", self)
        manage_isms_action.triggered.connect(self.manage_ism_directory)
        directory_menu.addAction(manage_isms_action)
        open_directory_action = QAction("Open Team Directory...", self)
        open_directory_action.triggered.connect(self.show_team_directory)
        directory_menu.addAction(open_directory_action)
        # Leave Approvals menu
        approvals_menu = menu_bar.addMenu("Approvals")
        leave_approvals_action = QAction("Leave Approvals...", self)
        leave_approvals_action.triggered.connect(self.open_leave_approvals_dialog)
        approvals_menu.addAction(leave_approvals_action)
        
        # Log Management Menu
        log_menu = menu_bar.addMenu("Logs")
        view_logs_action = QAction("View Activity Logs...", self)
        view_logs_action.triggered.connect(self.show_logs_dialog)
        log_menu.addAction(view_logs_action)
        
        reset_logs_action = QAction("Reset All Logs...", self)
        reset_logs_action.triggered.connect(self.reset_logging_data)
        log_menu.addAction(reset_logs_action)
        
        rotate_logs_action = QAction("Rotate Logs (Archive Old)", self)
        rotate_logs_action.triggered.connect(self._rotate_logs)
        log_menu.addAction(rotate_logs_action)
        
        log_menu.addSeparator()
        export_logs_action = QAction("Export Logs to CSV...", self)
        export_logs_action.triggered.connect(self.export_logs_to_csv)
        log_menu.addAction(export_logs_action)
        global_search_action = QAction("Global Search (Ctrl+K)", self)
        global_search_action.setShortcut("Ctrl+K")
        global_search_action.triggered.connect(self.open_global_search)
        view_menu.addAction(global_search_action)
        # Removed: Show Recent Activity (moved out of View menu)
        # Removed: Reset Logging Data (available under Logs menu)
        help_action = QAction("Help / Shortcuts", self)
        help_action.triggered.connect(self.show_help_dialog)
        view_menu.addAction(help_action)
        # Force-show tab actions removed from View menu

        # Tabs
        self.tabs = QTabWidget()
        self.tabs.setTabPosition(QTabWidget.TabPosition.North)
        
        # Wiki / Knowledge Base functionality has been removed
        
        try:
            # Ensure overflowed tabs are reachable
            tb = self.tabs.tabBar()
            if tb is not None:
                tb.setUsesScrollButtons(True)
                tb.setExpanding(False)
        except Exception as e:
            ErrorHandler.handle_ui_error("configure tab bar", e)
        self.root_layout.addWidget(self.tabs)
        
        # Home Page Tab (first tab)
        self.home_tab = QWidget()
        self._init_home_page()
        self.tabs.addTab(self.home_tab, "🏠 Home")

        # Global shortcuts for undo/redo
        try:
            QShortcut(QKeySequence.StandardKey.Undo, self, activated=self._undo_last_edit)
            QShortcut(QKeySequence.StandardKey.Redo, self, activated=self._redo_last_edit)
        except Exception as e:
            logger.debug(f"Failed to set keyboard shortcuts: {e}")

        # Dashboard tab
        self.dashboard_tab = QWidget()
        self.dashboard_layout = QVBoxLayout(self.dashboard_tab)
        self.cards_bar = QHBoxLayout()
        self.dashboard_layout.addLayout(self.cards_bar)
        # Dashboard filters
        dash_filters = QHBoxLayout()
        self.ism_filter = QComboBox()
        self.ism_filter.addItem("All ISMs")
        self.ism_filter.currentTextChanged.connect(lambda _=None: self.update_dashboard())
        dash_filters.addWidget(QLabel("ISM:"))
        dash_filters.addWidget(self.ism_filter)
        dash_filters.addStretch(1)
        self.dashboard_layout.addLayout(dash_filters)
        self.chart_holder = QWidget()
        self.chart_layout = QVBoxLayout(self.chart_holder)
        self.dashboard_layout.addWidget(self.chart_holder)
        # Status bar chart holder
        self.chart_holder2 = QWidget()
        self.chart_layout2 = QVBoxLayout(self.chart_holder2)
        self.dashboard_layout.addWidget(self.chart_holder2)
        # Overdue table (hidden by default as per feedback)
        self.overdue_table = None
        # Remove in-dashboard activity list; add under View menu
        # Client Visits / Audits consolidated table
        self.dashboard_layout.addWidget(QLabel("Client Visits / Audits - In Progress & Next 5 Days"))
        # Quick filters
        visits_filters = QHBoxLayout()
        self.visits_only_inprogress = QCheckBox("Only In Progress")
        self.visits_audit_type = QComboBox()
        self.visits_audit_type.addItems(["All Types", "Internal - Tech Scope", "Internal - No Scope", "External - Tech Scope", "External - No Scope", "Client Visit - Tech Scope", "Client Visit - No Scope"]) 
        visits_filters.addWidget(self.visits_only_inprogress)
        visits_filters.addWidget(QLabel("Audit Type:"))
        visits_filters.addWidget(self.visits_audit_type)
        visits_filters.addStretch(1)
        self.dashboard_layout.addLayout(visits_filters)
        self.visits_table = QTableWidget()
        self.visits_table.setColumnCount(7)
        self.visits_table.setHorizontalHeaderLabels(["Project Name", "Audit Type", "ISM Name", "Start Date", "End Date", "Status", "RAG"])
        self.visits_table.setSortingEnabled(True)
        self.visits_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.dashboard_layout.addWidget(self.visits_table)
        self.tabs.addTab(self.dashboard_tab, "📊 Dashboard")

        # Org Chart tab
        self.org_tab = QWidget()
        org_layout = QVBoxLayout(self.org_tab)
        org_layout.setContentsMargins(6, 6, 6, 6)
        org_layout.setSpacing(8)
        org_controls = QHBoxLayout()
        add_member_btn = QPushButton("Add Team Member")
        add_member_btn.setObjectName("primary")
        add_member_btn.clicked.connect(self.add_org_member)
        org_controls.addWidget(add_member_btn)
        
        # Add Sample Data button removed as requested
        
        # Color scheme controls
        org_controls.addWidget(QLabel("Color Scheme:"))
        self.color_scheme_combo = QComboBox()
        self.color_scheme_combo.addItems(["Level-based", "Designation-based", "Combined"])
        self.color_scheme_combo.setCurrentText("Level-based")
        self.color_scheme_combo.currentTextChanged.connect(self._on_color_scheme_changed)
        org_controls.addWidget(self.color_scheme_combo)
        
        # Color scheme info button
        color_info_btn = QPushButton("ℹ️")
        color_info_btn.setToolTip("Click to see color scheme information")
        color_info_btn.setStyleSheet("""
            QPushButton {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 4px 8px;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #e9ecef;
            }
        """)
        color_info_btn.clicked.connect(self._show_color_scheme_info)
        org_controls.addWidget(color_info_btn)
        
        # Add Export/Import buttons
        export_btn = QPushButton("📤 Export")
        export_btn.setObjectName("secondary")
        export_btn.clicked.connect(self._export_org_chart)
        org_controls.addWidget(export_btn)
        
        import_btn = QPushButton("📥 Import")
        import_btn.setObjectName("secondary")
        import_btn.clicked.connect(self._import_org_chart)
        org_controls.addWidget(import_btn)
        
        sample_btn = QPushButton("📋 Sample")
        sample_btn.setObjectName("secondary")
        sample_btn.clicked.connect(self._download_org_chart_sample)
        org_controls.addWidget(sample_btn)
        
        org_controls.addStretch()
        # Org chart toolbar (cleaned: no export/print)
        tb = QToolBar()
        zoom_in = QAction("Zoom +", self); zoom_out = QAction("Zoom -", self)
        search_lbl = QLabel("Find:"); search_edit = QLineEdit(); search_edit.setPlaceholderText("Name or ID")
        tb.addAction(zoom_in); tb.addAction(zoom_out)
        tb.addSeparator(); tb.addWidget(search_lbl); tb.addWidget(search_edit)
        def do_zoom(factor: float):
            self._org_zoom *= factor
            self.org_view.scale(factor, factor)
        zoom_in.triggered.connect(lambda: do_zoom(1.15))
        zoom_out.triggered.connect(lambda: do_zoom(1/1.15))
        def do_reset():
            self._org_zoom = 1.0
            self.org_view.resetTransform()
            # Auto-fit the org chart to the view size with intelligent padding
            if hasattr(self, 'org_scene') and self.org_scene.itemsBoundingRect().isValid():
                rect = self.org_scene.itemsBoundingRect()
                
                # Get view and screen dimensions
                view_size = self.org_view.size()
                screen = QApplication.primaryScreen()
                screen_geometry = screen.availableGeometry() if screen else QRect(0, 0, 1920, 1080)
                
                # Calculate optimal padding based on multiple factors
                min_view_dim = min(view_size.width(), view_size.height())
                max_view_dim = max(view_size.width(), view_size.height())
                
                # Determine screen size category with much smaller padding for better space utilization
                screen_area = screen_geometry.width() * screen_geometry.height()
                if screen_area >= 3840 * 2160:  # 4K and above
                    base_padding_percent = 0.05  # 5% for large screens - maximize content
                elif screen_area >= 2560 * 1440:  # 1440p
                    base_padding_percent = 0.04  # 4% for medium-large screens
                elif screen_area >= 1920 * 1080:  # 1080p
                    base_padding_percent = 0.03  # 3% for standard screens
                else:  # Smaller screens
                    base_padding_percent = 0.02  # 2% for small screens
                
                # Adjust padding based on view aspect ratio
                aspect_ratio = view_size.width() / view_size.height() if view_size.height() > 0 else 1.0
                if aspect_ratio > 2.0:  # Very wide view
                    base_padding_percent *= 0.8  # Reduce padding for wide views
                elif aspect_ratio < 0.5:  # Very tall view
                    base_padding_percent *= 0.8  # Reduce padding for tall views
                
                # Calculate dynamic padding with much smaller minimums for better space utilization
                dynamic_padding = max(5, int(min_view_dim * base_padding_percent))
                
                # Ensure reasonable bounds - much more generous maximums
                dynamic_padding = min(dynamic_padding, int(max_view_dim * 0.08))  # Max 8% of larger dimension
                dynamic_padding = max(dynamic_padding, 3)  # Minimum 3px for very small screens
                
                # Apply dynamic padding and fit
                padded = rect.adjusted(-dynamic_padding, -dynamic_padding, dynamic_padding, dynamic_padding)
                self.org_view.fitInView(padded, Qt.AspectRatioMode.KeepAspectRatio)
                # Center the view
                self.org_view.centerOn(rect.center())
        reset_zoom = QAction("Reset View", self)
        reset_zoom.triggered.connect(do_reset)
        tb.addAction(reset_zoom)
        # Removed PNG/SVG export and Print features per request
        def do_search():
            q = (search_edit.text() or "").lower().strip()
            if not q:
                return
            # find first match
            root = self.org_tree.invisibleRootItem()
            stack = [root]
            found = None
            while stack and not found:
                cur = stack.pop()
                for i in range(cur.childCount()):
                    ch = cur.child(i)
                    if q in ch.text(0).lower() or q in ch.text(2).lower():
                        found = ch; break
                    stack.append(ch)
            if found:
                self.render_org_chart()
                # Center graphics view on found node by locating its scene rect
                # Naively re-run layout to compute positions and find nearest item by tooltip id
                # Highlight by showing a temporary toast and centering view
                try:
                    name = found.text(0); ent = found.text(2)
                    tip = f"<b>{name}</b><br/>{found.text(1)}<br/>ID: {ent}"
                    for it in self.org_scene.items():
                        if hasattr(it, 'toolTip'):
                            if it.toolTip() and (f"ID: {ent}" in it.toolTip()):
                                rect = it.sceneBoundingRect()
                                self.org_view.centerOn(rect.center())
                                break
                except Exception as e:
                    ErrorHandler.handle_ui_error("org chart search", e)
                
                # Also show an avatar in a toast-like hint
                try:
                    avatar = self._default_avatar(name, 28)
                    icon = QIcon(avatar)
                    self._show_toast(f"Found: {name} ({ent})")
                except Exception as e:
                    ErrorHandler.handle_ui_error("show search result", e)
                QMessageBox.information(self, "Found", f"Found: {found.text(0)} ({found.text(2)})")
        search_edit.returnPressed.connect(do_search)
        org_controls.addWidget(tb)
        org_controls.addStretch(1)
        org_layout.addLayout(org_controls)
        
        # Main org chart view (no splitter initially)
        # Data-only tree for hierarchy; hidden
        self.org_tree = QTreeWidget()
        self.org_tree.setHeaderLabels(["Name", "Designation", "Enterprise ID", "Email ID", "Manager Enterprise ID", "Location"])
        self.org_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.org_tree.customContextMenuRequested.connect(self._open_org_context)
        self.org_tree.setVisible(False)
        # Enable drag-drop to reparent with guard
        try:
            # Ensure we're in main thread before setting up drag-drop
            self.org_tree.setDragDropMode(QTreeWidget.DragDropMode.InternalMove)
            # Note: Event filter removed to prevent threading issues
        except Exception as e:
            self.notifications.show_warning(f"Failed to setup org tree drag-drop: {str(e)}")
            self._log_change("Error", "Org Chart", f"Failed to setup drag-drop: {str(e)}")
        # Visual org chart view
        self.org_scene = QGraphicsScene(self)
        self.org_view = QGraphicsView(self.org_scene)
        # Ensure correct render hints (use QPainter, not QPixmap)
        self.org_view.setRenderHints(self.org_view.renderHints() | QPainter.RenderHint.SmoothPixmapTransform | QPainter.RenderHint.Antialiasing)
        self.org_view.setBackgroundBrush(QColor("#f8fafc"))
        self.org_view.setDragMode(QGraphicsView.DragMode.ScrollHandDrag)
        self.org_view.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.org_view.customContextMenuRequested.connect(self._open_org_graph_context)
        org_layout.addWidget(self.org_view)
        
        # Store reference to org layout for dynamic splitter management
        self.org_layout = org_layout
        self.org_chart_widget = None  # Will be set when splitter is created
        # Map graphics item -> tree item for interactions
        self._org_item_map: dict = {}
        # Make the Org Chart tab label more prominent
        self.tabs.addTab(self.org_tab, "Org Chart ✨")

        # Force-pin critical tabs at indices 0 and 1 before adding the rest
        try:
            self._pin_critical_tabs_first()
        except Exception as e:
            ErrorHandler.handle_ui_error("pin critical tabs", e)

        # Data tabs
        self.tables: dict[str, QTableWidget] = {}
        
        # Force immediate tab creation for new panes (after tables is initialized)
        try:
            self._rebuild_all_tabs()
        except Exception as e:
            ErrorHandler.handle_ui_error("immediate tab rebuild", e)
        
        for pane_name, columns in PANE_COLUMNS.items():
            if pane_name == "Leave Tracker":
                # Leave Tracker uses a custom UI, add after creation
                continue
            elif pane_name == "Project Details":
                # Project Details uses a custom UI with dashboard, add after creation
                continue
            self._add_pane_tab(pane_name, columns)

        # Enforce final tab order immediately after all tabs are added
        try:
            self._enforce_tab_order()
        except Exception:
            pass

        # Persistence paths - use script directory instead of current working directory
        script_dir = os.path.dirname(os.path.abspath(__file__ if '__file__' in globals() else '.'))
        # Use a temporary SQLite autosave file (no JSON)
        self.autosave_path = os.path.join(tempfile.gettempdir(), "tracker_autosave.sqlite")
        # Prefer SQLite backend by default, but respect existing Excel file
        sqlite_default = os.path.join(script_dir, "backend_data.sqlite")
        xlsx_default = os.path.join(script_dir, "backend_data.xlsx")
        if not getattr(self, 'backend_sqlite_path', None):
            self.backend_sqlite_path = sqlite_default if os.path.exists(sqlite_default) or not os.path.exists(xlsx_default) else xlsx_default
        
        # Load backend (SQLite or Excel) first, then fallback to autosave if backend doesn't exist
        backend_loaded = False
        if os.path.exists(self.backend_sqlite_path):
            try:
                self._load_backend_sqlite(self.backend_sqlite_path)
                self.__backend_path__ = self.backend_sqlite_path
                backend_loaded = True
            except Exception as e:
                print(f"Failed to load backend data: {e}")
        else:
            # Create new backend file if it doesn't exist
            try:
                if str(self.backend_sqlite_path).lower().endswith((".sqlite", ".db")):
                    self._create_new_backend_sqlite(self.backend_sqlite_path)
                else:
                    self._create_new_backend_file()
                self.__backend_path__ = self.backend_sqlite_path
                backend_loaded = True
            except Exception as e:
                print(f"Failed to create backend file: {e}")
        
        # Only load autosave if backend wasn't loaded
        if not backend_loaded:
            self._load_autosave()
        
        # Auto-load backend data if path is remembered and not already loaded
        if not backend_loaded and hasattr(self, '__backend_path__') and self.__backend_path__ and os.path.exists(self.__backend_path__):
            try:
                self._load_backend_sqlite(self.__backend_path__)
            except Exception as e:
                print(f"Failed to auto-load backend data: {e}")

        # Apply light palette
        self.apply_light_palette()

        # Update ageing for Potential Issues
        self.update_ageing_for_potential_issues()
        
        # Initial dashboard
        self.update_dashboard()
        self.refresh_ism_filter()
        # Populate Project Details UI at startup if data present
        try:
            if hasattr(self, 'projects_table') and self.projects_table is not None:
                self._load_projects_data()
                if hasattr(self, '_update_summary_metrics'):
                    self._update_summary_metrics()
        except Exception:
            pass
        # Populate Leave Tracker UI at startup if data present
        try:
            if hasattr(self, 'leave_tab') and self.leave_tab is not None:
                self.refresh_leave_tracker()
        except Exception:
            pass

        # Initial notification count update
        self._pending_notifications = []
        self._update_notification_count()

    def _update_notification_count(self) -> None:
        try:
            count = len(getattr(self, '_pending_notifications', []))
            if count > 0:
                self.notification_count_badge.setText(str(min(count, 99)))
                self.notification_count_badge.show()
            else:
                self.notification_count_badge.hide()
        except Exception as e:
            ErrorHandler.handle_ui_error("update notification count", e)

        # Wiki tab will be added at the very end of initialization

        # Initial org chart render
        self.render_org_chart()
        # Auto-fit the org chart to the view size with intelligent padding
        try:
            # Reset to 1:1 scale for full size display
            self._org_zoom = 1.0
            self.org_view.resetTransform()
            
            # Auto-fit with intelligent padding based on screen resolution
            rect = self.org_scene.itemsBoundingRect()
            if rect.isValid():
                # Get view and screen dimensions
                view_size = self.org_view.size()
                screen = QApplication.primaryScreen()
                screen_geometry = screen.availableGeometry() if screen else QRect(0, 0, 1920, 1080)
                
                # Calculate optimal padding based on multiple factors
                min_view_dim = min(view_size.width(), view_size.height())
                max_view_dim = max(view_size.width(), view_size.height())
                
                # Determine screen size category with much smaller padding for better space utilization
                screen_area = screen_geometry.width() * screen_geometry.height()
                if screen_area >= 3840 * 2160:  # 4K and above
                    base_padding_percent = 0.05  # 5% for large screens - maximize content
                elif screen_area >= 2560 * 1440:  # 1440p
                    base_padding_percent = 0.04  # 4% for medium-large screens
                elif screen_area >= 1920 * 1080:  # 1080p
                    base_padding_percent = 0.03  # 3% for standard screens
                else:  # Smaller screens
                    base_padding_percent = 0.02  # 2% for small screens
                
                # Adjust padding based on view aspect ratio
                aspect_ratio = view_size.width() / view_size.height() if view_size.height() > 0 else 1.0
                if aspect_ratio > 2.0:  # Very wide view
                    base_padding_percent *= 0.8  # Reduce padding for wide views
                elif aspect_ratio < 0.5:  # Very tall view
                    base_padding_percent *= 0.8  # Reduce padding for tall views
                
                # Calculate dynamic padding with much smaller minimums for better space utilization
                dynamic_padding = max(5, int(min_view_dim * base_padding_percent))
                
                # Ensure reasonable bounds - much more generous maximums
                dynamic_padding = min(dynamic_padding, int(max_view_dim * 0.08))  # Max 8% of larger dimension
                dynamic_padding = max(dynamic_padding, 3)  # Minimum 3px for very small screens
                
                # Apply dynamic padding and fit
                padded = rect.adjusted(-dynamic_padding, -dynamic_padding, dynamic_padding, dynamic_padding)
                self.org_view.fitInView(padded, Qt.AspectRatioMode.KeepAspectRatio)
                
                # Center the view on the content
                self.org_view.centerOn(rect.center())
        except Exception as e:
            ErrorHandler.handle_ui_error("org chart auto-fit", e)

        # Leave Tracker tab
        self._init_leave_tracker_tab()
        # Ensure the Leave tab is actually present (guard against earlier failures)
        try:
            has_leave = False
            for i in range(self.tabs.count()):
                try:
                    if self.tabs.tabText(i).startswith("Leave Tracker"):
                        has_leave = True
                        break
                except Exception:
                    continue
            if not has_leave:
                # If initialization created self.leave_tab, add it; otherwise create a minimal fallback tab
                try:
                    leave_widget = getattr(self, 'leave_tab', None)
                    if leave_widget is None:
                        from PyQt6.QtWidgets import QWidget, QVBoxLayout, QLabel
                        leave_widget = QWidget()
                        _lay = QVBoxLayout(leave_widget)
                        _lay.addWidget(QLabel("Leave Tracker unavailable due to initialization error."))
                    # Insert at the front for maximum visibility
                    insert_at = 0
                    self.tabs.insertTab(insert_at, leave_widget, "Leave Tracker 🗓️")
                    # Ensure the tab is visible and enabled
                    self.tabs.setTabEnabled(insert_at, True)
                    if hasattr(self.tabs, 'setTabVisible'):
                        self.tabs.setTabVisible(insert_at, True)
                except Exception as e:
                    ErrorHandler.handle_ui_error("add leave tracker tab (guard)", e)
        except Exception as e:
            ErrorHandler.handle_ui_error("verify leave tracker tab (guard)", e)
        # After ensuring presence, try focusing it once so it becomes visible in UI
        try:
            for i in range(self.tabs.count()):
                if self.tabs.tabText(i).startswith("Leave Tracker"):
                    self.tabs.setCurrentIndex(i)
                    break
        except Exception:
            pass
        
        # Project Details tab
        self._init_project_details_tab()
        
        # Add Project Details tab to the tab widget
        try:
            found = False
            for i in range(self.tabs.count()):
                try:
                    txt = self.tabs.tabText(i)
                    if txt.startswith("Project Details"):
                        found = True
                        break
                except Exception:
                    continue
            if not found:
                # Insert right after Leave Tracker
                insert_at = 1
                self.tabs.insertTab(insert_at, self.project_details_tab, "Project Details 📋")
                # Ensure the tab is visible and enabled
                self.tabs.setTabEnabled(insert_at, True)
                if hasattr(self.tabs, 'setTabVisible'):
                    self.tabs.setTabVisible(insert_at, True)
        except Exception as e:
            ErrorHandler.handle_ui_error("add project details tab", e)


        # Guard: ensure Project Details tab is present even if initialization had issues
        print("About to start Project Details guard")
        try:
            has_project = False
            for i in range(self.tabs.count()):
                try:
                    if self.tabs.tabText(i).startswith("Project Details"):
                        has_project = True
                        break
                except Exception:
                    continue
            if not has_project:
                try:
                    pd_widget = getattr(self, 'project_details_tab', None)
                    if pd_widget is None:
                        from PyQt6.QtWidgets import QWidget, QVBoxLayout, QLabel
                        pd_widget = QWidget()
                        _pdl = QVBoxLayout(pd_widget)
                        _pdl.addWidget(QLabel("Project Details unavailable due to initialization error."))
                    insert_at = 1
                    self.tabs.insertTab(insert_at, pd_widget, "Project Details 📋")
                    # Ensure the tab is visible and enabled
                    self.tabs.setTabEnabled(insert_at, True)
                    if hasattr(self.tabs, 'setTabVisible'):
                        self.tabs.setTabVisible(insert_at, True)
                except Exception as e:
                    ErrorHandler.handle_ui_error("add project details tab (guard)", e)
        except Exception as e:
            ErrorHandler.handle_ui_error("verify project details tab (guard)", e)
        # After ensuring presence, focus Project Details once to surface it
        print("About to focus Project Details tab")
        try:
            for i in range(self.tabs.count()):
                if self.tabs.tabText(i).startswith("Project Details"):
                    self.tabs.setCurrentIndex(i)
                    break
        except Exception:
            pass
        print("Finished focusing Project Details tab")

        # Ensure critical tabs are always visible after startup
        # Use QTimer to ensure this runs after all UI initialization is complete
        print("About to start tab repair process")
        try:
            QTimer.singleShot(200, lambda: self._show_leave_tab())
            QTimer.singleShot(250, lambda: self._show_project_details_tab())
            # Additional delayed fallback to ensure visibility
            QTimer.singleShot(500, lambda: self._ensure_critical_tabs_visible())
            # Final check to ensure tabs are visible
            QTimer.singleShot(1000, lambda: self._final_tab_visibility_check())
            # Diagnostic and final repair: log, ensure, and pin tabs at positions 0 and 1
            QTimer.singleShot(1300, lambda: self._diagnose_and_repair_critical_tabs())
            # Full rebuild as last resort to guarantee all tabs are present and visible
            QTimer.singleShot(1700, lambda: self._full_tab_repair())
            # Enforce final order
            QTimer.singleShot(2000, lambda: self._enforce_tab_order())
        except Exception:
            # Fallback: call directly if QTimer fails
            try:
                self._show_leave_tab()
                self._show_project_details_tab()
                self._ensure_critical_tabs_visible()
                self._final_tab_visibility_check()
                self._diagnose_and_repair_critical_tabs()
                self._full_tab_repair()
                self._enforce_tab_order()
            except Exception:
                pass

        # System tray functionality removed - notifications now handled by Bell icon in header
        
        # Setup automatic log cleanup timer (runs every 30 days)
        try:
            self._setup_log_cleanup_timer()
        except Exception as e:
            ErrorHandler.handle_ui_error("log cleanup timer setup", e)
        
        # Setup auto-save timer
        try:
            self._setup_auto_save_timer()
        except Exception as e:
            ErrorHandler.handle_ui_error("auto-save timer setup", e)
        
        # Setup window state management
        try:
            self._setup_window_state()
        except Exception as e:
            ErrorHandler.handle_ui_error("window state setup", e)
        
        # Realtime backend watcher to reflect changes immediately for all users
        try:
            self._setup_backend_file_watcher()
        except Exception as e:
            ErrorHandler.handle_ui_error("backend file watcher setup", e)
        
        # Auto-refresh all panes periodically
        try:
            self._setup_auto_refresh_timer()
        except Exception as e:
            ErrorHandler.handle_ui_error("auto-refresh timer setup", e)
        

    def _push_notification(self, recipient_eid: str | None, message: str, status_ms: int = 6000) -> None:
        """Lightweight in-app notification. If the recipient is the current user, surface immediately; always add to bell badge queue."""
        try:
            # Normalize
            recipient_eid = (recipient_eid or "").strip().lower()
            cur = str(getattr(self, 'logged_in_user', '')).strip().lower()
            # Show immediately for the current user
            if recipient_eid and cur and recipient_eid == cur:
                try:
                    self.statusBar().showMessage(message, status_ms)
                except Exception:
                    pass
            # Queue for bell badge
            try:
                if not hasattr(self, '_pending_notifications') or self._pending_notifications is None:
                    self._pending_notifications = []
                self._pending_notifications.append(message)
                self._update_notification_count()
            except Exception:
                pass
        except Exception as e:
            try:
                ErrorHandler.handle_ui_error("push notification", e)
            except Exception:
                pass

    def _notify_leave_logged(self, date_str: str, ism_name: str, approver_eid: str, requester_eid: str) -> None:
        """Notify both approver (manager) and the requester/selected ISM when a leave is logged."""
        try:
            # Notify manager/approver
            approver_msg = f"Leave request logged for {date_str} by {ism_name or requester_eid}. Please review in Approvals."
            self._push_notification(approver_eid, approver_msg)
            # Notify requester/ISM
            req_msg = f"Your leave request for {date_str} has been submitted to your manager for approval."
            self._push_notification(requester_eid, req_msg)
        except Exception as e:
            ErrorHandler.handle_ui_error("notify leave logged", e)

    def _notify_leave_decision_to_requester(self, requester_eid: str, status: str, date_str: str, approver_name: str) -> None:
        """Notify requester once a decision is recorded (Approved/Rejected). Safe to call from anywhere."""
        try:
            status_clean = (status or "").strip().title() or "Updated"
            msg = f"Your leave request for {date_str} was {status_clean} by {approver_name}."
            self._push_notification(requester_eid, msg, status_ms=8000)
        except Exception as e:
            ErrorHandler.handle_ui_error("notify leave decision", e)

    def _show_leave_tab(self):
        """Ensure Leave Tracker tab exists, insert if missing, and focus it."""
        try:
            # Ensure widget exists
            if not hasattr(self, 'leave_tab') or self.leave_tab is None:
                try:
                    self._init_leave_tracker_tab()
                except Exception:
                    pass
            widget = getattr(self, 'leave_tab', None)
            if widget is None:
                return False
            # Ensure tab present
            idx = self.tabs.indexOf(widget)
            if idx < 0:
                try:
                    self.tabs.insertTab(0, widget, "Leave Tracker 🗓️")
                    idx = self.tabs.indexOf(widget)
                except Exception:
                    idx = -1
            # Fallback: search by title if different widget instance is present
            if idx < 0:
                for i in range(self.tabs.count()):
                    try:
                        if "leave tracker" in self.tabs.tabText(i).lower():
                            idx = i; break
                    except Exception:
                        continue
            if idx >= 0:
                # Ensure enabled/visible
                try:
                    self.tabs.setTabEnabled(idx, True)
                    if hasattr(self.tabs, 'setTabVisible'):
                        self.tabs.setTabVisible(idx, True)
                except Exception:
                    pass
                try:
                    self.tabs.setCurrentIndex(idx)
                except Exception:
                    try:
                        self.tabs.setCurrentWidget(widget)
                    except Exception:
                        pass
                return True
            return False
        except Exception as e:
            ErrorHandler.handle_ui_error("show leave tab", e)
            return False

    def _show_project_details_tab(self):
        """Ensure Project Details tab exists, insert if missing, and focus it."""
        try:
            # Ensure widget exists
            if not hasattr(self, 'project_details_tab') or self.project_details_tab is None:
                try:
                    self._init_project_details_tab()
                except Exception:
                    pass
            widget = getattr(self, 'project_details_tab', None)
            if widget is None:
                return False
            # Ensure tab present
            idx = self.tabs.indexOf(widget)
            if idx < 0:
                try:
                    pos = min(1, self.tabs.count())
                    self.tabs.insertTab(pos, widget, "Project Details 📋")
                    idx = self.tabs.indexOf(widget)
                except Exception:
                    idx = -1
            # Fallback: search by title if different widget instance is present
            if idx < 0:
                for i in range(self.tabs.count()):
                    try:
                        if "project details" in self.tabs.tabText(i).lower():
                            idx = i; break
                    except Exception:
                        continue
            if idx >= 0:
                # Ensure enabled/visible
                try:
                    self.tabs.setTabEnabled(idx, True)
                    if hasattr(self.tabs, 'setTabVisible'):
                        self.tabs.setTabVisible(idx, True)
                except Exception:
                    pass
                try:
                    self.tabs.setCurrentIndex(idx)
                except Exception:
                    try:
                        self.tabs.setCurrentWidget(widget)
                    except Exception:
                        pass
                return True
            return False
        except Exception as e:
            ErrorHandler.handle_ui_error("show project details tab", e)
            return False

    def _ensure_critical_tabs_visible(self):
        """Final safeguard to ensure critical tabs are always visible and accessible"""
        try:
            # Check if Leave Tracker tab exists and is visible
            leave_found = False
            project_found = False
            
            for i in range(self.tabs.count()):
                try:
                    tab_text = self.tabs.tabText(i)
                    if tab_text.startswith("Leave Tracker"):
                        leave_found = True
                        # Ensure tab is enabled and visible
                        self.tabs.setTabEnabled(i, True)
                        if hasattr(self.tabs, 'setTabVisible'):
                            self.tabs.setTabVisible(i, True)
                    elif tab_text.startswith("Project Details"):
                        project_found = True
                        # Ensure tab is enabled and visible
                        self.tabs.setTabEnabled(i, True)
                        if hasattr(self.tabs, 'setTabVisible'):
                            self.tabs.setTabVisible(i, True)
                except Exception:
                    continue
            
            # If tabs are missing, try to add them again
            if not leave_found:
                try:
                    if hasattr(self, 'leave_tab') and self.leave_tab:
                        insert_at = min(0, self.tabs.count())
                        self.tabs.insertTab(insert_at, self.leave_tab, "Leave Tracker 🗓️")
                except Exception as e:
                    ErrorHandler.handle_ui_error("ensure leave tab visible", e)
            
            if not project_found:
                try:
                    if hasattr(self, 'project_details_tab') and self.project_details_tab:
                        insert_at = min(1, self.tabs.count())
                        self.tabs.insertTab(insert_at, self.project_details_tab, "Project Details 📋")
                except Exception as e:
                    ErrorHandler.handle_ui_error("ensure project details tab visible", e)
            
            # Finally, try to show Leave Tracker tab by default
            if leave_found or project_found:
                self._show_leave_tab()
                
        except Exception as e:
            ErrorHandler.handle_ui_error("ensure critical tabs visible", e)
        finally:
            try:
                self._enforce_tab_order()
            except Exception:
                pass

    def _final_tab_visibility_check(self):
        """Final check to ensure tabs are visible and properly displayed"""
        try:
            # Force refresh of tab visibility
            for i in range(self.tabs.count()):
                try:
                    tab_text = self.tabs.tabText(i)
                    if tab_text.startswith("Leave Tracker") or tab_text.startswith("Project Details"):
                        # Ensure tab is visible and enabled
                        self.tabs.setTabEnabled(i, True)
                        if hasattr(self.tabs, 'setTabVisible'):
                            self.tabs.setTabVisible(i, True)
                        # Force a repaint
                        self.tabs.update()
                except Exception:
                    continue
            
            # Try to show Leave Tracker tab
            self._show_leave_tab()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("final tab visibility check", e)
        finally:
            try:
                self._enforce_tab_order()
            except Exception:
                pass

    def _pin_critical_tabs_first(self):
        """Unconditionally create and insert Leave Tracker (0) and Project Details (1)."""
        try:
            # Ensure widgets exist
            try:
                if not hasattr(self, 'leave_tab') or self.leave_tab is None:
                    self._init_leave_tracker_tab()
            except Exception:
                pass
            try:
                if not hasattr(self, 'project_details_tab') or self.project_details_tab is None:
                    self._init_project_details_tab()
            except Exception:
                pass

            # Remove any existing instances to avoid duplicates
            try:
                to_remove = []
                for i in range(self.tabs.count()):
                    t = self.tabs.tabText(i)
                    if t.lower().startswith("leave tracker") or t.lower().startswith("project details"):
                        to_remove.append(i)
                for idx in reversed(to_remove):
                    self.tabs.removeTab(idx)
            except Exception:
                pass

            # Insert pinned (temporarily at the end; final order enforced below)
            if hasattr(self, 'leave_tab') and self.leave_tab is not None:
                pos_end = self.tabs.count()
                self.tabs.insertTab(pos_end, self.leave_tab, "Leave Tracker 🗓️")
            if hasattr(self, 'project_details_tab') and self.project_details_tab is not None:
                pos_end = self.tabs.count()
                self.tabs.insertTab(pos_end, self.project_details_tab, "Project Details 📋")

            # Ensure enabled/visible
            for i in range(self.tabs.count()):
                try:
                    self.tabs.setTabEnabled(i, True)
                    if hasattr(self.tabs, 'setTabVisible'):
                        self.tabs.setTabVisible(i, True)
                except Exception:
                    pass
            # Enforce final order now
            try:
                self._enforce_tab_order()
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("pin critical tabs first", e)

    def _find_tab_index_by_prefix(self, prefix: str) -> int:
        try:
            p = prefix.lower().strip()
            for i in range(self.tabs.count()):
                try:
                    txt = self.tabs.tabText(i)
                    # Match by containment to tolerate emojis/prefix icons
                    if p in (txt or "").lower():
                        return i
                except Exception:
                    continue
        except Exception:
            pass
        return -1

    def _enforce_tab_order(self):
        """Reorder tabs to the exact order requested by the user."""
        try:
            desired = [
                "Home",
                "Dashboard",
                "Org Chart",
                "Initiatives",
                "Potential Issues",
                "Activities",
                "Client Visits / Audits",
                "Accolades",
            ]
            # Compute widgets in desired order first
            # Build an ordered list of (title, widget)
            ordered_widgets = []
            seen = set()
            for name in desired:
                idx = self._find_tab_index_by_prefix(name)
                if idx >= 0:
                    try:
                        w = self.tabs.widget(idx)
                        ordered_widgets.append((name, w))
                        seen.add(idx)
                    except Exception:
                        pass
            # Append Leave Tracker and Project Details at the end if present
            for tail in ["Leave Tracker", "Project Details"]:
                idx = self._find_tab_index_by_prefix(tail)
                if idx >= 0 and idx not in seen:
                    try:
                        w = self.tabs.widget(idx)
                        ordered_widgets.append((tail, w))
                        seen.add(idx)
                    except Exception:
                        pass
            # Append any leftover tabs in their current order
            for i in range(self.tabs.count()):
                if i in seen:
                    continue
                try:
                    w = self.tabs.widget(i)
                    t = self.tabs.tabText(i)
                    ordered_widgets.append((t, w))
                except Exception:
                    continue
            # Rebuild in desired order with original labels (preserve emojis)
            try:
                original_titles = [self.tabs.tabText(i) for i in range(self.tabs.count())]
                original_widgets = [self.tabs.widget(i) for i in range(self.tabs.count())]
            except Exception:
                original_titles = []
                original_widgets = []
            try:
                while self.tabs.count() > 0:
                    self.tabs.removeTab(0)
            except Exception:
                pass
            for title, widget in ordered_widgets:
                try:
                    # If this widget existed before, reuse its original title to keep icons/emojis
                    label = title
                    try:
                        idx = original_widgets.index(widget) if widget in original_widgets else -1
                        if idx >= 0 and idx < len(original_titles):
                            label = original_titles[idx]
                    except Exception:
                        pass
                    self.tabs.addTab(widget, label)
                except Exception:
                    continue
        except Exception as e:
            ErrorHandler.handle_ui_error("enforce tab order", e)

    def showEvent(self, event) -> None:
        """Ensure critical tabs are visible once the window is shown."""
        try:
            super().showEvent(event)
        except Exception:
            try:
                QWidget.showEvent(self, event)
            except Exception:
                pass
        try:
            if not getattr(self, '_did_tab_autoshow', False):
                self._did_tab_autoshow = True
                # Run immediately after the window is visible
                try:
                    QTimer.singleShot(0, lambda: self._full_tab_repair())
                except Exception:
                    try:
                        self._full_tab_repair()
                    except Exception:
                        pass
        except Exception as e:
            ErrorHandler.handle_ui_error("showEvent autoshow tabs", e)

    def _diagnose_and_repair_critical_tabs(self):
        """Diagnose current tabs, log titles, and repair/pin critical tabs deterministically."""
        try:
            titles_before = []
            try:
                titles_before = [self.tabs.tabText(i) for i in range(self.tabs.count())]
            except Exception:
                pass

            # Ensure Leave Tracker widget exists
            leave_widget = getattr(self, 'leave_tab', None)
            if leave_widget is None:
                try:
                    self._init_leave_tracker_tab()
                    leave_widget = getattr(self, 'leave_tab', None)
                except Exception:
                    leave_widget = None

            # Ensure Project Details widget exists
            pd_widget = getattr(self, 'project_details_tab', None)
            if pd_widget is None:
                try:
                    self._init_project_details_tab()
                    pd_widget = getattr(self, 'project_details_tab', None)
                except Exception:
                    pd_widget = None

            # Remove any existing instances to avoid duplicates, then insert pinned
            try:
                # Collect indices to remove (from end to start)
                to_remove = []
                for i in range(self.tabs.count()):
                    txt = self.tabs.tabText(i)
                    if txt.startswith("Leave Tracker") or txt.startswith("Project Details"):
                        to_remove.append(i)
                for idx in reversed(to_remove):
                    self.tabs.removeTab(idx)
            except Exception:
                pass

            # Insert pinned at 0 and 1 if widgets exist
            try:
                if leave_widget is not None:
                    self.tabs.insertTab(0, leave_widget, "Leave Tracker 🗓️")
                    self.tabs.setTabEnabled(0, True)
                    if hasattr(self.tabs, 'setTabVisible'):
                        self.tabs.setTabVisible(0, True)
            except Exception as e:
                ErrorHandler.handle_ui_error("repair: insert leave tab", e)
            try:
                if pd_widget is not None:
                    pos = min(1, self.tabs.count())
                    self.tabs.insertTab(pos, pd_widget, "Project Details 📋")
                    self.tabs.setTabEnabled(pos, True)
                    if hasattr(self.tabs, 'setTabVisible'):
                        self.tabs.setTabVisible(pos, True)
            except Exception as e:
                ErrorHandler.handle_ui_error("repair: insert project tab", e)

            # Ensure scroll buttons on overflow
            try:
                tb = self.tabs.tabBar()
                if tb is not None:
                    tb.setUsesScrollButtons(True)
                    tb.setExpanding(False)
            except Exception:
                pass

            # Log after state
            try:
                titles_after = [self.tabs.tabText(i) for i in range(self.tabs.count())]
            except Exception:
                titles_after = []
            try:
                print(f"Tabs before repair: {titles_before}")
                print(f"Tabs after repair:  {titles_after}")
            except Exception:
                pass

            # Focus Leave Tracker by default, then bounce to Project Details to ensure both surface
            try:
                self._show_leave_tab()
                self._show_project_details_tab()
                self._show_leave_tab()
            except Exception:
                pass
            # Enforce final order post-repair
            try:
                self._enforce_tab_order()
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("diagnose and repair tabs", e)

    def _ensure_tab_widget_attached(self):
        """Ensure the QTabWidget is attached to the main layout and its tab bar is visible."""
        try:
            try:
                if self.tabs.parent() is None:
                    self.root_layout.addWidget(self.tabs)
            except Exception:
                # Fallback: always try to add; Qt will ignore if already added
                try:
                    self.root_layout.addWidget(self.tabs)
                except Exception:
                    pass
            try:
                tb = self.tabs.tabBar()
                if tb is not None:
                    tb.setVisible(True)
                    tb.setUsesScrollButtons(True)
                    tb.setExpanding(False)
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("ensure tab widget attached", e)

    def _rebuild_all_tabs(self):
        """Clear and re-add all tabs based on PANE_COLUMNS, pinning custom tabs first."""
        try:
            self._ensure_tab_widget_attached()
            # Remove everything
            try:
                while self.tabs.count() > 0:
                    self.tabs.removeTab(0)
            except Exception:
                pass

            # Create or reuse custom widgets
            try:
                if not hasattr(self, 'home_tab') or self.home_tab is None:
                    self._init_home_page()
            except Exception:
                pass
            try:
                if not hasattr(self, 'dashboard_tab') or self.dashboard_tab is None:
                    self._init_dashboard_tab()
            except Exception:
                pass
            try:
                if not hasattr(self, 'org_tab') or self.org_tab is None:
                    self._init_org_chart_tab()
            except Exception:
                pass
            try:
                if not hasattr(self, 'leave_tab') or self.leave_tab is None:
                    self._init_leave_tracker_tab()
            except Exception:
                pass
            try:
                if not hasattr(self, 'project_details_tab') or self.project_details_tab is None:
                    self._init_project_details_tab()
            except Exception:
                pass
            try:
                if not hasattr(self, 'kanban_tab') or self.kanban_tab is None:
                    self._init_kanban_tab()
            except Exception:
                pass
            try:
                if not hasattr(self, 'calendar_tab') or self.calendar_tab is None:
                    self._init_calendar_tab()
            except Exception:
                pass

            # Insert custom tabs first, pinned
            try:
                if hasattr(self, 'home_tab') and self.home_tab is not None:
                    self.tabs.insertTab(0, self.home_tab, "🏠 Home")
            except Exception as e:
                ErrorHandler.handle_ui_error("rebuild: insert home tab", e)
            try:
                if hasattr(self, 'dashboard_tab') and self.dashboard_tab is not None:
                    pos = min(1, self.tabs.count())
                    self.tabs.insertTab(pos, self.dashboard_tab, "📊 Dashboard")
            except Exception as e:
                ErrorHandler.handle_ui_error("rebuild: insert dashboard tab", e)
            try:
                if hasattr(self, 'org_tab') and self.org_tab is not None:
                    pos = min(2, self.tabs.count())
                    self.tabs.insertTab(pos, self.org_tab, "Org Chart ✨")
            except Exception as e:
                ErrorHandler.handle_ui_error("rebuild: insert org chart tab", e)
            try:
                if hasattr(self, 'leave_tab') and self.leave_tab is not None:
                    pos = min(3, self.tabs.count())
                    self.tabs.insertTab(pos, self.leave_tab, "Leave Tracker 🗓️")
            except Exception as e:
                ErrorHandler.handle_ui_error("rebuild: insert leave tab", e)
            try:
                if hasattr(self, 'project_details_tab') and self.project_details_tab is not None:
                    pos = min(4, self.tabs.count())
                    self.tabs.insertTab(pos, self.project_details_tab, "Project Details 📋")
            except Exception as e:
                ErrorHandler.handle_ui_error("rebuild: insert project tab", e)
            try:
                if hasattr(self, 'kanban_tab') and self.kanban_tab is not None:
                    pos = min(5, self.tabs.count())
                    self.tabs.insertTab(pos, self.kanban_tab, "Kanban Board 📋")
            except Exception as e:
                ErrorHandler.handle_ui_error("rebuild: insert kanban tab", e)
            try:
                if hasattr(self, 'calendar_tab') and self.calendar_tab is not None:
                    pos = min(6, self.tabs.count())
                    self.tabs.insertTab(pos, self.calendar_tab, "Calendar View 📅")
            except Exception as e:
                ErrorHandler.handle_ui_error("rebuild: insert calendar tab", e)

            # Note: Generic panes are handled by the regular initialization loop
            # This method only handles custom tabs (Leave Tracker, Project Details, Kanban, Calendar)

            # Ensure visibility and enablement for all tabs
            try:
                for i in range(self.tabs.count()):
                    self.tabs.setTabEnabled(i, True)
                    if hasattr(self.tabs, 'setTabVisible'):
                        self.tabs.setTabVisible(i, True)
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("rebuild all tabs", e)

    def _full_tab_repair(self):
        """Final, comprehensive repair: attach tab widget, rebuild all tabs, and focus critical ones."""
        try:
            self._ensure_tab_widget_attached()
            self._rebuild_all_tabs()
            # Focus both once
            try:
                self._show_leave_tab()
                self._show_project_details_tab()
            except Exception:
                pass
            # Log final tab titles
            try:
                titles = [self.tabs.tabText(i) for i in range(self.tabs.count())]
                print(f"Tabs after full repair: {titles}")
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("full tab repair", e)

    def _init_leave_tracker_tab(self):
        """Initialize the Leave Tracker tab with custom UI"""
        try:
            # Create the Leave Tracker tab widget
            self.leave_tab = QWidget()
            layout = QVBoxLayout(self.leave_tab)
            layout.setSpacing(10)
            layout.setContentsMargins(10, 10, 10, 10)
            
            # Add title with Refresh button on the right
            title_label = QLabel("Leave Tracker 🗓️")
            title_label.setStyleSheet("""
                QLabel {
                    font-size: 18px;
                    font-weight: bold;
                    color: #2c3e50;
                    margin: 10px 0;
                }
            """)
            header_row = QHBoxLayout()
            header_row.addWidget(title_label)
            header_row.addStretch()
            refresh_btn_top = QPushButton("Refresh")
            refresh_btn_top.clicked.connect(self.refresh_leave_tracker)
            header_row.addWidget(refresh_btn_top)
            layout.addLayout(header_row)
            
            # Add calendar widget
            self.leave_calendar = LeaveCalendar(self)
            layout.addWidget(self.leave_calendar)
            
            # Add controls
            controls_layout = QHBoxLayout()
            
            # ISM filter
            ism_label = QLabel("Filter by ISM:")
            controls_layout.addWidget(ism_label)
            
            self.leave_ism_combo = QComboBox()
            self.leave_ism_combo.addItem("All ISMs")
            controls_layout.addWidget(self.leave_ism_combo)
            
            # Add/Edit/Delete buttons
            add_btn = QPushButton("Add Leave")
            add_btn.clicked.connect(self._add_leave_entry)
            controls_layout.addWidget(add_btn)
            
            edit_btn = QPushButton("Edit Leave")
            edit_btn.clicked.connect(self._edit_leave_entry)
            controls_layout.addWidget(edit_btn)
            
            delete_btn = QPushButton("Delete Leave")
            delete_btn.clicked.connect(self._delete_leave_entry)
            controls_layout.addWidget(delete_btn)
            
            # Refresh button
            refresh_btn = QPushButton("Refresh")
            refresh_btn.clicked.connect(self.refresh_leave_tracker)
            controls_layout.addWidget(refresh_btn)
            
            controls_layout.addStretch()
            layout.addLayout(controls_layout)
            
            # Add table for leave entries
            self.leave_table = QTableWidget()
            self.leave_table.setColumnCount(5)
            self.leave_table.setHorizontalHeaderLabels(["Date", "Type", "Duration", "Description", "ISM Name"])
            self.leave_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            layout.addWidget(self.leave_table)
            
            # Populate ISM combo
            self._populate_leave_ism_combo()
            
            # Refresh the table with existing data
            self._refresh_leave_table()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("init leave tracker tab", e)
            # Create fallback tab
            self.leave_tab = QWidget()
            layout = QVBoxLayout(self.leave_tab)
            layout.addWidget(QLabel("Leave Tracker unavailable due to initialization error."))

    def refresh_leave_tracker(self) -> None:
        """Refresh Leave Tracker UI elements and data safely."""
        try:
            # Repopulate ISM filter list
            self._populate_leave_ism_combo()
        except Exception as e:
            ErrorHandler.handle_ui_error("refresh leave tracker (combo)", e)
        try:
            # Refresh table rows
            self._refresh_leave_table()
        except Exception as e:
            ErrorHandler.handle_ui_error("refresh leave tracker (table)", e)
        try:
            # If calendar supports refresh, invoke it
            if hasattr(self, 'leave_calendar') and hasattr(self.leave_calendar, 'refresh'):
                self.leave_calendar.refresh()
        except Exception as e:
            ErrorHandler.handle_ui_error("refresh leave tracker (calendar)", e)

    def _auto_refresh_all_panes(self) -> None:
        """Auto-refresh all panes, including custom tabs, in a lightweight way."""
        try:
            # Refresh generic data panes
            try:
                for pane_name in getattr(self, 'data', {}) or {}:
                    if pane_name in ("Leave Tracker", "Project Details"):
                        continue
                    try:
                        if pane_name in PANE_COLUMNS:
                            self.rebuild_table(pane_name)
                    except Exception as pane_err:
                        ErrorHandler.handle_ui_error(f"auto-refresh pane: {pane_name}", pane_err)
            except Exception:
                pass

            # Refresh custom tabs
            try:
                self.refresh_leave_tracker()
            except Exception as e:
                ErrorHandler.handle_ui_error("auto-refresh leave tracker", e)
            try:
                self._refresh_project_table()
            except Exception as e:
                ErrorHandler.handle_ui_error("auto-refresh project details", e)

            # Refresh dashboard and home stats if available
            try:
                if hasattr(self, 'update_dashboard'):
                    self.update_dashboard()
            except Exception:
                pass
            try:
                if hasattr(self, 'update_home_stats'):
                    self.update_home_stats()
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("auto-refresh all panes", e)

    def _setup_auto_refresh_timer(self) -> None:
        """Set up a periodic auto-refresh timer for all panes."""
        try:
            # Reuse existing timer if present
            if hasattr(self, 'auto_refresh_timer') and self.auto_refresh_timer:
                try:
                    self.auto_refresh_timer.stop()
                    self.auto_refresh_timer.deleteLater()
                except Exception:
                    pass
            self.auto_refresh_timer = QTimer(self)
            # Default interval: 60 seconds
            self.auto_refresh_timer.setInterval(60 * 1000)
            self.auto_refresh_timer.timeout.connect(self._auto_refresh_all_panes)
            self.auto_refresh_timer.start()
            # Track for cleanup
            try:
                self._active_timers.append(self.auto_refresh_timer)
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("setup auto refresh timer", e)

    def refresh_all_panes_now(self) -> None:
        """Manually refresh/rebuild all panes and custom tabs."""
        try:
            # Rebuild generic panes
            for pane_name in PANE_COLUMNS.keys():
                if pane_name in ("Leave Tracker", "Project Details"):
                    continue
                try:
                    self.rebuild_table(pane_name)
                except Exception as pane_err:
                    ErrorHandler.handle_ui_error(f"refresh all: {pane_name}", pane_err)
            # Refresh custom tabs
            try:
                if hasattr(self, 'projects_table') and self.projects_table is not None:
                    self._load_projects_data()
            except Exception as e:
                ErrorHandler.handle_ui_error("refresh all: project details", e)
            try:
                self.refresh_leave_tracker()
            except Exception as e:
                ErrorHandler.handle_ui_error("refresh all: leave tracker", e)
            # Update dashboard and filters
            try:
                self.update_dashboard()
            except Exception:
                pass
            try:
                self.refresh_ism_filter()
            except Exception:
                pass
            try:
                if hasattr(self, '_update_summary_metrics'):
                    self._update_summary_metrics()
            except Exception:
                pass
            # Status
            try:
                self._show_toast("All panes refreshed", level="SUCCESS")
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("refresh all panes now", e)

    def _init_project_details_tab(self):
        """Initialize the Project Details tab with custom UI"""
        try:
            # Create the Project Details tab widget
            self.project_details_tab = QWidget()
            layout = QVBoxLayout(self.project_details_tab)
            layout.setSpacing(10)
            layout.setContentsMargins(10, 10, 10, 10)
            
            # Add title
            title_label = QLabel("Project Details 📋")
            title_label.setStyleSheet("""
                QLabel {
                    font-size: 18px;
                    font-weight: bold;
                    color: #2c3e50;
                    margin: 10px 0;
                }
            """)
            layout.addWidget(title_label)
            
            # Add controls
            controls_layout = QHBoxLayout()
            
            # ISM filter
            ism_label = QLabel("Filter by ISM:")
            controls_layout.addWidget(ism_label)
            
            self.project_ism_combo = QComboBox()
            self.project_ism_combo.addItem("All ISMs")
            controls_layout.addWidget(self.project_ism_combo)
            
            # Add/Edit/Delete buttons
            add_btn = QPushButton("Add Project")
            add_btn.clicked.connect(self._add_project)
            controls_layout.addWidget(add_btn)
            
            edit_btn = QPushButton("Edit Project")
            edit_btn.clicked.connect(self._edit_project)
            controls_layout.addWidget(edit_btn)
            
            delete_btn = QPushButton("Delete Project")
            delete_btn.clicked.connect(self._delete_project)
            controls_layout.addWidget(delete_btn)
            
            controls_layout.addStretch()
            layout.addLayout(controls_layout)
            
            # Add table for project details
            self.projects_table = QTableWidget()
            self.projects_table.setColumnCount(17)
            self.projects_table.setHorizontalHeaderLabels([
                "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM", 
                "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type", 
                "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
            ])
            self.projects_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            layout.addWidget(self.projects_table)
            
            # Populate ISM combo
            self._populate_project_ism_combo()
            
            # Refresh the table with existing data
            self._refresh_project_table()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("init project details tab", e)
            # Create fallback tab
            self.project_details_tab = QWidget()
            layout = QVBoxLayout(self.project_details_tab)
            layout.addWidget(QLabel("Project Details unavailable due to initialization error."))

    def _populate_leave_ism_combo(self):
        """Populate the Leave Tracker ISM combo box"""
        try:
            if not hasattr(self, 'leave_ism_combo'):
                return
            self.leave_ism_combo.clear()
            self.leave_ism_combo.addItem("All ISMs")
            
            # Get ISM names from org directory
            ism_names = set()
            try:
                org_rows = self._collect_org_directory_rows()
                for (name, _, _, _, _, _) in org_rows:
                    if name:
                        ism_names.add(name)
            except Exception:
                pass
            
            # Add ISM names to combo
            for name in sorted(ism_names):
                self.leave_ism_combo.addItem(name)
                
        except Exception as e:
            ErrorHandler.handle_ui_error("populate leave ism combo", e)

    def _populate_project_ism_combo(self):
        """Populate the Project Details ISM combo box"""
        try:
            if not hasattr(self, 'project_ism_combo'):
                return
            self.project_ism_combo.clear()
            self.project_ism_combo.addItem("All ISMs")
            
            # Get ISM names from org directory
            ism_names = set()
            try:
                org_rows = self._collect_org_directory_rows()
                for (name, _, _, _, _, _) in org_rows:
                    if name:
                        ism_names.add(name)
            except Exception:
                pass
            
            # Add ISM names to combo
            for name in sorted(ism_names):
                self.project_ism_combo.addItem(name)
                
        except Exception as e:
            ErrorHandler.handle_ui_error("populate project ism combo", e)

    def _init_kanban_tab(self):
        """Initialize the Kanban Board tab"""
        try:
            self.kanban_tab = KanbanPane(self, self.logged_in_user)
        except Exception as e:
            ErrorHandler.handle_ui_error("init kanban tab", e)
            # Create fallback tab
            self.kanban_tab = QWidget()
            layout = QVBoxLayout(self.kanban_tab)
            layout.addWidget(QLabel("Kanban Board unavailable due to initialization error."))

    def _init_calendar_tab(self):
        """Initialize the Calendar View tab"""
        try:
            self.calendar_tab = CalendarPane(self, self.logged_in_user)
        except Exception as e:
            ErrorHandler.handle_ui_error("init calendar tab", e)
            # Create fallback tab
            self.calendar_tab = QWidget()
            layout = QVBoxLayout(self.calendar_tab)
            layout.addWidget(QLabel("Calendar View unavailable due to initialization error."))

    def _init_dashboard_tab(self):
        """Initialize the Dashboard tab"""
        try:
            self.dashboard_tab = QWidget()
            self.dashboard_layout = QVBoxLayout(self.dashboard_tab)
            self.cards_bar = QHBoxLayout()
            self.dashboard_layout.addLayout(self.cards_bar)
            # Dashboard filters
            dash_filters = QHBoxLayout()
            self.ism_filter = QComboBox()
            self.ism_filter.addItem("All ISMs")
            self.ism_filter.currentTextChanged.connect(lambda _=None: self.update_dashboard())
            dash_filters.addWidget(QLabel("ISM:"))
            dash_filters.addWidget(self.ism_filter)
            dash_filters.addStretch(1)
            self.dashboard_layout.addLayout(dash_filters)
            self.chart_holder = QWidget()
            self.chart_layout = QVBoxLayout(self.chart_holder)
            self.dashboard_layout.addWidget(self.chart_holder)
            # Status bar chart holder
            self.chart_holder2 = QWidget()
            self.chart_layout2 = QVBoxLayout(self.chart_holder2)
            self.dashboard_layout.addWidget(self.chart_holder2)
            # Overdue table (hidden by default as per feedback)
            self.overdue_table = None
            # Remove in-dashboard activity list; add under View menu
            # Client Visits / Audits consolidated table
            self.dashboard_layout.addWidget(QLabel("Client Visits / Audits - In Progress & Next 5 Days"))
            # Quick filters
            visits_filters = QHBoxLayout()
            self.visits_only_inprogress = QCheckBox("Only In Progress")
            self.visits_audit_type = QComboBox()
            self.visits_audit_type.addItems(["All Types", "Internal - Tech Scope", "Internal - No Scope", "External - Tech Scope", "External - No Scope", "Client Visit - Tech Scope", "Client Visit - No Scope"]) 
            visits_filters.addWidget(self.visits_only_inprogress)
            visits_filters.addWidget(QLabel("Audit Type:"))
            visits_filters.addWidget(self.visits_audit_type)
            visits_filters.addStretch(1)
            self.dashboard_layout.addLayout(visits_filters)
            self.visits_table = QTableWidget()
            self.visits_table.setColumnCount(7)
            self.visits_table.setHorizontalHeaderLabels(["Project Name", "Audit Type", "ISM Name", "Start Date", "End Date", "Status", "RAG"])
            self.visits_table.setSortingEnabled(True)
            self.visits_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            self.dashboard_layout.addWidget(self.visits_table)
        except Exception as e:
            ErrorHandler.handle_ui_error("init dashboard tab", e)
            # Create fallback tab
            self.dashboard_tab = QWidget()
            layout = QVBoxLayout(self.dashboard_tab)
            layout.addWidget(QLabel("Dashboard unavailable due to initialization error."))

    def _init_org_chart_tab(self):
        """Initialize the Org Chart tab"""
        try:
            self.org_tab = QWidget()
            org_layout = QVBoxLayout(self.org_tab)
            org_layout.setContentsMargins(6, 6, 6, 6)
            org_layout.setSpacing(8)
            org_controls = QHBoxLayout()
            add_member_btn = QPushButton("Add Team Member")
            add_member_btn.setObjectName("primary")
            add_member_btn.clicked.connect(self.add_org_member)
            org_controls.addWidget(add_member_btn)
            
            # Color scheme controls
            org_controls.addWidget(QLabel("Color Scheme:"))
            self.color_scheme_combo = QComboBox()
            self.color_scheme_combo.addItems(["Level-based", "Designation-based", "Combined"])
            self.color_scheme_combo.setCurrentText("Level-based")
            self.color_scheme_combo.currentTextChanged.connect(self._on_color_scheme_changed)
            org_controls.addWidget(self.color_scheme_combo)
            
            # Color scheme info button
            color_info_btn = QPushButton("ℹ️")
            color_info_btn.setToolTip("Click to see color scheme information")
            color_info_btn.setStyleSheet("""
                QPushButton {
                    background-color: #f8f9fa;
                    border: 1px solid #dee2e6;
                    border-radius: 4px;
                    padding: 4px 8px;
                    font-size: 12px;
                }
                QPushButton:hover {
                    background-color: #e9ecef;
                }
            """)
            color_info_btn.clicked.connect(self._show_color_scheme_info)
            org_controls.addWidget(color_info_btn)
            
            # Add Export/Import buttons
            export_btn = QPushButton("📤 Export")
            export_btn.setObjectName("secondary")
            export_btn.clicked.connect(self._export_org_chart)
            org_controls.addWidget(export_btn)
            
            import_btn = QPushButton("📥 Import")
            import_btn.setObjectName("secondary")
            import_btn.clicked.connect(self._import_org_chart)
            org_controls.addWidget(import_btn)
            
            org_controls.addStretch(1)
            org_layout.addLayout(org_controls)
            
            # Splitter for tree and chart
            self.org_splitter = QSplitter(Qt.Orientation.Horizontal)
            org_layout.addWidget(self.org_splitter)
            
            # Tree view for org structure
            self.org_tree = QTreeWidget()
            self.org_tree.setHeaderLabels(["Name", "Designation", "Enterprise ID", "Manager EID", "Level", "Team"])
            self.org_tree.setAlternatingRowColors(True)
            self.org_tree.setRootIsDecorated(True)
            self.org_tree.setSortingEnabled(True)
            self.org_tree.itemDoubleClicked.connect(self._on_org_item_double_clicked)
            self.org_splitter.addWidget(self.org_tree)
            
            # Note: Event filter removed to prevent threading issues
            try:
                self.notifications.show_warning(f"Failed to setup org tree drag-drop: {str(e)}")
                self._log_change("Error", "Org Chart", f"Failed to setup drag-drop: {str(e)}")
            except Exception:
                pass
            # Visual org chart view
            self.org_scene = QGraphicsScene(self)
            self.org_view = QGraphicsView(self.org_scene)
            self.org_view.setRenderHint(QPainter.RenderHint.Antialiasing)
            self.org_view.setDragMode(QGraphicsView.DragMode.RubberBandDrag)
            self.org_view.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
            self.org_view.customContextMenuRequested.connect(self._on_org_chart_context_menu)
            self.org_splitter.addWidget(self.org_view)
            
            # Set splitter proportions (tree: 40%, chart: 60%)
            self.org_splitter.setSizes([400, 600])
            
            # Store reference for later use
            self.org_chart_widget = None  # Will be set when splitter is created
            # Map graphics item -> tree item for interactions
            self._org_item_map: dict = {}
        except Exception as e:
            ErrorHandler.handle_ui_error("init org chart tab", e)
            # Create fallback tab
            self.org_tab = QWidget()
            layout = QVBoxLayout(self.org_tab)
            layout.addWidget(QLabel("Org Chart unavailable due to initialization error."))

    def _add_leave_entry(self):
        """Add a new leave entry"""
        try:
            # Simple dialog for adding leave entry
            dialog = QDialog(self)
            dialog.setWindowTitle("Add Leave Entry")
            layout = QVBoxLayout(dialog)
            
            # Date input
            date_layout = QHBoxLayout()
            date_layout.addWidget(QLabel("Date:"))
            date_edit = QLineEdit()
            date_edit.setPlaceholderText("YYYY-MM-DD")
            date_layout.addWidget(date_edit)
            layout.addLayout(date_layout)
            
            # Type input
            type_layout = QHBoxLayout()
            type_layout.addWidget(QLabel("Type:"))
            type_combo = QComboBox()
            type_combo.addItems(["WFH", "Planned Leave", "Public Holiday", "Earned Leave", "Casual Leave"])
            type_layout.addWidget(type_combo)
            layout.addLayout(type_layout)
            
            # Duration input
            duration_layout = QHBoxLayout()
            duration_layout.addWidget(QLabel("Duration:"))
            duration_combo = QComboBox()
            duration_combo.addItems(["Full Day", "Half Day - AM", "Half Day - PM"])
            duration_layout.addWidget(duration_combo)
            layout.addLayout(duration_layout)
            
            # Description input
            desc_layout = QHBoxLayout()
            desc_layout.addWidget(QLabel("Description:"))
            desc_edit = QLineEdit()
            desc_edit.setPlaceholderText("Leave description")
            desc_layout.addWidget(desc_edit)
            layout.addLayout(desc_layout)
            
            # ISM input
            ism_layout = QHBoxLayout()
            ism_layout.addWidget(QLabel("ISM Name:"))
            ism_combo = QComboBox()
            ism_combo.addItem("All ISMs")
            # Populate with ISM names
            try:
                org_rows = self._collect_org_directory_rows()
                for (name, _, _, _, _, _) in org_rows:
                    if name:
                        ism_combo.addItem(name)
            except Exception:
                pass
            ism_layout.addWidget(ism_combo)
            layout.addLayout(ism_layout)
            
            # Buttons
            button_layout = QHBoxLayout()
            ok_btn = QPushButton("OK")
            cancel_btn = QPushButton("Cancel")
            button_layout.addWidget(ok_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)
            
            ok_btn.clicked.connect(dialog.accept)
            cancel_btn.clicked.connect(dialog.reject)
            
            if dialog.exec() == QDialog.DialogCode.Accepted:
                # Add to data
                if "Leave Tracker" not in self.data:
                    self.data["Leave Tracker"] = []
                
                new_entry = [
                    date_edit.text(),
                    type_combo.currentText(),
                    duration_combo.currentText(),
                    desc_edit.text(),
                    ism_combo.currentText()
                ]
                self.data["Leave Tracker"].append(new_entry)
                
                # Refresh the table
                self._refresh_leave_table()
                
        except Exception as e:
            ErrorHandler.handle_ui_error("add leave entry", e)

    def _edit_leave_entry(self):
        """Edit selected leave entry"""
        try:
            if not hasattr(self, 'leave_table') or not self.leave_table.currentRow():
                return
            
            # Get selected row
            row = self.leave_table.currentRow()
            if row < 0 or row >= len(self.data.get("Leave Tracker", [])):
                return
            
            # Edit the entry (simplified for now)
            QMessageBox.information(self, "Edit Leave", "Edit functionality will be implemented")
            
        except Exception as e:
            ErrorHandler.handle_ui_error("edit leave entry", e)

    def _delete_leave_entry(self):
        """Delete selected leave entry"""
        try:
            if not hasattr(self, 'leave_table') or not self.leave_table.currentRow():
                return
            
            # Get selected row
            row = self.leave_table.currentRow()
            if row < 0 or row >= len(self.data.get("Leave Tracker", [])):
                return
            
            # Confirm deletion
            reply = QMessageBox.question(self, "Delete Leave Entry", 
                                       "Are you sure you want to delete this leave entry?",
                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            
            if reply == QMessageBox.StandardButton.Yes:
                # Remove from data
                self.data["Leave Tracker"].pop(row)
                
                # Refresh the table
                self._refresh_leave_table()
                
        except Exception as e:
            ErrorHandler.handle_ui_error("delete leave entry", e)

    def _refresh_leave_table(self):
        """Refresh the leave table display"""
        try:
            if not hasattr(self, 'leave_table'):
                return
            
            # Clear table
            self.leave_table.setRowCount(0)
            
            # Add data
            for entry in self.data.get("Leave Tracker", []):
                row = self.leave_table.rowCount()
                self.leave_table.insertRow(row)
                
                for col, value in enumerate(entry):
                    if col < self.leave_table.columnCount():
                        self.leave_table.setItem(row, col, QTableWidgetItem(str(value)))
                        
        except Exception as e:
            ErrorHandler.handle_ui_error("refresh leave table", e)

    def _refresh_project_table(self):
        """Refresh the project table display"""
        try:
            if not hasattr(self, 'projects_table'):
                return
            
            # Clear table
            self.projects_table.setRowCount(0)
            
            # Add data
            for entry in self.data.get("Project Details", []):
                row = self.projects_table.rowCount()
                self.projects_table.insertRow(row)
                
                for col, value in enumerate(entry):
                    if col < self.projects_table.columnCount():
                        self.projects_table.setItem(row, col, QTableWidgetItem(str(value)))
                        
        except Exception as e:
            ErrorHandler.handle_ui_error("refresh project table", e)

    def _init_ui_helpers(self):
        """Initialize helper methods for UI components"""
        # This method is called early in __init__ to ensure all UI helper methods are available
        # Initialize thread-safe UI operation timer
        self._setup_ui_operation_timer()
    
    def _setup_ui_operation_timer(self):
        """Set up the timer for processing UI operations in the main thread"""
        try:
            if QApplication.instance() is not None:
                self._ui_operation_timer = QTimer()
                self._ui_operation_timer.timeout.connect(self._process_ui_operations)
                self._ui_operation_timer.setSingleShot(False)
                self._ui_operation_timer.start(50)  # Process every 50ms
        except Exception as e:
            ErrorHandler.handle_ui_error("setup ui operation timer", e)

    def _show_leave_tab(self):
        """Ensure Leave Tracker tab exists and focus it."""
        try:
            # Create if needed
            self._init_leave_tracker_tab()
        except Exception:
            pass
        try:
            # Find existing
            for i in range(self.tabs.count()):
                if self.tabs.tabText(i).startswith("Leave Tracker"):
                    self.tabs.setCurrentIndex(i)
                    return
            # Not found; add fallback
            leave_widget = getattr(self, 'leave_tab', None)
            if leave_widget is None:
                from PyQt6.QtWidgets import QWidget, QVBoxLayout, QLabel
                leave_widget = QWidget()
                _lay = QVBoxLayout(leave_widget)
                _lay.addWidget(QLabel("Leave Tracker unavailable due to initialization error."))
            insert_at = min(3, self.tabs.count())
            idx = self.tabs.insertTab(insert_at, leave_widget, "Leave Tracker 🗓️")
            self.tabs.setCurrentIndex(idx)
        except Exception as e:
            ErrorHandler.handle_ui_error("show leave tab", e)

    def _show_project_details_tab(self):
        """Ensure Project Details tab exists and focus it."""
        try:
            self._init_project_details_tab()
        except Exception:
            pass
        try:
            for i in range(self.tabs.count()):
                if self.tabs.tabText(i).startswith("Project Details"):
                    self.tabs.setCurrentIndex(i)
                    return
            pd_widget = getattr(self, 'project_details_tab', None)
            if pd_widget is None:
                from PyQt6.QtWidgets import QWidget, QVBoxLayout, QLabel
                pd_widget = QWidget()
                _pdl = QVBoxLayout(pd_widget)
                _pdl.addWidget(QLabel("Project Details unavailable due to initialization error."))
            insert_at = min(4, self.tabs.count())
            idx = self.tabs.insertTab(insert_at, pd_widget, "Project Details 📋")
            self.tabs.setCurrentIndex(idx)
        except Exception as e:
            ErrorHandler.handle_ui_error("show project details tab", e)
    
    def _schedule_ui_operation(self, operation, *args, **kwargs):
        """Schedule a UI operation to be executed in the main thread"""
        try:
            self._ui_operation_queue.append((operation, args, kwargs))
        except Exception as e:
            ErrorHandler.handle_ui_error("schedule ui operation", e)
    
    def _process_ui_operations(self):
        """Process queued UI operations in the main thread"""
        try:
            # Process up to 10 operations per timer tick to avoid blocking
            processed = 0
            while self._ui_operation_queue and processed < 10:
                operation, args, kwargs = self._ui_operation_queue.pop(0)
                try:
                    operation(*args, **kwargs)
                except Exception as e:
                    ErrorHandler.handle_ui_error("process ui operation", e)
                processed += 1
        except Exception as e:
            ErrorHandler.handle_ui_error("process ui operations", e)
    
    def _ensure_main_thread(self, operation, *args, **kwargs):
        """Run the given callable on the Qt main thread; call directly if already on it."""
        try:
            import threading
            if threading.current_thread() is threading.main_thread() or QApplication.instance() is None:
                operation(*args, **kwargs)
            else:
                # Post to the main event loop
                QTimer.singleShot(0, lambda: operation(*args, **kwargs))
        except Exception as e:
            ErrorHandler.handle_ui_error("execute operation", e)

    def _validate_cell_input(self, column_name: str, value: str) -> tuple[bool, str]:
        """Enhanced cell input validation using InputValidator with feedback"""
        try:
            # Use the existing validation engine as fallback
            is_valid, error_msg = self.validation_engine.validate_field(column_name, value)
            if not is_valid:
                # Show validation feedback to user
                self._show_validation_feedback(column_name, error_msg, False)
                return False, error_msg
            
            # Enhanced validation using InputValidator
            if "Date" in column_name:
                return InputValidator.validate_date(value, column_name)
            elif "RAG" in column_name:
                return InputValidator.validate_rag_status(value)
            elif "Status" in column_name:
                return InputValidator.validate_status(value)
            elif "Priority" in column_name:
                return InputValidator.validate_priority(value)
            elif "Hours" in column_name:
                return InputValidator.validate_hours(value)
            elif "Email" in column_name or "@" in value:
                return InputValidator.validate_email(value)
            elif "Project ID" in column_name:
                return InputValidator.validate_project_id(value)
            elif column_name in ["Project Name", "ISM Name", "Ownership"]:
                # Required fields
                return InputValidator.validate_required_field(value, column_name)
            else:
                # General text length validation
                return InputValidator.validate_text_length(value, column_name, 255)
                
        except Exception as e:
            ErrorHandler.handle_validation_error(column_name, value, f"Validation error: {e}")
            return False, f"Validation error: {e}"

    def _recover_from_error(self, operation: str, error: Exception, context: dict = None) -> bool:
        """Comprehensive error recovery mechanism"""
        try:
            logger.info(f"Attempting error recovery for operation: {operation}")
            
            if "Excel" in operation:
                return self._recover_excel_error(operation, error, context)
            elif "UI" in operation:
                return self._recover_ui_error(operation, error, context)
            elif "Data" in operation:
                return self._recover_data_error(operation, error, context)
            else:
                return self._recover_generic_error(operation, error, context)
                
        except Exception as recovery_error:
            logger.error(f"Error recovery failed for {operation}: {recovery_error}")
            return False

    def _recover_excel_error(self, operation: str, error: Exception, context: dict = None) -> bool:
        """Recover from Excel-related errors"""
        try:
            if "File is not a zip file" in str(error):
                # Try to create a new Excel file
                backend_path = context.get('backend_path', self.backend_sqlite_path) if context else self.backend_sqlite_path
                self._create_new_backend_file(backend_path)
                ErrorHandler.log_operation("Excel file recovery", True, "Created new file")
                return True
            elif "Permission denied" in str(error):
                # Try to close and reopen the file
                ErrorHandler.log_operation("Excel file recovery", False, "Permission denied - file may be in use")
                return False
            else:
                # Generic Excel recovery
                ErrorHandler.log_operation("Excel file recovery", False, f"Unknown error: {error}")
                return False
        except Exception as e:
            logger.error(f"Excel recovery failed: {e}")
            return False

    def _recover_ui_error(self, operation: str, error: Exception, context: dict = None) -> bool:
        """Recover from UI-related errors"""
        try:
            if "commitData" in str(error):
                # Try to close all editors and refresh
                self._close_all_editors()
                ErrorHandler.log_operation("UI recovery", True, "Closed all editors")
                return True
            elif "widget" in str(error).lower():
                # Try to refresh the UI
                self.update_dashboard()
                ErrorHandler.log_operation("UI recovery", True, "Refreshed dashboard")
                return True
            else:
                ErrorHandler.log_operation("UI recovery", False, f"Unknown UI error: {error}")
                return False
        except Exception as e:
            logger.error(f"UI recovery failed: {e}")
            return False
    def _recover_data_error(self, operation: str, error: Exception, context: dict = None) -> bool:
        """Recover from data-related errors"""
        try:
            if "index" in str(error).lower():
                # Try to rebuild the data structure
                self._rebuild_data_structures()
                ErrorHandler.log_operation("Data recovery", True, "Rebuilt data structures")
                return True
            elif "validation" in str(error).lower():
                # Try to validate and clean data
                self._clean_invalid_data()
                ErrorHandler.log_operation("Data recovery", True, "Cleaned invalid data")
                return True
            else:
                ErrorHandler.log_operation("Data recovery", False, f"Unknown data error: {error}")
                return False
        except Exception as e:
            logger.error(f"Data recovery failed: {e}")
            return False

    def _recover_generic_error(self, operation: str, error: Exception, context: dict = None) -> bool:
        """Generic error recovery"""
        try:
            # Try to save current state
            self._save_autosave()
            ErrorHandler.log_operation("Generic recovery", True, "Saved current state")
            return True
        except Exception as e:
            logger.error(f"Generic recovery failed: {e}")
            return False

    def _close_all_editors(self):
        """Close all active editors to prevent commitData errors"""
        try:
            # Ensure this runs on the main thread
            try:
                import threading
                if QApplication.instance() is not None and threading.current_thread() is not threading.main_thread():
                    QTimer.singleShot(0, self._close_all_editors)
                    return
            except Exception:
                pass
            for table_name, table in self.tables.items():
                if hasattr(table, 'closePersistentEditor'):
                    # Close all persistent editors
                    for row in range(table.rowCount()):
                        for col in range(table.columnCount()):
                            item = table.item(row, col)
                            if item:
                                table.closePersistentEditor(item)
                    # Commit any pending data
                    current_item = table.currentItem()
                    if current_item:
                        try:
                            # Get the editor for the current item
                            editor = table.indexWidget(table.currentIndex())
                            if editor is not None:
                                table.commitData(editor)
                        except Exception as e:
                            ErrorHandler.handle_ui_error("commit data", e)
        except Exception as e:
            logger.warning(f"Failed to close all editors: {e}")

    def _rebuild_data_structures(self):
        """Rebuild data structures to recover from corruption"""
        try:
            # Rebuild all tables
            for pane_name in PANE_COLUMNS.keys():
                if pane_name in self.tables:
                    self.rebuild_table(pane_name)
            ErrorHandler.log_operation("Data structure rebuild", True)
        except Exception as e:
            logger.error(f"Failed to rebuild data structures: {e}")

    def _clean_invalid_data(self):
        """Clean invalid data entries"""
        try:
            for pane_name, data in self.data.items():
                if isinstance(data, list):
                    # Remove empty or invalid rows
                    self.data[pane_name] = [row for row in data if row and len(row) > 0]
            ErrorHandler.log_operation("Data cleaning", True)
        except Exception as e:
            logger.error(f"Failed to clean invalid data: {e}")

    def _setup_log_cleanup_timer(self):
        """Setup automatic log cleanup timer that runs every 30 days"""
        try:
            # Setup timer directly
            
            # Create a timer that runs every 24 hours and checks if 30 days have passed
            self.log_cleanup_timer = QTimer()
            self.log_cleanup_timer.timeout.connect(self._check_and_cleanup_logs)
            # Start the timer to run every 24 hours (24 * 60 * 60 * 1000 milliseconds)
            self.log_cleanup_timer.start(86400000)
            
            # Track timer for cleanup
            self._active_timers.append(self.log_cleanup_timer)
            
            # Also run cleanup immediately on startup
            self._cleanup_old_logs()
            
        except Exception as e:
            self._log_change("Error", "System", f"Failed to setup log cleanup timer: {str(e)}")
            print(f"Error setting up log cleanup timer: {e}")
    
    def _setup_auto_save_timer(self):
        """Setup auto-save timer based on user preferences"""
        try:
            # Ensure we're in the main thread for timer creation
            
            if self.preferences.get("auto_save", True):
                interval = self.preferences.get("auto_save_interval", 30) * 1000  # Convert to milliseconds
                
                self.auto_save_timer = QTimer()
                self.auto_save_timer.timeout.connect(self._perform_auto_save)
                self.auto_save_timer.start(interval)
                
                # Track timer for cleanup
                self._active_timers.append(self.auto_save_timer)
                
                self._show_toast(f"Auto-save enabled (every {interval//1000}s)", level="INFO")
                print(f"Auto-save timer started: {interval//1000} seconds")
            else:
                print("Auto-save disabled by user preferences")
                
        except Exception as e:
            ErrorHandler.handle_ui_error("setup auto-save timer", e)
            print(f"Error setting up auto-save timer: {e}")

    def _setup_backend_file_watcher(self):
        """Watch the backend file and autosave SQLite for changes; reload on change and refresh UI."""
        try:
            self._fs_watcher = QFileSystemWatcher(self)
            watch_paths = []
            if hasattr(self, '__backend_path__') and self.__backend_path__ and os.path.exists(self.__backend_path__):
                watch_paths.append(self.__backend_path__)
            # Also watch autosave JSON if present
            if hasattr(self, 'autosave_path') and self.autosave_path and os.path.exists(self.autosave_path):
                watch_paths.append(self.autosave_path)
            if watch_paths:
                self._fs_watcher.addPaths(watch_paths)
                def on_change(_):
                    try:
                        # Reload backend if exists; otherwise reload autosave
                        if hasattr(self, '__backend_path__') and self.__backend_path__ and os.path.exists(self.__backend_path__):
                            self._load_backend_sqlite(self.__backend_path__)
                        else:
                            self._load_autosave()
                        # Refresh key UI elements so approvers see new requests immediately
                        self.update_dashboard()
                        self._refresh_calendar_decorations()
                        
                        # Refresh calendar if it's visible
                        try:
                            if hasattr(self, 'calendar_tab') and self.calendar_tab:
                                self.calendar_tab.refresh_calendar_if_visible()
                        except Exception:
                            pass
                        # Ensure Project Details view reflects external changes
                        try:
                            if hasattr(self, 'projects_table') and self.projects_table is not None:
                                self._load_projects_data()
                                if hasattr(self, '_update_summary_metrics'):
                                    self._update_summary_metrics()
                        except Exception:
                            pass
                        try:
                            # Leave tracker list refresh if tab is initialized
                            if hasattr(self, 'leave_calendar'):
                                # Trigger the leave list to rebuild using existing refresh_for_date linkage
                                self.leave_calendar.selectionChanged.emit()
                        except Exception:
                            pass
                        # Update bell badge count if we have pending approvals
                        self._update_notification_count()
                    except Exception as e:
                        ErrorHandler.handle_ui_error("reload on file change", e)
                self._fs_watcher.fileChanged.connect(on_change)
        except Exception as e:
            ErrorHandler.handle_ui_error("setup backend file watcher", e)
    
    def _perform_auto_save(self):
        """Perform auto-save operation"""
        try:
            if self.preferences.get("auto_save", True):
                # Show progress indicator
                if self.preferences.get("show_progress_bars", True):
                    self.progress_manager.show_progress("Auto-saving...", 100)
                
                # Save to backend Excel
                self._save_backend_sqlite()
                
                # Update progress
                if self.preferences.get("show_progress_bars", True):
                    self.progress_manager.update_progress(50, "Creating backup...")
                
                # Create backup if enabled
                backup_count = self.preferences.get("backup_count", 5)
                if backup_count > 0:
                    self._create_auto_save_backup()
                
                # Hide progress
                if self.preferences.get("show_progress_bars", True):
                    self.progress_manager.update_progress(100, "Auto-save completed")
                    QTimer.singleShot(2000, self.progress_manager.hide_progress)
                
                # Show subtle notification
                current_time = datetime.now().strftime("%H:%M:%S")
                self.statusBar().showMessage(f"Auto-saved at {current_time}", 3000)
                
        except Exception as e:
            ErrorHandler.handle_ui_error("auto-save", e)
            if self.preferences.get("show_progress_bars", True):
                self.progress_manager.hide_progress()
    
    def _create_auto_save_backup(self):
        """Create numbered backup files for the current backend.
        
        - For SQLite backends: create a .sqlite backup using the SQLite backup API
          and honor the configured backup directory and count.
        - For Excel backends: keep existing .xlsx file copy behavior.
        """
        try:
            backend_path = getattr(self, '__backend_path__', None) or self.backend_sqlite_path
            is_sqlite = bool(backend_path) and str(backend_path).lower().endswith((".sqlite", ".db"))

            # Determine default backup directory
            script_dir = os.path.dirname(os.path.abspath(__file__))
            default_backup_dir = os.path.join(script_dir, "backups")
            if not os.path.exists(default_backup_dir):
                os.makedirs(default_backup_dir)

            if is_sqlite:
                # Prefer user-selected backup directory; fall back to default backups folder
                preferred_dir = self.preferences.get("backup_directory", "")
                target_dir = preferred_dir if preferred_dir and os.path.exists(preferred_dir) else default_backup_dir
                # Use the central backup routine (handles cleanup by backup_count)
                self._backup_sqlite_database(target_dir)
                return

            # Excel fallback (legacy behavior)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_filename = f"backup_{timestamp}.xlsx"
            backup_path = os.path.join(default_backup_dir, backup_filename)

            # Copy current backend Excel file to backup
            import shutil
            if backend_path and os.path.exists(backend_path):
                shutil.copy2(backend_path, backup_path)

            # Clean up old Excel backups (keep only the configured number)
            backup_count = self.preferences.get("backup_count", 5)
            excel_backups = [f for f in os.listdir(default_backup_dir) if f.startswith("backup_") and f.endswith(".xlsx")]
            excel_backups.sort(reverse=True)
            for old_backup in excel_backups[backup_count:]:
                try:
                    os.remove(os.path.join(default_backup_dir, old_backup))
                except Exception:
                    pass

        except Exception as e:
            logger.warning(f"Failed to create auto-save backup: {e}")
    
    def _setup_window_state(self):
        """Setup window state management for better display handling"""
        try:
            # Load saved window geometry if available
            saved_geometry = self.preferences.get("window_geometry")
            if saved_geometry:
                if not self._restore_window_geometry(saved_geometry):
                    # Clear invalid geometry
                    self.preferences.set("window_geometry", None)
            
            # Connect window state change events
            self.resizeEvent = self._on_window_resize
            self.moveEvent = self._on_window_move
            
            # Set window properties for better multi-monitor support
            self.setAttribute(Qt.WidgetAttribute.WA_DeleteOnClose, True)
            
            # Ensure window is visible and properly positioned
            self._ensure_window_visible()
            self.show()
            self.raise_()
            self.activateWindow()
            
        except Exception as e:
            logger.warning(f"Failed to setup window state: {e}")
    
    def _save_window_geometry(self):
        """Save window geometry as base64 string to preferences"""
        try:
            geometry = self.saveGeometry()
            import base64
            geometry_b64 = base64.b64encode(geometry.data()).decode('utf-8')
            self.preferences.set("window_geometry", geometry_b64)
        except Exception as e:
            logger.debug(f"Error saving window geometry: {e}")
    
    def _restore_window_geometry(self, saved_geometry):
        """Restore window geometry from base64 string"""
        try:
            import base64
            if isinstance(saved_geometry, str):
                geometry_bytes = base64.b64decode(saved_geometry)
                from PyQt6.QtCore import QByteArray
                geometry_bytearray = QByteArray(geometry_bytes)
                self.restoreGeometry(geometry_bytearray)
                return True
            return False
        except Exception as e:
            logger.debug(f"Error restoring window geometry: {e}")
            return False
    
    def _on_window_resize(self, event):
        """Handle window resize events"""
        try:
            # Save window geometry to preferences (with debouncing)
            if not hasattr(self, '_geometry_save_timer'):
                self._geometry_save_timer = QTimer()
                self._geometry_save_timer.setSingleShot(True)
                self._geometry_save_timer.timeout.connect(self._save_window_geometry)
            
            # Restart timer to debounce rapid resize events
            self._geometry_save_timer.stop()
            self._geometry_save_timer.start(500)  # Save after 500ms of no changes
            
            # Call the original resize event
            super().resizeEvent(event)
            
        except Exception as e:
            logger.debug(f"Error in window resize handler: {e}")
            super().resizeEvent(event)
    
    def _on_window_move(self, event):
        """Handle window move events"""
        try:
            # Save window geometry to preferences (with debouncing)
            if not hasattr(self, '_geometry_save_timer'):
                self._geometry_save_timer = QTimer()
                self._geometry_save_timer.setSingleShot(True)
                self._geometry_save_timer.timeout.connect(self._save_window_geometry)
            
            # Restart timer to debounce rapid move events
            self._geometry_save_timer.stop()
            self._geometry_save_timer.start(500)  # Save after 500ms of no changes
            
            # Call the original move event
            super().moveEvent(event)
            
        except Exception as e:
            logger.debug(f"Error in window move handler: {e}")
            super().moveEvent(event)
    
    def _ensure_window_visible(self):
        """Ensure window is visible and properly sized for current display"""
        try:
            screen = QApplication.primaryScreen()
            if not screen:
                return
            
            screen_geometry = screen.availableGeometry()
            current_geometry = self.geometry()
            
            # Check if window is completely outside the screen
            if (current_geometry.right() < screen_geometry.left() or 
                current_geometry.left() > screen_geometry.right() or
                current_geometry.bottom() < screen_geometry.top() or
                current_geometry.top() > screen_geometry.bottom()):
                
                # Move window to center of screen
                self.move(
                    screen_geometry.center().x() - self.width() // 2,
                    screen_geometry.center().y() - self.height() // 2
                )
            
            # Ensure window is not too large for screen
            if (self.width() > screen_geometry.width() or 
                self.height() > screen_geometry.height()):
                
                self.resize(
                    min(self.width(), screen_geometry.width() - 50),
                    min(self.height(), screen_geometry.height() - 50)
                )
                
        except Exception as e:
            logger.debug(f"Error ensuring window visibility: {e}")
    
    def _show_validation_feedback(self, field_name: str, error_msg: str, is_valid: bool):
        """Show validation feedback to user"""
        try:
            if is_valid:
                # Show success feedback
                self._show_toast(f"✅ {field_name} is valid", level="SUCCESS")
            else:
                # Show error feedback with styling
                self._show_toast(f"❌ {error_msg}", level="ERROR")
                
                # Also show in status bar for persistence
                self.statusBar().showMessage(f"Validation Error: {error_msg}", 10000)
                
                # If preferences allow, play error sound
                if self.preferences.get("enable_sounds", False):
                    # Could add sound here if available
                    pass
                    
        except Exception as e:
            logger.warning(f"Failed to show validation feedback: {e}")
        
        # Load backend data after UI initialization
        print("About to load backend data")
        try:
            if hasattr(self, 'backend_sqlite_path') and self.backend_sqlite_path:
                print(f"Loading backend data from: {self.backend_sqlite_path}")
                self._load_backend_sqlite(self.backend_sqlite_path)
                print("Backend data loaded successfully")
            else:
                print("No backend path available")
        except Exception as e:
            print(f"Error loading backend data: {e}")
            ErrorHandler.handle_ui_error("load backend data on startup", e)
        print("Finished loading backend data")
    
    def _show_preferences_dialog(self):
        """Show user preferences dialog"""
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle("Preferences")
            dialog.setModal(True)
            dialog.resize(500, 600)
            
            layout = QVBoxLayout(dialog)
            layout.setSpacing(20)
            layout.setContentsMargins(20, 20, 20, 20)
            
            # Title
            title = QLabel("⚙️ Application Preferences")
            title.setStyleSheet("""
                QLabel {
                    font-size: 18px;
                    font-weight: bold;
                    color: #2c3e50;
                    padding: 10px 0;
                }
            """)
            layout.addWidget(title)
            
            # Auto-save section
            auto_save_group = QGroupBox("Auto-Save Settings")
            auto_save_layout = QVBoxLayout(auto_save_group)
            
            auto_save_enabled = QCheckBox("Enable auto-save")
            auto_save_enabled.setChecked(self.preferences.get("auto_save", True))
            auto_save_layout.addWidget(auto_save_enabled)
            
            interval_layout = QHBoxLayout()
            interval_layout.addWidget(QLabel("Auto-save interval (seconds):"))
            interval_spin = QSpinBox()
            interval_spin.setRange(10, 3600)
            interval_spin.setValue(self.preferences.get("auto_save_interval", 30))
            interval_layout.addWidget(interval_spin)
            auto_save_layout.addLayout(interval_layout)
            
            backup_layout = QHBoxLayout()
            backup_layout.addWidget(QLabel("Keep backup count:"))
            backup_spin = QSpinBox()
            backup_spin.setRange(0, 50)
            backup_spin.setValue(self.preferences.get("backup_count", 5))
            backup_layout.addWidget(backup_spin)
            auto_save_layout.addLayout(backup_layout)
            
            # Backup directory setting
            backup_dir_layout = QHBoxLayout()
            backup_dir_layout.addWidget(QLabel("Backup directory:"))
            backup_dir_edit = QLineEdit()
            backup_dir_edit.setText(self.preferences.get("backup_directory", ""))
            backup_dir_edit.setPlaceholderText("Select directory for SQLite backups...")
            backup_dir_layout.addWidget(backup_dir_edit)
            
            browse_backup_btn = QPushButton("Browse...")
            browse_backup_btn.clicked.connect(lambda: self._browse_backup_directory(backup_dir_edit))
            backup_dir_layout.addWidget(browse_backup_btn)
            auto_save_layout.addLayout(backup_dir_layout)
            
            # Backup button
            backup_btn_layout = QHBoxLayout()
            backup_now_btn = QPushButton("💾 Backup SQLite Database Now")
            backup_now_btn.setStyleSheet("""
                QPushButton {
                    background-color: #3498db;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #2980b9;
                }
                QPushButton:pressed {
                    background-color: #21618c;
                }
            """)
            backup_now_btn.clicked.connect(lambda: self._backup_sqlite_database(backup_dir_edit.text() if backup_dir_edit.text() else None))
            backup_btn_layout.addWidget(backup_now_btn)
            backup_btn_layout.addStretch()
            auto_save_layout.addLayout(backup_btn_layout)
            
            layout.addWidget(auto_save_group)
            
            # UI Settings section
            ui_group = QGroupBox("UI Settings")
            ui_layout = QVBoxLayout(ui_group)
            
            show_tooltips = QCheckBox("Show tooltips")
            show_tooltips.setChecked(self.preferences.get("show_tooltips", True))
            ui_layout.addWidget(show_tooltips)
            
            tooltip_delay_layout = QHBoxLayout()
            tooltip_delay_layout.addWidget(QLabel("Tooltip delay (ms):"))
            tooltip_delay_spin = QSpinBox()
            tooltip_delay_spin.setRange(100, 5000)
            tooltip_delay_spin.setValue(self.preferences.get("tooltip_delay", 1000))
            tooltip_delay_layout.addWidget(tooltip_delay_spin)
            ui_layout.addLayout(tooltip_delay_layout)
            
            show_progress = QCheckBox("Show progress indicators")
            show_progress.setChecked(self.preferences.get("show_progress_bars", True))
            ui_layout.addWidget(show_progress)
            
            confirm_deletions = QCheckBox("Confirm deletions")
            confirm_deletions.setChecked(self.preferences.get("confirm_deletions", True))
            ui_layout.addWidget(confirm_deletions)
            
            enable_sounds = QCheckBox("Enable sounds")
            enable_sounds.setChecked(self.preferences.get("enable_sounds", False))
            ui_layout.addWidget(enable_sounds)
            
            layout.addWidget(ui_group)
            
            # Font settings
            font_group = QGroupBox("Font Settings")
            font_layout = QVBoxLayout(font_group)
            
            font_size_layout = QHBoxLayout()
            font_size_layout.addWidget(QLabel("Font size:"))
            font_size_spin = QSpinBox()
            font_size_spin.setRange(8, 20)
            font_size_spin.setValue(self.preferences.get("font_size", 10))
            font_size_layout.addWidget(font_size_spin)
            font_layout.addLayout(font_size_layout)
            
            layout.addWidget(font_group)
            
            # Buttons
            button_layout = QHBoxLayout()
            
            reset_btn = QPushButton("Reset to Defaults")
            reset_btn.clicked.connect(lambda: self._reset_preferences(dialog))
            button_layout.addWidget(reset_btn)
            
            button_layout.addStretch()
            
            cancel_btn = QPushButton("Cancel")
            cancel_btn.clicked.connect(dialog.reject)
            button_layout.addWidget(cancel_btn)
            
            ok_btn = QPushButton("OK")
            ok_btn.setDefault(True)
            ok_btn.clicked.connect(lambda: self._save_preferences(
                dialog, auto_save_enabled.isChecked(), interval_spin.value(),
                backup_spin.value(), backup_dir_edit.text(), show_tooltips.isChecked(), 
                tooltip_delay_spin.value(), show_progress.isChecked(),
                confirm_deletions.isChecked(), enable_sounds.isChecked(),
                font_size_spin.value()
            ))
            button_layout.addWidget(ok_btn)
            
            layout.addLayout(button_layout)
            
            dialog.exec()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("show preferences dialog", e)

    # ====== Wiki / Knowledge Base components removed ======
    def _save_preferences(self, dialog, auto_save: bool, interval: int, backup_count: int, backup_directory: str,
                         show_tooltips: bool, tooltip_delay: int, show_progress: bool,
                         confirm_deletions: bool, enable_sounds: bool, font_size: int):
        """Save preferences and apply changes"""
        try:
            # Save preferences
            self.preferences.set("auto_save", auto_save)
            self.preferences.set("auto_save_interval", interval)
            self.preferences.set("backup_count", backup_count)
            self.preferences.set("backup_directory", backup_directory)
            self.preferences.set("show_tooltips", show_tooltips)
            self.preferences.set("tooltip_delay", tooltip_delay)
            self.preferences.set("show_progress_bars", show_progress)
            self.preferences.set("confirm_deletions", confirm_deletions)
            self.preferences.set("enable_sounds", enable_sounds)
            self.preferences.set("font_size", font_size)
            
            # Apply auto-save changes
            if auto_save:
                if hasattr(self, 'auto_save_timer'):
                    self.auto_save_timer.stop()
                self._setup_auto_save_timer()
            else:
                if hasattr(self, 'auto_save_timer'):
                    self.auto_save_timer.stop()
            
            # Apply font size changes (would need font update logic here)
            # Apply other UI changes...
            
            self._show_toast("Preferences saved successfully", level="SUCCESS")
            dialog.accept()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("save preferences", e)
    
    def _reset_preferences(self, dialog):
        """Reset preferences to defaults"""
        try:
            reply = QMessageBox.question(
                dialog, "Reset Preferences",
                "Are you sure you want to reset all preferences to default values?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                self.preferences.reset_to_defaults()
                self._show_toast("Preferences reset to defaults", level="INFO")
                dialog.accept()
                
        except Exception as e:
            ErrorHandler.handle_ui_error("reset preferences", e)
    
    def _browse_backup_directory(self, backup_dir_edit):
        """Browse for backup directory"""
        try:
            current_dir = backup_dir_edit.text() if backup_dir_edit.text() else ""
            backend_path = getattr(self, '__backend_path__', None) or self.backend_sqlite_path
            if not current_dir and backend_path:
                current_dir = os.path.dirname(backend_path)
            
            selected_dir = QFileDialog.getExistingDirectory(
                self,
                "Select Backup Directory",
                current_dir,
                QFileDialog.Option.ShowDirsOnly
            )
            
            if selected_dir:
                backup_dir_edit.setText(selected_dir)
                
        except Exception as e:
            ErrorHandler.handle_ui_error("browse backup directory", e)
    
    def _duplicate_row(self, pane_name: str, row_index: int):
        """Duplicate a row in the specified pane"""
        try:
            if pane_name not in self.data or row_index >= len(self.data[pane_name]):
                return
            
            # Get the row data
            original_row = self.data[pane_name][row_index].copy()
            
            # Modify some fields to indicate it's a duplicate
            cols = PANE_COLUMNS[pane_name]
            if "Description" in cols:
                desc_idx = cols.index("Description")
                if len(original_row) > desc_idx:
                    original_row[desc_idx] = f"[COPY] {original_row[desc_idx]}"
            
            # Add the duplicated row
            self.data[pane_name].append(original_row)
            
            # Rebuild the table
            self.rebuild_table(pane_name)
            
            # Add to undo stack
            self.undo_redo_manager.add_operation(
                "duplicate_row", pane_name, len(self.data[pane_name])-1, -1,
                "", str(original_row), f"Duplicate row in {pane_name}"
            )
            
            self._show_toast(f"Row duplicated in {pane_name}", level="SUCCESS")
            
        except Exception as e:
            ErrorHandler.handle_ui_error("duplicate row", e)
    
    def show_help_dialog(self):
        """Show comprehensive help dialog with context-sensitive help"""
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle("Help & Shortcuts")
            dialog.setModal(True)
            dialog.resize(800, 700)
            
            layout = QVBoxLayout(dialog)
            layout.setSpacing(15)
            layout.setContentsMargins(20, 20, 20, 20)
            
            # Title
            title = QLabel("📚 Help & Shortcuts")
            title.setStyleSheet("""
                QLabel {
                    font-size: 20px;
                    font-weight: bold;
                    color: #2c3e50;
                    padding: 10px 0;
                }
            """)
            layout.addWidget(title)
            
            # Tab widget for different help sections
            help_tabs = QTabWidget()
            help_tabs.setStyleSheet("""
                QTabWidget::pane {
                    border: 1px solid #ddd;
                    background: white;
                }
                QTabBar::tab {
                    padding: 8px 16px;
                    margin-right: 2px;
                }
                QTabBar::tab:selected {
                    background: #3498db;
                    color: white;
                }
            """)
            
            # Keyboard Shortcuts tab
            shortcuts_widget = QWidget()
            shortcuts_layout = QVBoxLayout(shortcuts_widget)
            
            shortcuts_text = QLabel("""
<h3>🔤 Keyboard Shortcuts</h3>

<h4>📁 File Operations</h4>
• <b>Ctrl+N</b> - Add new project<br>
• <b>Ctrl+O</b> - Load data file<br>
• <b>Ctrl+S</b> - Save data<br>
• <b>Ctrl+E</b> - Export current pane<br>
• <b>Ctrl+I</b> - Import to current pane<br>

<h4>✏️ Edit Operations</h4>
• <b>Ctrl+Z</b> - Undo last action<br>
• <b>Ctrl+Y</b> - Redo last action<br>
• <b>Delete</b> - Delete selected row<br>
• <b>Ctrl+D</b> - Duplicate selected row<br>
• <b>Ctrl+A</b> - Select all rows<br>

<h4>🧭 Navigation</h4>
• <b>Ctrl+F</b> or <b>Ctrl+K</b> - Global search<br>
• <b>F1</b> or <b>Ctrl+H</b> - Show this help<br>
• <b>F5</b> - Refresh current pane<br>
• <b>Ctrl+Tab</b> - Next tab<br>
• <b>Ctrl+Shift+Tab</b> - Previous tab<br>
• <b>Ctrl+1-5</b> - Go to specific tab<br>

<h4>⚙️ Settings</h4>
• <b>Ctrl+,</b> - Open preferences<br>
• <b>Ctrl+Alt+S</b> - Toggle auto-save<br>
            """)
            shortcuts_text.setWordWrap(True)
            shortcuts_text.setStyleSheet("font-size: 12px; line-height: 1.4;")
            shortcuts_layout.addWidget(shortcuts_text)
            
            # Features tab
            features_widget = QWidget()
            features_layout = QVBoxLayout(features_widget)
            
            features_text = QLabel("""
<h3>✨ Application Features</h3>

<h4>🏠 Home Dashboard</h4>
• Quick navigation to all panes<br>
• Real-time statistics and metrics<br>
• Project overview and status<br>

<h4>📊 Data Management</h4>
• Multiple data panes (Initiatives, Issues, Activities, etc.)<br>
• Excel import/export functionality<br>
• Auto-save with backup creation<br>
• Data validation and error checking<br>

<h4>🔍 Search & Filter</h4>
• Global search across all data (Ctrl+K)<br>
• Per-pane filtering and sorting<br>
• Quick filter chips for common searches<br>

<h4>👥 Organization Chart</h4>
• Visual team hierarchy<br>
• Drag-and-drop member management<br>
• Color-coded by level or designation<br>

<h4>📋 Project Details</h4>
• Comprehensive project tracking<br>
• ISM hours and progress monitoring<br>
• Voice solutions and audit requirements<br>

<h4>🔔 Notifications</h4>
• Due date reminders<br>
• Overdue item alerts<br>
• System status notifications<br>
            """)
            features_text.setWordWrap(True)
            features_text.setStyleSheet("font-size: 12px; line-height: 1.4;")
            features_layout.addWidget(features_text)
            
            # Tips & Tricks tab
            tips_widget = QWidget()
            tips_layout = QVBoxLayout(tips_widget)
            
            tips_text = QLabel("""
<h3>💡 Tips & Tricks</h3>

<h4>⚡ Quick Actions</h4>
• Double-click cells to edit inline<br>
• Right-click for context menus<br>
• Use search boxes to filter data quickly<br>
• Click column headers to sort<br>

<h4>📈 Data Entry</h4>
• Use dropdown menus for consistent data<br>
• Date format: YYYY-MM-DD<br>
• RAG status: Red, Amber, Green<br>
• Priority: Critical, High, Medium, Low<br>

<h4>💾 Data Safety</h4>
• Auto-save runs every 30 seconds by default<br>
• Backups are created automatically<br>
• Use Ctrl+S to save manually<br>
• Check the status bar for save confirmations<br>

<h4>🎨 Customization</h4>
• Use Ctrl+, to open preferences<br>
• Adjust auto-save intervals<br>
• Configure tooltips and notifications<br>
• Set font sizes and UI preferences<br>

<h4>🔧 Troubleshooting</h4>
• Check the status bar for error messages<br>
• Use F5 to refresh if data seems stale<br>
• Check Activity Logs for detailed information<br>
• Use Ctrl+Z to undo accidental changes<br>
            """)
            tips_text.setWordWrap(True)
            tips_text.setStyleSheet("font-size: 12px; line-height: 1.4;")
            tips_layout.addWidget(tips_text)
            
            # Getting Started tab
            getting_started_widget = QWidget()
            getting_started_layout = QVBoxLayout(getting_started_widget)
            
            getting_started_text = QLabel("""
<h3>🚀 Getting Started</h3>

<h4>1. 📁 Load Your Data</h4>
• Click "Load Data" or use Ctrl+O<br>
• Select your Excel file<br>
• Data will be imported automatically<br>

<h4>2. 🏠 Explore the Dashboard</h4>
• Start with the Home tab for overview<br>
• Check the Dashboard for key metrics<br>
• Review notifications for urgent items<br>
<h4>4. 🔍 Find Information</h4>
• Use Ctrl+K for global search<br>
• Use pane-specific search boxes<br>
• Click column headers to sort<br>

<h4>5. 💾 Save Your Work</h4>
• Auto-save is enabled by default<br>
• Use Ctrl+S to save manually<br>
• Export specific panes with Ctrl+E<br>

<h4>6. ⚙️ Customize Settings</h4>
• Use Ctrl+, to open preferences<br>
• Adjust auto-save frequency<br>
• Configure UI preferences<br>
            """)
            getting_started_text.setWordWrap(True)
            getting_started_text.setStyleSheet("font-size: 12px; line-height: 1.4;")
            getting_started_layout.addWidget(getting_started_text)
            
            # Add tabs
            help_tabs.addTab(getting_started_widget, "🚀 Getting Started")
            help_tabs.addTab(shortcuts_widget, "🔤 Shortcuts")
            help_tabs.addTab(features_widget, "✨ Features")
            help_tabs.addTab(tips_widget, "💡 Tips")
            
            layout.addWidget(help_tabs)
            
            # Buttons
            button_layout = QHBoxLayout()
            
            preferences_btn = QPushButton("⚙️ Open Preferences")
            preferences_btn.clicked.connect(lambda: (dialog.accept(), self._show_preferences_dialog()))
            button_layout.addWidget(preferences_btn)
            
            button_layout.addStretch()
            
            close_btn = QPushButton("✖️ Close")
            close_btn.setDefault(True)
            close_btn.clicked.connect(dialog.accept)
            button_layout.addWidget(close_btn)
            
            layout.addLayout(button_layout)
            
            # Style the dialog
            dialog.setStyleSheet("""
                QDialog {
                    background: #f8fafc;
                }
                QPushButton {
                    background: #3498db;
                    color: white;
                    padding: 8px 16px;
                    border: none;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background: #2980b9;
                }
                QPushButton:default {
                    background: #27ae60;
                }
                QPushButton:default:hover {
                    background: #229954;
                }
            """)
            
            dialog.exec()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("show help dialog", e)
    
    def export_data_multiple_formats(self, pane_name: str = None):
        """Export data in multiple formats with user choice"""
        try:
            # Create export format dialog
            dialog = QDialog(self)
            dialog.setWindowTitle("Export Data")
            dialog.setModal(True)
            dialog.resize(400, 300)
            
            layout = QVBoxLayout(dialog)
            layout.setSpacing(15)
            layout.setContentsMargins(20, 20, 20, 20)
            
            # Title
            title = QLabel("📤 Export Data")
            title.setStyleSheet("""
                QLabel {
                    font-size: 16px;
                    font-weight: bold;
                    color: #2c3e50;
                    padding: 10px 0;
                }
            """)
            layout.addWidget(title)
            
            # Data source selection
            source_group = QGroupBox("Data to Export")
            source_layout = QVBoxLayout(source_group)
            
            source_combo = QComboBox()
            if pane_name:
                source_combo.addItem(f"Current Pane: {pane_name}", pane_name)
            else:
                source_combo.addItem("All Data", "all")
                for pane in PANE_COLUMNS.keys():
                    source_combo.addItem(f"Pane: {pane}", pane)
            source_layout.addWidget(source_combo)
            layout.addWidget(source_group)
            
            # Format selection
            format_group = QGroupBox("Export Format")
            format_layout = QVBoxLayout(format_group)
            
            format_combo = QComboBox()
            format_combo.addItem("📊 Excel (.xlsx)", "xlsx")
            format_combo.addItem("📄 CSV (.csv)", "csv")
            format_combo.addItem("🌐 JSON (.json)", "json")
            format_combo.addItem("📋 TSV (.tsv)", "tsv")
            if PANDAS_AVAILABLE:
                format_combo.addItem("🗄️ Parquet (.parquet)", "parquet")
                format_combo.addItem("🐍 Pickle (.pkl)", "pickle")
            format_layout.addWidget(format_combo)
            layout.addWidget(format_group)
            
            # Options
            options_group = QGroupBox("Export Options")
            options_layout = QVBoxLayout(options_group)
            
            include_headers = QCheckBox("Include column headers")
            include_headers.setChecked(True)
            options_layout.addWidget(include_headers)
            
            include_metadata = QCheckBox("Include metadata")
            include_metadata.setChecked(True)
            options_layout.addWidget(include_metadata)
            
            filter_empty = QCheckBox("Filter out empty rows")
            filter_empty.setChecked(False)
            options_layout.addWidget(filter_empty)
            
            layout.addWidget(options_group)
            
            # Buttons
            button_layout = QHBoxLayout()
            
            preview_btn = QPushButton("👁️ Preview")
            preview_btn.clicked.connect(lambda: self._preview_export_data(
                source_combo.currentData(), format_combo.currentData()
            ))
            button_layout.addWidget(preview_btn)
            
            button_layout.addStretch()
            
            cancel_btn = QPushButton("❌ Cancel")
            cancel_btn.clicked.connect(dialog.reject)
            button_layout.addWidget(cancel_btn)
            
            export_btn = QPushButton("📤 Export")
            export_btn.setDefault(True)
            export_btn.clicked.connect(lambda: self._perform_data_export(
                dialog, source_combo.currentData(), format_combo.currentData(),
                include_headers.isChecked(), include_metadata.isChecked(),
                filter_empty.isChecked()
            ))
            button_layout.addWidget(export_btn)
            
            layout.addLayout(button_layout)
            
            # Style the dialog
            dialog.setStyleSheet("""
                QDialog {
                    background: #f8fafc;
                }
                QGroupBox {
                    font-weight: bold;
                    border: 2px solid #cbd5e0;
                    border-radius: 8px;
                    margin-top: 1ex;
                    padding-top: 10px;
                }
                QGroupBox::title {
                    subcontrol-origin: margin;
                    left: 10px;
                    padding: 0 5px 0 5px;
                }
                QPushButton {
                    background: #3498db;
                    color: white;
                    padding: 8px 16px;
                    border: none;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background: #2980b9;
                }
                QPushButton:default {
                    background: #27ae60;
                }
            """)
            
            dialog.exec()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("export data multiple formats", e)
    
    def _preview_export_data(self, source: str, format_type: str):
        """Preview export data before actual export"""
        try:
            # Get data to export
            if source == "all":
                data_preview = {}
                for pane_name, pane_data in self.data.items():
                    if pane_data:  # Only include non-empty panes
                        data_preview[pane_name] = pane_data[:5]  # First 5 rows
            else:
                data_preview = {source: self.data.get(source, [])[:5]}
            
            # Create preview dialog
            preview_dialog = QDialog(self)
            preview_dialog.setWindowTitle(f"Export Preview - {format_type.upper()}")
            preview_dialog.resize(600, 400)
            
            layout = QVBoxLayout(preview_dialog)
            
            info_label = QLabel(f"Preview of data to be exported as {format_type.upper()} (showing first 5 rows)")
            info_label.setStyleSheet("font-weight: bold; padding: 10px;")
            layout.addWidget(info_label)
            
            # Show preview in a text widget
            preview_text = QTextEdit()
            preview_text.setReadOnly(True)
            
            preview_content = ""
            for pane_name, pane_data in data_preview.items():
                preview_content += f"\n=== {pane_name} ===\n"
                if pane_data:
                    headers = PANE_COLUMNS.get(pane_name, [])
                    preview_content += f"Columns: {', '.join(headers)}\n"
                    preview_content += f"Rows: {len(pane_data)} (showing first 5)\n\n"
                    for i, row in enumerate(pane_data[:5]):
                        preview_content += f"Row {i+1}: {', '.join(str(cell)[:50] + '...' if len(str(cell)) > 50 else str(cell) for cell in row)}\n"
                else:
                    preview_content += "No data\n"
                preview_content += "\n"
            
            preview_text.setText(preview_content)
            layout.addWidget(preview_text)
            
            close_btn = QPushButton("Close Preview")
            close_btn.clicked.connect(preview_dialog.accept)
            layout.addWidget(close_btn)
            
            preview_dialog.exec()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("preview export data", e)
    
    def _perform_data_export(self, dialog, source: str, format_type: str, 
                           include_headers: bool, include_metadata: bool, filter_empty: bool):
        """Perform the actual data export"""
        try:
            dialog.accept()
            
            # Show progress
            if self.preferences.get("show_progress_bars", True):
                self.progress_manager.show_progress("Preparing export...", 100)
            
            # Get export path
            last_export_path = self.preferences.get("last_export_path", "")
            
            # File extensions
            extensions = {
                "xlsx": "Excel Files (*.xlsx)",
                "csv": "CSV Files (*.csv)",
                "json": "JSON Files (*.json)",
                "tsv": "TSV Files (*.tsv)",
                "parquet": "Parquet Files (*.parquet)",
                "pickle": "Pickle Files (*.pkl)"
            }
            
            default_name = f"export_{source}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format_type}"
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Export Data", 
                os.path.join(last_export_path, default_name),
                extensions.get(format_type, f"{format_type.upper()} Files (*.{format_type})")
            )
            
            if not file_path:
                if self.preferences.get("show_progress_bars", True):
                    self.progress_manager.hide_progress()
                return
            
            # Save the export path
            self.preferences.set("last_export_path", os.path.dirname(file_path))
            
            # Update progress
            if self.preferences.get("show_progress_bars", True):
                self.progress_manager.update_progress(25, "Collecting data...")
            
            # Collect data to export
            export_data = {}
            if source == "all":
                for pane_name, pane_data in self.data.items():
                    if not filter_empty or pane_data:
                        export_data[pane_name] = pane_data
            else:
                export_data[source] = self.data.get(source, [])
            
            # Update progress
            if self.preferences.get("show_progress_bars", True):
                self.progress_manager.update_progress(50, f"Exporting to {format_type.upper()}...")
            
            # Export based on format
            if format_type == "csv":
                self._export_to_csv(file_path, export_data, include_headers)
            elif format_type == "json":
                self._export_to_json(file_path, export_data, include_metadata)
            elif format_type == "tsv":
                self._export_to_tsv(file_path, export_data, include_headers)
            elif format_type == "parquet" and PANDAS_AVAILABLE:
                self._export_to_parquet(file_path, export_data, include_headers)
            elif format_type == "pickle" and PANDAS_AVAILABLE:
                self._export_to_pickle(file_path, export_data, include_metadata)
            
            # Update progress
            if self.preferences.get("show_progress_bars", True):
                self.progress_manager.update_progress(100, "Export completed!")
                QTimer.singleShot(2000, self.progress_manager.hide_progress)
            
            # Show success message
            self._show_toast(f"Data exported successfully to {format_type.upper()}", level="SUCCESS")
            
            # Ask if user wants to open the file
            reply = QMessageBox.question(
                self, "Export Complete",
                f"Data exported successfully to:\n{file_path}\n\nWould you like to open the file?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                import subprocess
                import platform
                
                try:
                    if platform.system() == "Windows":
                        os.startfile(file_path)
                    elif platform.system() == "Darwin":  # macOS
                        subprocess.run(["open", file_path])
                    else:  # Linux
                        subprocess.run(["xdg-open", file_path])
                except Exception as e:
                    self._show_toast(f"Could not open file: {e}", level="WARNING")
            
        except Exception as e:
            if self.preferences.get("show_progress_bars", True):
                self.progress_manager.hide_progress()
            ErrorHandler.handle_ui_error("perform data export", e)
    
    
    def _export_to_csv(self, file_path: str, export_data: dict, include_headers: bool):
        """Export data to CSV format"""
        try:
            import csv
            
            if len(export_data) == 1:
                # Single pane - save directly as CSV
                pane_name, pane_data = next(iter(export_data.items()))
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f)
                    
                    if include_headers:
                        headers = PANE_COLUMNS.get(pane_name, [])
                        writer.writerow(headers)
                    
                    writer.writerows(pane_data)
            else:
                # Multiple panes - create a folder with multiple CSV files
                import os
                base_path = file_path.rsplit('.', 1)[0]
                os.makedirs(base_path, exist_ok=True)
                
                for pane_name, pane_data in export_data.items():
                    pane_file = os.path.join(base_path, f"{pane_name}.csv")
                    with open(pane_file, 'w', newline='', encoding='utf-8') as f:
                        writer = csv.writer(f)
                        
                        if include_headers:
                            headers = PANE_COLUMNS.get(pane_name, [])
                            writer.writerow(headers)
                        
                        writer.writerows(pane_data)
                        
        except Exception as e:
            raise Exception(f"Failed to export to CSV: {e}")
    
    def _export_to_json(self, file_path: str, export_data: dict, include_metadata: bool):
        """Export data to JSON format"""
        try:
            output = {}
            
            for pane_name, pane_data in export_data.items():
                headers = PANE_COLUMNS.get(pane_name, [])
                pane_records = []
                
                for row in pane_data:
                    record = {}
                    for i, value in enumerate(row):
                        if i < len(headers):
                            record[headers[i]] = value
                        else:
                            record[f"Column_{i+1}"] = value
                    pane_records.append(record)
                
                output[pane_name] = pane_records
            
            if include_metadata:
                output["_metadata"] = {
                    "export_date": datetime.now().isoformat(),
                    "export_user": self.logged_in_user,
                    "application_version": "1.0.0",
                    "total_panes": len(export_data),
                    "total_rows": sum(len(data) for data in export_data.values())
                }
            
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(output, f, indent=2, ensure_ascii=False)
                
        except Exception as e:
            raise Exception(f"Failed to export to JSON: {e}")
    
    def _export_to_tsv(self, file_path: str, export_data: dict, include_headers: bool):
        """Export data to TSV format"""
        try:
            import csv
            
            if len(export_data) == 1:
                # Single pane - save directly as TSV
                pane_name, pane_data = next(iter(export_data.items()))
                with open(file_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.writer(f, delimiter='\t')
                    
                    if include_headers:
                        headers = PANE_COLUMNS.get(pane_name, [])
                        writer.writerow(headers)
                    
                    writer.writerows(pane_data)
            else:
                # Multiple panes - create a folder with multiple TSV files
                base_path = file_path.rsplit('.', 1)[0]
                os.makedirs(base_path, exist_ok=True)
                
                for pane_name, pane_data in export_data.items():
                    pane_file = os.path.join(base_path, f"{pane_name}.tsv")
                    with open(pane_file, 'w', newline='', encoding='utf-8') as f:
                        writer = csv.writer(f, delimiter='\t')
                        
                        if include_headers:
                            headers = PANE_COLUMNS.get(pane_name, [])
                            writer.writerow(headers)
                        
                        writer.writerows(pane_data)
                        
        except Exception as e:
            raise Exception(f"Failed to export to TSV: {e}")
    
    def _export_to_parquet(self, file_path: str, export_data: dict, include_headers: bool):
        """Export data to Parquet format"""
        try:
            if not PANDAS_AVAILABLE:
                raise Exception("Pandas is required for Parquet export")
            
            if len(export_data) == 1:
                # Single pane
                pane_name, pane_data = next(iter(export_data.items()))
                headers = PANE_COLUMNS.get(pane_name, []) if include_headers else [f"Column_{i+1}" for i in range(len(pane_data[0]) if pane_data else 0)]
                
                df = pd.DataFrame(pane_data, columns=headers)
                df.to_parquet(file_path, index=False)
            else:
                # Multiple panes - save as a single file with multi-index
                all_data = []
                for pane_name, pane_data in export_data.items():
                    headers = PANE_COLUMNS.get(pane_name, []) if include_headers else [f"Column_{i+1}" for i in range(len(pane_data[0]) if pane_data else 0)]
                    
                    for row in pane_data:
                        row_dict = {'_pane': pane_name}
                        for i, value in enumerate(row):
                            if i < len(headers):
                                row_dict[headers[i]] = value
                        all_data.append(row_dict)
                
                df = pd.DataFrame(all_data)
                df.to_parquet(file_path, index=False)
                
        except Exception as e:
            raise Exception(f"Failed to export to Parquet: {e}")
    
    def _export_to_pickle(self, file_path: str, export_data: dict, include_metadata: bool):
        """Export data to Pickle format"""
        try:
            if not PANDAS_AVAILABLE:
                raise Exception("Pandas is required for Pickle export")
            
            output = {}
            
            for pane_name, pane_data in export_data.items():
                headers = PANE_COLUMNS.get(pane_name, [])
                df = pd.DataFrame(pane_data, columns=headers)
                output[pane_name] = df
            
            if include_metadata:
                output["_metadata"] = {
                    "export_date": datetime.now().isoformat(),
                    "export_user": self.logged_in_user,
                    "application_version": "1.0.0",
                    "total_panes": len(export_data),
                    "total_rows": sum(len(data) for data in export_data.values())
                }
            
            import pickle
            with open(file_path, 'wb') as f:
                pickle.dump(output, f)
                
        except Exception as e:
            raise Exception(f"Failed to export to Pickle: {e}")

    def _check_and_cleanup_logs(self):
        """Check if 30 days have passed since last cleanup and run cleanup if needed"""
        try:
            # Check if we have a last cleanup date stored
            last_cleanup_key = "__last_log_cleanup__"
            current_date = datetime.now().strftime("%Y-%m-%d")
            
            # Get last cleanup date from autosave or use a default
            last_cleanup = getattr(self, '_last_log_cleanup_date', None)
            
            if not last_cleanup:
                # First time running, set today as last cleanup date
                self._last_log_cleanup_date = current_date
                return
            
            # Check if 30 days have passed
            from datetime import datetime, timedelta
            last_cleanup_dt = datetime.strptime(last_cleanup, "%Y-%m-%d")
            current_dt = datetime.strptime(current_date, "%Y-%m-%d")
            
            if (current_dt - last_cleanup_dt).days >= 30:
                self._cleanup_old_logs()
                self._last_log_cleanup_date = current_date
                
        except Exception as e:
            print(f"Error checking log cleanup: {e}")

    def _create_bell_icon(self):
        """Create a compact bell icon for notifications with subtle appearance"""
        pixmap = QPixmap(36, 36)  # Compact size to match icon size
        pixmap.fill(Qt.GlobalColor.transparent)
        painter = QPainter(pixmap)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        painter.setRenderHint(QPainter.RenderHint.TextAntialiasing)
        
        # Compact bell shape with subtle styling
        painter.setPen(QPen(QColor("#2563eb"), 2))  # Clean thickness
        painter.setBrush(QBrush(QColor("#f8fafc")))  # Light background
        
        # Draw bell shape - scaled for 36x36 space
        bell_path = QPainterPath()
        bell_path.moveTo(18, 5)  # Start from top center
        bell_path.arcTo(5, 5, 26, 18, 0, 180)  # Main bell body
        bell_path.lineTo(9, 23)  # Left side
        bell_path.arcTo(9, 23, 18, 9, 0, 180)  # Bottom curve
        bell_path.lineTo(18, 32)  # Center bottom
        bell_path.lineTo(27, 32)  # Right bottom
        bell_path.arcTo(9, 23, 18, 9, 180, 180)  # Bottom curve - right side
        bell_path.lineTo(31, 23)  # Right side
        bell_path.closeSubpath()
        
        painter.drawPath(bell_path)
        
        # Bell clapper with accent color - compact and clean
        painter.setPen(QPen(QColor("#2563eb"), 1))
        painter.setBrush(QBrush(QColor("#2563eb")))
        painter.drawEllipse(16, 30, 3, 3)  # Compact clapper
        
        # Add a subtle highlight for depth
        painter.setPen(QPen(QColor("#ffffff"), 1))
        painter.drawLine(11, 9, 25, 9)  # Top highlight
        
        painter.end()
        return pixmap

    def show_notifications_dialog(self):
        """Show notifications dialog with pending items"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Notifications")
        dialog.setModal(True)
        dialog.resize(500, 400)
        dialog.setStyleSheet("""
            QDialog {
                background-color: #f8f9fa;
            }
        """)
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(10)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Title
        title_label = QLabel("🔔 Notifications")
        title_label.setStyleSheet("""
            font-size: 18px;
            font-weight: bold;
            color: #2c3e50;
            padding: 10px 0;
        """)
        layout.addWidget(title_label)
        
        # Get pending notifications
        notifications = self._get_pending_notifications()
        
        if not notifications:
            # No notifications
            empty_label = QLabel("✅ No pending notifications")
            empty_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            empty_label.setStyleSheet("""
                font-size: 14px;
                color: #7f8c8d;
                padding: 40px;
            """)
            layout.addWidget(empty_label)
        else:
            # Notifications list
            notifications_list = QListWidget()
            notifications_list.setStyleSheet("""
                QListWidget {
                    background-color: white;
                    border: 1px solid #dee2e6;
                    border-radius: 6px;
                    padding: 5px;
                }
                QListWidget::item {
                    padding: 10px;
                    border-bottom: 1px solid #f1f3f4;
                    border-radius: 4px;
                    margin: 2px;
                }
                QListWidget::item:hover {
                    background-color: #e3f2fd;
                }
                QListWidget::item:selected {
                    background-color: #2196f3;
                    color: white;
                }
            """)
            
            for notification in notifications:
                item = QListWidgetItem(notification['text'])
                item.setData(Qt.ItemDataRole.UserRole, notification)
                notifications_list.addItem(item)
            
            notifications_list.itemClicked.connect(lambda item: self._handle_notification_click(item.data(Qt.ItemDataRole.UserRole)))
            layout.addWidget(notifications_list)
        
        # Buttons
        open_btn = QPushButton("Open Selected")
        open_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a7de1;
                color: white;
                padding: 10px 20px;
                border: none;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2563eb;
            }
            QPushButton:disabled {
                background-color: #6c757d;
                color: #adb5bd;
            }
        """)
        # Open the currently selected notification and close the dialog
        def _open_selected():
            try:
                if 'notifications_list' in locals():
                    it = notifications_list.currentItem()
                    if it is None and notifications:
                        notifications_list.setCurrentRow(0)
                        it = notifications_list.currentItem()
                    if it is not None:
                        self._handle_notification_click(it.data(Qt.ItemDataRole.UserRole))
                        dialog.accept()
            except Exception as e:
                ErrorHandler.handle_ui_error("open selected notification", e)
        try:
            # If list exists in this context, wire the button
            open_btn.clicked.connect(_open_selected)
        except Exception:
            pass

        clear_btn = QPushButton("🗑️ Clear All")
        clear_btn.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                padding: 10px 20px;
                border: none;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
            QPushButton:disabled {
                background-color: #6c757d;
                color: #adb5bd;
            }
        """)
        clear_btn.setEnabled(len(notifications) > 0)
        clear_btn.clicked.connect(lambda: self._clear_all_notifications(dialog))
        
        close_btn = QPushButton("Close")
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #6c757d;
                color: white;
                padding: 10px 20px;
                border: none;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
        """)
        close_btn.clicked.connect(dialog.accept)
        
        button_layout = QHBoxLayout()
        button_layout.addWidget(open_btn)
        button_layout.addWidget(clear_btn)
        button_layout.addStretch()
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)
        
        dialog.exec()
    def _get_pending_notifications(self):
        """Get all pending notifications from data scoped to the logged-in ISM and honoring dismissed IDs."""
        notifications = []
        today = datetime.now().date()
        
        # Check for overdue items in Potential Issues
        if "Potential Issues" in self.data:
            cols_pi = PANE_COLUMNS["Potential Issues"]
            ism_idx = cols_pi.index("ISM Name") if "ISM Name" in cols_pi else None
            for row_idx, row in enumerate(self.data["Potential Issues"]):
                if len(row) > 4:  # Ensure Due Date column exists
                    # Scope to logged in user if ISM column available
                    if ism_idx is not None and (ism_idx >= len(row) or str(row[ism_idx]).strip() != str(self.logged_in_user)):
                        continue
                    due_date_str = row[4]  # Due Date column
                    if due_date_str:
                        try:
                            due_date = datetime.strptime(due_date_str, "%Y-%m-%d").date()
                            if due_date < today:
                                project_name = str(row[5]) if len(row) > 5 else "Unknown Project"
                                # Stable ID based on project and date (row index can change across loads)
                                notification_id = f"overdue_{project_name}_{due_date_str}"
                                if notification_id not in self._dismissed_notifications:
                                    notifications.append({
                                        'type': 'overdue',
                                        'pane': 'Potential Issues',
                                        'row': row_idx,
                                        'text': f"⚠️ Overdue: {project_name} (Due: {due_date_str})",
                                        'project': project_name,
                                        'id': notification_id,
                                        'match': {'project': project_name, 'date': due_date_str}
                                    })
                        except ValueError:
                            pass
        
        # Check for upcoming due dates (next 3 days)
        for days_ahead in range(1, 4):
            check_date = today + timedelta(days=days_ahead)
            if "Potential Issues" in self.data:
                cols_pi = PANE_COLUMNS["Potential Issues"]
                ism_idx = cols_pi.index("ISM Name") if "ISM Name" in cols_pi else None
                for row_idx, row in enumerate(self.data["Potential Issues"]):
                    if len(row) > 4:
                        if ism_idx is not None and (ism_idx >= len(row) or str(row[ism_idx]).strip() != str(self.logged_in_user)):
                            continue
                        due_date_str = row[4]
                        if due_date_str:
                            try:
                                due_date = datetime.strptime(due_date_str, "%Y-%m-%d").date()
                                if due_date == check_date:
                                    project_name = str(row[5]) if len(row) > 5 else "Unknown Project"
                                    notification_id = f"upcoming_{project_name}_{due_date_str}"
                                    if notification_id not in self._dismissed_notifications:
                                        notifications.append({
                                            'type': 'upcoming',
                                            'pane': 'Potential Issues',
                                            'row': row_idx,
                                            'text': f"📅 Due Soon: {project_name} (Due: {due_date_str})",
                                            'project': project_name,
                                            'id': notification_id,
                                            'match': {'project': project_name, 'date': due_date_str}
                                        })
                            except ValueError:
                                pass
        
        # Check for activities with upcoming target dates
        if "Activities" in self.data:
            cols_act = PANE_COLUMNS["Activities"]
            ism_idx = cols_act.index("ISM Name") if "ISM Name" in cols_act else None
            act_idx = cols_act.index("Activity/Issue") if "Activity/Issue" in cols_act else None
            proj_idx = cols_act.index("Project Name") if "Project Name" in cols_act else None
            for row_idx, row in enumerate(self.data["Activities"]):
                if len(row) > 4:  # Target Date column
                    if ism_idx is not None and (ism_idx >= len(row) or str(row[ism_idx]).strip() != str(self.logged_in_user)):
                        continue
                    target_date_str = row[4]
                    if target_date_str:
                        try:
                            target_date = datetime.strptime(target_date_str, "%Y-%m-%d").date()
                            if target_date <= today + timedelta(days=3):
                                project_name = str(row[proj_idx]) if proj_idx is not None and proj_idx < len(row) else (str(row[1]) if len(row) > 1 else "Unknown Project")
                                activity = str(row[act_idx]) if act_idx is not None and act_idx < len(row) else (str(row[2]) if len(row) > 2 else "Activity")
                                notification_id = f"activity_{project_name}_{activity}_{target_date_str}"
                                if notification_id not in self._dismissed_notifications:
                                    notifications.append({
                                        'type': 'activity',
                                        'pane': 'Activities',
                                        'row': row_idx,
                                        'text': f"📋 Activity Due: {activity} - {project_name} (Target: {target_date_str})",
                                        'project': project_name,
                                        'id': notification_id,
                                        'match': {'project': project_name, 'activity': activity, 'date': target_date_str}
                                    })
                        except ValueError:
                            pass
        
        return notifications
    
    def _update_notification_count(self):
        """Update the notification count badge on the bell icon"""
        try:
            notifications = self._get_pending_notifications()
            count = len(notifications)
            
            if count > 0:
                # Show the badge with count
                self.notification_count_badge.setText(str(count) if count < 100 else "99+")
                self.notification_count_badge.show()
                
                # Update tooltip to include count
                self.notification_bell.setToolTip(f"Notifications ({count})")
            else:
                # Hide the badge
                self.notification_count_badge.hide()
                self.notification_bell.setToolTip("Notifications")
                
        except Exception as e:
            ErrorHandler.handle_ui_error("update notification count", e)

    def _clear_all_notifications(self, dialog):
        """Clear all notifications by marking them as dismissed"""
        try:
            # Get current notifications to dismiss them
            current_notifications = self._get_pending_notifications()
            
            if not current_notifications:
                self._show_toast("No notifications to clear", level="INFO")
                return
            
            # Show confirmation dialog
            reply = QMessageBox.question(
                dialog,
                "Clear All Notifications",
                f"Are you sure you want to clear all {len(current_notifications)} notifications?\n\nNote: Notifications will reappear if items are still overdue or due soon.",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                # Add all current notification IDs to dismissed set
                for notification in current_notifications:
                    if 'id' in notification:
                        self._dismissed_notifications.add(notification['id'])
                
                # Update notification count immediately
                self._update_notification_count()
                
                # Close the current dialog
                dialog.accept()
                
                # Show success message
                self._show_toast(f"Cleared {len(current_notifications)} notifications", level="SUCCESS")
                
                # Log the action
                self._log(f"Cleared {len(current_notifications)} notifications", "INFO")
                
        except Exception as e:
            ErrorHandler.handle_ui_error("clear all notifications", e)

    def _handle_notification_click(self, notification):
        """Handle clicking on a notification item"""
        if not notification:
            return
        
        # Navigate to the relevant pane
        pane_name = notification['pane']
        row_idx = notification.get('row', 0)
        
        # Find the tab index for the pane (tab text may include emoji/prefix)
        for i in range(self.tabs.count()):
            try:
                if pane_name.lower() in self.tabs.tabText(i).lower():
                    self.tabs.setCurrentIndex(i)
                    break
            except Exception:
                continue
        
        # Find table for the pane and highlight
        table = self.tables.get(pane_name)
        if table is None:
            # Try to resolve by partial match on keys
            for key, t in self.tables.items():
                if pane_name.lower() in key.lower():
                    table = t
                    break
        if table:
            # If stored index is invalid (due to filtering or order changes), try to locate by match fields
            if row_idx >= table.rowCount() and 'match' in notification:
                match = notification['match']
                candidate = None
                for r in range(table.rowCount()):
                    first_col = table.item(r, 0).text() if table.item(r, 0) else ""
                    row_text = " ".join((table.item(r, c).text() if table.item(r, c) else "") for c in range(min(5, table.columnCount())))
                    if all(str(v) in row_text for v in match.values()) or (match.get('project') and match['project'] in row_text):
                        candidate = r
                        break
                if candidate is not None:
                    row_idx = candidate
            if 0 <= row_idx < table.rowCount():
                table.selectRow(row_idx)
                table.scrollToItem(table.item(row_idx, 0), QAbstractItemView.ScrollHint.PositionAtCenter)
                self._start_row_blink_animation(table, row_idx)

    def _start_row_blink_animation(self, table, row_idx):
        """Start blinking animation for a table row"""
        if not hasattr(self, '_blink_timers'):
            self._blink_timers = {}
        
        # Stop any existing animation for this row
        timer_key = f"{id(table)}_{row_idx}"
        if timer_key in self._blink_timers:
            self._blink_timers[timer_key].stop()
        
        # Create new blinking animation
        blink_count = 0
        max_blinks = 6
        
        def blink():
            nonlocal blink_count
            if blink_count >= max_blinks:
                # Reset to normal color and stop
                for col in range(table.columnCount()):
                    item = table.item(row_idx, col)
                    if item:
                        item.setBackground(QColor(255, 255, 255))
                if timer_key in self._blink_timers:
                    self._blink_timers[timer_key].stop()
                    del self._blink_timers[timer_key]
                return
            
            # Alternate between highlight and normal color
            if blink_count % 2 == 0:
                color = QColor(255, 255, 0, 100)  # Yellow highlight
            else:
                color = QColor(255, 255, 255)  # Normal white
            
            for col in range(table.columnCount()):
                item = table.item(row_idx, col)
                if item:
                    item.setBackground(color)
            
            blink_count += 1
        
        # Start the timer
        timer = QTimer()
        timer.timeout.connect(blink)
        timer.start(500)  # Blink every 500ms
        self._blink_timers[timer_key] = timer

    def show_profile_dialog(self):
        """Show user profile dialog with user information"""
        dialog = QDialog(self)
        dialog.setWindowTitle("User Profile")
        dialog.setModal(True)
        dialog.resize(400, 300)
        dialog.setStyleSheet("""
            QDialog {
                background-color: #f8f9fa;
            }
        """)
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Profile header
        header_layout = QHBoxLayout()
        
        # Avatar - Enhanced with larger size and better styling
        avatar_label = QLabel()
        avatar_pixmap = self._default_avatar(self.logged_in_user, 100)  # Increased from 80 to 100
        avatar_label.setPixmap(avatar_pixmap)
        avatar_label.setStyleSheet("""
            QLabel {
                border: 4px solid #e5e7eb;
                border-radius: 50px;
                padding: 6px;
                background-color: #f9fafb;
            }
        """)
        header_layout.addWidget(avatar_label)
        
        # User info
        user_info_layout = QVBoxLayout()
        user_info_layout.setSpacing(5)
        
        # Get user display name (ISM name) instead of EID
        display_name = self._get_user_display_name(self.logged_in_user)
        name_label = QLabel(display_name)
        name_label.setStyleSheet("""
            font-size: 18px;
            font-weight: bold;
            color: #2c3e50;
        """)
        user_info_layout.addWidget(name_label)
        
        # Show EID as secondary info
        eid_label = QLabel(f"EID: {self.logged_in_user}")
        eid_label.setStyleSheet("""
            font-size: 12px;
            color: #6c757d;
            font-style: italic;
        """)
        user_info_layout.addWidget(eid_label)
        
        # Get user role and email from org directory
        role, email = self._get_user_info(self.logged_in_user)
        
        role_label = QLabel(f"Role: {role}")
        role_label.setStyleSheet("""
            font-size: 14px;
            color: #6c757d;
        """)
        user_info_layout.addWidget(role_label)
        
        email_label = QLabel(f"Email: {email}")
        email_label.setStyleSheet("""
            font-size: 14px;
            color: #6c757d;
        """)
        user_info_layout.addWidget(email_label)
        
        header_layout.addLayout(user_info_layout)
        header_layout.addStretch()
        layout.addLayout(header_layout)
        
        # User statistics
        stats_group = QGroupBox("Statistics")
        stats_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #dee2e6;
                border-radius: 6px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 5px 0 5px;
            }
        """)
        stats_layout = QVBoxLayout(stats_group)
        
        # Calculate user statistics
        user_stats = self._calculate_user_statistics()
        
        for stat_name, stat_value in user_stats.items():
            stat_layout = QHBoxLayout()
            stat_label = QLabel(f"{stat_name}:")
            stat_label.setStyleSheet("font-size: 12px; color: #6c757d;")
            stat_value_label = QLabel(str(stat_value))
            stat_value_label.setStyleSheet("font-size: 12px; font-weight: bold; color: #2c3e50;")
            stat_layout.addWidget(stat_label)
            stat_layout.addStretch()
            stat_layout.addWidget(stat_value_label)
            stats_layout.addLayout(stat_layout)
        
        layout.addWidget(stats_group)
        layout.addStretch()
        
        # Buttons
        settings_btn = QPushButton("⚙️ Settings")
        settings_btn.setStyleSheet("""
            QPushButton {
                background-color: #2a7de1;
                color: white;
                padding: 10px 20px;
                border: none;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #1e6bb8;
            }
        """)
        settings_btn.clicked.connect(lambda: (dialog.accept(), self._show_preferences_dialog()))
        
        close_btn = QPushButton("Close")
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #6c757d;
                color: white;
                padding: 10px 20px;
                border: none;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
        """)
        close_btn.clicked.connect(dialog.accept)
        
        button_layout = QHBoxLayout()
        button_layout.addWidget(settings_btn)
        button_layout.addStretch()
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)
        
        dialog.exec()

    def _get_user_info(self, username):
        """Get user role and email from org directory"""
        try:
            org_rows = self._collect_org_directory_rows()
            for name, designation, enterprise_id, email, level, _ in org_rows:
                if name == username or enterprise_id == username:
                    return designation or "ISM", email or f"{username}@accenture.com"
        except Exception:
            pass
        return "ISM", f"{username}@accenture.com"
    
    def _get_user_display_name(self, username):
        """Get user display name (ISM name) from org directory or return username"""
        try:
            org_rows = self._collect_org_directory_rows()
            for name, designation, enterprise_id, email, level, _ in org_rows:
                if enterprise_id == username:
                    return name or username
        except Exception:
            pass
        return username
    
    def _refresh_profile_icon(self):
        """Refresh the profile icon tooltip with current user display name"""
        try:
            if hasattr(self, 'profile_icon') and self.profile_icon:
                display_name = self._get_user_display_name(self.logged_in_user)
                self.profile_icon.setToolTip(f"Profile: {display_name}")
        except Exception as e:
            # Don't show error dialog for profile icon refresh failures
            pass

    def _calculate_user_statistics(self):
        """Calculate user statistics for profile dialog"""
        stats = {}
        
        # Count items assigned to user
        for pane_name, rows in self.data.items():
            if pane_name in PANE_COLUMNS:
                columns = PANE_COLUMNS[pane_name]
                if "ISM Name" in columns:
                    ism_idx = columns.index("ISM Name")
                    count = sum(1 for row in rows if len(row) > ism_idx and row[ism_idx] == self.logged_in_user)
                    if count > 0:
                        stats[f"{pane_name} Items"] = count
        
        # Count overdue items
        today = datetime.now().date()
        overdue_count = 0
        if "Potential Issues" in self.data:
            for row in self.data["Potential Issues"]:
                if len(row) > 4:  # Due Date column
                    due_date_str = row[4]
                    if due_date_str:
                        try:
                            due_date = datetime.strptime(due_date_str, "%Y-%m-%d").date()
                            if due_date < today:
                                overdue_count += 1
                        except ValueError:
                            pass
        stats["Overdue Items"] = overdue_count
        
        return stats

    def show_save_status(self, message="Data Saved", duration=3000):
        """Show save status near the Load Data button"""
        if hasattr(self, 'save_status_label'):
            self.save_status_label.setText(message)
            # Reset to success styling
            self.save_status_label.setStyleSheet("""
                QLabel {
                    color: #27ae60;
                    font-size: 12px;
                    font-weight: bold;
                    padding: 4px 8px;
                    background-color: rgba(39, 174, 96, 0.1);
                    border: 1px solid #27ae60;
                    border-radius: 4px;
                    min-width: 80px;
                    text-align: center;
                }
            """)
            self.save_status_label.setVisible(True)
            
            # Auto-hide after duration - use a proper timer instance
            if not hasattr(self, '_save_status_timer'):
                self._save_status_timer = QTimer()
                self._save_status_timer.setSingleShot(True)
                self._save_status_timer.timeout.connect(self.hide_save_status)
            self._save_status_timer.start(duration)
            
            # Toast notification removed to avoid duplicate "Data Saved" messages

    def hide_save_status(self):
        """Hide the save status indicator"""
        if hasattr(self, 'save_status_label'):
            self.save_status_label.setVisible(False)
    
    def _ensure_main_thread(self, func, *args, **kwargs):
        """Execute function directly (no threading)"""
        try:
            func(*args, **kwargs)
        except Exception as e:
            print(f"Error executing function: {e}")

    def _safe_timer_single_shot(self, msec: int, slot) -> None:
        """Thread-safe timer single shot that works across threads"""
        try:
            QTimer.singleShot(msec, slot)
        except Exception as e:
            print(f"Error setting timer: {e}")
            # Fallback: try to execute immediately
            try:
                slot()
            except Exception as fallback_error:
                print(f"Fallback execution also failed: {fallback_error}")

    def _safe_set_badge_text(self, text: str) -> None:
        """Thread-safe method to set badge text"""
        try:
            if hasattr(self, '_save_badge') and self._save_badge is not None:
                self._save_badge.setText(text)
        except Exception as e:
            print(f"Error setting badge text: {e}")

    def _safe_reset_load_button(self) -> None:
        """Thread-safe method to reset load button"""
        return

    def _safe_hide_toast(self) -> None:
        """Thread-safe method to hide toast notification"""
        try:
            if hasattr(self, '_toast') and self._toast is not None:
                self._toast.hide()
        except Exception as e:
            print(f"Error hiding toast: {e}")


    def _add_pane_tab(self, pane_name: str, columns: list[str]) -> None:
        tab = QWidget()
        vbox = QVBoxLayout(tab)

        controls = ControlsBar(pane_name, self)
        vbox.addWidget(controls)
        try:
            # Visual separator to emphasize sticky header
            sep = QFrame(); sep.setFrameShape(QFrame.Shape.HLine); sep.setStyleSheet("color:#e2e8f0;")
            vbox.addWidget(sep)
        except Exception as e:
            ErrorHandler.handle_ui_error("add separator", e)

        table = QTableWidget()
        table.setColumnCount(len(columns))
        table.setHorizontalHeaderLabels(columns)
        table.setAlternatingRowColors(True)
        table.setSortingEnabled(True)
        table.setWordWrap(True)
        table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        table.customContextMenuRequested.connect(lambda pos, p=pane_name: self._open_context_menu(p, pos))
        # Inline combo editors via delegate
        try:
            table.setItemDelegate(InlineComboDelegate(self, pane_name))
            table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.SelectedClicked | QAbstractItemView.EditTrigger.EditKeyPressed)
        except Exception as e:
            ErrorHandler.handle_ui_error("set table delegate", e)
        # Install multi-line text editors for specific columns across panes, except Activity Brief Update which stays popup
        try:
            for col_index, col_name in enumerate(columns):
                multiline_targets = {"Remarks", "Additional Remarks", "Description", "Briefupdate", "Brief Update"}
                if col_name in multiline_targets:
                    if not (pane_name == "Activities" and col_name in {"Brief Update", "Briefupdate"}):
                        table.setItemDelegateForColumn(col_index, MultiLineTextDelegate(self))
        except Exception as e:
            ErrorHandler.handle_ui_error("set multiline delegates", e)
        # Table visuals
        table.verticalHeader().setDefaultSectionSize(28)
        table.horizontalHeader().setStretchLastSection(True)
        header_font = table.horizontalHeader().font()
        header_font.setPointSize(10)
        header_font.setWeight(QFont.Weight.DemiBold)
        table.horizontalHeader().setFont(header_font)
        
        # Add tooltips to column headers
        tooltips = TooltipManager.get_common_tooltips()
        for col, column_name in enumerate(columns):
            if table.horizontalHeaderItem(col):
                TooltipManager.set_tooltip(
                    table.horizontalHeaderItem(col), 
                    tooltips.get("sort_column", f"Click to sort by {column_name}")
                )
        # Quick filter chips label
        chips = QLabel("")
        chips.setStyleSheet("color:#475569; padding:4px; font-size:11px;")
        vbox.addWidget(chips)
        vbox.addWidget(table)
        try:
            # Ensure controls and chips keep fixed height; table expands and scrolls
            from PyQt6.QtWidgets import QSizePolicy
            table.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        except Exception as e:
            ErrorHandler.handle_ui_error("set table size policy", e)

        self.tables[pane_name] = table
        
        # Add symbols to pane tabs
        pane_symbols = {
            "Initiatives": "🎯",
            "Potential Issues": "⚠️", 
            "Activities": "📋",
            "Client Visits / Audits": "🔍",
            "Accolades": "🏆"
        }
        tab_name = f"{pane_symbols.get(pane_name, '📄')} {pane_name}"
        self.tabs.addTab(tab, tab_name)

        # Add tooltips to control buttons
        tooltips = TooltipManager.get_common_tooltips()
        TooltipManager.set_tooltip(controls.add_btn_primary, tooltips.get("add_row", "Add new row (Ctrl+N)"))
        TooltipManager.set_tooltip(controls.btn_del, tooltips.get("delete_row", "Delete selected row (Delete)"))
        TooltipManager.set_tooltip(controls.btn_imp, "Import data from Excel/CSV")
        TooltipManager.set_tooltip(controls.btn_exp, tooltips.get("export", "Export all data to Excel (Ctrl+E)"))
        TooltipManager.set_tooltip(controls.btn_export_view, "Export current filtered view to Excel")
        TooltipManager.set_tooltip(controls.btn_sample, "Download sample template with required headers")
        TooltipManager.set_tooltip(controls.search, "Type to filter rows by any column")
        
        # Wire controls
        # Wire add button
        controls.add_btn_primary.clicked.connect(lambda: self.add_row(pane_name))
        controls.btn_del.clicked.connect(lambda: self.delete_selected(pane_name))
        controls.btn_imp.clicked.connect(lambda: self._show_import_guide_and_import(pane_name))
        controls.btn_exp.clicked.connect(lambda: self.export_pane(pane_name))
        controls.btn_export_view.clicked.connect(lambda: self.export_current_view(pane_name))
        controls.btn_sample.clicked.connect(lambda: self.show_import_tutorial(pane_name))
        controls.btn_search.clicked.connect(lambda: self.filter_table(pane_name, controls.search.text()))
        controls.btn_clear.clicked.connect(lambda: (controls.search.clear(), self.filter_table(pane_name, "")))

        # Cell edit
        table.cellDoubleClicked.connect(lambda r, c, p=pane_name: self._handle_cell_double_click(p, r, c))
        
        # Cell change events for real-time calendar updates
        if pane_name in ["Potential Issues", "Activities", "Leave Tracker"]:
            table.cellChanged.connect(lambda r, c, p=pane_name: self._on_cell_changed_for_calendar(p, r, c))
        # Hook filters
        def _apply_filters():
            self.filter_pane_rows(pane_name, controls.status_combo.currentText(), controls.rag_combo.currentText(), controls.search.text())
            parts = []
            if controls.search.text():
                parts.append(f"Search='{controls.search.text()}'")
            if controls.status_combo.currentText() != "All":
                parts.append(f"Status={controls.status_combo.currentText()}")
            if controls.rag_combo.currentText() != "All":
                parts.append(f"RAG={controls.rag_combo.currentText()}")
            chips.setText(" | ".join(parts))
        controls.status_combo.currentTextChanged.connect(lambda _=None: _apply_filters())
        controls.rag_combo.currentTextChanged.connect(lambda _=None: _apply_filters())
        controls.btn_search.clicked.connect(lambda _=False: _apply_filters())
        controls.btn_clear.clicked.connect(lambda _=False: (controls.search.clear(), controls.status_combo.setCurrentText("All"), controls.rag_combo.setCurrentText("All"), _apply_filters()))
        _apply_filters()

        # Keyboard-first flow enhancements
        try:
            table.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked | QAbstractItemView.EditTrigger.SelectedClicked | QAbstractItemView.EditTrigger.EditKeyPressed)
            # Note: Event filter removed to prevent threading issues
        except Exception as e:
            self.notifications.show_warning(f"Failed to setup table edit triggers: {str(e)}")
            self._log_change("Error", "Table Setup", f"Failed to setup edit triggers: {str(e)}")

        # Header context menu for show/hide columns
        try:
            header = table.horizontalHeader()
            header.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
            def open_header_menu(pos):
                menu = QMenu(self)
                for i, name in enumerate(columns):
                    act = QAction(name, self)
                    act.setCheckable(True)
                    act.setChecked(not table.isColumnHidden(i))
                    def toggle(idx=i, p=pane_name, t=table):
                        t.setColumnHidden(idx, not t.isColumnHidden(idx))
                        self._persist_hidden_columns(p, t)
                    act.toggled.connect(lambda _=False, fn=toggle: fn())
                    menu.addAction(act)
                menu.exec(table.mapToGlobal(pos))
            header.customContextMenuRequested.connect(open_header_menu)
        except Exception as e:
            ErrorHandler.handle_ui_error("setup header context menu", e)
    # --- Home Page Tab ---
    def _init_home_page(self) -> None:
        """Initialize the modern home page with SVG logo and navigation buttons"""
        layout = QVBoxLayout(self.home_tab)
        layout.setSpacing(30)
        layout.setContentsMargins(40, 40, 40, 40)
        
        # Welcome section with logo
        welcome_frame = QFrame()
        welcome_frame.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1, 
                    stop:0 #667eea, stop:1 #764ba2);
                border-radius: 20px;
                padding: 30px;
            }
        """)
        welcome_layout = QVBoxLayout(welcome_frame)
        
        # Static heading (same size regardless of screen)
        self.home_heading_label = QLabel()
        self.home_heading_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.home_heading_label.setTextFormat(Qt.TextFormat.RichText)
        # Fixed sizes chosen for clarity across screens
        main_size = 34
        sub_size = 20
        line_height = 1.15
        heading_html = (
            f"<div style='color:#ffffff; text-align:center; line-height:{line_height};'>"
            f"<div style='font-weight:800; font-size:{main_size}px; letter-spacing:1px;'>ACCENTURE ISMT</div>"
            f"<div style='font-weight:600; font-size:{sub_size}px;'>Task Management Dashboard</div>"
            f"</div>"
        )
        self.home_heading_label.setText(heading_html)
        welcome_layout.addWidget(self.home_heading_label)
        
        # Welcome text
        welcome_text = QLabel("Welcome to your Task Management Dashboard")
        welcome_text.setStyleSheet("""
            QLabel {
                color: white;
                font-size: 18px;
                font-weight: bold;
                margin-top: 20px;
            }
        """)
        welcome_text.setAlignment(Qt.AlignmentFlag.AlignCenter)
        welcome_layout.addWidget(welcome_text)
        
        layout.addWidget(welcome_frame)
        
        # Navigation buttons grid
        nav_frame = QFrame()
        nav_frame.setStyleSheet("""
            QFrame {
                background: white;
                border-radius: 15px;
                border: 1px solid #e2e8f0;
            }
        """)
        nav_layout = QGridLayout(nav_frame)
        nav_layout.setSpacing(20)
        nav_layout.setContentsMargins(30, 30, 30, 30)
        
        # Define navigation buttons with their corresponding tab names
        nav_buttons = [
            ("📊 Dashboard", "Dashboard", "#4CAF50"),
            ("🎯 Initiatives", "Initiatives", "#2196F3"),
            ("⚠️ Potential Issues", "Potential Issues", "#FF9800"),
            ("📋 Activities", "Activities", "#9C27B0"),
            ("🔍 Client Visits / Audits", "Client Visits / Audits", "#607D8B"),
            ("🏆 Accolades", "Accolades", "#FFC107"),
            ("🗓️ Leave Tracker", "Leave Tracker", "#795548"),
            ("📁 Project Details", "Project Details", "#E91E63")
        ]
        
        # Create navigation buttons
        for i, (text, tab_name, color) in enumerate(nav_buttons):
            btn = QPushButton(text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: {color};
                    color: white;
                    border: none;
                    border-radius: 12px;
                    padding: 20px;
                    font-size: 14px;
                    font-weight: bold;
                    min-height: 40px;
                }}
                QPushButton:hover {{
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                        stop:0 {color}, stop:1 rgba(255,255,255,0.1));
                }}
                QPushButton:pressed {{
                    background: rgba(0,0,0,0.1);
                }}
            """)
            
            # Connect button to switch to corresponding tab
            def make_switch_tab(tab_name):
                def switch_tab():
                    for i in range(self.tabs.count()):
                        if tab_name.lower() in self.tabs.tabText(i).lower():
                            self.tabs.setCurrentIndex(i)
                            break
                return switch_tab
            
            btn.clicked.connect(make_switch_tab(tab_name))
            
            # Add shadow effect
            shadow = QGraphicsDropShadowEffect()
            shadow.setBlurRadius(15)
            shadow.setXOffset(0)
            shadow.setYOffset(5)
            shadow.setColor(QColor(0, 0, 0, 30))
            btn.setGraphicsEffect(shadow)
            
            # Position buttons in grid (4 columns)
            row = i // 4
            col = i % 4
            nav_layout.addWidget(btn, row, col)
        
        layout.addWidget(nav_frame)
        
        # Quick stats section
        stats_frame = QFrame()
        stats_frame.setStyleSheet("""
            QFrame {
                background: #f8fafc;
                border-radius: 15px;
                border: 1px solid #e2e8f0;
            }
        """)
        stats_layout = QHBoxLayout(stats_frame)
        stats_layout.setContentsMargins(20, 20, 20, 20)
        
        # Add some quick stats - create labels that can be updated
        self.stats_labels = {}
        stats_info = [
            ("Total Tasks", "0"),
            ("Active Projects", "0"),
            ("Today's Visits/Audits", "0"),
            ("Current User", self.logged_in_user)
        ]
        
        for title, initial_value in stats_info:
            stat_widget = QWidget()
            stat_layout = QVBoxLayout()
            stat_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
            stat_widget.setLayout(stat_layout)
            
            value_label = QLabel(initial_value)
            value_label.setStyleSheet("""
                QLabel {
                    font-size: 24px;
                    font-weight: bold;
                    color: #2563eb;
                }
            """)
            value_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            title_label = QLabel(title)
            title_label.setStyleSheet("""
                QLabel {
                    font-size: 12px;
                    color: #64748b;
                    margin-top: 5px;
                }
            """)
            title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            stat_layout.addWidget(value_label)
            stat_layout.addWidget(title_label)
            stats_layout.addWidget(stat_widget)
            
            # Store the value label for updating
            self.stats_labels[title] = value_label
        
        # Update stats initially
        self.update_home_stats()
        
        layout.addWidget(stats_frame)
        layout.addStretch()

    def update_home_stats(self):
        """Update the home page statistics"""
        if hasattr(self, 'stats_labels'):
            # Calculate total tasks only from Potential Issues and Activities panes
            task_panes = ["Potential Issues", "Activities"]
            total_tasks = sum(len(self.data.get(pane, [])) for pane in task_panes)
            self.stats_labels["Total Tasks"].setText(str(total_tasks))
            
            # Update active projects count
            self.stats_labels["Active Projects"].setText(str(len(self.projects)))
            
            # Calculate today's client visits/audits
            today_visits = self._count_todays_visits_audits()
            self.stats_labels["Today's Visits/Audits"].setText(str(today_visits))
            
            # Current user doesn't need updating

    def _wrap_resize(self, original_handler, callback):
        # Keep helper in case other parts reuse it, but no-op callback here
        def _handler(event):
            if callable(original_handler):
                return original_handler(event)
            try:
                return QWidget.resizeEvent(self, event)
            except Exception:
                return None
        return _handler

    def _update_home_heading_scale(self) -> None:
        # Fixed, static heading; do nothing (left for compatibility)
        try:
            if hasattr(self, 'home_heading_label') and self.home_heading_label is not None:
                # Ensure fixed sizes are kept
                main_size = 34
                sub_size = 20
                line_height = 1.15
                html = (
                    f"<div style='color:#ffffff; text-align:center; line-height:{line_height};'>"
                    f"<div style='font-weight:800; font-size:{main_size}px; letter-spacing:1px;'>ACCENTURE ISMT</div>"
                    f"</div>"
                )
                self.home_heading_label.setText(html)
        except Exception as e:
            ErrorHandler.handle_ui_error("update home heading static", e)

    def _count_todays_visits_audits(self):
        """Count client visits/audits scheduled for today"""
        try:
            from datetime import datetime
            today = datetime.now().strftime("%Y-%m-%d")
            count = 0
            
            # Check Client Visits / Audits pane
            if "Client Visits / Audits" in self.data:
                cols = PANE_COLUMNS["Client Visits / Audits"]
                
                # Find the date columns
                start_date_idx = None
                end_date_idx = None
                
                for i, col in enumerate(cols):
                    if "Start Date" in col:
                        start_date_idx = i
                    elif "End Date" in col:
                        end_date_idx = i
                
                # Count visits/audits that are scheduled for today
                for row in self.data["Client Visits / Audits"]:
                    if len(row) > max(start_date_idx or 0, end_date_idx or 0):
                        # Check if today falls within the start and end date range
                        start_date = row[start_date_idx] if start_date_idx is not None else ""
                        end_date = row[end_date_idx] if end_date_idx is not None else ""
                        
                        # If today matches start date or end date, or falls between them
                        if start_date == today or end_date == today:
                            count += 1
                        elif start_date and end_date:
                            # Check if today falls between start and end date
                            try:
                                start_dt = datetime.strptime(start_date, "%Y-%m-%d")
                                end_dt = datetime.strptime(end_date, "%Y-%m-%d")
                                today_dt = datetime.strptime(today, "%Y-%m-%d")
                                
                                if start_dt <= today_dt <= end_dt:
                                    count += 1
                            except ValueError:
                                # Skip invalid date formats
                                continue
            
            return count
            
        except Exception as e:
            print(f"Error counting today's visits/audits: {e}")
            return 0

    def sync_projects_from_details(self):
        """Sync the projects list from Project Details pane data"""
        try:
            if "Project Details" in self.data:
                project_details = self.data["Project Details"]
                # Extract Project Name and Project ID from Project Details data
                new_projects = []
                for row in project_details:
                    if len(row) >= 2 and row[0] and row[1]:  # Project Name and Project ID exist
                        project_name = str(row[0]).strip()
                        project_id = str(row[1]).strip()
                        if project_name and project_id:
                            new_projects.append((project_name, project_id))
                
                # Update the projects list
                self.projects = new_projects
                print(f"Synced {len(self.projects)} projects from Project Details")
        except Exception as e:
            print(f"Failed to sync projects from details: {e}")

    def get_project_dropdown_data(self):
        """Get project data for dropdowns with name-to-id and id-to-name mappings"""
        self.sync_projects_from_details()
        project_names = [name for name, _ in self.projects]
        project_ids = [pid for _, pid in self.projects]
        name_to_id = {name: pid for name, pid in self.projects}
        id_to_name = {pid: name for name, pid in self.projects}
        return project_names, project_ids, name_to_id, id_to_name

    # --- Leave Tracker Tab ---
    def _init_leave_tracker_tab(self) -> None:
        self.leave_tab = QWidget()
        lay = QVBoxLayout(self.leave_tab)

        controls = QHBoxLayout()
        ism_combo = QComboBox(); ism_combo.setEditable(False)
        ism_combo.addItem("All ISMs")
        ism_combo.addItem(self.logged_in_user)
        for n in sorted(self._collect_all_isms() | set(self.ism_directory)):
            if ism_combo.findText(n) < 0:
                ism_combo.addItem(n)
        controls.addWidget(QLabel("ISM:")); controls.addWidget(ism_combo)
        type_filter = QComboBox(); type_filter.addItems(["All Types", "WFH", "Planned Leave", "Public Holiday", "Earned Leave", "Casual Leave"]) 
        search_box = QLineEdit(); search_box.setPlaceholderText("Search description or name")
        controls.addWidget(QLabel("Type:")); controls.addWidget(type_filter)
        controls.addWidget(search_box)
        add_btn = QPushButton("Add Leave"); add_range_btn = QPushButton("Add Range"); edit_btn = QPushButton("Edit Selected"); del_btn = QPushButton("Delete Selected")
        add_btn.setObjectName("primary")
        month_prev = QPushButton("◀ Previous Month"); month_next = QPushButton("Next Month ▶")
        import_btn = QPushButton("📥 Import"); sample_btn = QPushButton("📋 Sample")
        export_btn = QPushButton("Export Month (XLSX)"); export_csv_btn = QPushButton("Export Month (CSV)"); export_ics_btn = QPushButton("Export Month (ICS)")
        controls.addStretch(1); controls.addWidget(month_prev); controls.addWidget(month_next)
        controls.addWidget(add_btn); controls.addWidget(add_range_btn); controls.addWidget(edit_btn); controls.addWidget(del_btn)
        controls.addWidget(import_btn); controls.addWidget(sample_btn)
        controls.addWidget(export_btn); controls.addWidget(export_csv_btn); controls.addWidget(export_ics_btn)
        lay.addLayout(controls)

        cal = LeaveCalendar(self); cal.setGridVisible(True)
        # Hide week numbers to avoid confusion with dates > 31
        try:
            cal.setVerticalHeaderFormat(QCalendarWidget.VerticalHeaderFormat.NoVerticalHeader)
        except Exception as e:
            ErrorHandler.handle_ui_error("set calendar header format", e)
        lay.addWidget(cal)
        # Legend
        legend = QHBoxLayout()
        def chip(text, color):
            lbl = QLabel(text)
            lbl.setStyleSheet(f"QLabel{{background:{color};border-radius:8px;padding:4px 8px;color:#0f172a;}}")
            return lbl
        legend.addWidget(chip("WFH", "#e0f2fe"))
        legend.addWidget(chip("Planned", "#dcfce7"))
        legend.addWidget(chip("Public Holiday", "#fef3c7"))
        legend.addWidget(chip("Earned", "#ffedd5"))
        legend.addWidget(chip("Casual", "#fee2e2"))
        legend.addStretch(1)
        lay.addLayout(legend)
        summary = QLabel("")
        table = QTableWidget(); table.setColumnCount(7); table.setHorizontalHeaderLabels(["ISM", "Type", "Duration", "Description", "Approval Status", "Approver Name", "Approval Comments"])
        table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        table.setSortingEnabled(True)
        lay.addWidget(summary); lay.addWidget(table)

        self.leave_calendar = cal
        self.leave_list = table
        self.leave_ism_combo = ism_combo
        # Build per-ISM color map (stable but derived)
        self.leave_ism_colors: dict[str, QColor] = {}

        def refresh_for_date():
            date_str = cal.selectedDate().toString("yyyy-MM-dd")
            table.setRowCount(0)
            rows = self.data.get("Leave Tracker", [])
            cols = PANE_COLUMNS["Leave Tracker"]
            didx = cols.index("Date"); tidx = cols.index("Type"); duridx = cols.index("Duration"); descidx = cols.index("Description"); ismidx = cols.index("ISM Name")
            stidx = cols.index("Approval Status") if "Approval Status" in cols else None
            anidx = cols.index("Approver Name") if "Approver Name" in cols else None
            cmidx = cols.index("Approval Comments") if "Approval Comments" in cols else None
            sel_ism = ism_combo.currentText()
            sel_type = type_filter.currentText()
            query = (search_box.text() or "").lower().strip()
            day_rows = []
            for r in rows:
                if didx < len(r) and r[didx] == date_str:
                    if sel_ism != "All ISMs" and (ismidx >= len(r) or r[ismidx] != sel_ism):
                        continue
                    if sel_type != "All Types" and r[tidx] != sel_type:
                        continue
                    desc_l = (r[descidx] if descidx < len(r) else "")
                    ism_l = (r[ismidx] if ismidx < len(r) else "")
                    if query and query not in (str(desc_l) or "").lower() and query not in (str(ism_l) or "").lower():
                        continue
                    status_v = (r[stidx] if (stidx is not None and stidx < len(r)) else "")
                    appr_name_v = (r[anidx] if (anidx is not None and anidx < len(r)) else "")
                    comments_v = (r[cmidx] if (cmidx is not None and cmidx < len(r)) else "")
                    day_rows.append((ism_l, r[tidx], r[duridx] if duridx < len(r) else "Full Day", r[descidx] if descidx < len(r) else "", status_v, appr_name_v, comments_v))
            table.setRowCount(len(day_rows))
            for i, (ism, typ, dur, desc, status_v, appr_name_v, comments_v) in enumerate(day_rows):
                table.setItem(i, 0, QTableWidgetItem(str(ism)))
                it_type = QTableWidgetItem(str(typ))
                it_dur = QTableWidgetItem(str(dur))
                it_desc = QTableWidgetItem(str(desc))
                # Color background by type
                color = {
                    "WFH": "#e0f2fe",
                    "Planned Leave": "#dcfce7",
                    "Public Holiday": "#fef3c7",
                    "Earned Leave": "#ffedd5",
                    "Casual Leave": "#fee2e2",
                }.get(typ, "#f1f5f9")
                it_type.setBackground(QColor(color)); it_desc.setBackground(QColor(color))
                table.setItem(i, 1, it_type)
                table.setItem(i, 2, it_dur)
                table.setItem(i, 3, it_desc)
                table.setItem(i, 4, QTableWidgetItem(str(status_v or "")))
                table.setItem(i, 5, QTableWidgetItem(str(appr_name_v or "")))
                table.setItem(i, 6, QTableWidgetItem(str(comments_v or "")))
            # Update monthly summary
            try:
                y = cal.selectedDate().year(); m = cal.selectedDate().month()
                counts = {k: 0 for k in ["WFH", "Planned Leave", "Public Holiday", "Earned Leave", "Casual Leave"]}
                total = 0
                for r in rows:
                    if didx < len(r) and r[didx]:
                        dt = datetime.strptime(r[didx], "%Y-%m-%d")
                        if dt.year == y and dt.month == m and (sel_ism == "All ISMs" or (ismidx < len(r) and r[ismidx] == sel_ism)):
                            t = r[tidx]
                            if t in counts:
                                counts[t] += 1; total += 1
                summary.setText(f"This month — Total: {total} | WFH: {counts['WFH']} | Planned: {counts['Planned Leave']} | Public Hol: {counts['Public Holiday']} | Earned: {counts['Earned Leave']} | Casual: {counts['Casual Leave']}")
            except Exception:
                summary.setText("")
        cal.selectionChanged.connect(refresh_for_date)
        ism_combo.currentTextChanged.connect(lambda _=None: (self._refresh_calendar_decorations(), refresh_for_date()))
        type_filter.currentTextChanged.connect(lambda _=None: refresh_for_date())
        search_box.textChanged.connect(lambda _=None: refresh_for_date())
        def nav_prev():
            d = cal.selectedDate(); cal.setSelectedDate(d.addMonths(-1))
            # Jump to first day of new month
            cal.setSelectedDate(QDate(cal.selectedDate().year(), cal.selectedDate().month(), 1))
            refresh_for_date(); self._refresh_calendar_decorations()
        def nav_next():
            d = cal.selectedDate(); cal.setSelectedDate(d.addMonths(1))
            cal.setSelectedDate(QDate(cal.selectedDate().year(), cal.selectedDate().month(), 1))
            refresh_for_date(); self._refresh_calendar_decorations()
        month_prev.clicked.connect(nav_prev); month_next.clicked.connect(nav_next)

        def add_or_edit(existing_idx: int | None = None):
            dlg = QDialog(self)
            dlg.setWindowTitle("Leave Details")
            v = QVBoxLayout(dlg)
            date_edit = QLineEdit(cal.selectedDate().toString("yyyy-MM-dd"))
            type_cb = QComboBox(); type_cb.addItems(["WFH", "Planned Leave", "Public Holiday", "Earned Leave", "Casual Leave"]) 
            dur_cb = QComboBox(); dur_cb.addItems(["Full Day", "Half Day - AM", "Half Day - PM"]) 
            desc_edit = QLineEdit()
            ism_cb = QComboBox();
            # Populate ISM dropdown from Org Directory with fallback
            dir_rows = self._collect_org_directory_rows()
            ism_names = sorted({n for (n, _, _, _, _, _) in dir_rows if n})
            if not ism_names:
                ism_names = sorted(self._collect_all_isms())
            if ism_combo.currentText() != "All ISMs" and ism_combo.currentText() not in ism_names:
                ism_names = [ism_combo.currentText()] + ism_names
            for n in (["All ISMs"] + (ism_names or [self.logged_in_user])):
                ism_cb.addItem(n)
            # Preselect current filter if not All ISMs
            if ism_combo.currentText() != "All ISMs":
                idx = ism_cb.findText(ism_combo.currentText())
                if idx >= 0:
                    ism_cb.setCurrentIndex(idx)
            # Enforce Public Holiday for All ISMs only
            def _enforce_public_holiday_rule():
                if type_cb.currentText() == "Public Holiday":
                    i = ism_cb.findText("All ISMs")
                    if i >= 0:
                        ism_cb.setCurrentIndex(i)
                    try:
                        ism_cb.setDisabled(True)
                    except Exception as e:
                        ErrorHandler.handle_ui_error("excel schema check", e)
                else:
                    try:
                        ism_cb.setDisabled(False)
                    except Exception as e:
                        ErrorHandler.handle_ui_error("build ISM directory from Org Directory", e)
            type_cb.currentTextChanged.connect(lambda _=None: _enforce_public_holiday_rule())
            _enforce_public_holiday_rule()
            grid = QGridLayout();
            grid.addWidget(QLabel("Date (YYYY-MM-DD)"), 0, 0); grid.addWidget(date_edit, 0, 1)
            grid.addWidget(QLabel("Type"), 1, 0); grid.addWidget(type_cb, 1, 1)
            grid.addWidget(QLabel("Duration"), 2, 0); grid.addWidget(dur_cb, 2, 1)
            grid.addWidget(QLabel("Description"), 3, 0); grid.addWidget(desc_edit, 3, 1)
            grid.addWidget(QLabel("ISM Name"), 4, 0); grid.addWidget(ism_cb, 4, 1)
            v.addLayout(grid)
            btns = QHBoxLayout(); ok = QPushButton("Save"); cancel = QPushButton("Cancel")
            ok.setObjectName("primary"); btns.addStretch(1); btns.addWidget(ok); btns.addWidget(cancel); v.addLayout(btns)
            ok.clicked.connect(dlg.accept); cancel.clicked.connect(dlg.reject)
            # preload if editing
            if existing_idx is not None:
                row = self.data["Leave Tracker"][existing_idx]
                try:
                    date_edit.setText(row[0]); type_cb.setCurrentText(row[1]); dur_cb.setCurrentText(row[2] if len(row) > 2 else "Full Day"); desc_edit.setText(row[3 if len(row) > 3 else 2]);
                    # set ISM dropdown
                    cur_ism = row[4] if len(row) > 4 else (row[3] if len(row) > 3 else self.logged_in_user)
                    idx = ism_cb.findText(cur_ism)
                    if idx >= 0:
                        ism_cb.setCurrentIndex(idx)
                except Exception as e:
                    ErrorHandler.handle_ui_error("preload leave data", e)
            if dlg.exec() != QDialog.DialogCode.Accepted:
                return
            # validate date
            try:
                datetime.strptime(date_edit.text().strip(), "%Y-%m-%d")
            except Exception:
                QMessageBox.warning(self, "Invalid Date", "Please enter date as YYYY-MM-DD")
                return
            # Validation: Public Holiday must be for All ISMs
            if type_cb.currentText() == "Public Holiday" and (ism_cb.currentText() != "All ISMs"):
                QMessageBox.warning(self, "Public Holiday", "'Public Holiday' can only be logged for 'All ISMs'.")
                return
            new_row = [date_edit.text().strip(), type_cb.currentText(), dur_cb.currentText(), desc_edit.text().strip(), ism_cb.currentText() or self.logged_in_user]
            # Fill approval routing based on selected ISM's manager
            try:
                cols_l = PANE_COLUMNS.get("Leave Tracker", [])
                # Ensure row length
                new_row = list(new_row) + [""] * max(0, len(cols_l) - len(new_row))
                if "Approval Status" in cols_l:
                    new_row[cols_l.index("Approval Status")] = "Pending"
                # Route to the manager of the selected ISM
                selected_ism_name = new_row[cols_l.index("ISM Name")] if "ISM Name" in cols_l else (ism_cb.currentText() or "")
                approver_eid = self._resolve_manager_eid_for_user_name(selected_ism_name) or ""
                approver_name = self._resolve_name_from_eid(approver_eid) if approver_eid else ""
                if "Approver Enterprise ID" in cols_l:
                    new_row[cols_l.index("Approver Enterprise ID")] = approver_eid
                if "Approver Name" in cols_l:
                    new_row[cols_l.index("Approver Name")] = approver_name
                # Requested by reflects the submitter (current user)
                if "Requested By Enterprise ID" in cols_l:
                    new_row[cols_l.index("Requested By Enterprise ID")] = self.logged_in_user
                if "Requested By Name" in cols_l:
                    new_row[cols_l.index("Requested By Name")] = self._resolve_name_from_eid(self.logged_in_user) or self.logged_in_user
            except Exception:
                pass
            if existing_idx is None:
                self.data["Leave Tracker"].append(new_row)
                self._log("Added leave entry")
            else:
                self.data["Leave Tracker"][existing_idx] = new_row
                self._log("Updated leave entry")
            # Notify approver of new request
            try:
                cols_l = PANE_COLUMNS.get("Leave Tracker", [])
                date_val = new_row[cols_l.index("Date")] if "Date" in cols_l else ""
                approver_eid = new_row[cols_l.index("Approver Enterprise ID")] if "Approver Enterprise ID" in cols_l else ""
                requester_name = new_row[cols_l.index("Requested By Name")] if "Requested By Name" in cols_l else (self._resolve_name_from_eid(self.logged_in_user) or self.logged_in_user)
                self._notify_approver_new_leave(approver_eid, date_val, requester_name)
            except Exception:
                pass
            self._save_autosave(); self._save_backend_sqlite(); refresh_for_date(); self._refresh_calendar_decorations()

        def on_add():
            # Add new entry and notify manager/requester
            add_or_edit(None)
            try:
                # Determine date and ISM for notification
                date_str = cal.selectedDate().toString("yyyy-MM-dd")
                cols_l = PANE_COLUMNS.get("Leave Tracker", [])
                ism_name = ""
                try:
                    ism_name = ism_combo.currentText() or ""
                except Exception:
                    pass
                # Resolve participants
                requester_eid = str(getattr(self, 'logged_in_user', '')).strip()
                eff_ism_name = ism_name or (self._resolve_name_from_eid(requester_eid) or requester_eid)
                approver_eid = self._resolve_manager_eid_for_user_name(eff_ism_name) or ""
                self._notify_leave_logged(date_str, eff_ism_name, approver_eid, requester_eid)
            except Exception:
                pass
        def on_add_range():
            dlg = QDialog(self)
            dlg.setWindowTitle("Add Leave Range")
            v = QVBoxLayout(dlg)
            start_edit = QLineEdit(cal.selectedDate().toString("yyyy-MM-dd"))
            end_edit = QLineEdit(cal.selectedDate().toString("yyyy-MM-dd"))
            type_cb = QComboBox(); type_cb.addItems(["WFH", "Planned Leave", "Public Holiday", "Earned Leave", "Casual Leave"]) 
            dur_cb = QComboBox(); dur_cb.addItems(["Full Day", "Half Day - AM", "Half Day - PM"]) 
            desc_edit = QLineEdit(); ism_cb2 = QComboBox()
            # Populate ISM dropdown for range
            dir_rows = self._collect_org_directory_rows()
            ism_names = sorted({n for (n, _, _, _, _, _) in dir_rows if n}) or sorted(self._collect_all_isms()) or [self.logged_in_user]
            for n in (["All ISMs"] + ism_names):
                ism_cb2.addItem(n)
            if ism_combo.currentText() != "All ISMs":
                idx = ism_cb2.findText(ism_combo.currentText())
                if idx >= 0:
                    ism_cb2.setCurrentIndex(idx)
            # Enforce Public Holiday for All ISMs only
            def _enforce_public_holiday_range():
                if type_cb.currentText() == "Public Holiday":
                    i = ism_cb2.findText("All ISMs")
                    if i >= 0:
                        ism_cb2.setCurrentIndex(i)
                    try:
                        ism_cb2.setDisabled(True)
                    except Exception as e:
                        ErrorHandler.handle_ui_error("save badge timer", e)
                else:
                    try:
                        ism_cb2.setDisabled(False)
                    except Exception as e:
                        ErrorHandler.handle_ui_error("save badge timer", e)
            type_cb.currentTextChanged.connect(lambda _=None: _enforce_public_holiday_range())
            _enforce_public_holiday_range()
            grid = QGridLayout();
            grid.addWidget(QLabel("Start Date (YYYY-MM-DD)"), 0, 0); grid.addWidget(start_edit, 0, 1)
            grid.addWidget(QLabel("End Date (YYYY-MM-DD)"), 1, 0); grid.addWidget(end_edit, 1, 1)
            grid.addWidget(QLabel("Type"), 2, 0); grid.addWidget(type_cb, 2, 1)
            grid.addWidget(QLabel("Duration"), 3, 0); grid.addWidget(dur_cb, 3, 1)
            grid.addWidget(QLabel("Description"), 4, 0); grid.addWidget(desc_edit, 4, 1)
            grid.addWidget(QLabel("ISM Name"), 5, 0); grid.addWidget(ism_cb2, 5, 1)
            v.addLayout(grid)
            btns = QHBoxLayout(); ok = QPushButton("Add Range"); cancel = QPushButton("Cancel")
            ok.setObjectName("primary"); btns.addStretch(1); btns.addWidget(ok); btns.addWidget(cancel); v.addLayout(btns)
            ok.clicked.connect(dlg.accept); cancel.clicked.connect(dlg.reject)
            if dlg.exec() != QDialog.DialogCode.Accepted:
                return
            try:
                sdt = datetime.strptime(start_edit.text().strip(), "%Y-%m-%d").date()
                edt = datetime.strptime(end_edit.text().strip(), "%Y-%m-%d").date()
            except Exception:
                QMessageBox.warning(self, "Invalid Date", "Please enter valid start/end dates as YYYY-MM-DD")
                return
            if edt < sdt:
                QMessageBox.warning(self, "Invalid Range", "End date must be on/after start date")
                return
            cur = sdt
            count = 0
            # Validation: Public Holiday must be for All ISMs
            if type_cb.currentText() == "Public Holiday" and (ism_cb2.currentText() != "All ISMs"):
                QMessageBox.warning(self, "Public Holiday", "'Public Holiday' can only be logged for 'All ISMs'.")
                return
            while cur <= edt:
                base = [
                    cur.strftime("%Y-%m-%d"), type_cb.currentText(), dur_cb.currentText(), desc_edit.text().strip(), ism_cb2.currentText() or self.logged_in_user
                ]
                try:
                    cols_l = PANE_COLUMNS.get("Leave Tracker", [])
                    base = list(base) + [""] * max(0, len(cols_l) - len(base))
                    if "Approval Status" in cols_l:
                        base[cols_l.index("Approval Status")] = "Pending"
                    # Route to manager of selected ISM for each day
                    selected_ism_name = base[cols_l.index("ISM Name")] if "ISM Name" in cols_l else (ism_cb2.currentText() or "")
                    approver_eid = self._resolve_manager_eid_for_user_name(selected_ism_name) or ""
                    approver_name = self._resolve_name_from_eid(approver_eid) if approver_eid else ""
                    if "Approver Enterprise ID" in cols_l:
                        base[cols_l.index("Approver Enterprise ID")] = approver_eid
                    if "Approver Name" in cols_l:
                        base[cols_l.index("Approver Name")] = approver_name
                    if "Requested By Enterprise ID" in cols_l:
                        base[cols_l.index("Requested By Enterprise ID")] = self.logged_in_user
                    if "Requested By Name" in cols_l:
                        base[cols_l.index("Requested By Name")] = self._resolve_name_from_eid(self.logged_in_user) or self.logged_in_user
                except Exception:
                    pass
                self.data["Leave Tracker"].append(base)
                cur += timedelta(days=1)
                count += 1
            self._log(f"Added {count} leave day(s)")
            self._save_autosave(); self._save_backend_sqlite(); refresh_for_date(); self._refresh_calendar_decorations()
            try:
                # Notify both approver (manager) and requester/ISM for the starting date
                start_date_str = sdt.strftime("%Y-%m-%d")
                # Resolve participants
                requester_eid = str(getattr(self, 'logged_in_user', '')).strip()
                # ISM name selected in the dialog, fallback to requester name
                selected_ism_name = ism_cb2.currentText() or self._resolve_name_from_eid(requester_eid) or requester_eid
                approver_eid = self._resolve_manager_eid_for_user_name(selected_ism_name) or ""
                self._notify_leave_logged(start_date_str, selected_ism_name, approver_eid, requester_eid)
            except Exception:
                pass
        def on_edit():
            # Find selected in list against underlying data for selected date
            sel = table.currentRow()
            if sel < 0:
                return
            # Map to index in data
            date_str = cal.selectedDate().toString("yyyy-MM-dd")
            cols = PANE_COLUMNS["Leave Tracker"]; didx = cols.index("Date"); ismidx = cols.index("ISM Name")
            # Determine ISM of selected row
            sel_ism = table.item(sel, 0).text() if table.item(sel, 0) else ism_combo.currentText()
            counter = -1
            for i, r in enumerate(self.data.get("Leave Tracker", [])):
                if didx < len(r) and r[didx] == date_str and (not sel_ism or r[ismidx] == sel_ism):
                    counter += 1
                    if counter == sel:
                        add_or_edit(i)
                        return
        def on_delete():
            sel = table.currentRow()
            if sel < 0:
                return
            date_str = cal.selectedDate().toString("yyyy-MM-dd")
            cols = PANE_COLUMNS["Leave Tracker"]; didx = cols.index("Date"); ismidx = cols.index("ISM Name")
            sel_ism = table.item(sel, 0).text() if table.item(sel, 0) else ism_combo.currentText()
            counter = -1
            for i in range(len(self.data.get("Leave Tracker", [])) - 1, -1, -1):
                r = self.data["Leave Tracker"][i]
                if didx < len(r) and r[didx] == date_str and (not sel_ism or r[ismidx] == sel_ism):
                    counter += 1
                    if counter == sel:
                        self.data["Leave Tracker"].pop(i)
                        self._log("Deleted leave entry")
                        self._save_autosave(); self._save_backend_sqlite(); refresh_for_date(); self._refresh_calendar_decorations()
                        return

        add_btn.clicked.connect(on_add)
        edit_btn.clicked.connect(on_edit)
        del_btn.clicked.connect(on_delete)
        add_range_btn.clicked.connect(on_add_range)

        # Approve/Reject controls for approvers (acts on selected row in the list for chosen date)
        def _decide_selected(approved: bool):
            try:
                sel = table.currentRow()
                if sel < 0:
                    return
                dstr = cal.selectedDate().toString("yyyy-MM-dd")
                cols = PANE_COLUMNS.get("Leave Tracker", [])
                didx = cols.index("Date") if "Date" in cols else None
                ismidx = cols.index("ISM Name") if "ISM Name" in cols else None
                appr_name_idx = cols.index("Approver Name") if "Approver Name" in cols else None
                appr_eid_idx = cols.index("Approver Enterprise ID") if "Approver Enterprise ID" in cols else None
                req_eid_idx = cols.index("Requested By Enterprise ID") if "Requested By Enterprise ID" in cols else None
                status_idx = cols.index("Approval Status") if "Approval Status" in cols else None
                dec_date_idx = cols.index("Decision Date") if "Decision Date" in cols else None
                # Find matching underlying row by (date, ism filter)
                sel_ism = table.item(sel, 0).text() if table.item(sel, 0) else ism_combo.currentText()
                counter = -1
                data_idx = None
                for i, r in enumerate(self.data.get("Leave Tracker", [])):
                    ok = True
                    if didx is not None and didx < len(r):
                        ok = ok and str(r[didx]) == dstr
                    if ismidx is not None and ismidx < len(r) and sel_ism:
                        ok = ok and str(r[ismidx]) == sel_ism
                    if ok:
                        counter += 1
                        if counter == sel:
                            data_idx = i
                            break
                if data_idx is None:
                    return
                row = self.data["Leave Tracker"][data_idx]
                # Write decision fields
                approver_name = self._resolve_name_from_eid(self.logged_in_user) or str(self.logged_in_user)
                if status_idx is not None:
                    while len(row) <= status_idx: row.append("")
                    row[status_idx] = "Approved" if approved else "Rejected"
                if dec_date_idx is not None:
                    while len(row) <= dec_date_idx: row.append("")
                    row[dec_date_idx] = datetime.now().strftime("%Y-%m-%d")
                if appr_name_idx is not None:
                    while len(row) <= appr_name_idx: row.append("")
                    row[appr_name_idx] = approver_name
                if appr_eid_idx is not None:
                    while len(row) <= appr_eid_idx: row.append("")
                    row[appr_eid_idx] = str(self.logged_in_user)
                # Persist and refresh UI
                self._save_autosave(); self._save_backend_sqlite(); refresh_for_date(); self._refresh_calendar_decorations()
                # Notify requester
                try:
                    requester_eid = row[req_eid_idx] if (req_eid_idx is not None and req_eid_idx < len(row)) else str(getattr(self, 'logged_in_user', ''))
                    self._notify_leave_decision_to_requester(requester_eid, "Approved" if approved else "Rejected", dstr, approver_name)
                except Exception:
                    pass
            except Exception as e:
                ErrorHandler.handle_ui_error("leave decision", e)

        approve_btn = QPushButton("Approve"); reject_btn = QPushButton("Reject")
        approve_btn.setObjectName("primary")
        approve_btn.clicked.connect(lambda _=False: _decide_selected(True))
        reject_btn.clicked.connect(lambda _=False: _decide_selected(False))
        # Place decision buttons below the add/edit/delete buttons
        btn_row2 = QHBoxLayout(); btn_row2.addStretch(1); btn_row2.addWidget(approve_btn); btn_row2.addWidget(reject_btn)
        # Add to the main Leave tab layout
        try:
            if self.leave_tab and self.leave_tab.layout():
                self.leave_tab.layout().addLayout(btn_row2)
        except Exception:
            pass

        # Context menu helpers for LeaveCalendar
        def _ensure_leave_tab_visible(date_str: str):
            try:
                # Focus Leave Tracker tab
                for i in range(self.tabs.count()):
                    if self.tabs.tabText(i).startswith("Leave Tracker"):
                        self.tabs.setCurrentIndex(i)
                        break
                # Select date in calendar
                try:
                    dt = datetime.strptime(date_str, "%Y-%m-%d")
                    cal.setSelectedDate(QDate(dt.year, dt.month, dt.day))
                except Exception as e:
                    ErrorHandler.handle_ui_error("set calendar date", e)
            except Exception as e:
                ErrorHandler.handle_ui_error("ensure leave tab visible", e)

        def _find_leave_indices_for_date(dstr: str) -> list[tuple[int, str]]:
            out: list[tuple[int, str]] = []
            rows = self.data.get("Leave Tracker", [])
            cols = PANE_COLUMNS["Leave Tracker"]; didx = cols.index("Date"); ismidx = cols.index("ISM Name"); tidx = cols.index("Type"); duridx = cols.index("Duration"); descidx = cols.index("Description")
            sel_ism = ism_combo.currentText()
            for idx, r in enumerate(rows):
                if didx < len(r) and r[didx] == dstr:
                    ism = r[ismidx] if ismidx < len(r) else ""
                    # Apply ISM filter unless All ISMs is selected
                    if sel_ism != "All ISMs" and ism != sel_ism:
                        continue
                    typ = r[tidx] if tidx < len(r) else ""
                    dur = r[duridx] if duridx < len(r) else "Full Day"
                    desc = r[descidx] if descidx < len(r) else ""
                    label = f"{ism} — {typ} ({dur}) — {desc}".strip()
                    out.append((idx, label))
            return out

        def _pick_leave_for_date(dstr: str) -> int | None:
            items = _find_leave_indices_for_date(dstr)
            if not items:
                QMessageBox.information(self, "No Leave Entries", f"No leave entries found for {dstr}")
                return None
            if len(items) == 1:
                return items[0][0]
            dlg = QDialog(self)
            dlg.setWindowTitle(f"Select Leave — {dstr}")
            dlg.resize(400, 300)
            v = QVBoxLayout(dlg)
            lst = QListWidget()
            for _, label in items:
                lst.addItem(label)
            v.addWidget(lst)
            btns = QHBoxLayout(); ok = QPushButton("Select"); cancel = QPushButton("Cancel"); ok.setObjectName("primary"); btns.addStretch(1); btns.addWidget(ok); btns.addWidget(cancel)
            v.addLayout(btns)
            ok.clicked.connect(dlg.accept); cancel.clicked.connect(dlg.reject)
            if dlg.exec() != QDialog.DialogCode.Accepted:
                return None
            row = lst.currentRow()
            if row < 0 or row >= len(items):
                return None
            return items[row][0]

        self._open_leave_add_dialog_for_date = lambda d: (_ensure_leave_tab_visible(d), on_add())
        def _add_range_for_date(dstr: str):
            try:
                # Pre-fill the range dialog to the chosen date for both start and end
                # Reuse on_add_range flow by setting calendar selection first
                _ensure_leave_tab_visible(dstr)
                on_add_range()
            except Exception as e:
                ErrorHandler.handle_ui_error("add range for date", e)
        self._open_leave_add_range_for_date = lambda d: _add_range_for_date(d)

        def _edit_for_date(dstr: str):
            idx = _pick_leave_for_date(dstr)
            if idx is not None:
                add_or_edit(idx)
        self._open_leave_edit_dialog_for_date = lambda d: (_ensure_leave_tab_visible(d), _edit_for_date(d))

        def _delete_for_date(dstr: str):
            idx = _pick_leave_for_date(dstr)
            if idx is None:
                return
            try:
                # Get the leave entry details before deleting for confirmation
                leave_entry = self.data["Leave Tracker"][idx]
                ism = leave_entry[4] if len(leave_entry) > 4 else "Unknown"
                typ = leave_entry[1] if len(leave_entry) > 1 else "Unknown"
                desc = leave_entry[3] if len(leave_entry) > 3 else "No description"
                
                # Ask for confirmation
                reply = QMessageBox.question(self, "Confirm Delete", 
                    f"Are you sure you want to delete this leave entry?\n\n"
                    f"Date: {dstr}\n"
                    f"ISM: {ism}\n"
                    f"Type: {typ}\n"
                    f"Description: {desc}",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                
                if reply == QMessageBox.StandardButton.Yes:
                    self.data["Leave Tracker"].pop(idx)
                    self._log("Deleted leave entry")
                    self._save_autosave(); self._save_backend_sqlite(); refresh_for_date(); self._refresh_calendar_decorations()
                    QMessageBox.information(self, "Success", "Leave entry deleted successfully")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to delete leave entry: {str(e)}")
        self._open_leave_delete_for_date = lambda d: (_ensure_leave_tab_visible(d), _delete_for_date(d))

        def _export_month(to_csv: bool = False) -> None:
            y = cal.selectedDate().year(); m = cal.selectedDate().month()
            rows = self.data.get("Leave Tracker", [])
            cols = PANE_COLUMNS["Leave Tracker"]
            didx = cols.index("Date"); ismidx = cols.index("ISM Name")
            sel_ism = ism_combo.currentText()
            month_rows = []
            for r in rows:
                try:
                    dt = datetime.strptime(r[didx], "%Y-%m-%d")
                except Exception:
                    continue
                if dt.year == y and dt.month == m and (not sel_ism or r[ismidx] == sel_ism):
                    month_rows.append(r)
            if to_csv:
                path, _ = QFileDialog.getSaveFileName(self, "Export Leave Month (CSV)", f"leave_{y}_{m:02d}.csv", "CSV (*.csv)")
                if not path:
                    return
                try:
                    pd.DataFrame(month_rows, columns=cols).to_csv(path, index=False)
                    QMessageBox.information(self, "Export", "CSV exported")
                except Exception as e:
                    QMessageBox.critical(self, "Export", str(e))
            else:
                path, _ = QFileDialog.getSaveFileName(self, "Export Leave Month (CSV)", f"leave_{y}_{m:02d}.csv", "CSV (*.csv)")
                if not path:
                    return
                try:
                    pd.DataFrame(month_rows, columns=cols).to_csv(path, index=False)
                    QMessageBox.information(self, "Export", "CSV exported")
                except Exception as e:
                    QMessageBox.critical(self, "Export", str(e))

        def _export_ics() -> None:
            try:
                from datetime import date as _date
                y = cal.selectedDate().year(); m = cal.selectedDate().month()
                rows = self.data.get("Leave Tracker", [])
                cols = PANE_COLUMNS["Leave Tracker"]
                didx = cols.index("Date"); tidx = cols.index("Type"); duridx = cols.index("Duration"); descidx = cols.index("Description"); ismidx = cols.index("ISM Name")
                sel_ism = ism_combo.currentText()
                events = []
                for r in rows:
                    try:
                        dt = datetime.strptime(r[didx], "%Y-%m-%d")
                    except Exception:
                        continue
                    if dt.year == y and dt.month == m and (not sel_ism or r[ismidx] == sel_ism):
                        # Build simple VEVENT
                        dstr = dt.strftime("%Y%m%d")
                        summary_txt = f"{r[tidx]} ({r[duridx]}) — {r[descidx]}"
                        events.append("".join([
                            "BEGIN:VEVENT\n",
                            f"DTSTART;VALUE=DATE:{dstr}\n",
                            f"DTEND;VALUE=DATE:{dstr}\n",
                            f"SUMMARY:{summary_txt}\n",
                            "END:VEVENT\n",
                        ]))
                cal_str = "BEGIN:VCALENDAR\nVERSION:2.0\nPRODID:-//LeaveTracker//EN\n" + "".join(events) + "END:VCALENDAR\n"
                path, _ = QFileDialog.getSaveFileName(self, "Export Leave Month (ICS)", f"leave_{y}_{m:02d}.ics", "ICS (*.ics)")
                if not path:
                    return
                with open(path, "w", encoding="utf-8") as f:
                    f.write(cal_str)
                QMessageBox.information(self, "Export", "ICS exported")
            except Exception as e:
                QMessageBox.critical(self, "Export", str(e))

        export_btn.clicked.connect(lambda: _export_month(False))
        export_csv_btn.clicked.connect(lambda: _export_month(True))
        export_ics_btn.clicked.connect(_export_ics)
        # Import CSV for bulk add
        import_btn = QPushButton("Import CSV")
        controls.addWidget(import_btn)
        def do_import_csv():
            path, _ = QFileDialog.getOpenFileName(self, "Import Leave CSV", "", "CSV (*.csv)")
            if not path:
                return
            try:
                df = pd.read_csv(path)
                # Expect columns: Date, Type, Duration, Description, ISM Name (Duration optional -> defaults Full Day)
                cols = {c.lower(): c for c in df.columns}
                required = ["date", "type", "description", "ism name"]
                for rc in required:
                    if rc not in cols:
                        raise ValueError("CSV must include Date, Type, Description, ISM Name")
                # Build rows in new schema order
                add_rows = []
                for _, row in df.iterrows():
                    date_v = str(row[cols['date']])
                    type_v = str(row[cols['type']])
                    dur_v = str(row[cols['duration']]) if 'duration' in cols else 'Full Day'
                    desc_v = str(row[cols['description']])
                    ism_v = str(row[cols['ism name']])
                    # Build with approval defaults routed to selected ISM's manager
                    req_eid = self.logged_in_user
                    req_name = self._resolve_name_from_eid(req_eid) or self.logged_in_user
                    mgr_eid = self._resolve_manager_eid_for_user_name(ism_v)
                    mgr_name = self._resolve_name_from_eid(mgr_eid) if mgr_eid else ""
                    add_rows.append([
                        date_v, type_v, dur_v, desc_v, ism_v,
                        "Pending", (mgr_eid or ""), (mgr_name or ""),
                        req_eid, req_name, "", ""
                    ])
                self.data["Leave Tracker"].extend(add_rows)
                self._save_autosave(); self._save_backend_sqlite()
                self._log(f"Imported {len(add_rows)} leave rows")
                self._refresh_calendar_decorations(); refresh_for_date()
            except Exception as e:
                QMessageBox.critical(self, "Import", str(e))
        import_btn.clicked.connect(do_import_csv)
        
        # Sample download functionality
        def do_download_sample():
            try:
                path, _ = QFileDialog.getSaveFileName(self, "Save Leave Tracker Sample", "leave_tracker_sample.csv", "CSV (*.csv);;Excel (*.xlsx)")
                if not path:
                    return
                
                # Create sample data
                sample_data = [
                    ["2024-01-15", "WFH", "Full Day", "Working from home", "john.smith"],
                    ["2024-01-20", "Earned Leave", "Full Day", "Personal vacation", "sarah.johnson"],
                    ["2024-01-25", "Casual Leave", "Half Day", "Medical appointment", "mike.wilson"]
                ]
                
                if path.lower().endswith('.csv'):
                    import csv
                    with open(path, 'w', newline='', encoding='utf-8') as csvfile:
                        writer = csv.writer(csvfile)
                        writer.writerow(["Date", "Type", "Duration", "Description", "ISM Name"])
                        writer.writerows(sample_data)
                elif path.lower().endswith('.xlsx'):
                    if PANDAS_AVAILABLE:
                        import pandas as pd
                        df = pd.DataFrame(sample_data, columns=["Date", "Type", "Duration", "Description", "ISM Name"])
                        df.to_excel(path, index=False)
                    else:
                        QMessageBox.warning(self, "Sample Download", "Pandas not available for Excel export. Please use CSV format.")
                        return
                
                QMessageBox.information(self, "Sample Download", f"Sample template saved to {path}")
            except Exception as e:
                QMessageBox.critical(self, "Sample Download", str(e))
        
        sample_btn.clicked.connect(do_download_sample)

        # Ensure data structure exists and normalize to new schema
        if "Leave Tracker" not in self.data:
            self.data["Leave Tracker"] = []
        else:
            # Normalize existing rows to include approval fields
            cols_lt = PANE_COLUMNS.get("Leave Tracker", [])
            target_len = len(cols_lt)
            normalized = []
            for r in self.data["Leave Tracker"]:
                row = list(r) + [""] * max(0, target_len - len(r))
                row = row[:target_len]
                # If missing approver/requester, set sensible defaults
                try:
                    # Requested By (current user context)
                    if "Requested By Enterprise ID" in cols_lt:
                        ridx = cols_lt.index("Requested By Enterprise ID")
                        if not row[ridx]:
                            row[ridx] = self.logged_in_user
                    if "Requested By Name" in cols_lt:
                        rnidx = cols_lt.index("Requested By Name")
                        if not row[rnidx]:
                            row[rnidx] = self._resolve_name_from_eid(self.logged_in_user) or self.logged_in_user
                    if "Approver Enterprise ID" in cols_lt:
                        aidx = cols_lt.index("Approver Enterprise ID")
                        if not row[aidx]:
                            # Route to manager of the ISM in the row, if present
                            try:
                                ism_name = row[cols_lt.index("ISM Name")] if "ISM Name" in cols_lt else None
                            except Exception:
                                ism_name = None
                            row[aidx] = (self._resolve_manager_eid_for_user_name(ism_name) if ism_name else None) or ""
                    if "Approver Name" in cols_lt:
                        anidx = cols_lt.index("Approver Name")
                        if not row[anidx]:
                            app_eid = row[cols_lt.index("Approver Enterprise ID")] if "Approver Enterprise ID" in cols_lt else ""
                            row[anidx] = self._resolve_name_from_eid(app_eid) or ""
                    if "Approval Status" in cols_lt:
                        sidx = cols_lt.index("Approval Status")
                        if not row[sidx]:
                            row[sidx] = "Pending"
                except Exception:
                    pass
                normalized.append(row)
            self.data["Leave Tracker"] = normalized

        self._refresh_calendar_decorations()
        refresh_for_date()
        # Add to tabs if not already present
        try:
            found = False
            for i in range(self.tabs.count()):
                if self.tabs.tabText(i) == "Leave Tracker":
                    found = True; break
            if not found:
                idx = self.tabs.addTab(self.leave_tab, "Leave Tracker 🗓️")
                # Place the tab near data tabs (after Dashboard and Org Chart)
                # PyQt6 QTabWidget doesn't support reordering directly; we accept default append
        except Exception as e:
            ErrorHandler.handle_ui_error("add leave tracker tab", e)

    def _init_project_details_tab(self) -> None:
        """Initialize the Project Details tab with comprehensive functionality"""
        # Ensure we're in the main thread
        # Initialize project details tab directly
        
        try:
            self.project_details_tab = QWidget()
            layout = QVBoxLayout(self.project_details_tab)
        except Exception as e:
            ErrorHandler.handle_ui_error("init project details tab", e)
            return
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Header with title and add button
        header_layout = QHBoxLayout()
        title_label = QLabel("Project Details")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(title_label)
        header_layout.addStretch()
        
        # Button container for equal sizing and spacing
        button_container = QHBoxLayout()
        button_container.setSpacing(8)  # Small gap between buttons
        
        self.add_project_btn = QPushButton("Add Project")
        self.add_project_btn.setFixedSize(120, 36)  # Equal size
        self.add_project_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                border-radius: 4px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        self.add_project_btn.clicked.connect(self._add_project)
        button_container.addWidget(self.add_project_btn)
        
        # Export button
        self.export_projects_btn = QPushButton("Export Projects")
        self.export_projects_btn.setFixedSize(120, 36)  # Equal size
        self.export_projects_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                border-radius: 4px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        self.export_projects_btn.clicked.connect(self._export_projects)
        button_container.addWidget(self.export_projects_btn)
        
        # Import button
        self.import_projects_btn = QPushButton("📥 Import")
        self.import_projects_btn.setFixedSize(120, 36)
        self.import_projects_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                border-radius: 4px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        self.import_projects_btn.clicked.connect(self._import_projects)
        button_container.addWidget(self.import_projects_btn)
        
        # Sample button
        self.sample_projects_btn = QPushButton("📋 Sample")
        self.sample_projects_btn.setFixedSize(120, 36)
        self.sample_projects_btn.setStyleSheet("""
            QPushButton {
                background-color: #9b59b6;
                color: white;
                border: none;
                border-radius: 4px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #8e44ad;
            }
        """)
        self.sample_projects_btn.clicked.connect(self._download_projects_sample)
        button_container.addWidget(self.sample_projects_btn)
        
        # Add button container to header
        header_layout.addLayout(button_container)
        layout.addLayout(header_layout)
        
        # Filter controls
        filter_layout = QHBoxLayout()
        
        # ISM Filter
        ism_filter_label = QLabel("Filter by ISM:")
        ism_filter_label.setStyleSheet("font-weight: bold; color: #34495e;")
        filter_layout.addWidget(ism_filter_label)
        
        self.ism_filter_combo = QComboBox()
        self.ism_filter_combo.setStyleSheet("""
            QComboBox {
                background-color: white;
                border: 1px solid #bdc3c7;
                padding: 5px;
                border-radius: 3px;
                min-width: 150px;
            }
            QComboBox:focus {
                border-color: #3498db;
            }
        """)
        self.ism_filter_combo.addItem("All ISMs")
        self.ism_filter_combo.currentTextChanged.connect(self._filter_by_ism)
        filter_layout.addWidget(self.ism_filter_combo)
        
        # Search box
        search_label = QLabel("Search:")
        search_label.setStyleSheet("font-weight: bold; color: #34495e; margin-left: 20px;")
        filter_layout.addWidget(search_label)
        
        self.projects_search = QLineEdit()
        self.projects_search.setPlaceholderText("Search projects...")
        self.projects_search.setStyleSheet("""
            QLineEdit {
                background-color: white;
                border: 1px solid #bdc3c7;
                padding: 5px;
                border-radius: 3px;
                min-width: 200px;
            }
            QLineEdit:focus {
                border-color: #3498db;
            }
        """)
        self.projects_search.textChanged.connect(self._filter_projects_table)
        filter_layout.addWidget(self.projects_search)
        
        filter_layout.addStretch()
        layout.addLayout(filter_layout)
        
        # Projects Table
        self.projects_table = QTableWidget()
        self.projects_table.setColumnCount(17)
        self.projects_table.setHorizontalHeaderLabels([
            "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM", 
            "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type", 
            "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
        ])
        
        # Table styling
        self.projects_table.setStyleSheet("""
            QTableWidget {
                background-color: white;
                border: 1px solid #bdc3c7;
                border-radius: 5px;
                gridline-color: #ecf0f1;
                alternate-background-color: #f8f9fa;
            }
            QHeaderView::section {
                background-color: #34495e;
                color: white;
                padding: 10px 8px;
                border: none;
                border-right: 1px solid #2c3e50;
                font-weight: bold;
                font-size: 12px;
            }
            QHeaderView::section:last {
                border-right: none;
            }
            QTableWidget::item {
                padding: 8px;
                border: none;
                border-bottom: 1px solid #ecf0f1;
                color: #2c3e50;
                font-size: 13px;
            }
            QTableWidget::item:selected {
                background-color: #f8f9fa;
                color: #2c3e50;
            }
            QLineEdit {
                background-color: white;
                border: 1px solid #bdc3c7;
                padding: 4px;
                border-radius: 3px;
            }
            QComboBox {
                background-color: white;
                border: 1px solid #bdc3c7;
                padding: 4px;
                border-radius: 3px;
            }
        """)
        
        # Table configuration
        self.projects_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.projects_table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.projects_table.setSortingEnabled(True)
        self.projects_table.setAlternatingRowColors(True)
        self.projects_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.projects_table.customContextMenuRequested.connect(self._show_projects_context_menu)
        
        # Enable text wrapping
        self.projects_table.setWordWrap(True)
        
        # Set row height for better text wrapping
        self.projects_table.verticalHeader().setDefaultSectionSize(50)
        self.projects_table.verticalHeader().setVisible(False)
        
        # Configure dropdown delegates and columns map
        try:
            # Update global columns map if available so rebuild/export use the new schema
            try:
                cols_pd = [
                    "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM",
                    "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type",
                    "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
                ]
                if 'PANE_COLUMNS' in globals() and isinstance(PANE_COLUMNS, dict):
                    PANE_COLUMNS["Project Details"] = cols_pd
            except Exception:
                pass
            self._setup_project_delegates()
        except Exception as e:
            ErrorHandler.handle_ui_error("setup project delegates", e)
        
        self.projects_table.setEditTriggers(
            QAbstractItemView.EditTrigger.DoubleClicked | 
            QAbstractItemView.EditTrigger.SelectedClicked | 
            QAbstractItemView.EditTrigger.EditKeyPressed
        )
        
        # Connect signals
        # Use debounced handler to avoid recursive edits and crashes
        try:
            self.projects_table.blockSignals(False)
        except Exception:
            pass
        self.projects_table.cellChanged.connect(self._on_project_cell_changed)
        self._updating_table = False
        
        layout.addWidget(self.projects_table)
        
        # Summary Metrics Section
        summary_group = QGroupBox("Project Summary Metrics")
        summary_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                font-size: 14px;
                color: #2c3e50;
                border: 2px solid #bdc3c7;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
                background-color: #f8f9fa;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 8px 0 8px;
                color: #2c3e50;
                background-color: #f8f9fa;
            }
        """)
        summary_layout = QVBoxLayout(summary_group)
        
        # Metrics grid
        metrics_layout = QGridLayout()
        
        # Total ISM Hours
        self.total_hours_label = QLabel("Total ISM Hours: 0")
        self.total_hours_label.setStyleSheet(
            "QLabel { font-size: 14px; color: #27ae60; font-weight: bold; border: 2px solid #27ae60; border-radius: 6px; padding: 6px 10px; background-color: rgba(39,174,96,0.08); }"
        )
        metrics_layout.addWidget(self.total_hours_label, 0, 0)
        
        # Total Audit Deals
        self.total_audit_label = QLabel("Total Audit Deals: 0")
        self.total_audit_label.setStyleSheet(
            "QLabel { font-size: 14px; color: #e74c3c; font-weight: bold; border: 2px solid #e74c3c; border-radius: 6px; padding: 6px 10px; background-color: rgba(231,76,60,0.08); }"
        )
        metrics_layout.addWidget(self.total_audit_label, 0, 1)
        
        # Total Voice Deals
        self.total_voice_label = QLabel("Total Voice Deals: 0")
        self.total_voice_label.setStyleSheet(
            "QLabel { font-size: 14px; color: #9b59b6; font-weight: bold; border: 2px solid #9b59b6; border-radius: 6px; padding: 6px 10px; background-color: rgba(155,89,182,0.08); }"
        )
        metrics_layout.addWidget(self.total_voice_label, 0, 2)
        
        # Total Projects
        self.total_projects_label = QLabel("Total Projects: 0")
        self.total_projects_label.setStyleSheet(
            "QLabel { font-size: 14px; color: #3498db; font-weight: bold; border: 2px solid #3498db; border-radius: 6px; padding: 6px 10px; background-color: rgba(52,152,219,0.08); }"
        )
        metrics_layout.addWidget(self.total_projects_label, 0, 3)
        
        # Unique Primary/Secondary ISM metrics removed as requested
        
        # Per-ISM Audit Deals removed as requested
        
        summary_layout.addLayout(metrics_layout)
        
        # Per-ISM sections (side-by-side, responsive)
        per_ism_row = QHBoxLayout()
        per_ism_row.setSpacing(12)
        
        # Primary ISM Project Counts
        self.primary_ism_group = QGroupBox("Primary ISM Project Counts")
        self.primary_ism_group.setStyleSheet("QGroupBox { font-size: 12px; font-weight: bold; color: #2c3e50; }")
        primary_ism_layout = QVBoxLayout(self.primary_ism_group)
        self.primary_counts_label = QLabel("")
        self.primary_counts_label.setWordWrap(True)
        self.primary_counts_label.setStyleSheet("font-size: 11px; color: #7f8c8d;")
        primary_ism_layout.addWidget(self.primary_counts_label)
        try:
            from PyQt6.QtWidgets import QSizePolicy
            self.primary_ism_group.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        except Exception:
            pass
        per_ism_row.addWidget(self.primary_ism_group)
        
        # Secondary ISM Project Counts
        self.secondary_ism_group = QGroupBox("Secondary ISM Project Counts")
        self.secondary_ism_group.setStyleSheet("QGroupBox { font-size: 12px; font-weight: bold; color: #2c3e50; }")
        secondary_ism_layout = QVBoxLayout(self.secondary_ism_group)
        self.secondary_counts_label = QLabel("")
        self.secondary_counts_label.setWordWrap(True)
        self.secondary_counts_label.setStyleSheet("font-size: 11px; color: #7f8c8d;")
        secondary_ism_layout.addWidget(self.secondary_counts_label)
        try:
            from PyQt6.QtWidgets import QSizePolicy
            self.secondary_ism_group.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        except Exception:
            pass
        per_ism_row.addWidget(self.secondary_ism_group)
        
        # ISM Hours per Primary ISM
        self.ism_hours_group = QGroupBox("ISM Hours per Primary ISM")
        self.ism_hours_group.setStyleSheet("QGroupBox { font-size: 12px; font-weight: bold; color: #2c3e50; }")
        ism_hours_layout = QVBoxLayout(self.ism_hours_group)
        self.ism_hours_details = QLabel("")
        self.ism_hours_details.setWordWrap(True)
        self.ism_hours_details.setStyleSheet("font-size: 11px; color: #7f8c8d;")
        ism_hours_layout.addWidget(self.ism_hours_details)
        try:
            from PyQt6.QtWidgets import QSizePolicy
            self.ism_hours_group.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        except Exception:
            pass
        per_ism_row.addWidget(self.ism_hours_group)
        
        # Ensure equal stretch for responsive alignment
        per_ism_row.setStretch(0, 1)
        per_ism_row.setStretch(1, 1)
        per_ism_row.setStretch(2, 1)
        summary_layout.addLayout(per_ism_row)
        
        layout.addWidget(summary_group)
        
        # Load data and populate ISM filter
        try:
            self._load_projects_data()
            self._populate_ism_filter()
            self._update_summary_metrics()
        except Exception as e:
            ErrorHandler.handle_ui_error("project details initialization", e)
    
    def _setup_project_delegates(self):
        """Setup project table delegates"""
        try:
            if hasattr(self, 'projects_table') and self.projects_table:
                # Check if delegates are already set to avoid conflicts
                if hasattr(self, '_project_delegates_setup') and self._project_delegates_setup:
                    return
                
                # Install base inline combo delegate (if available in the app)
                try:
                    inline_delegate = InlineComboDelegate(self, "Project Details")
                    self.projects_table.setItemDelegate(inline_delegate)
                except Exception:
                    pass
                
                # Column indexes with new schema
                vertical_col = 3
                complexity_col = 4
                headcount_col = 15
                seatcount_col = 16

                # Vertical dropdown (fixed options)
                class VerticalDelegate(QStyledItemDelegate):
                    def createEditor(self, parent, option, index):
                        combo = QComboBox(parent)
                        combo.addItems([
                            "", "Health", "Pharma", "F&A", "Support", "CMT", "Digital", "F&A-Mortgage", "SCM", "Resources", "Utilities Back office"
                        ])
                        return combo
                    def setEditorData(self, editor, index):
                        val = index.data() or ""
                        i = editor.findText(val)
                        editor.setCurrentIndex(i if i >= 0 else 0)
                    def setModelData(self, editor, model, index):
                        model.setData(index, editor.currentText())

                self.projects_table.setItemDelegateForColumn(vertical_col, VerticalDelegate(self.projects_table))

                # Complexity dropdown if you want to restrict values (keep existing free text otherwise)
                class ComplexityDelegate(QStyledItemDelegate):
                    def createEditor(self, parent, option, index):
                        combo = QComboBox(parent)
                        combo.addItems(["", "High", "Medium", "Low"])
                        return combo
                    def setEditorData(self, editor, index):
                        val = (index.data() or "").strip()
                        i = editor.findText(val)
                        editor.setCurrentIndex(i if i >= 0 else 0)
                    def setModelData(self, editor, model, index):
                        model.setData(index, editor.currentText())

                self.projects_table.setItemDelegateForColumn(complexity_col, ComplexityDelegate(self.projects_table))

                # ISM Hours (integer-only) and Head/Seat Count numeric delegates
                class IntSpinDelegate(QStyledItemDelegate):
                    def __init__(self, parent=None):
                        super().__init__(parent)
                    def createEditor(self, parent, option, index):
                        spin = QSpinBox(parent)
                        spin.setRange(0, 9999)
                        return spin
                    def setEditorData(self, editor, index):
                        try:
                            val = int(str(index.data() or "0").strip() or "0")
                        except Exception:
                            val = 0
                        editor.setValue(max(0, min(9999, val)))
                    def setModelData(self, editor, model, index):
                        model.setData(index, str(editor.value()))

                int_delegate = IntSpinDelegate(self.projects_table)
                # ISM Hours column index = 8 in the 17-col schema
                self.projects_table.setItemDelegateForColumn(8, int_delegate)
                self.projects_table.setItemDelegateForColumn(headcount_col, int_delegate)
                self.projects_table.setItemDelegateForColumn(seatcount_col, int_delegate)
                
                # Mark as setup to prevent duplicate setup
                self._project_delegates_setup = True
        except Exception as e:
            ErrorHandler.handle_ui_error("setup project delegates", e)
    
    def _load_projects_data(self):
        """Load projects data into the table"""
        try:
            self._updating_table = True
            if "Project Details" not in self.data:
                self.data["Project Details"] = []
            
            rows = self.data["Project Details"]
            self.projects_table.setRowCount(len(rows))
            
            # Normalize rows to new column count (17)
            target_cols = 17
            for row_idx, row_data in enumerate(rows):
                norm = list(row_data) + [""] * max(0, target_cols - len(row_data))
                norm = norm[:target_cols]
                for col_idx, value in enumerate(norm):
                    if col_idx < self.projects_table.columnCount():
                        item = QTableWidgetItem(str(value) if value is not None else "")
                        self.projects_table.setItem(row_idx, col_idx, item)
            # Apply validators and dropdowns for new columns
            try:
                # Vertical dropdown values
                vertical_options = [
                    "Health", "Pharma", "F&A", "Support", "CMT", "Digital", "F&A-Mortgage", "SCM", "Resources", "Utilities Back office"
                ]
                # Install a combo delegate for Vertical and Complexity columns if using inline editing
                if hasattr(self, 'projects_table'):
                    from PyQt6.QtWidgets import QComboBox
                    # Pre-populate current rows with valid defaults where empty
                    for r in range(self.projects_table.rowCount()):
                        # Vertical (col 3)
                        it_v = self.projects_table.item(r, 3)
                        if it_v and not it_v.text().strip():
                            it_v.setText("")
                        # Complexity (col 4) keep existing values
                        # Head Count (col 15) and Seat Count (col 16) ensure numeric
                        for c in (15, 16):
                            it_n = self.projects_table.item(r, c)
                            if it_n and not it_n.text().strip():
                                it_n.setText("0")
            except Exception:
                pass
            
            self._updating_table = False
        except Exception as e:
            self._updating_table = False
            ErrorHandler.handle_ui_error("load projects data", e)
    
    def _add_project(self) -> None:
        """Add a new project entry to Project Details and refresh the table."""
        try:
            # Basic dialog to capture at least Project Name and Project ID
            dlg = QDialog(self)
            dlg.setWindowTitle("Add Project")
            lay = QVBoxLayout(dlg)
            grid = QGridLayout()
            name_edit = QLineEdit(); id_edit = QLineEdit()
            name_edit.setPlaceholderText("Project Name")
            id_edit.setPlaceholderText("Project ID")
            grid.addWidget(QLabel("Project Name:"), 0, 0); grid.addWidget(name_edit, 0, 1)
            grid.addWidget(QLabel("Project ID:"), 1, 0); grid.addWidget(id_edit, 1, 1)
            # Optional quick fields
            loc_edit = QLineEdit(); vert_edit = QLineEdit(); prim_ism = QLineEdit(); sec_ism = QLineEdit(); hours_edit = QLineEdit()
            grid.addWidget(QLabel("Location:"), 2, 0); grid.addWidget(loc_edit, 2, 1)
            grid.addWidget(QLabel("Vertical:"), 3, 0); grid.addWidget(vert_edit, 3, 1)
            grid.addWidget(QLabel("Primary ISM:"), 4, 0); grid.addWidget(prim_ism, 4, 1)
            grid.addWidget(QLabel("Secondary ISM:"), 5, 0); grid.addWidget(sec_ism, 5, 1)
            grid.addWidget(QLabel("ISM Hours:"), 6, 0); grid.addWidget(hours_edit, 6, 1)
            lay.addLayout(grid)
            btns = QHBoxLayout(); ok = QPushButton("Add"); cancel = QPushButton("Cancel")
            ok.setObjectName("primary"); btns.addStretch(1); btns.addWidget(ok); btns.addWidget(cancel); lay.addLayout(btns)
            ok.clicked.connect(dlg.accept); cancel.clicked.connect(dlg.reject)
            if dlg.exec() != QDialog.DialogCode.Accepted:
                return
            name = name_edit.text().strip(); pid = id_edit.text().strip()
            if not name or not pid:
                QMessageBox.warning(self, "Add Project", "Project Name and Project ID are required.")
                return
            # Ensure data structure exists
            if "Project Details" not in self.data:
                self.data["Project Details"] = []
            # Build row in the 17-column schema
            row = [""] * 17
            # Columns: [0] Project Name, [1] Project ID, [2] Location, [3] Vertical, [4] Complexity,
            # [5] Complexity Details, [6] Primary ISM, [7] Secondary ISM, [8] ISM Hours,
            # [9] CXL Name, [10] SDL Name, [11] Connectivity Type, [12] Audits in Deal,
            # [13] Voice Solution, [14] Contact Center, [15] Head Count, [16] Seat Count
            row[0] = name
            row[1] = pid
            row[2] = loc_edit.text().strip()
            row[3] = vert_edit.text().strip()
            row[6] = prim_ism.text().strip()
            row[7] = sec_ism.text().strip()
            row[8] = (hours_edit.text().strip() or "0")
            # Append and refresh
            self.data["Project Details"].append(row)
            # Persist and refresh UI
            self._save_autosave()
            self._save_backend_sqlite()
            self._load_projects_data()
            try:
                self._populate_ism_filter()
            except Exception:
                pass
            # Also refresh via simple table refresher if that variant is active
            try:
                self._refresh_project_table()
            except Exception:
                pass
            try:
                self._populate_ism_filter()
            except Exception:
                pass
            try:
                if hasattr(self, '_update_summary_metrics'):
                    self._update_summary_metrics()
            except Exception:
                pass
            # Select the newly added row
            try:
                r = max(0, self.projects_table.rowCount() - 1)
                self.projects_table.selectRow(r)
                self.projects_table.scrollToItem(self.projects_table.item(r, 0), QAbstractItemView.ScrollHint.PositionAtCenter)
            except Exception:
                pass
            self._show_toast("Project added", level="SUCCESS")
        except Exception as e:
            ErrorHandler.handle_ui_error("add project", e)
    
    def _populate_ism_filter(self):
        """Populate ISM filter"""
        try:
            if hasattr(self, 'ism_filter_combo') and self.ism_filter_combo:
                # Clear existing items
                self.ism_filter_combo.clear()
                
                # Add "All ISMs" option
                self.ism_filter_combo.addItem("All ISMs")
                
                # Get unique ISMs from Project Details
                isms = set()
                if "Project Details" in self.data:
                    for row in self.data["Project Details"]:
                        if len(row) > 4:  # Primary ISM column
                            primary_ism = row[4] if row[4] else ""
                            if primary_ism:
                                isms.add(primary_ism)
                        if len(row) > 5:  # Secondary ISM column
                            secondary_ism = row[5] if row[5] else ""
                            if secondary_ism:
                                isms.add(secondary_ism)
                
                # Add sorted ISMs
                for ism in sorted(isms):
                    self.ism_filter_combo.addItem(ism)
        except Exception as e:
            ErrorHandler.handle_ui_error("populate ism filter", e)
    
    def _update_summary_metrics(self):
        """Update summary metrics"""
        try:
            if not hasattr(self, 'total_hours_label'):
                return
            
            # Calculate metrics
            total_hours = 0
            total_audit_deals = 0
            total_voice_deals = 0
            total_projects = 0
            
            if "Project Details" in self.data:
                # Use column-name lookups to be robust to schema changes
                cols = [
                    "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM",
                    "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type",
                    "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
                ]
                # Best effort: fall back to PANE_COLUMNS if available
                try:
                    from_cols = PANE_COLUMNS.get("Project Details", cols)
                    if isinstance(from_cols, list) and from_cols:
                        cols = from_cols
                except Exception:
                    pass
                idx_hours = cols.index("ISM Hours") if "ISM Hours" in cols else None
                idx_audit = cols.index("Audits in Deal") if "Audits in Deal" in cols else None
                idx_voice = cols.index("Voice Solution") if "Voice Solution" in cols else None
                for row in self.data["Project Details"]:
                    total_projects += 1
                    # ISM Hours
                    if idx_hours is not None and idx_hours < len(row):
                        try:
                            hours = int(str(row[idx_hours]).strip() or "0")
                            total_hours += hours
                        except Exception:
                            pass
                    # Audit deals
                    if idx_audit is not None and idx_audit < len(row):
                        if str(row[idx_audit]).strip().lower() == "yes":
                            total_audit_deals += 1
                    # Voice deals
                    if idx_voice is not None and idx_voice < len(row):
                        if str(row[idx_voice]).strip().lower() == "yes":
                            total_voice_deals += 1
            
            # Update labels
            self.total_hours_label.setText(f"Total ISM Hours: {total_hours}")
            self.total_audit_label.setText(f"Total Audit Deals: {total_audit_deals}")
            self.total_voice_label.setText(f"Total Voice Deals: {total_voice_deals}")
            self.total_projects_label.setText(f"Total Projects: {total_projects}")
            
        except Exception as e:
            ErrorHandler.handle_ui_error("update summary metrics", e)
    
    def _populate_ism_filter(self):
        """Populate ISM filter with org chart members and ISM directory; fallback to named indices."""
        try:
            if hasattr(self, 'ism_filter_combo') and self.ism_filter_combo:
                current_selection = self.ism_filter_combo.currentText()
                self.ism_filter_combo.clear()
                self.ism_filter_combo.addItem("All ISMs")
                
                isms: set[str] = set()
                # 1) Org chart members (primary)
                try:
                    org_rows = self._collect_org_directory_rows()
                    for name, _, _, _, _, _ in org_rows:
                        name_clean = (name or "").strip()
                        if name_clean:
                            isms.add(name_clean)
                except Exception as e:
                    ErrorHandler.handle_ui_error("collect org chart ISMs", e)
                
                # 2) ISM directory (user-managed)
                try:
                    for ism in getattr(self, 'ism_directory', []) or []:
                        ism_clean = (ism or "").strip()
                        if ism_clean:
                            isms.add(ism_clean)
                except Exception:
                    pass
                
                # 3) Fallback: scan data using named indices only (avoid hardcoded columns)
                try:
                    # From Project Details primary/secondary ISM by name
                    cols_pd = PANE_COLUMNS.get("Project Details", [])
                    idx_p = cols_pd.index("Primary ISM") if "Primary ISM" in cols_pd else None
                    idx_s = cols_pd.index("Secondary ISM") if "Secondary ISM" in cols_pd else None
                    for project in self.data.get("Project Details", []) or []:
                        for col in [idx_p, idx_s]:
                            if col is None:
                                continue
                            if col < len(project):
                                val = (project[col] or "").strip()
                                if val:
                                    isms.add(val)
                    # From other panes' "ISM Name"
                    for pane_name, rows in self.data.items():
                        cols = PANE_COLUMNS.get(pane_name, [])
                        if "ISM Name" not in cols:
                            continue
                        idx = cols.index("ISM Name")
                        for r in rows:
                            if idx < len(r):
                                val = (r[idx] or "").strip()
                                if val:
                                    isms.add(val)
                except Exception:
                    pass
                
                # Populate combo
                for ism in sorted(isms):
                    self.ism_filter_combo.addItem(ism)
                
                # Restore selection
                idx_restore = self.ism_filter_combo.findText(current_selection)
                if idx_restore >= 0:
                    self.ism_filter_combo.setCurrentIndex(idx_restore)
        except Exception as e:
            ErrorHandler.handle_ui_error("populate ism filter", e)
    
    def refresh_ism_filter(self):
        """Refresh ISM filter"""
        try:
            self._populate_ism_filter()
        except Exception as e:
            ErrorHandler.handle_ui_error("refresh ism filter", e)
    
    def _resolve_ism_name_from_eid(self, enterprise_id: str) -> Optional[str]:
        """Lookup ISM name from Org Chart by Enterprise ID.
        Returns name if found, otherwise None.
        """
        try:
            rows = self._collect_org_directory_rows()
            for (name, _, eid, _, _, _) in rows:
                if (eid or "").strip().lower() == str(enterprise_id or "").strip().lower():
                    return (name or "").strip()
        except Exception as e:
            ErrorHandler.handle_ui_error("resolve ISM from EID", e, {"eid": enterprise_id})
        return None

    def _resolve_manager_eid_for_user_eid(self, enterprise_id: str) -> Optional[str]:
        """Given a user's Enterprise ID, return their manager's Enterprise ID from Org Chart."""
        try:
            rows = self._collect_org_directory_rows()
            for (_, _, eid, _, manager_eid, _) in rows:
                if (eid or "").strip().lower() == str(enterprise_id or "").strip().lower():
                    return (manager_eid or "").strip() or None
        except Exception as e:
            ErrorHandler.handle_ui_error("resolve manager EID", e, {"eid": enterprise_id})
        return None

    def _resolve_name_from_eid(self, enterprise_id: str) -> Optional[str]:
        """Generic EID->Name resolver from Org Chart."""
        return self._resolve_ism_name_from_eid(enterprise_id)

    def _resolve_eid_from_name(self, name: str) -> Optional[str]:
        """Lookup Enterprise ID from Org Chart by ISM name."""
        try:
            nm = (name or "").strip().lower()
            if not nm:
                return None
            for (n, _, eid, _, _, _) in self._collect_org_directory_rows():
                if (n or "").strip().lower() == nm:
                    return (eid or "").strip() or None
        except Exception as e:
            ErrorHandler.handle_ui_error("resolve EID from name", e, {"name": name})
        return None

    def _resolve_manager_eid_for_user_name(self, name: str) -> Optional[str]:
        """Given a user's name, return their manager's EID via Org Chart."""
        eid = self._resolve_eid_from_name(name)
        if not eid:
            return None
        return self._resolve_manager_eid_for_user_eid(eid)

    def _notify_approver_new_leave(self, approver_eid: str, for_date: str, requester_name: str) -> None:
        """Lightweight notification to approver about a new leave request."""
        try:
            if not approver_eid:
                return
            msg = f"New leave request from {requester_name} for {for_date}. Open Approvals to review."
            # If approver is current user, surface a toast immediately
            if str(approver_eid).strip().lower() == str(self.logged_in_user).strip().lower():
                self._show_toast(msg, level="INFO")
                self.statusBar().showMessage(msg, 6000)
            # Update notification bell badge
            try:
                if hasattr(self, '_pending_notifications'):
                    self._pending_notifications.append(msg)
                else:
                    self._pending_notifications = [msg]
                self._update_notification_count()
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("notify approver new leave", e)

    def _notify_leave_decision(self, requester_eid: str, status: str, date_str: str, approver_name: str) -> None:
        """Notify requester about decision outcome."""
        try:
            msg = f"Your leave request for {date_str} was {status} by {approver_name}."
            # If requester is current user, show toast immediately
            if str(requester_eid or "").strip().lower() == str(self.logged_in_user).strip().lower():
                self._show_toast(msg, level="SUCCESS" if status == "Approved" else "WARNING")
                self.statusBar().showMessage(msg, 8000)
            # Update bell badge for requester context as well
            try:
                if hasattr(self, '_pending_notifications'):
                    self._pending_notifications.append(msg)
                else:
                    self._pending_notifications = [msg]
                self._update_notification_count()
            except Exception:
                pass
        except Exception as e:
            ErrorHandler.handle_ui_error("notify leave decision", e)

    def open_leave_approvals_dialog(self) -> None:
        """Dialog for approvers to approve/reject pending leave requests assigned to them."""
        try:
            pane = "Leave Tracker"
            if pane not in self.data:
                QMessageBox.information(self, "Leave Approvals", "No leave data available.")
                return
            cols = PANE_COLUMNS.get(pane, [])
            required_cols = [
                "Date", "Type", "Duration", "Description", "ISM Name",
                "Approval Status", "Approver Enterprise ID", "Approver Name",
                "Requested By Enterprise ID", "Requested By Name", "Decision Date", "Approval Comments"
            ]
            # Build list of pending rows for current approver
            try:
                idx = {c: cols.index(c) for c in required_cols if c in cols}
            except ValueError:
                QMessageBox.warning(self, "Leave Approvals", "Leave Tracker columns are not configured as expected.")
                return
            my_eid = str(self.logged_in_user).strip().lower()
            pending = []
            for i, r in enumerate(self.data.get(pane, [])):
                status = (r[idx["Approval Status"]] if idx.get("Approval Status") is not None and idx["Approval Status"] < len(r) else "").strip()
                app_eid = (r[idx["Approver Enterprise ID"]] if idx.get("Approver Enterprise ID") is not None and idx["Approver Enterprise ID"] < len(r) else "").strip().lower()
                if status in ("", "Pending") and app_eid == my_eid:
                    pending.append((i, r))

            dlg = QDialog(self)
            dlg.setWindowTitle("Leave Approvals")
            lay = QVBoxLayout(dlg)
            tbl = QTableWidget()
            hdrs = ["Date", "ISM Name", "Type", "Duration", "Description", "Requested By", "Status"]
            tbl.setColumnCount(len(hdrs))
            tbl.setHorizontalHeaderLabels(hdrs)
            tbl.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            tbl.setRowCount(len(pending))
            for row_idx, (data_idx, row_vals) in enumerate(pending):
                def g(name):
                    ci = idx.get(name)
                    return row_vals[ci] if ci is not None and ci < len(row_vals) else ""
                vals = [g("Date"), g("ISM Name"), g("Type"), g("Duration"), g("Description"), g("Requested By Name"), g("Approval Status") or "Pending"]
                for c, v in enumerate(vals):
                    tbl.setItem(row_idx, c, QTableWidgetItem(str(v)))
                # store original row index
                tbl.setVerticalHeaderItem(row_idx, QTableWidgetItem(str(data_idx)))
            lay.addWidget(tbl)

            comments = QLineEdit(); comments.setPlaceholderText("Optional comments...")
            lay.addWidget(comments)

            btns = QHBoxLayout()
            approve_btn = QPushButton("Approve")
            reject_btn = QPushButton("Reject")
            close_btn = QPushButton("Close")
            btns.addWidget(approve_btn); btns.addWidget(reject_btn); btns.addStretch(1); btns.addWidget(close_btn)
            lay.addLayout(btns)

            def decide(approved: bool):
                sel = tbl.currentRow()
                if sel < 0:
                    QMessageBox.information(dlg, "Leave Approvals", "Please select a request.")
                    return
                data_row_idx_item = tbl.verticalHeaderItem(sel)
                if not data_row_idx_item:
                    return
                data_idx = int(data_row_idx_item.text())
                r = self.data[pane][data_idx]
                now_str = datetime.now().strftime("%Y-%m-%d")
                # Write back status, decision date, approver name, comments
                def setf(field, value):
                    if field in cols:
                        ci = cols.index(field)
                        while len(r) <= ci:
                            r.append("")
                        r[ci] = value
                setf("Approval Status", "Approved" if approved else "Rejected")
                setf("Decision Date", now_str)
                setf("Approver Name", self._resolve_name_from_eid(self.logged_in_user) or str(self.logged_in_user))
                setf("Approval Comments", comments.text().strip())
                # Persist
                self._save_autosave(); self._save_backend_sqlite()
                # Notify requester
                requester_eid = r[cols.index("Requested By Enterprise ID")] if "Requested By Enterprise ID" in cols and cols.index("Requested By Enterprise ID") < len(r) else ""
                date_str = r[cols.index("Date")] if "Date" in cols and cols.index("Date") < len(r) else ""
                self._notify_leave_decision(requester_eid, "Approved" if approved else "Rejected", date_str, self._resolve_name_from_eid(self.logged_in_user) or str(self.logged_in_user))
                # Refresh UI and dialog
                try:
                    self.rebuild_table(pane)
                except Exception:
                    pass
                QMessageBox.information(dlg, "Leave Approvals", f"Request {'approved' if approved else 'rejected'}.")
                dlg.accept()

            approve_btn.clicked.connect(lambda: decide(True))
            reject_btn.clicked.connect(lambda: decide(False))
            close_btn.clicked.connect(dlg.reject)
            dlg.resize(900, 420)
            dlg.exec()
        except Exception as e:
            ErrorHandler.handle_ui_error("open leave approvals dialog", e)

    def _filter_by_ism(self, ism_text: str):
        """Filter Project Details table by selected ISM using correct column indices."""
        try:
            if not hasattr(self, 'projects_table') or not self.projects_table:
                return
            # Show all rows if 'All ISMs'
            if ism_text == "All ISMs":
                for row in range(self.projects_table.rowCount()):
                    self.projects_table.setRowHidden(row, False)
                return
            # Resolve column indices by header names
            cols = PANE_COLUMNS.get("Project Details", [])
            idx_primary = cols.index("Primary ISM") if "Primary ISM" in cols else None
            idx_secondary = cols.index("Secondary ISM") if "Secondary ISM" in cols else None
            for row in range(self.projects_table.rowCount()):
                row_visible = False
                for col in [idx_primary, idx_secondary]:
                    if col is None or col < 0:
                        continue
                    item = self.projects_table.item(row, col)
                    if item and item.text().strip() == ism_text:
                        row_visible = True
                        break
                self.projects_table.setRowHidden(row, not row_visible)
        except Exception as e:
            ErrorHandler.handle_ui_error("filter by ism", e)
    


    

    def _update_summary_metrics(self):
        """Update all summary metrics based on current data (robust to schema changes)"""
        projects = self.data.get("Project Details", [])
        cols = PANE_COLUMNS.get("Project Details", [])
        # Resolve indices by name to avoid breakage when columns are added/reordered
        idx_primary = cols.index("Primary ISM") if "Primary ISM" in cols else None
        idx_secondary = cols.index("Secondary ISM") if "Secondary ISM" in cols else None
        idx_hours = cols.index("ISM Hours") if "ISM Hours" in cols else None
        idx_audit = cols.index("Audits in Deal") if "Audits in Deal" in cols else None
        idx_voice = cols.index("Voice Solution") if "Voice Solution" in cols else None

        total_hours = 0
        total_audit_deals = 0
        total_voice_deals = 0
        primary_ism_project_counts: dict[str,int] = {}
        secondary_ism_project_counts: dict[str,int] = {}
        ism_audit_counts: dict[str,int] = {}
        ism_voice_counts: dict[str,int] = {}
        ism_hours_counts: dict[str,int] = {}

        for row in projects:
            # Hours
            if idx_hours is not None and idx_hours < len(row):
                try:
                    h = int(str(row[idx_hours]).strip() or "0")
                    total_hours += h
                except Exception:
                    h = 0
            else:
                h = 0
            # Audit deals
            if idx_audit is not None and idx_audit < len(row):
                if str(row[idx_audit]).strip().lower() == "yes":
                    total_audit_deals += 1
            # Voice deals
            if idx_voice is not None and idx_voice < len(row):
                if str(row[idx_voice]).strip().lower() == "yes":
                    total_voice_deals += 1
            # Primary ISM aggregations
            if idx_primary is not None and idx_primary < len(row):
                p = (row[idx_primary] or "").strip()
                if p:
                    primary_ism_project_counts[p] = primary_ism_project_counts.get(p, 0) + 1
                    ism_hours_counts[p] = ism_hours_counts.get(p, 0) + h
                    if idx_audit is not None and idx_audit < len(row) and str(row[idx_audit]).strip().lower() == "yes":
                        ism_audit_counts[p] = ism_audit_counts.get(p, 0) + 1
                    if idx_voice is not None and idx_voice < len(row) and str(row[idx_voice]).strip().lower() == "yes":
                        ism_voice_counts[p] = ism_voice_counts.get(p, 0) + 1
            # Secondary ISM counts
            if idx_secondary is not None and idx_secondary < len(row):
                s = (row[idx_secondary] or "").strip()
                if s:
                    secondary_ism_project_counts[s] = secondary_ism_project_counts.get(s, 0) + 1

        # Update labels
        try:
            self.total_hours_label.setText(f"Total ISM Hours: {total_hours:,}")
            self.total_audit_label.setText(f"Total Audit Deals: {total_audit_deals}")
            self.total_voice_label.setText(f"Total Voice Deals: {total_voice_deals}")
            self.total_projects_label.setText(f"Total Projects: {len(projects)}")
        except Exception:
            pass

        # Fill Primary ISM Project Counts
        if primary_ism_project_counts:
            prim_lines = []
            for ism in sorted(primary_ism_project_counts):
                prim_lines.append(
                    f"• {ism}: {primary_ism_project_counts[ism]} total, "
                    f"{ism_audit_counts.get(ism, 0)} audit, {ism_voice_counts.get(ism, 0)} voice"
                )
            self.primary_counts_label.setText("\n".join(prim_lines))
        else:
            self.primary_counts_label.setText("No primary ISM data available")

        # Fill Secondary ISM Project Counts
        if secondary_ism_project_counts:
            sec_lines = []
            for ism in sorted(secondary_ism_project_counts):
                sec_lines.append(f"• {ism}: {secondary_ism_project_counts[ism]} total")
            self.secondary_counts_label.setText("\n".join(sec_lines))
        else:
            self.secondary_counts_label.setText("No secondary ISM data available")

        # Fill ISM Hours per Primary ISM
        if ism_hours_counts:
            hours_lines = [f"• {ism}: {hours:,} hours" for ism, hours in sorted(ism_hours_counts.items())]
            self.ism_hours_details.setText("\n".join(hours_lines))
        else:
            self.ism_hours_details.setText("No ISM hours data available")

    def _export_projects(self):
        """Export project data to Excel file"""
        try:
            from PyQt6.QtWidgets import QFileDialog
            import pandas as pd
            
            # Get file path from user
            file_path, _ = QFileDialog.getSaveFileName(
                self, 
                "Export Projects", 
                "projects_export.xlsx", 
                "Excel Files (*.xlsx);;All Files (*)"
            )
            
            if not file_path:
                return
            
            # Get project data
            projects = self.data.get("Project Details", [])
            if not projects:
                QMessageBox.information(self, "Export", "No project data to export.")
                return
            
            # Create DataFrame with proper column names
            columns = [
                "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM", 
                "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type", 
                "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
            ]
            
            # Ensure all rows have the same length
            max_cols = len(columns)
            normalized_projects = []
            for project in projects:
                # Pad with empty strings if project has fewer columns
                normalized_project = list(project) + [''] * (max_cols - len(project))
                # Truncate if project has more columns
                normalized_project = normalized_project[:max_cols]
                normalized_projects.append(normalized_project)
            
            df = pd.DataFrame(normalized_projects, columns=columns)
            
            # Export to CSV
            df.to_csv(file_path, index=False)
            
            QMessageBox.information(
                self, 
                "Export Successful", 
                f"Project data exported successfully to:\n{file_path}\n\nExported {len(projects)} projects."
            )
            
            self._log_change("Export", "Project Details", f"Exported {len(projects)} projects to {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export projects:\n{str(e)}")
            self._log_change("Error", "Project Details", f"Export failed: {str(e)}")
    
    def _import_projects(self):
        """Import project data from file"""
        try:
            from PyQt6.QtWidgets import QFileDialog
            import pandas as pd
            
            # Get file path from user
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Import Projects",
                "",
                "Excel Files (*.xlsx);;CSV Files (*.csv);;All Files (*)"
            )
            
            if not file_path:
                return
            
            # Read the file
            if file_path.lower().endswith('.xlsx'):
                df = pd.read_excel(file_path)
            else:
                df = pd.read_csv(file_path)
            
            # Expected columns
            expected_columns = [
                "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM", 
                "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type", 
                "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
            ]
            
            # Check if required columns exist
            missing_columns = [col for col in expected_columns if col not in df.columns]
            if missing_columns:
                QMessageBox.warning(
                    self, 
                    "Import Error", 
                    f"Missing required columns:\n{', '.join(missing_columns)}\n\nPlease use the Sample template to get the correct format."
                )
                return
            
            # Convert to list of lists
            imported_projects = []
            for _, row in df.iterrows():
                project_row = []
                for col in expected_columns:
                    project_row.append(str(row.get(col, '')))
                imported_projects.append(project_row)
            
            # Clear existing data and add imported data
            self.data["Project Details"] = imported_projects
            
            # Save and refresh
            self._save_autosave()
            self._save_backend_sqlite()
            self._refresh_project_details_table()
            
            QMessageBox.information(
                self, 
                "Import Successful", 
                f"Successfully imported {len(imported_projects)} projects."
            )
            
            self._log_change("Import", "Project Details", f"Imported {len(imported_projects)} projects from {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"Failed to import projects:\n{str(e)}")
            self._log_change("Error", "Project Details", f"Import failed: {str(e)}")
    
    def _download_projects_sample(self):
        """Download sample project template"""
        try:
            from PyQt6.QtWidgets import QFileDialog
            import pandas as pd
            
            # Get file path from user
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Project Details Sample Template",
                "project_details_sample.xlsx",
                "Excel Files (*.xlsx);;CSV Files (*.csv)"
            )
            
            if not file_path:
                return
            
            # Create sample data
            sample_data = [
                [
                    "Sample Project 1", "PRJ001", "New York", "Financial Services", "Medium", 
                    "Standard implementation", "john.smith", "sarah.johnson", "40", 
                    "Client CXL", "Internal SDL", "Dedicated", "Yes", "Yes", "Yes", "50", "25"
                ],
                [
                    "Sample Project 2", "PRJ002", "San Francisco", "Technology", "High", 
                    "Complex integration", "mike.wilson", "lisa.brown", "60", 
                    "Client CXL", "Internal SDL", "Shared", "No", "No", "No", "100", "50"
                ]
            ]
            
            # Expected columns
            columns = [
                "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM", 
                "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type", 
                "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
            ]
            
            # Create DataFrame
            df = pd.DataFrame(sample_data, columns=columns)
            
            # Export based on file type
            if file_path.lower().endswith('.xlsx'):
                df.to_excel(file_path, index=False)
            else:
                df.to_csv(file_path, index=False)
            
            QMessageBox.information(self, "Sample Download", f"Sample template saved to {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "Sample Download Error", f"Failed to create sample template:\n{str(e)}")

    def resizeEvent(self, event):
        """Handle window resize events to auto-fit org chart"""
        super().resizeEvent(event)
        # Auto-fit org chart when window is resized
        if hasattr(self, 'org_view') and hasattr(self, 'org_scene'):
            try:
                rect = self.org_scene.itemsBoundingRect()
                if rect.isValid():
                    # Reset transform first
                    try:
                        self.org_view.resetTransform()
                    except Exception:
                        pass
                    
                    # Get view and screen dimensions
                    view_size = self.org_view.size()
                    screen = QApplication.primaryScreen()
                    screen_geometry = screen.availableGeometry() if screen else QRect(0, 0, 1920, 1080)
                    
                    # Calculate optimal padding based on multiple factors
                    min_view_dim = min(view_size.width(), view_size.height())
                    max_view_dim = max(view_size.width(), view_size.height())
                    
                    # Determine screen size category
                    screen_area = screen_geometry.width() * screen_geometry.height()
                    if screen_area >= 3840 * 2160:  # 4K and above
                        base_padding_percent = 0.12  # 12% for large screens
                    elif screen_area >= 2560 * 1440:  # 1440p
                        base_padding_percent = 0.10  # 10% for medium-large screens
                    elif screen_area >= 1920 * 1080:  # 1080p
                        base_padding_percent = 0.08  # 8% for standard screens
                    else:  # Smaller screens
                        base_padding_percent = 0.06  # 6% for small screens
                    
                    # Adjust padding based on view aspect ratio
                    aspect_ratio = view_size.width() / view_size.height() if view_size.height() > 0 else 1.0
                    if aspect_ratio > 2.0:  # Very wide view
                        base_padding_percent *= 0.8  # Reduce padding for wide views
                    elif aspect_ratio < 0.5:  # Very tall view
                        base_padding_percent *= 0.8  # Reduce padding for tall views
                    
                    # Calculate dynamic padding
                    dynamic_padding = max(15, int(min_view_dim * base_padding_percent))
                    
                    # Ensure reasonable bounds
                    dynamic_padding = min(dynamic_padding, int(max_view_dim * 0.15))  # Max 15% of larger dimension
                    dynamic_padding = max(dynamic_padding, 10)  # Minimum 10px
                    
                    # Apply dynamic padding and fit
                    padded = rect.adjusted(-dynamic_padding, -dynamic_padding, dynamic_padding, dynamic_padding)
                    self.org_view.fitInView(padded, Qt.AspectRatioMode.KeepAspectRatio)
                    
                    # Center the view
                    try:
                        self.org_view.centerOn(rect.center())
                    except Exception:
                        pass
            except Exception as e:
                ErrorHandler.handle_ui_error("org chart auto-fit", e)

    def showEvent(self, event):
        """Auto-fit org chart to screen on first show."""
        try:
            super().showEvent(event)
        except Exception:
            try:
                QMainWindow.showEvent(self, event)
            except Exception:
                pass
        # Ensure View menu contains action to show Wiki / Knowledge Base
        try:
            if not getattr(self, '_wiki_view_menu_installed', False):
                self._install_wiki_view_menu()
        except Exception:
            pass
        # Defer fit until after the window is shown to ensure correct geometry
        try:
            QTimer.singleShot(0, self._auto_fit_org_chart)
        except Exception as e:
            ErrorHandler.handle_ui_error("schedule org chart auto-fit", e)

    def _auto_fit_org_chart(self) -> None:
        """Fit the entire org chart into the view based on current screen/window size."""
        try:
            if not hasattr(self, 'org_view') or not hasattr(self, 'org_scene'):
                return
            if self.org_scene is None or self.org_view is None:
                return
            rect = self.org_scene.itemsBoundingRect()
            if rect.isValid() and rect.width() > 0 and rect.height() > 0:
                # Reset transform first
                try:
                    self.org_view.resetTransform()
                except Exception:
                    pass
                
                # Get view and screen dimensions
                view_size = self.org_view.size()
                screen = QApplication.primaryScreen()
                screen_geometry = screen.availableGeometry() if screen else QRect(0, 0, 1920, 1080)
                
                # Calculate optimal padding based on multiple factors
                min_view_dim = min(view_size.width(), view_size.height())
                max_view_dim = max(view_size.width(), view_size.height())
                
                # Determine screen size category with much smaller padding for better space utilization
                screen_area = screen_geometry.width() * screen_geometry.height()
                if screen_area >= 3840 * 2160:  # 4K and above
                    base_padding_percent = 0.05  # 5% for large screens - maximize content
                elif screen_area >= 2560 * 1440:  # 1440p
                    base_padding_percent = 0.04  # 4% for medium-large screens
                elif screen_area >= 1920 * 1080:  # 1080p
                    base_padding_percent = 0.03  # 3% for standard screens
                else:  # Smaller screens
                    base_padding_percent = 0.02  # 2% for small screens
                
                # Adjust padding based on view aspect ratio
                aspect_ratio = view_size.width() / view_size.height() if view_size.height() > 0 else 1.0
                if aspect_ratio > 2.0:  # Very wide view
                    base_padding_percent *= 0.8  # Reduce padding for wide views
                elif aspect_ratio < 0.5:  # Very tall view
                    base_padding_percent *= 0.8  # Reduce padding for tall views
                
                # Calculate dynamic padding with much smaller minimums for better space utilization
                dynamic_padding = max(5, int(min_view_dim * base_padding_percent))
                
                # Ensure reasonable bounds - much more generous maximums
                dynamic_padding = min(dynamic_padding, int(max_view_dim * 0.08))  # Max 8% of larger dimension
                dynamic_padding = max(dynamic_padding, 3)  # Minimum 3px for very small screens
                
                # Apply dynamic padding
                padded = rect.adjusted(-dynamic_padding, -dynamic_padding, dynamic_padding, dynamic_padding)
                
                # Ensure minimum chart size - if chart would be too small, reduce padding further
                chart_width = padded.width()
                chart_height = padded.height()
                min_chart_size = min(min_view_dim * 0.8, 400)  # At least 80% of view or 400px minimum
                
                if chart_width < min_chart_size or chart_height < min_chart_size:
                    # Reduce padding to ensure minimum chart size
                    scale_factor = min_chart_size / max(chart_width, chart_height, 1)
                    reduced_padding = max(2, int(dynamic_padding * scale_factor * 0.5))
                    padded = rect.adjusted(-reduced_padding, -reduced_padding, reduced_padding, reduced_padding)
                
                # Fit with aspect ratio preservation
                self.org_view.fitInView(padded, Qt.AspectRatioMode.KeepAspectRatio)
                
                # Center the view
                try:
                    self.org_view.centerOn(rect.center())
                except Exception:
                    pass
                    
                # Store the current zoom level for reference
                self._org_zoom = 1.0
                
        except Exception as e:
            ErrorHandler.handle_ui_error("org chart fit-in-view", e)

    def _ensure_wiki_tab_visible(self) -> None:
        """Ensure the Wiki / Knowledge Base pane exists and is selected in tabs."""
        try:
            # Ensure tab widget exists
            if not hasattr(self, 'tabs') or self.tabs is None:
                return
            # Try to find existing Wiki tab by title or object name
            target_titles = {"Wiki", "Knowledge Base", "Wiki / Knowledge Base"}
            target_widget = None
            for i in range(self.tabs.count()):
                w = self.tabs.widget(i)
                title = self.tabs.tabText(i) or ""
                if title in target_titles or (hasattr(w, 'objectName') and w.objectName() in target_titles):
                    target_widget = w
                    break
            # Create if missing
            if target_widget is None:
                try:
                    # Wiki functionality has been removed
                    self.wiki_pane = QWidget()
                    layout = QVBoxLayout(self.wiki_pane)
                    layout.addWidget(QLabel("Wiki functionality has been removed."))
                except Exception:
                    # Fallback: create a minimal placeholder widget
                    placeholder = QWidget(self)
                    lay = QVBoxLayout(placeholder)
                    lay.addWidget(QLabel("Wiki / Knowledge Base is unavailable."))
                    self.wiki_pane = placeholder
                target_widget = self.wiki_pane
                self.tabs.addTab(target_widget, "Wiki / Knowledge Base")
            # Focus the tab
            idx = self.tabs.indexOf(target_widget)
            if idx >= 0:
                self.tabs.setCurrentIndex(idx)
        except Exception as e:
            ErrorHandler.handle_ui_error("show wiki pane", e)

    def _install_wiki_view_menu(self) -> None:
        """Add a View menu action to show the Wiki / Knowledge Base pane."""
        try:
            mb = self.menuBar() if hasattr(self, 'menuBar') else None
            if mb is None:
                return
            # Find or create View menu
            view_menu = None
            try:
                for act in mb.actions():
                    try:
                        if act.text() and 'View' in act.text():
                            view_menu = act.menu()
                            break
                    except Exception:
                        continue
            except Exception:
                pass
            if view_menu is None:
                view_menu = QMenu("View", self)
                mb.addMenu(view_menu)
            # Avoid duplicate action
            for a in view_menu.actions():
                if a.text() and 'Wiki' in a.text():
                    self._wiki_view_menu_installed = True
                    return
            act_show_wiki = QAction("Show Wiki / Knowledge Base", self)
            try:
                act_show_wiki.setShortcut(QKeySequence("Ctrl+Alt+W"))
            except Exception:
                pass
            act_show_wiki.triggered.connect(self._ensure_wiki_tab_visible)
            view_menu.addAction(act_show_wiki)
            self._wiki_view_menu_installed = True
        except Exception as e:
            ErrorHandler.handle_ui_error("install wiki view menu", e)

    def _populate_ism_dropdowns(self):
        """Populate ISM dropdowns with Org Chart members - thread-safe version"""
        try:
            # Get ISMs from Org Chart directory
            org_rows = self._collect_org_directory_rows()
            isms = sorted({name for (name, _, _, _, _, _) in org_rows if name})
            
            # Fallback to collected ISMs if no org data
            if not isms:
                isms = sorted(self._collect_all_isms())
            
            # Populate dropdowns directly
            try:
                for combo in [self.primary_ism_combo, self.secondary_ism_combo]:
                    if combo:
                        combo.clear()
                        combo.addItem("")
                        for ism in isms:
                            combo.addItem(ism)
            except Exception as e:
                ErrorHandler.handle_ui_error("populate ism dropdowns", e)
        except Exception as e:
            ErrorHandler.handle_ui_error("populate ism dropdowns", e)

    def _on_voice_solution_changed(self):
        """Handle voice solution selection change"""
        try:
            is_yes = self.voice_solution_combo.currentText() == "Yes"
            self.contact_center_combo.setEnabled(is_yes)
            if not is_yes:
                self.contact_center_combo.setCurrentIndex(0)
        except Exception as e:
            ErrorHandler.handle_ui_error("voice solution changed", e)

    def _create_metric_card(self, title: str, value: str, color: str, bg_color: str) -> QFrame:
        """Create a styled metric card with enhanced visual design - improved visibility"""
        card = QFrame()
        
        # Create a more visible background with better contrast
        card.setStyleSheet(f"""
            QFrame {{
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #ffffff, stop:1 #f8fafc);
                border: 2px solid {color};
                border-radius: 10px;
                padding: 10px;
                margin: 3px;
            }}
            QFrame:hover {{
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #f0f9ff, stop:1 #e0f2fe);
                border: 2px solid {color};
            }}
        """)
        card.setFixedHeight(70)  # Slightly increased for better readability
        card.setMinimumWidth(110)  # Increased for better text display
        card.setMaximumWidth(140)  # Increased maximum width
        
        layout = QVBoxLayout(card)
        layout.setContentsMargins(8, 6, 8, 6)  # Better margins
        layout.setSpacing(4)  # Better spacing
        
        # Title with improved styling
        title_label = QLabel(title)
        title_label.setStyleSheet(f"""
            color: #374151; 
            font-size: 10px; 
            font-weight: 700; 
            background: transparent;
            border: none;
            padding: 2px;
        """)
        title_label.setWordWrap(True)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)
        
        # Value with enhanced styling and better contrast
        value_label = QLabel(value)
        value_label.setStyleSheet(f"""
            color: {color}; 
            font-size: 18px; 
            font-weight: 900; 
            background: transparent;
            border: none;
            padding: 2px;
            text-shadow: 0px 1px 2px rgba(0,0,0,0.1);
        """)
        value_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(value_label)
        
        # Add a subtle shadow effect
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(8)
        shadow.setXOffset(0)
        shadow.setYOffset(2)
        shadow.setColor(QColor(0, 0, 0, 25))
        card.setGraphicsEffect(shadow)
        
        # Store reference for updating
        card.value_label = value_label
        
        return card
    def _add_project(self):
        """Add a new project row to the table"""
        try:
            self._updating_table = True
            
            # Add new row to table
            row_count = self.projects_table.rowCount()
            self.projects_table.insertRow(row_count)
            
            # Add default values for all 17 columns
            default_values = default_row_for_columns(PANE_COLUMNS["Project Details"])
            default_values = (list(default_values) + [""] * max(0, 17 - len(default_values)))[:17]
            for col, value in enumerate(default_values):
                item = QTableWidgetItem(str(value))
                # Enable text wrapping for better visibility
                item.setTextAlignment(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft)
                self.projects_table.setItem(row_count, col, item)
            
            # Add to data
            if "Project Details" not in self.data:
                self.data["Project Details"] = []
            self.data["Project Details"].append(default_values)
            
            # Save and update
            self._save_autosave()
            self._save_backend_sqlite()
            self._update_summary_metrics()
            self._populate_ism_filter()
            
            self._updating_table = False
            
            # Focus on the new row for editing
            self.projects_table.setCurrentCell(row_count, 0)
            self.projects_table.edit(self.projects_table.model().index(row_count, 0))
            
            # Show success notification
            self.notifications.show_success("Project added successfully!")
            
        except Exception as e:
            self._updating_table = False
            self.notifications.show_error(f"Failed to add project: {str(e)}")
            self._log_change("Error", "Project Details", f"Failed to add project: {str(e)}")

    def _edit_project(self):
        """Edit selected project - start inline editing"""
        current_row = self.projects_table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "No Selection", "Please select a project to edit")
            return
        
        # Start editing the first cell (Project Name)
        self.projects_table.setCurrentCell(current_row, 0)
        self.projects_table.edit(self.projects_table.model().index(current_row, 0))
        
        self._log(f"Started editing project at row {current_row + 1}")

    def _update_project(self):
        """Update the project being edited - no longer needed with inline editing"""
        # This method is kept for compatibility but does nothing
        # since we now use inline table editing with real-time updates
        pass

    def _show_projects_context_menu(self, position):
        """Show context menu for projects table"""
        menu = QMenu(self)
        
        # Add Row action (always available)
        add_action = QAction("➕ Add Row", self)
        add_action.triggered.connect(self._add_project)
        menu.addAction(add_action)
        
        menu.addSeparator()
        
        # Check if we're right-clicking on an item
        item = self.projects_table.itemAt(position)
        if item is not None:
            # Edit action (only when right-clicking on an item)
            edit_action = QAction("✏️ Edit Selected", self)
            edit_action.triggered.connect(self._edit_project)
            menu.addAction(edit_action)
            
            # Delete action (only when right-clicking on an item)
            delete_action = QAction("🗑️ Delete Selected", self)
            delete_action.triggered.connect(self._delete_project)
            menu.addAction(delete_action)
            
            menu.addSeparator()
            
            # Export selected action
            export_selected_action = QAction("📊 Export Selected", self)
            export_selected_action.triggered.connect(self._export_selected_projects)
            menu.addAction(export_selected_action)
        
        # Show context menu
        menu.exec(self.projects_table.mapToGlobal(position))

    def _export_selected_projects(self):
        """Export selected projects to Excel"""
        selected_rows = []
        for item in self.projects_table.selectedItems():
            row = item.row()
            if row not in selected_rows:
                selected_rows.append(row)
        
        if not selected_rows:
            QMessageBox.warning(self, "No Selection", "Please select projects to export")
            return
        
        try:
            path, _ = QFileDialog.getSaveFileName(self, "Export Selected Projects", "selected_projects.csv", "CSV (*.csv)")
            if not path:
                return
            
            # Get selected project data
            selected_projects = []
            for row in selected_rows:
                if row < len(self.data["Project Details"]):
                    selected_projects.append(self.data["Project Details"][row])
            
            # Create DataFrame and export
            df = pd.DataFrame(selected_projects, columns=[
                "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM", 
                "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type", 
                "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
            ])
            df.to_csv(path, index=False)
            
            QMessageBox.information(self, "Export", f"Exported {len(selected_projects)} projects successfully!")
            
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export: {str(e)}")

    def _export_all_projects(self):
        """Export all projects to CSV"""
        try:
            path, _ = QFileDialog.getSaveFileName(self, "Export All Projects", "all_projects.csv", "CSV (*.csv)")
            if not path:
                return
            
            projects = self.data.get("Project Details", [])
            if not projects:
                QMessageBox.information(self, "No Data", "No projects to export")
                return
            
            # Create DataFrame and export
            df = pd.DataFrame(projects, columns=[
                "Project Name", "Project ID", "Location", "Vertical", "Complexity", "Complexity Details", "Primary ISM", 
                "Secondary ISM", "ISM Hours", "CXL Name", "SDL Name", "Connectivity Type", 
                "Audits in Deal", "Voice Solution", "Contact Center", "Head Count", "Seat Count"
            ])
            df.to_csv(path, index=False)
            
            QMessageBox.information(self, "Export", f"Exported {len(projects)} projects successfully!")
            
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export: {str(e)}")

    def _delete_project(self):
        """Delete selected project"""
        current_row = self.projects_table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "No Selection", "Please select a project to delete")
            return
        
        project_name = self.projects_table.item(current_row, 0).text() if self.projects_table.item(current_row, 0) else "Unknown"
        reply = QMessageBox.question(self, "Confirm Delete", 
            f"Are you sure you want to delete project '{project_name}'?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        
        if reply == QMessageBox.StandardButton.Yes:
            # Remove from data
            if current_row < len(self.data["Project Details"]):
                self.data["Project Details"].pop(current_row)
            
            # Remove from table
            self.projects_table.removeRow(current_row)
            
            # Sync projects list
            self.sync_projects_from_details()
            
            # Save changes
            self._save_autosave()
            self._save_backend_sqlite()
            
            # Update simple stats
            self._update_simple_stats()
            
            self._log("Deleted project")
            self._log_change("Delete Project", "Project Details", f"Deleted project: {project_name}")


    def _edit_project(self):
        """Edit selected project - focus on the first cell for editing"""
        current_row = self.projects_table.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "No Selection", "Please select a project to edit")
            return
        
        # Focus on the first cell (Project Name) for editing
        self.projects_table.setCurrentCell(current_row, 0)
        self.projects_table.edit(self.projects_table.currentIndex())

    def _on_project_cell_changed(self, row, column):
        """Handle cell changes in the projects table"""
        # Reentrancy guard to prevent recursive updates/crashes during rapid edits
        if getattr(self, "_pd_edit_in_progress", False):
            return
        self._pd_edit_in_progress = True
        # Ensure we run on the main Qt thread to avoid event filter errors when edits come from worker callbacks
        try:
            import threading
            if QApplication.instance() is not None and threading.current_thread() is not threading.main_thread():
                QTimer.singleShot(0, lambda r=row, c=column: self._on_project_cell_changed(r, c))
                self._pd_edit_in_progress = False
                return
        except Exception:
            pass
        if self._updating_table or row < 0 or column < 0:
            self._pd_edit_in_progress = False
            return
        
        item = self.projects_table.item(row, column)
        if not item:
            self._pd_edit_in_progress = False
            return
        
        new_value = item.text()
        column_name = PANE_COLUMNS["Project Details"][column]
        
        # Enhanced validation using InputValidator + numeric constraints for Head/Seat Count
        if column_name in ("Head Count", "Seat Count"):
            try:
                iv = int(new_value.strip() or "0")
                if iv < 0 or iv > 9999:
                    raise ValueError("out of range")
                new_value = str(iv)
            except Exception:
                ErrorHandler.handle_validation_error(column_name, new_value, "Must be an integer between 0 and 9999")
                item.setBackground(QColor("#ffebee"))
                self._pd_edit_in_progress = False
                return
        is_valid, error_msg = self._validate_cell_input(column_name, new_value)
        if not is_valid:
            ErrorHandler.handle_validation_error(column_name, new_value, error_msg)
            # Highlight the cell with error
            item.setBackground(QColor("#ffebee"))  # Light red background
            return
        else:
            # Clear error highlighting
            item.setBackground(QColor("#ffffff"))  # White background
        
        # Update data
        if "Project Details" not in self.data:
            self.data["Project Details"] = []
        
        while len(self.data["Project Details"]) <= row:
            self.data["Project Details"].append(default_row_for_columns(PANE_COLUMNS["Project Details"]))
        
        if column < len(self.data["Project Details"][row]):
            self.data["Project Details"][row][column] = new_value
        
        # Handle Project Name/ID synchronization
        self._handle_project_sync(row, column, new_value)
        
        # Avoid full rebuilds during edit; just refresh viewport and metrics
        try:
            self.projects_table.viewport().update()
        except Exception:
            pass
        self._update_summary_metrics()
        self._populate_ism_filter()
        self.sync_projects_from_details()  # Sync global projects list
        # Autosave behavior consistent with other panes
        try:
            if self.preferences.get("auto_save", True):
                # Debounce saves to prevent thrashing when typing
                if self._pd_save_timer is None:
                    self._pd_save_timer = QTimer(self)
                    self._pd_save_timer.setSingleShot(True)
                    self._pd_save_timer.timeout.connect(lambda: (self._save_autosave(), self._save_backend_sqlite()))
                try:
                    self._pd_save_timer.stop()
                except Exception:
                    pass
                self._pd_save_timer.start(800)
        except Exception as e:
            ErrorHandler.handle_ui_error("autosave Project Details", e)
        finally:
            self._pd_edit_in_progress = False

    def _handle_project_sync(self, row, column, new_value):
        """Handle Project Name/ID synchronization in Project Details"""
        if self._updating_table:
            return
        
        try:
            # Get project data for mapping
            project_names, project_ids, name_to_id, id_to_name = self.get_project_dropdown_data()
            
            # If Project Name changed (column 0), update Project ID (column 1)
            if column == 0 and new_value in name_to_id:
                project_id = name_to_id[new_value]
                if len(self.data["Project Details"]) > row and len(self.data["Project Details"][row]) > 1:
                    self.data["Project Details"][row][1] = project_id
                    # Reflect UI directly without rebuild
                    try:
                        if hasattr(self, 'projects_table'):
                            self.projects_table.blockSignals(True)
                            existing = self.projects_table.item(row, 1)
                            if existing is not None:
                                existing.setText(project_id)
                            else:
                                self.projects_table.setItem(row, 1, QTableWidgetItem(project_id))
                    except Exception as e:
                        ErrorHandler.handle_ui_error("sync Project ID cell", e)
                    finally:
                        try:
                            self.projects_table.blockSignals(False)
                        except Exception:
                            pass
            
            # If Project ID changed (column 1), update Project Name (column 0)
            elif column == 1 and new_value in id_to_name:
                project_name = id_to_name[new_value]
                if len(self.data["Project Details"]) > row and len(self.data["Project Details"][row]) > 0:
                    self.data["Project Details"][row][0] = project_name
                    # Reflect UI directly without rebuild
                    try:
                        if hasattr(self, 'projects_table'):
                            self.projects_table.blockSignals(True)
                            existing = self.projects_table.item(row, 0)
                            if existing is not None:
                                existing.setText(project_name)
                            else:
                                self.projects_table.setItem(row, 0, QTableWidgetItem(project_name))
                    except Exception as e:
                        ErrorHandler.handle_ui_error("sync Project Name cell", e)
                    finally:
                        try:
                            self.projects_table.blockSignals(False)
                        except Exception:
                            pass
        except Exception as e:
            print(f"Error in project sync: {e}")

    def _on_cell_changed_for_calendar(self, pane_name: str, row: int, column: int):
        """Handle cell changes in panes that affect calendar data"""
        try:
            # Check if the changed cell affects calendar data (date fields, status, etc.)
            if pane_name not in self.data or row >= len(self.data[pane_name]):
                return
                
            cols = PANE_COLUMNS.get(pane_name, [])
            if column >= len(cols):
                return
                
            column_name = cols[column]
            
            # Check if this is a date field or status field that affects calendar
            date_fields = ["Due Date", "Target Date", "Leave Date", "Date", "End Date", "Audit End Date"]
            status_fields = ["Status", "Approval Status", "Task Type"]
            
            if column_name in date_fields or column_name in status_fields:
                # Refresh calendar if it's visible
                if hasattr(self, 'calendar_tab') and self.calendar_tab:
                    self.calendar_tab.refresh_calendar_if_visible()
                    
        except Exception as e:
            print(f"Calendar cell change error: {e}")

    def _load_projects_data(self):
        """Load projects data into the table - thread-safe version"""
        try:
            self._updating_table = True
            
            projects = self.data.get("Project Details", [])
            self.projects_table.setRowCount(len(projects))
            
            # Render all 17 columns
            for row, project in enumerate(projects):
                normalized = list(project) + [""] * max(0, 17 - len(project))
                normalized = normalized[:17]
                for col, value in enumerate(normalized):
                    item = QTableWidgetItem(str(value))
                    item.setTextAlignment(Qt.AlignmentFlag.AlignTop | Qt.AlignmentFlag.AlignLeft)
                    self.projects_table.setItem(row, col, item)
            
            self._updating_table = False
            
            # Update summary metrics after loading
            self._update_summary_metrics()
        except Exception as e:
            self._updating_table = False
            ErrorHandler.handle_ui_error("load projects data", e)

    def _show_projects_context_menu(self, position):
        """Show context menu for projects table"""
        menu = QMenu(self)
        
        # Add Row action
        add_action = QAction("➕ Add Project", self)
        add_action.triggered.connect(self._add_project)
        menu.addAction(add_action)
        
        # Delete Row action (if row selected)
        current_row = self.projects_table.currentRow()
        if current_row >= 0:
            menu.addSeparator()
            delete_action = QAction("🗑️ Delete Project", self)
            delete_action.triggered.connect(self._delete_project)
            menu.addAction(delete_action)
        
        menu.exec(self.projects_table.mapToGlobal(position))

    def _delete_project(self):
        """Delete selected project"""
        try:
            current_row = self.projects_table.currentRow()
            if current_row < 0:
                self.notifications.show_warning("Please select a project to delete")
                return
            
            # Get project name for confirmation
            project_name = ""
            name_item = self.projects_table.item(current_row, 0)
            if name_item:
                project_name = name_item.text()
            
            # Confirm deletion
            reply = QMessageBox.question(
                self, "Confirm Delete", 
                f"Are you sure you want to delete project '{project_name}'?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                # Remove from table
                self.projects_table.removeRow(current_row)
                
                # Remove from data
                if "Project Details" in self.data and current_row < len(self.data["Project Details"]):
                    del self.data["Project Details"][current_row]
                
                # Save and update
                self._save_autosave()
                self._save_backend_sqlite()
                self._update_summary_metrics()
                self._populate_ism_filter()
                
                # Show success notification
                self.notifications.show_success(f"Project '{project_name}' deleted successfully!")
                
        except Exception as e:
            self.notifications.show_error(f"Failed to delete project: {str(e)}")
            self._log_change("Error", "Project Details", f"Failed to delete project: {str(e)}")

    def _filter_projects_table(self):
        """Filter projects table based on search text"""
        search_text = self.projects_search.text().lower()
        for row in range(self.projects_table.rowCount()):
            row_visible = not search_text
            if search_text:
                for col in range(self.projects_table.columnCount()):
                    item = self.projects_table.item(row, col)
                    if item and search_text in item.text().lower():
                        row_visible = True
                        break
            self.projects_table.setRowHidden(row, not row_visible)





    def _sync_ism_directory_with_org(self) -> None:
        try:
            if not hasattr(self, 'org_tree'):
                return
            names = set(self.ism_directory or [])
            root = self.org_tree.invisibleRootItem()
            stack = [root]
            while stack:
                cur = stack.pop()
                for i in range(cur.childCount()):
                    ch = cur.child(i)
                    nm = ch.text(0)
                    if nm:
                        names.add(nm)
                    stack.append(ch)
            self.ism_directory = sorted(names)
            self.refresh_ism_filter()
            # Refresh profile icon to update display name if org directory changed
            self._refresh_profile_icon()
        except Exception as e:
            ErrorHandler.handle_ui_error("sync ISM directory", e)

    def _refresh_calendar_decorations(self) -> None:
        try:
            # Check if leave_calendar exists (Leave Tracker tab might not be initialized yet)
            if not hasattr(self, 'leave_calendar'):
                return
            cal = self.leave_calendar
            if not cal:
                return
            # reset text formats
            default_fmt = cal.weekdayTextFormat(Qt.DayOfWeek.Monday)
            year = cal.selectedDate().year(); month = cal.selectedDate().month()
            for d in range(1, 32):
                try:
                    qd = QDate(year, month, d)
                    cal.setDateTextFormat(qd, default_fmt)
                except Exception:
                    continue
            # apply for each leave entry
            colors = {
                "WFH": QColor("#60a5fa"),
                "Planned Leave": QColor("#34d399"),
                "Public Holiday": QColor("#f59e0b"),
                "Earned Leave": QColor("#f97316"),
                "Casual Leave": QColor("#ef4444"),
            }
            # Merge colors for multiple leaves; add tiny dot indicators for ISMs
            daily_types: dict[tuple[int,int,int], set[str]] = {}
            daily_isms: dict[tuple[int,int,int], set[str]] = {}
            for r in self.data.get("Leave Tracker", []):
                try:
                    dt = datetime.strptime(r[0], "%Y-%m-%d")
                except Exception:
                    continue
                if dt.year != year or dt.month != month:
                    continue
                key = (dt.year, dt.month, dt.day)
                daily_types.setdefault(key, set()).add(r[1])
                daily_isms.setdefault(key, set()).add(r[3] if len(r) > 3 else "")
            for (yy, mm, dd), types in daily_types.items():
                qd = QDate(yy, mm, dd)
                fmt = cal.dateTextFormat(qd)
                if len(types) == 1:
                    t = next(iter(types))
                    fmt.setBackground(QBrush(colors.get(t, QColor("#94a3b8"))))
                else:
                    # For multiple leave types, use a neutral highlight and bold font
                    fmt.setBackground(QBrush(QColor("#e2e8f0")))
                    f = fmt.font(); f.setBold(True); fmt.setFont(f)
                cal.setDateTextFormat(qd, fmt)
        except Exception as e:
            ErrorHandler.handle_ui_error("refresh calendar decorations", e)

    def _open_context_menu(self, pane_name: str, pos) -> None:
        table = self.tables[pane_name]
        menu = QMenu(self)
        add_act = QAction("Add Row", self, triggered=lambda: self.add_row(pane_name))
        del_act = QAction("Delete Selected", self, triggered=lambda: self.delete_selected(pane_name))
        hide_rows_act = QAction("Hide Selected Rows", self)
        show_all_rows_act = QAction("Show All Rows", self)
        dup_act = QAction("Duplicate Selected", self)
        def do_dup():
            rows = sorted({idx.row() for idx in table.selectedIndexes()})
            if not rows:
                return
            for r in rows:
                if 0 <= r < len(self.data[pane_name]):
                    new = list(self.data[pane_name][r])
                    self.data[pane_name].insert(r+1, new)
            self.rebuild_table(pane_name)
            self.update_dashboard()
            self._save_backend_sqlite()
            self._log(f"Duplicated {len(rows)} row(s) in {pane_name}")
        dup_act.triggered.connect(do_dup)
        def do_hide_rows():
            rows = sorted({idx.row() for idx in table.selectedIndexes()})
            if not rows:
                return
            for r in rows:
                table.setRowHidden(r, True)
        def do_show_all_rows():
            for r in range(table.rowCount()):
                table.setRowHidden(r, False)
        hide_rows_act.triggered.connect(do_hide_rows)
        show_all_rows_act.triggered.connect(do_show_all_rows)
        menu.addAction(add_act)
        menu.addAction(del_act)
        menu.addAction(dup_act)
        menu.addSeparator(); menu.addAction(hide_rows_act); menu.addAction(show_all_rows_act)
        # If an Updates cell is selected in Potential Issues, offer 'Edit Updates...'
        try:
            if pane_name == "Potential Issues":
                sel = table.selectedIndexes()
                if sel:
                    r = sel[0].row(); c = sel[0].column()
                    cols = PANE_COLUMNS[pane_name]
                    if c < len(cols) and cols[c] == "Updates":
                        edit_upd = QAction("Edit Updates...", self, triggered=lambda: self.edit_cell(pane_name, r, c))
                        menu.addSeparator()
                        menu.addAction(edit_upd)
        except Exception as e:
            ErrorHandler.handle_ui_error("context menu setup", e)
        menu.exec(table.viewport().mapToGlobal(pos))

    def open_rag_report(self, rag: str) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle(f"{rag} Tasks Report")
        lay = QVBoxLayout(dlg)
        tbl = QTableWidget()
        headers = ["Pane", "ISM Name", "Project Name", "Summary", "Start Date", "Due/End Date", "Status", "RAG"]
        tbl.setColumnCount(len(headers))
        tbl.setHorizontalHeaderLabels(headers)
        tbl.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        rows_to_add = []
        ism_filter = self.ism_filter.currentText() if hasattr(self, 'ism_filter') else "All ISMs"
        
        # Only include Potential Issues and Activities panes for Red Tasks (matching dashboard calculation)
        target_panes = ["Potential Issues", "Activities"]
        
        for pane, rows in self.data.items():
            # Skip panes that are not in our target list
            if pane not in target_panes:
                continue
                
            cols = PANE_COLUMNS[pane]
            
            # For these panes, look for RAG field
            rag_idx = next((i for i, c in enumerate(cols) if "RAG" in c), None)
            if rag_idx is None:
                continue
            def map_priority_to_rag(val):
                return val  # Return as-is for regular RAG fields
            
            ism_idx = cols.index("ISM Name") if "ISM Name" in cols else None
            proj_idx = cols.index("Project Name") if "Project Name" in cols else None
            summary_idx = 0 if cols else 0
            start_idx = next((cols.index(c) for c in ("Start Date","Created Date","Audit Start Date") if c in cols), None)
            due_idx = next((cols.index(c) for c in ("Due Date","Target Date","End Date","Audit End Date") if c in cols), None)
            status_idx = cols.index("Status") if "Status" in cols else None
            for r in rows:
                raw_val = r[rag_idx] if rag_idx < len(r) else ""
                rag_val = map_priority_to_rag(raw_val)
                if rag_val != rag:
                    continue
                if ism_filter and ism_filter != "All ISMs":
                    ism_val = r[ism_idx] if (ism_idx is not None and ism_idx < len(r)) else ""
                    if ism_val != ism_filter:
                        continue
                row_tuple = (
                    pane,
                    r[ism_idx] if (ism_idx is not None and ism_idx < len(r)) else "",
                    r[proj_idx] if (proj_idx is not None and proj_idx < len(r)) else "",
                    r[summary_idx] if summary_idx < len(r) else "",
                    r[start_idx] if (start_idx is not None and start_idx < len(r)) else "",
                    r[due_idx] if (due_idx is not None and due_idx < len(r)) else "",
                    r[status_idx] if (status_idx is not None and status_idx < len(r)) else "",
                    rag_val
                )
                rows_to_add.append(row_tuple)
        tbl.setRowCount(len(rows_to_add))
        for i, row in enumerate(rows_to_add):
            for j, val in enumerate(row):
                tbl.setItem(i, j, QTableWidgetItem(str(val)))
        lay.addWidget(tbl)
        btn = QPushButton("Close")
        btn.clicked.connect(dlg.accept)
        lay.addWidget(btn)
        dlg.resize(900, 500)
        dlg.exec()

    # Data ops
    def add_row(self, pane_name: str) -> None:
        cols = PANE_COLUMNS[pane_name]
        row = default_row_for_columns(cols)
        # Default ISM Name to logged-in user if present
        if "ISM Name" in cols:
            idx = cols.index("ISM Name")
            mapped_name = self._resolve_ism_name_from_eid(self.logged_in_user)
            row[idx] = mapped_name if mapped_name else self.logged_in_user
        # Leave Tracker approval defaults
        if pane_name == "Leave Tracker":
            try:
                if "Approval Status" in cols:
                    row[cols.index("Approval Status")] = "Pending"
                req_eid = self.logged_in_user
                req_name = self._resolve_name_from_eid(req_eid) or self.logged_in_user
                mgr_eid = self._resolve_manager_eid_for_user_eid(req_eid) or ""
                mgr_name = self._resolve_name_from_eid(mgr_eid) if mgr_eid else ""
                if "Requested By Enterprise ID" in cols:
                    row[cols.index("Requested By Enterprise ID")] = req_eid
                if "Requested By Name" in cols:
                    row[cols.index("Requested By Name")] = req_name
                if "Approver Enterprise ID" in cols:
                    row[cols.index("Approver Enterprise ID")] = mgr_eid
                if "Approver Name" in cols:
                    row[cols.index("Approver Name")] = mgr_name or ""
                # Notify approver if current user is approver
                try:
                    date_val = row[cols.index("Date")] if "Date" in cols else ""
                    self._notify_approver_new_leave(mgr_eid, date_val, req_name)
                except Exception:
                    pass
            except Exception:
                pass
        # Client Visits & Audits defaults
        if pane_name == "Client Visits / Audits":
            if "Status" in cols:
                sidx = cols.index("Status")
                row[sidx] = AUDIT_STATUS_OPTIONS[0]
            if "Detailed Action Plan/Status" in cols:
                aidx = cols.index("Detailed Action Plan/Status")
                row[aidx] = "No action plan provided"
        # Accolades defaults and auto-fill
        if pane_name == "Accolades":
            if "ISMT Enterprise ID" in cols:
                eidx = cols.index("ISMT Enterprise ID")
                row[eidx] = self.logged_in_user
            # Auto Month as MMM-YY and Fortnight based on today's date
            try:
                today = datetime.today().date()
                mon_val = today.strftime("%b-%y")
                if "Month" in cols:
                    midx = cols.index("Month")
                    row[midx] = mon_val
                if "Fortnight" in cols:
                    fidx = cols.index("Fortnight")
                    row[fidx] = "1st Fortnight" if today.day <= 15 else "2nd Fortnight"
            except Exception as e:
                ErrorHandler.handle_ui_error("set date fields", e)
        
        # Potential Issues ageing calculation
        if pane_name == "Potential Issues":
            # Set Created Date to today if not already set
            if "Created Date" in cols:
                created_idx = cols.index("Created Date")
                if not row[created_idx]:  # If Created Date is empty
                    row[created_idx] = datetime.today().strftime("%Y-%m-%d")
            
            # Calculate Ageing based on Created Date
            if "Ageing" in cols and "Created Date" in cols:
                ageing_idx = cols.index("Ageing")
                created_idx = cols.index("Created Date")
                created_date_str = row[created_idx]
                
                if created_date_str:
                    try:
                        # Parse the created date
                        created_date = datetime.strptime(created_date_str, "%Y-%m-%d").date()
                        today = datetime.today().date()
                        
                        # Calculate days difference
                        days_diff = (today - created_date).days
                        row[ageing_idx] = str(days_diff)
                        
                        # Set Ageing RAG based on days
                        if "Ageing RAG" in cols:
                            ageing_rag_idx = cols.index("Ageing RAG")
                            if days_diff <= 7:
                                row[ageing_rag_idx] = "Green"
                            elif days_diff <= 14:
                                row[ageing_rag_idx] = "Amber"
                            else:
                                row[ageing_rag_idx] = "Red"
                    except ValueError:
                        # If date format is invalid, set ageing to 0
                        row[ageing_idx] = "0"
                        if "Ageing RAG" in cols:
                            ageing_rag_idx = cols.index("Ageing RAG")
                            row[ageing_rag_idx] = "Green"
        self.data[pane_name].append(row)
        # Refresh UI for the affected pane
        if pane_name == "Project Details" and hasattr(self, 'projects_table') and self.projects_table is not None:
            try:
                self._load_projects_data()
                if hasattr(self, '_update_summary_metrics'):
                    self._update_summary_metrics()
            except Exception:
                pass
        else:
            self.rebuild_table(pane_name)
        self.update_dashboard()
        self.update_home_stats()  # Update home page stats
        
        # Refresh calendar if the added row affects calendar data
        if pane_name in ["Potential Issues", "Activities", "Leave Tracker"]:
            try:
                if hasattr(self, 'calendar_tab') and self.calendar_tab:
                    self.calendar_tab.refresh_calendar_if_visible()
            except Exception:
                pass
        
        self._save_backend_sqlite()
        self._log(f"Added row to {pane_name}", "SUCCESS")
        self._log_change("Add Row", pane_name, f"Added new row with {len(row)} columns", "SUCCESS")
        # Show congratulations for Accolades
        if pane_name == "Accolades":
            try:
                dlg = QDialog(self)
                dlg.setWindowTitle("Congratulations!")
                lay = QVBoxLayout(dlg)
                title = QLabel("🎉 Congratulations!")
                tf = QFont(); tf.setPointSize(20); tf.setWeight(QFont.Weight.Bold)
                title.setFont(tf)
                title.setAlignment(Qt.AlignmentFlag.AlignCenter)
                msg = QLabel("Your accolade has been added.")
                msg.setAlignment(Qt.AlignmentFlag.AlignCenter)
                msg.setStyleSheet("color:#0f172a; font-size:14px;")
                lay.addWidget(title)
                lay.addWidget(msg)
                btn = QPushButton("Nice!")
                btn.setObjectName("primary")
                btn.clicked.connect(dlg.accept)
                lay.addWidget(btn, alignment=Qt.AlignmentFlag.AlignCenter)
                dlg.setStyleSheet("QDialog{background:#ffffff;} QPushButton#primary{background:#2a7de1;color:#fff;border:none;padding:8px 16px;border-radius:6px;}")
                dlg.resize(360, 180)
                dlg.exec()
            except Exception as e:
                ErrorHandler.handle_ui_error("show congratulations dialog", e)

    def update_ageing_for_potential_issues(self):
        """Update ageing calculation for all Potential Issues rows with enhanced logging"""
        if "Potential Issues" not in self.data:
            self._log("No Potential Issues data found for ageing calculation", "WARNING")
            return
        
        cols = PANE_COLUMNS["Potential Issues"]
        if "Ageing" not in cols or "Created Date" not in cols:
            self._log("Ageing or Created Date columns not found in Potential Issues", "WARNING")
            return
        
        ageing_idx = cols.index("Ageing")
        created_idx = cols.index("Created Date")
        ageing_rag_idx = cols.index("Ageing RAG") if "Ageing RAG" in cols else None
        
        today = datetime.today().date()
        updated_count = 0
        error_count = 0
        
        for row_idx, row in enumerate(self.data["Potential Issues"]):
            if len(row) > max(ageing_idx, created_idx):
                created_date_str = row[created_idx] if created_idx < len(row) else ""
                
                if created_date_str:
                    try:
                        # Parse the created date
                        created_date = datetime.strptime(created_date_str, "%Y-%m-%d").date()
                        
                        # Calculate days difference
                        days_diff = (today - created_date).days
                        old_ageing = row[ageing_idx] if ageing_idx < len(row) else "0"
                        row[ageing_idx] = str(days_diff)
                        
                        # Set Ageing RAG based on days
                        if ageing_rag_idx is not None and ageing_rag_idx < len(row):
                            old_rag = row[ageing_rag_idx] if ageing_rag_idx < len(row) else "Green"
                            if days_diff <= 7:
                                row[ageing_rag_idx] = "Green"
                            elif days_diff <= 14:
                                row[ageing_rag_idx] = "Amber"
                            else:
                                row[ageing_rag_idx] = "Red"
                            
                            # Log significant changes
                            if old_rag != row[ageing_rag_idx]:
                                self._log_change("Ageing RAG Update", "Potential Issues", 
                                               f"Row {row_idx + 1}: {old_rag} → {row[ageing_rag_idx]} ({days_diff} days)", "INFO")
                        
                        updated_count += 1
                        
                    except ValueError as e:
                        # If date format is invalid, set ageing to 0
                        row[ageing_idx] = "0"
                        if ageing_rag_idx is not None and ageing_rag_idx < len(row):
                            row[ageing_rag_idx] = "Green"
                        error_count += 1
                        self._log_change("Ageing Calculation Error", "Potential Issues", 
                                       f"Row {row_idx + 1}: Invalid date format '{created_date_str}' - {str(e)}", "WARNING")
        
        # Log summary
        if updated_count > 0:
            self._log(f"Updated ageing for {updated_count} Potential Issues entries", "SUCCESS")
        if error_count > 0:
            self._log(f"Encountered {error_count} errors during ageing calculation", "WARNING")

    def auto_update_ageing(self):
        """Automatically update ageing calculations for all relevant panes"""
        self._log("Starting automatic ageing update", "INFO")
        
        # Update Potential Issues ageing
        self.update_ageing_for_potential_issues()
        
        # Update other panes that might have ageing calculations
        for pane_name in ["Activities", "Client Visits / Audits"]:
            if pane_name in self.data and self.data[pane_name]:
                self._update_ageing_for_pane(pane_name)
        # Also auto-advance Client Visits / Audits status based on start date
        try:
            if "Client Visits / Audits" in self.data:
                cols_cv = PANE_COLUMNS["Client Visits / Audits"]
                start_idx = cols_cv.index("Audit Start Date") if "Audit Start Date" in cols_cv else (cols_cv.index("Start Date") if "Start Date" in cols_cv else None)
                status_idx = cols_cv.index("Status") if "Status" in cols_cv else None
                if start_idx is not None and status_idx is not None:
                    today = datetime.today().date()
                    changed = 0
                    for r in self.data["Client Visits / Audits"]:
                        if start_idx < len(r) and status_idx < len(r):
                            s = str(r[start_idx]).strip()
                            st = str(r[status_idx]).strip().lower()
                            if s and st != "in progress":
                                try:
                                    sd = datetime.strptime(s, "%Y-%m-%d").date()
                                    if sd <= today:
                                        r[status_idx] = "In Progress"
                                        changed += 1
                                except Exception:
                                    # ignore unparsable
                                    pass
                    if changed:
                        try:
                            self.rebuild_table("Client Visits / Audits")
                        except Exception as e:
                            ErrorHandler.handle_ui_error("rebuild Client Visits / Audits", e)
                        self._log_change("Auto Status", "Client Visits / Audits", f"Marked {changed} row(s) as In Progress", "INFO")
        except Exception as e:
            ErrorHandler.handle_ui_error("auto update client visits status", e)
        
        self._log("Automatic ageing update completed", "SUCCESS")
    def _update_ageing_for_pane(self, pane_name: str):
        """Update ageing for a specific pane if it has ageing columns"""
        if pane_name not in self.data or not self.data[pane_name]:
            return
        
        cols = PANE_COLUMNS.get(pane_name, [])
        if "Ageing" not in cols:
            return
        
        ageing_idx = cols.index("Ageing")
        ageing_rag_idx = cols.index("Ageing RAG") if "Ageing RAG" in cols else None
        
        # Find start date column
        start_col_candidates = [c for c in ("Start Date", "Created Date", "Audit Start Date") if c in cols]
        if not start_col_candidates:
            return
        
        start_idx = cols.index(start_col_candidates[0])
        today = datetime.today().date()
        updated_count = 0
        
        for row_idx, row in enumerate(self.data[pane_name]):
            if len(row) > max(ageing_idx, start_idx):
                try:
                    start_date_str = str(row[start_idx])
                    if start_date_str:
                        start_date = datetime.strptime(start_date_str, "%Y-%m-%d").date()
                        days_diff = (today - start_date).days
                        row[ageing_idx] = str(days_diff)
                        
                        # Update RAG if present
                        if ageing_rag_idx is not None and ageing_rag_idx < len(row):
                            if days_diff <= 7:
                                row[ageing_rag_idx] = "Green"
                            elif days_diff <= 14:
                                row[ageing_rag_idx] = "Amber"
                            else:
                                row[ageing_rag_idx] = "Red"
                        
                        updated_count += 1
                except (ValueError, IndexError):
                    continue
        
        if updated_count > 0:
            self._log(f"Updated ageing for {updated_count} {pane_name} entries", "INFO")

    def delete_selected(self, pane_name: str) -> None:
        table = self.tables[pane_name]
        rows = sorted({idx.row() for idx in table.selectedIndexes()}, reverse=True)
        if not rows:
            QMessageBox.warning(self, "Delete Row", "Please select at least one row")
            return
        for r in rows:
            if 0 <= r < len(self.data[pane_name]):
                self.data[pane_name].pop(r)
        self.rebuild_table(pane_name)
        self.update_dashboard()
        self.update_home_stats()  # Update home page stats
        self._update_notification_count()  # Update notification count
        self._save_backend_sqlite()
        self._log(f"Deleted {len(rows)} row(s) from {pane_name}", "WARNING")
        self._log_change("Delete Rows", pane_name, f"Deleted {len(rows)} row(s) at positions: {sorted(rows)}", "WARNING")

    def import_pane(self, pane_name: str) -> None:
        path, _ = QFileDialog.getOpenFileName(self, f"Import {pane_name}", "", "Excel (*.xlsx);;CSV (*.csv)")
        if not path:
            return
        try:
            if os.path.basename(path).lower() == "sample":
                self.show_import_tutorial(pane_name)
                return
            if path.lower().endswith(".csv"):
                # CSV path: allow import even without pandas (use Python csv fallback)
                try:
                    import pandas as pd  # type: ignore
                    df = pd.read_csv(path)
                except Exception:
                    # Fallback minimal CSV reader into DataFrame-like structure using headers
                    import csv
                    rows = []
                    headers = []
                    with open(path, newline='', encoding='utf-8') as f:
                        reader = csv.reader(f)
                        for i, row in enumerate(reader):
                            if i == 0:
                                headers = row
                            else:
                                rows.append(row)
                    # Build a minimal dict-of-lists to align with existing code expectations
                    class _SimpleDF:
                        def __init__(self, headers, rows):
                            self.columns = headers
                            self._rows = rows
                        def __getitem__(self, cols):
                            # Return list-of-lists with selected columns
                            idxs = [self.columns.index(c) for c in cols if c in self.columns]
                            return [[r[i] if i < len(r) else '' for i in idxs] for r in self._rows]
                    df = _SimpleDF(headers, rows)
            else:
                # Excel path: require pandas; if missing, show clear prompt
                try:
                    import pandas as pd  # type: ignore
                    df = pd.read_csv(path)
                except Exception:
                    QMessageBox.critical(
                        self,
                        "Excel Import Requires pandas",
                        "Excel import needs the 'pandas' package.\n\n"
                        "Install with:\n"
                        "  pip install pandas openpyxl\n\n"
                        "Alternatively, save your file as CSV and import that instead."
                    )
                    return
            cols = PANE_COLUMNS[pane_name]
            # Handle both pandas DataFrame and _SimpleDF
            df_columns = list(getattr(df, 'columns', []))
            if all(c in df_columns for c in cols):
                values = df[cols]
            else:
                # Open mapping wizard
                # If not a real pandas DataFrame, try to coerce using pandas for the mapping UI
                try:
                    import pandas as pd  # type: ignore
                    if not hasattr(df, 'iloc'):
                        # Convert _SimpleDF into DataFrame
                        data = {c: [] for c in df_columns}
                        for row in getattr(df, '_rows', []):
                            for i, c in enumerate(df_columns):
                                data[c].append(row[i] if i < len(row) else '')
                        df_for_map = pd.DataFrame(data)
                    else:
                        df_for_map = df
                except Exception:
                    QMessageBox.critical(
                        self,
                        "Import Mapping Requires pandas",
                        "Column mapping requires the 'pandas' package.\n\n"
                        "Install with:\n"
                        "  pip install pandas\n\n"
                        "Alternatively, rename your CSV headers to exactly match required columns and retry."
                    )
                    return
                mapped = self.open_import_mapping_wizard(pane_name, df_for_map)
                if mapped is None:
                    return
                values = mapped
            # Convert to list of lists regardless of DataFrame-like
            if hasattr(values, 'astype') and hasattr(values, 'values'):
                imported = values.astype(str).values.tolist()
            else:
                # values came from _SimpleDF __getitem__
                imported = [[str(v) for v in row] for row in values]
            if "ISM Name" in cols:
                idx = cols.index("ISM Name")
                for row in imported:
                    if not row[idx]:
                        mapped_name = self._resolve_ism_name_from_eid(self.logged_in_user)
                        row[idx] = mapped_name if mapped_name else self.logged_in_user
            self.data[pane_name] = imported
            self.rebuild_table(pane_name)
            self.update_dashboard()
            self.update_home_stats()  # Update home page stats
            self._save_backend_sqlite()
            self._log(f"Imported data into {pane_name} from file")
        except Exception as e:
            QMessageBox.critical(self, "Import Failed", str(e))

    def open_import_mapping_wizard(self, pane_name: str, df: pd.DataFrame) -> pd.DataFrame | None:
        dlg = QDialog(self)
        dlg.setWindowTitle(f"Map Columns for {pane_name}")
        lay = QVBoxLayout(dlg)
        lay.addWidget(QLabel("Map source file columns to required columns"))
        combos: list[QComboBox] = []
        grid = QGridLayout()
        src_cols = [str(c) for c in df.columns]
        required = PANE_COLUMNS[pane_name]
        for i, col in enumerate(required):
            grid.addWidget(QLabel(col), i, 0)
            cb = QComboBox(); cb.setEditable(True)
            cb.addItem("<None>")
            cb.addItems(src_cols)
            # Preselect case-insensitive match
            matches = [c for c in src_cols if c.lower() == col.lower()]
            if matches:
                cb.setCurrentText(matches[0])
            combos.append(cb)
            grid.addWidget(cb, i, 1)
        lay.addLayout(grid)
        mode_row = QHBoxLayout(); mode_row.addWidget(QLabel("On import:"))
        mode = QComboBox(); mode.addItems(["Replace existing data", "Append to existing data"]) 
        mode_row.addWidget(mode); mode_row.addStretch(1)
        lay.addLayout(mode_row)
        btns = QHBoxLayout(); ok = QPushButton("Import"); cancel = QPushButton("Cancel")
        ok.setObjectName("primary"); btns.addStretch(1); btns.addWidget(ok); btns.addWidget(cancel)
        lay.addLayout(btns)
        ok.clicked.connect(dlg.accept); cancel.clicked.connect(dlg.reject)
        dlg.resize(600, 400)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return None
        # Build mapped dataframe with required columns
        out = pd.DataFrame()
        for req, cb in zip(required, combos):
            src = cb.currentText()
            if src and src != "<None>" and src in df.columns:
                out[req] = df[src]
            else:
                out[req] = ""
        if mode.currentText().startswith("Append") and pane_name in self.data and self.data[pane_name]:
            existing = pd.DataFrame(self.data[pane_name], columns=PANE_COLUMNS[pane_name])
            out = pd.concat([existing, out], ignore_index=True)
        return out

    def _show_import_guide_and_import(self, pane_name: str) -> None:
        """Show import guide dialog and then proceed with import"""
        try:
            # Show import guide dialog first
            result = self._show_pane_import_guide(pane_name)
            
            # Only proceed with import if user clicked "Continue to Import"
            if result:
                self.import_pane(pane_name)
            
        except Exception as e:
            ErrorHandler.handle_ui_error("show import guide and import", e)
    
    def _show_pane_import_guide(self, pane_name: str) -> bool:
        """Show import guide dialog for any pane. Returns True if user wants to continue with import."""
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle(f"{pane_name} Import Guide")
            dialog.setModal(True)
            dialog.resize(500, 400)
            dialog.setStyleSheet("""
                QDialog {
                    background-color: #f8f9fa;
                }
            """)
            
            layout = QVBoxLayout(dialog)
            layout.setSpacing(15)
            layout.setContentsMargins(20, 20, 20, 20)
            
            # Title
            title = QLabel(f"📋 {pane_name} Import Guide")
            title.setStyleSheet("""
                font-size: 18px;
                font-weight: bold;
                color: #2c3e50;
                margin-bottom: 10px;
            """)
            layout.addWidget(title)
            
            # Get pane columns for this pane
            columns = PANE_COLUMNS.get(pane_name, [])
            
            # Import steps
            steps_text = QLabel(f"""
<b>How to Import {pane_name} Data:</b><br><br>
1. <b>Download Sample:</b> Click "📋 Sample" to get the template with proper headers<br><br>
2. <b>Fill Data:</b> Add your data with the required fields:<br>
   {self._format_columns_for_guide(columns)}<br>
3. <b>Save File:</b> Save as .csv or .xlsx format<br><br>
4. <b>Import:</b> Click "📥 Import" and select your file<br><br>
5. <b>Verify:</b> Check the data displays correctly in the table<br><br>
<b>Important Notes:</b><br>
• Make sure all required fields are filled<br>
• Use the exact column names from the sample template<br>
• Data will be validated during import
            """)
            steps_text.setWordWrap(True)
            steps_text.setStyleSheet("""
                font-size: 12px;
                color: #495057;
                line-height: 1.4;
                padding: 15px;
                background-color: #ffffff;
                border: 1px solid #dee2e6;
                border-radius: 8px;
            """)
            layout.addWidget(steps_text)
            
            # Buttons
            button_layout = QHBoxLayout()
            
            # Download sample button
            sample_btn = QPushButton("📋 Download Sample")
            sample_btn.setStyleSheet("""
                QPushButton {
                    background-color: #9b59b6;
                    color: white;
                    padding: 10px 20px;
                    border: none;
                    border-radius: 6px;
                    font-size: 14px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #8e44ad;
                }
            """)
            sample_btn.clicked.connect(lambda: (dialog.accept(), self.export_sample_file(pane_name)))
            button_layout.addWidget(sample_btn)
            
            button_layout.addStretch()
            
            # Continue button
            continue_btn = QPushButton("Continue to Import")
            continue_btn.setStyleSheet("""
                QPushButton {
                    background-color: #2a7de1;
                    color: white;
                    padding: 10px 20px;
                    border: none;
                    border-radius: 6px;
                    font-size: 14px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #1e6bb8;
                }
            """)
            continue_btn.clicked.connect(dialog.accept)
            button_layout.addWidget(continue_btn)
            
            # Cancel button
            cancel_btn = QPushButton("Cancel")
            cancel_btn.setStyleSheet("""
                QPushButton {
                    background-color: #6c757d;
                    color: white;
                    padding: 10px 20px;
                    border: none;
                    border-radius: 6px;
                    font-size: 14px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #5a6268;
                }
            """)
            cancel_btn.clicked.connect(dialog.reject)
            button_layout.addWidget(cancel_btn)
            
            layout.addLayout(button_layout)
            
            # Return True if user clicked "Continue to Import", False if "Cancel" or "Download Sample"
            result = dialog.exec()
            return result == QDialog.DialogCode.Accepted
            
        except Exception as e:
            ErrorHandler.handle_ui_error("show pane import guide", e)
            return False
    
    def _format_columns_for_guide(self, columns: list) -> str:
        """Format column names for the import guide"""
        if not columns:
            return "• No specific columns defined"
        
        # Format columns in a readable way
        formatted = []
        for i, col in enumerate(columns, 1):
            formatted.append(f"   • <b>{col}:</b> {self._get_column_description(col)}")
        
        return "<br>".join(formatted)
    
    def _get_column_description(self, column_name: str) -> str:
        """Get a description for a column name"""
        descriptions = {
            "Action": "Description of the action to be taken",
            "Description": "Detailed description of the item",
            "Tracker": "Person tracking this item",
            "Ownership": "Person responsible for this item",
            "Remarks": "Additional comments or notes",
            "Start Date": "Start date (YYYY-MM-DD format)",
            "End Date": "End date (YYYY-MM-DD format)",
            "Status": "Current status of the item",
            "Priority": "Priority level (High/Medium/Low)",
            "Efforts": "Effort required or hours spent",
            "Additional Remarks": "Extra comments or notes",
            "Task Type": "Type of task or issue",
            "Created Date": "Date when created (YYYY-MM-DD)",
            "Due Date": "Due date (YYYY-MM-DD)",
            "Project Name": "Name of the project",
            "Project ID": "Unique project identifier",
            "RAG Status": "Red/Amber/Green status",
            "ISM Name": "Name of the ISM",
            "Action Owner": "Person responsible for the action",
            "Ageing": "Age of the item in days",
            "Ageing RAG": "RAG status based on ageing",
            "Added in PI Tool": "Whether added in PI tool (Yes/No)",
            "Leads Attention Required": "Whether leads attention is needed (Yes/No)",
            "Updates": "Latest updates or progress",
            "Activity/Issue": "Description of the activity or issue",
            "Target Date": "Target completion date (YYYY-MM-DD)",
            "Support Required": "Whether support is needed (Yes/No)",
            "RAG": "Red/Amber/Green status",
            "Brief Update": "Brief update on progress",
            "Audit Scope": "Scope of the audit",
            "Audit Type": "Type of audit",
            "Audit Start Date": "Audit start date (YYYY-MM-DD)",
            "Audit End Date": "Audit end date (YYYY-MM-DD)",
            "Detailed Action Plan/Status": "Detailed action plan or status",
            "Month": "Month (MM/YYYY format)",
            "ISMT Enterprise ID": "ISMT person's enterprise ID",
            "Appreciator Enterprise ID": "Appreciator's enterprise ID",
            "Appreciator Designation": "Appreciator's job title",
            "Accolades Description": "Description of the accolade",
            "Fortnight": "Fortnight period",
            "Date": "Date (YYYY-MM-DD format)",
            "Type": "Type of leave or activity",
            "Duration": "Duration of the activity",
            "Approval Status": "Status of approval (Pending/Approved/Rejected)",
            "Approver Enterprise ID": "Approver's enterprise ID",
            "Approver Name": "Name of the approver",
            "Requested By Enterprise ID": "Requester's enterprise ID",
            "Requested By Name": "Name of the requester",
            "Decision Date": "Date of decision (YYYY-MM-DD)",
            "Approval Comments": "Comments from approver"
        }
        
        return descriptions.get(column_name, "Data for this field")

    def show_import_tutorial(self, pane_name: str) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle(f"How to import data for {pane_name}")
        lay = QVBoxLayout(dlg)
        steps = QLabel(
            "1) Download sample file with required headers.\n"
            "2) Fill your data exactly under those headers.\n"
            "3) Save as .xlsx or .csv.\n"
            "4) Use Import and select the filled file." )
        lay.addWidget(steps)
        sample_btn = QPushButton("Download Sample File")
        sample_btn.clicked.connect(lambda: self.export_sample_file(pane_name))
        lay.addWidget(sample_btn)
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dlg.accept)
        lay.addWidget(close_btn)
        dlg.exec()

    def export_sample_file(self, pane_name: str) -> None:
        path, _ = QFileDialog.getSaveFileName(self, f"Save Sample for {pane_name}", f"{pane_name}_sample.xlsx", "Excel (*.xlsx);;CSV (*.csv)")
        if not path:
            return
        try:
            cols = PANE_COLUMNS[pane_name]
            sample_df = pd.DataFrame(columns=cols)
            sample_df.to_csv(path, index=False)
            QMessageBox.information(self, "Sample", "Sample file saved.")
        except Exception as e:
            QMessageBox.critical(self, "Sample Failed", str(e))

    def export_pane(self, pane_name: str) -> None:
        path, _ = QFileDialog.getSaveFileName(self, f"Export {pane_name}", f"{pane_name}.xlsx", "Excel (*.xlsx);;CSV (*.csv)")
        if not path:
            return
        try:
            df = pd.DataFrame(self.data[pane_name], columns=PANE_COLUMNS[pane_name])
            if path.lower().endswith(".csv"):
                df.to_csv(path, index=False)
            else:
                df.to_csv(path, index=False)
            QMessageBox.information(self, "Export", "Export successful")
            self._log(f"Exported {pane_name} to {os.path.basename(path)}")
        except Exception as e:
            QMessageBox.critical(self, "Export Failed", str(e))

    def export_current_view(self, pane_name: str) -> None:
        table = self.tables.get(pane_name)
        if not table:
            return
        # Collect only visible rows and current column order
        cols = [table.horizontalHeaderItem(i).text() for i in range(table.columnCount())]
        visible_rows: list[list[str]] = []
        for r in range(table.rowCount()):
            if table.isRowHidden(r):
                continue
            row_vals = []
            for c in range(table.columnCount()):
                it = table.item(r, c)
                row_vals.append(it.text() if it else "")
            visible_rows.append(row_vals)
        if not visible_rows:
            QMessageBox.information(self, "Export", "No visible rows to export.")
            return
        path, _ = QFileDialog.getSaveFileName(self, f"Export Current View — {pane_name}", f"{pane_name}_view.xlsx", "Excel (*.xlsx);;CSV (*.csv)")
        if not path:
            return
        try:
            df = pd.DataFrame(visible_rows, columns=cols)
            if path.lower().endswith(".csv"):
                df.to_csv(path, index=False)
            else:
                df.to_csv(path, index=False)
            QMessageBox.information(self, "Export", "Current view exported")
            self._log(f"Exported current view of {pane_name}")
            ErrorHandler.handle_success("export current view", f"Exported {len(visible_rows)} rows from {pane_name}")
        except Exception as e:
            ErrorHandler.handle_ui_error("export current view", e, {"pane_name": pane_name, "file_path": path})
            QMessageBox.critical(self, "Export Failed", str(e))

    def edit_cell(self, pane_name: str, row: int, col: int) -> None:
        table = self.tables[pane_name]
        col_name = PANE_COLUMNS[pane_name][col]
        item = table.item(row, col)
        text_val = item.text() if item else ""
        # Date columns: open date picker dialog (pane-specific)
        try:
            date_cols = {
                "initiatives": {"start date", "end date"},
                "potential issues": {"created date", "due date"},
                "activities": {"start date", "target date"},
                "client visits / audits": {"audit start date", "audit end date"},
            }
            pkey = pane_name.lower()
            if pkey in date_cols and col_name.lower() in date_cols[pkey]:
                try:
                    # parse to QDate if possible
                    qd = None
                    if text_val:
                        try:
                            dt = datetime.strptime(text_val, "%Y-%m-%d")
                            qd = QDate(dt.year, dt.month, dt.day)
                        except Exception:
                            qd = QDate.currentDate()
                    else:
                        qd = QDate.currentDate()
                    dlg = DatePickerDialog(self, qd)
                    if dlg.exec():
                        val = dlg.selected_date()
                        self._set_cell_value(pane_name, row, col, val)
                except Exception as e:
                    ErrorHandler.handle_ui_error("open date picker", e)
                return
        except Exception:
            pass
        # Special-case exceptions that should remain popup-based
        if pane_name == "Potential Issues" and col_name == "Updates":
            # Rich Updates dialog retained
            from PyQt6.QtWidgets import QTextEdit
            dlg = QDialog(self)
            dlg.setWindowTitle("Updates - View, Edit and Append")
            lay = QVBoxLayout(dlg)
            dlg.setStyleSheet(
                """
                QDialog { background:#ffffff; }
                QLabel { color:#0f172a; }
                QListWidget { background:#f8fafc; border:1px solid #dbe5f0; border-radius:8px; }
                QPushButton { background:#f1f5f9; border:1px solid #d4dbe6; border-radius:6px; padding:6px 10px; }
                QPushButton:hover { background:#e6effa; }
                QPushButton#primary { background:#2a7de1; color:#ffffff; border:none; }
                QPushButton#danger { background:#f55f77; color:#ffffff; border:none; }
                """
            )
            # Old updates as list
            old_hdr = QLabel("Old Updates"); old_hdr.setStyleSheet("font-weight:700; color:#0f172a; font-size:13px;")
            lay.addWidget(old_hdr)
            old_list = QListWidget(); old_list.setSelectionMode(QListWidget.SelectionMode.SingleSelection)
            old_entries = [line for line in (text_val or "").split('\n') if line.strip()]
            def render_old_list_from_entries(entries: list[str]):
                from PyQt6.QtWidgets import QLabel
                import re
                old_list.clear()
                for line in entries:
                    m = re.match(r"\s*(\[[^\]]+\])(.*)$", line)
                    if m:
                        head, rest = m.group(1), m.group(2)
                        html = f"<span style='font-weight:700;color:#0f172a'>{head}</span><span style='color:#334155'>{rest}</span>"
                    else:
                        html = f"<span style='color:#334155'>{line}</span>"
                    lbl = QLabel(); lbl.setTextFormat(Qt.TextFormat.RichText); lbl.setText(html); lbl.setWordWrap(True)
                    it = QListWidgetItem(old_list); it.setSizeHint(lbl.sizeHint()); old_list.addItem(it); old_list.setItemWidget(it, lbl)
            render_old_list_from_entries(old_entries)
            old_btns = QHBoxLayout(); btn_edit = QPushButton("Edit Selected"); btn_edit.setObjectName("primary"); btn_delete = QPushButton("Delete Selected"); btn_delete.setObjectName("danger")
            old_btns.addWidget(btn_edit); old_btns.addWidget(btn_delete); old_btns.addStretch(1)
            lay.addWidget(old_list); lay.addLayout(old_btns)
            new_hdr = QLabel("New Update (appends with date & user)"); new_hdr.setStyleSheet("font-weight:700; color:#0f172a; font-size:13px;")
            lay.addWidget(new_hdr)
            new_te = QTextEdit(); new_te.setPlaceholderText("Type new update here and click 'Append New Update'"); new_te.setStyleSheet("QTextEdit { background:#fff8e6; border:1px solid #f6a23c; border-radius:8px; padding:8px; }"); new_te.setMinimumHeight(80)
            lay.addWidget(new_te)
            mention_row = QHBoxLayout(); mention_lbl = QLabel("@Team:"); mention_input = QComboBox(); mention_input.setEditable(True); mention_input.addItems(sorted(self.team_history)); mention_btn = QPushButton("Insert @Team")
            mention_row.addWidget(mention_lbl); mention_row.addWidget(mention_input); mention_row.addWidget(mention_btn)
            lay.addLayout(mention_row)
            btn_append = QPushButton("Append New Update"); btn_append.setObjectName("primary"); btn_append.setMinimumHeight(32); lay.addWidget(btn_append)
            footer = QHBoxLayout(); save_all = QPushButton("Save All Changes"); cancel = QPushButton("Close"); save_all.setObjectName("primary"); footer.addStretch(1); footer.addWidget(save_all); footer.addWidget(cancel); lay.addLayout(footer)
            def refresh_cell_from_list():
                combined_items = []
                for i in range(old_list.count()):
                    it = old_list.item(i); w = old_list.itemWidget(it)
                    if w is not None: combined_items.append(w.text().replace("<b>", "").replace("</b>", ""))
                    else: combined_items.append(it.text())
                combined = "\n".join(combined_items)
                self._set_cell_value_internal(pane_name, row, col, combined, raw_override=True)
            def on_edit():
                it = old_list.currentItem()
                if not it:
                    return
                raw = None
                w = old_list.itemWidget(it)
                if w is not None:
                    raw = w.text()
                    raw = raw.replace("<span style='font-weight:700;color:#0f172a'>", "").replace("</span>", "")
                    raw = raw.replace("<span style='color:#334155'>", "")
                else:
                    raw = it.text()
                txt, ok = QInputDialog.getText(self, "Edit Update", "Update text:", text=raw)
                if ok:
                    entries = []
                    for j in range(old_list.count()):
                        if j == old_list.currentRow():
                            entries.append(txt)
                        else:
                            jt = old_list.item(j)
                            jw = old_list.itemWidget(jt)
                            if jw is not None:
                                entries.append(jw.text().replace("<span style='font-weight:700;color:#0f172a'>", "").replace("</span>", "").replace("<span style='color:#334155'>", ""))
                            else:
                                entries.append(jt.text())
                    render_old_list_from_entries(entries)
            def on_delete():
                r = old_list.currentRow()
                if r >= 0:
                    old_list.takeItem(r)
            def on_append():
                txt = new_te.toPlainText().strip()
                if not txt:
                    return
                self._set_cell_value(pane_name, row, col, txt)
                updated_val = self.data[pane_name][row][col]
                entries = [l for l in (updated_val or "").split('\n') if l.strip()]
                render_old_list_from_entries(entries)
                new_te.clear()
            def on_save_all():
                refresh_cell_from_list()
                dlg.accept()
            btn_edit.clicked.connect(on_edit); btn_delete.clicked.connect(on_delete)
            try:
                def create_shortcuts():
                    ctrl_enter_shortcut = QShortcut(QKeySequence("Ctrl+Return"), new_te); ctrl_enter_shortcut.activated.connect(on_append)
                    ctrl_s_shortcut = QShortcut(QKeySequence("Ctrl+S"), new_te); ctrl_s_shortcut.activated.connect(on_save_all)
                self._ensure_main_thread(create_shortcuts)
            except Exception as e:
                ErrorHandler.handle_ui_error("create keyboard shortcuts", e)
            def insert_mention():
                t = mention_input.currentText().strip()
                if t:
                    cursor = new_te.textCursor()
                    cursor.insertText(f"@{t} ")
                    new_te.setTextCursor(cursor)
            mention_btn.clicked.connect(insert_mention)
            btn_append.clicked.connect(on_append); save_all.clicked.connect(on_save_all); cancel.clicked.connect(dlg.reject)
            dlg.resize(700, 500); dlg.exec()
        elif pane_name == "Activities" and col_name in ("Brief Update", "Briefupdate"):
            # Keep popup editing for Brief Update in Activities
            val, ok = QInputDialog.getText(self, "Edit Brief Update", f"Enter value for {col_name}:", text=text_val)
            if ok:
                self._set_cell_value(pane_name, row, col, val)
        elif pane_name == "Potential Issues" and col_name == "Action Owner":
            # Multi-select teams with persistent history
            dlg = QDialog(self)
            dlg.setWindowTitle("Select Action Owner Teams")
            lay = QVBoxLayout(dlg)
            input_line = QLineEdit()
            input_line.setPlaceholderText("Type a team and press Add")
            add_btn = QPushButton("Add Team")
            add_row = QHBoxLayout(); add_row.addWidget(input_line); add_row.addWidget(add_btn)
            lay.addLayout(add_row)
            lw = QListWidget(); lw.setSelectionMode(QListWidget.SelectionMode.MultiSelection)
            # Seed from history and any values in other rows
            history = set(self.team_history)
            # Also scrape existing Action Owner values across Potential Issues
            try:
                pi_rows = self.data.get("Potential Issues", [])
                pi_cols = PANE_COLUMNS["Potential Issues"]
                ao_idx = pi_cols.index("Action Owner")
                for rvals in pi_rows:
                    if ao_idx < len(rvals) and rvals[ao_idx]:
                        for part in str(rvals[ao_idx]).split(','):
                            t = part.strip()
                            if t:
                                history.add(t)
            except Exception as e:
                ErrorHandler.handle_ui_error("collect action owner history", e)
            existing_sel = set([s.strip() for s in (text_val or "").split(',') if s.strip()])
            for t in sorted(history):
                it = QListWidgetItem(t)
                if t in existing_sel:
                    it.setSelected(True)
                lw.addItem(it)
            lay.addWidget(lw)
            def add_team():
                t = input_line.text().strip()
                if not t:
                    return
                # add to list if not present
                found = [lw.item(i).text() for i in range(lw.count())]
                if t not in found:
                    it = QListWidgetItem(t)
                    it.setSelected(True)
                    lw.addItem(it)
                if t not in self.team_history:
                    self.team_history.append(t)
                input_line.clear()
            add_btn.clicked.connect(add_team)
            btns = QHBoxLayout()
            ok = QPushButton("Select"); cancel = QPushButton("Cancel")
            ok.clicked.connect(dlg.accept); cancel.clicked.connect(dlg.reject)
            btns.addStretch(1); btns.addWidget(ok); btns.addWidget(cancel)
            lay.addLayout(btns)
            if dlg.exec():
                selected = ", ".join([lw.item(i).text() for i in range(lw.count()) if lw.item(i).isSelected()])
                self._set_cell_value(pane_name, row, col, selected)
        else:
            # Everything else: inline editing
            if item is None:
                table.setItem(row, col, QTableWidgetItem(text_val))
            try:
                table.editItem(table.item(row, col))
            except Exception as e:
                ErrorHandler.handle_ui_error("edit cell", e, {"pane_name": pane_name, "row": row, "col": col, "col_name": col_name})
            return
        

    def _handle_cell_double_click(self, pane_name: str, row: int, col: int) -> None:
        table = self.tables[pane_name]
        col_name = PANE_COLUMNS[pane_name][col]
        # Exceptions that must stay popup-based
        if (pane_name == "Potential Issues" and col_name == "Updates") or (pane_name == "Activities" and col_name in ("Brief Update", "Briefupdate")):
            self.edit_cell(pane_name, row, col)
            return
        # Route date columns to date picker
        try:
            date_cols = {
                "initiatives": {"start date", "end date"},
                "potential issues": {"created date", "due date"},
                "activities": {"start date", "target date"},
                "client visits / audits": {"audit start date", "audit end date"},
            }
            if pane_name.lower() in date_cols and col_name.lower() in date_cols[pane_name.lower()]:
                self.edit_cell(pane_name, row, col)
                return
        except Exception:
            pass
        # Default to inline editing for everything else
        if table.item(row, col) is None:
            table.setItem(row, col, QTableWidgetItem(""))
        try:
            table.editItem(table.item(row, col))
        except Exception as e:
            ErrorHandler.handle_ui_error("edit cell", e, {"pane_name": pane_name, "row": row, "col": col, "col_name": col_name})

    def _set_cell_value(self, pane_name: str, row: int, col: int, value: str) -> None:
        return self._set_cell_value_internal(pane_name, row, col, value, raw_override=False)

    def _set_cell_value_internal(self, pane_name: str, row: int, col: int, value: str, raw_override: bool = False) -> None:
        # CRITICAL FIX: Add bounds checking to prevent crashes
        if pane_name not in PANE_COLUMNS:
            self.notifications.show_warning(f"Invalid pane name: {pane_name}")
            return
            
        cols = PANE_COLUMNS[pane_name]
        
        # Bounds checking for row and column
        if row < 0 or col < 0:
            self.notifications.show_warning(f"Invalid row/column index: row={row}, col={col}")
            return
            
        if col >= len(cols):
            self.notifications.show_warning(f"Column index {col} out of range for pane {pane_name}")
            return
            
        # Ensure data structure exists
        if pane_name not in self.data:
            self.data[pane_name] = []
            
        # Ensure row exists
        while len(self.data[pane_name]) <= row:
            self.data[pane_name].append([""] * len(cols))
            
        # Ensure row has correct number of columns
        while len(self.data[pane_name][row]) < len(cols):
            self.data[pane_name][row].append("")
        
        # Validate the input if not raw override
        if not raw_override and col < len(cols):
            column_name = cols[col]
            is_valid, error_msg = self.validation_engine.validate_field(column_name, value)
            if not is_valid:
                self.notifications.show_warning(f"Validation Error in {pane_name}: {error_msg}")
                self._log_change("Warning", pane_name, f"Validation failed for {column_name}: {error_msg}")
                return
        
        # Push undo info before change
        try:
            old = self.data[pane_name][row][col] if (row < len(self.data[pane_name]) and col < len(self.data[pane_name][row])) else ""
        except Exception:
            old = ""
        # Special rule: Potential Issues > Updates cell appends with date and ISM
        if not raw_override and pane_name == "Potential Issues" and cols[col] == "Updates":
            prev = self.data[pane_name][row][col] if col < len(self.data[pane_name][row]) else ""
            stamp = datetime.now().strftime("%Y-%m-%d")
            ism = self.logged_in_user
            new_val = prev.strip()
            entry = f"[{stamp} - {ism}] {value.strip()}"
            value_to_set = (entry if not new_val else f"{new_val}\n{entry}")
            self.data[pane_name][row][col] = value_to_set
        elif pane_name == "Client Visits & Audits" and cols[col] == "Detailed Action Plan/Status":
            if not value.strip():
                self.data[pane_name][row][col] = "No action plan provided"
            else:
                self.data[pane_name][row][col] = value
        # Handle Voice Solution and Contact Center logic
        elif pane_name == "Project Details" and cols[col] == "Voice Solution":
            self.data[pane_name][row][col] = value
            # If Voice Solution is set to "No", automatically set Contact Center to "NA"
            if value == "No" and "Contact Center" in cols:
                contact_center_col = cols.index("Contact Center")
                self.data[pane_name][row][contact_center_col] = "NA"
                # Update the table item for Contact Center as well
                if hasattr(self, 'projects_table') and row < self.projects_table.rowCount() and contact_center_col < self.projects_table.columnCount():
                    cc_item = self.projects_table.item(row, contact_center_col)
                    if cc_item:
                        cc_item.setText("NA")
                    else:
                        new_cc_item = QTableWidgetItem("NA")
                        self.projects_table.setItem(row, contact_center_col, new_cc_item)
                self._log(f"Auto-set Contact Center to NA for row {row+1} (Voice Solution = No)", show_toast=False)
        # Handle Client Visits / Audits auto-status change when start date is mentioned
        elif pane_name == "Client Visits / Audits" and cols[col] == "Audit Start Date":
            self.data[pane_name][row][col] = value
            # If a start date is provided, automatically set status to "In Progress"
            if value.strip() and "Status" in cols:
                status_col = cols.index("Status")
                current_status = self.data[pane_name][row][status_col] if status_col < len(self.data[pane_name][row]) else ""
                # Only change status if it's empty or "Yet to Start"
                if not current_status or current_status in ["", "Yet to Start"]:
                    self.data[pane_name][row][status_col] = "In Progress"
                    self._log(f"Auto-set Status to 'In Progress' for row {row+1} (Audit Start Date provided)", show_toast=False)
        else:
            self.data[pane_name][row][col] = value
        try:
            self._edit_undo_stack.append((pane_name, row, col, old, self.data[pane_name][row][col]))
            self._edit_redo_stack.clear()
        except Exception as e:
            ErrorHandler.handle_ui_error("update undo stack", e)
        # Handle Project Details specially since it doesn't use rebuild_table
        if pane_name == "Project Details":
            # Update the table item directly for Project Details
            if hasattr(self, 'projects_table') and row < self.projects_table.rowCount() and col < self.projects_table.columnCount():
                item = self.projects_table.item(row, col)
                if item:
                    item.setText(value)
                    item.setForeground(QColor("#2c3e50"))
                    item.setBackground(QColor("#ffffff"))
                else:
                    # Create new item if it doesn't exist
                    new_item = QTableWidgetItem(value)
                    new_item.setForeground(QColor("#2c3e50"))
                    new_item.setBackground(QColor("#ffffff"))
                    self.projects_table.setItem(row, col, new_item)
                # Refresh the table viewport to ensure changes are visible
                self.projects_table.viewport().update()
        else:
            self.rebuild_table(pane_name)
        
        self.update_dashboard()
        self.update_home_stats()  # Update home page stats
        self._save_backend_sqlite()
        self._log(f"Edited {pane_name} r{row+1} c{col+1}", show_toast=False)
        # After edits, run reminder check lightly (ensure main thread)
        try:
            QTimer.singleShot(0, self.check_due_date_reminders)
        except Exception as e:
            print(f"Timer error: {e}")
            # Fallback: run directly if timer fails
            try:
                self.check_due_date_reminders()
            except Exception as e:
                ErrorHandler.handle_ui_error("check due date reminders", e)
    # Table helpers
    def rebuild_table(self, pane_name: str) -> None:
        # Skip panes that don't use generic table renderer (Leave Tracker & Project Details use custom UI)
        if pane_name in ["Leave Tracker", "Project Details"]:
            return
        if pane_name not in self.tables:
            print(f"Table for {pane_name} does not exist!")
            return
        if pane_name not in self.data:
            print(f"Data for {pane_name} does not exist!")
            return
            
        table = self.tables[pane_name]
        data = self.data[pane_name]
        cols = PANE_COLUMNS[pane_name]
        
        # Close any active editor before rebuilding to prevent commitData errors
        try:
            # Close all persistent editors first
            for row in range(table.rowCount()):
                for col in range(table.columnCount()):
                    item = table.item(row, col)
                    if item:
                        table.closePersistentEditor(item)
            
            # Commit data if there's a current item with an editor
            current_item = table.currentItem()
            if current_item is not None:
                try:
                    # Get the editor for the current item
                    editor = table.indexWidget(table.currentIndex())
                    if editor is not None:
                        table.commitData(editor)
                except Exception as e:
                    ErrorHandler.handle_ui_error("commit data", e)  # Ignore commitData errors
        except Exception as e:
            ErrorHandler.handle_ui_error("close editors", e)  # Ignore all editor-related errors
        
        # Clear the table safely
        table.setRowCount(0)
        table.setColumnCount(len(cols))
        table.setHorizontalHeaderLabels(cols)
        # Restore column order if stored
        try:
            if pane_name in self._column_orders:
                order = self._column_orders[pane_name]
                header = table.horizontalHeader()
                for visual, logical in enumerate(order):
                    cur_visual = header.visualIndex(logical)
                    if cur_visual != -1 and cur_visual != visual:
                        header.moveSection(cur_visual, visual)
        except Exception as e:
            ErrorHandler.handle_ui_error("restore column order", e)
        # Re-apply hidden columns
        try:
            if pane_name in self._hidden_columns:
                for idx in self._hidden_columns[pane_name]:
                    if 0 <= idx < len(cols):
                        table.setColumnHidden(idx, True)
        except Exception as e:
            self._log_change("Warning", "UI", f"Failed to restore column visibility for {pane_name}: {str(e)}")
        for row_vals in data:
            # Auto-calc Ageing and Ageing RAG if present
            if "Ageing" in cols:
                ageing_idx = cols.index("Ageing")
                # Determine start date column
                start_col_candidates = [c for c in ("Start Date", "Created Date", "Audit Start Date") if c in cols]
                sd_idx = cols.index(start_col_candidates[0]) if start_col_candidates else None
                if sd_idx is not None and sd_idx < len(row_vals):
                    try:
                        start_dt = datetime.strptime(str(row_vals[sd_idx]), "%Y-%m-%d").date()
                        days = (datetime.today().date() - start_dt).days
                        # Ensure row has enough columns
                        if ageing_idx < len(row_vals):
                            row_vals[ageing_idx] = str(days)
                        # Ageing RAG rule if present
                        if "Ageing RAG" in cols:
                            rag_idx = cols.index("Ageing RAG")
                            rag = "Green" if days <= 5 else ("Amber" if days <= 10 else "Red")
                            if rag_idx < len(row_vals):
                                row_vals[rag_idx] = rag
                    except ValueError as e:
                        # Handle date parsing errors specifically
                        self._log_change("Warning", "Data", f"Invalid date format in row: {str(e)}")
                        if ageing_idx < len(row_vals):
                            row_vals[ageing_idx] = "Invalid Date"
                        if "Ageing RAG" in cols:
                            rag_idx = cols.index("Ageing RAG")
                            if rag_idx < len(row_vals):
                                row_vals[rag_idx] = "Error"
                    except Exception as e:
                        # Handle other unexpected errors
                        self._log_change("Error", "Data", f"Ageing calculation failed: {str(e)}")
                        if ageing_idx < len(row_vals):
                            row_vals[ageing_idx] = "Error"
            # Auto-set Status for Client Visits / Audits based on Audit Start Date
            try:
                if pane_name in ("Client Visits / Audits", "Client Visits & Audits"):
                    if "Status" in cols:
                        status_idx = cols.index("Status")
                    else:
                        status_idx = None
                    # Prefer explicit Audit Start Date, fallback to Start Date
                    start_col_name = "Audit Start Date" if "Audit Start Date" in cols else ("Start Date" if "Start Date" in cols else None)
                    if start_col_name is not None:
                        start_idx = cols.index(start_col_name)
                    else:
                        start_idx = None
                    if status_idx is not None and start_idx is not None and start_idx < len(row_vals):
                        start_val = str(row_vals[start_idx]).strip()
                        if start_val:
                            try:
                                start_dt = datetime.strptime(start_val, "%Y-%m-%d").date()
                                today = datetime.today().date()
                                new_status = "In Progress" if start_dt <= today else "Yet to Start"
                                if status_idx < len(row_vals):
                                    row_vals[status_idx] = new_status
                            except Exception:
                                # Leave status unchanged on parse error
                                pass
            except Exception as e:
                ErrorHandler.handle_ui_error("auto status for client visits/audits", e)

            r = table.rowCount()
            table.insertRow(r)
            for c, val in enumerate(row_vals):
                item = QTableWidgetItem(str(val))
                # If this column is ISM Name, add an avatar/icon
                try:
                    if cols[c] == "ISM Name":
                        name = str(val)
                        # Attempt to find a photo from org directory (not stored here), fallback to default avatar
                        avatar = self._default_avatar(name, 24)
                        item.setIcon(QIcon(avatar))
                except Exception as e:
                    ErrorHandler.handle_ui_error("set ISM avatar", e)
                table.setItem(r, c, item)
        self.apply_rag_format(pane_name)
        # Additionally, apply complexity-based coloring for Project Details when using generic rebuild
        if pane_name == "Project Details":
            try:
                cols_pd = PANE_COLUMNS["Project Details"]
                if "Complexity" in cols_pd:
                    cx_idx = cols_pd.index("Complexity")
                    for r in range(table.rowCount()):
                        it = table.item(r, cx_idx)
                        val = (it.text().strip().lower() if it else "")
                        base = None; text_color = QColor("#000000")
                        if val == "high":
                            base = QColor("#ffebee"); text_color = QColor("#c62828")
                        elif val == "medium":
                            base = QColor("#fff8e1"); text_color = QColor("#f57f17")
                        elif val == "low":
                            base = QColor("#e8f5e8"); text_color = QColor("#2e7d32")
                        if base:
                            for c in range(table.columnCount()):
                                cell = table.item(r, c)
                                if cell:
                                    cell.setBackground(base)
                                    cell.setForeground(text_color)
            except Exception as e:
                ErrorHandler.handle_ui_error("apply complexity colors", e)
        self.refresh_ism_filter()
        
        # Refresh calendar if it's visible and the pane affects calendar data
        if pane_name in ["Potential Issues", "Activities", "Leave Tracker"]:
            try:
                if hasattr(self, 'calendar_tab') and self.calendar_tab:
                    self.calendar_tab.refresh_calendar_if_visible()
            except Exception:
                pass

        # Empty state helper
        try:
            if len(data) == 0:
                empty = self.empty_states.get(pane_name)
                if empty is None:
                    empty = QLabel(f"No {pane_name} yet — Use '+ Add' or Import to get started.")
                    empty.setStyleSheet("color:#64748b; padding:10px; font-style:italic;")
                    self.empty_states[pane_name] = empty
                if table.parentWidget() is not None:
                    table.parentWidget().layout().addWidget(empty)
                empty.setVisible(True)
            else:
                if self.empty_states.get(pane_name):
                    self.empty_states[pane_name].setVisible(False)
        except Exception as e:
            ErrorHandler.handle_ui_error("show empty state", e)

    def _get_date_from_row(self, cols: list[str], row_vals: list[str]) -> tuple[datetime | None, str | None]:
        # Returns (due_or_end_date, label)
        try:
            for key in ("Due Date", "Target Date", "End Date", "Audit End Date"):
                if key in cols:
                    idx = cols.index(key)
                    if idx < len(row_vals) and row_vals[idx]:
                        return datetime.strptime(str(row_vals[idx]), "%Y-%m-%d"), key
        except Exception:
            return None, None
        return None, None

    def check_due_date_reminders(self) -> None:
        # System tray reminders removed - notifications now handled by Bell icon
            pass

    def apply_rag_format(self, pane_name: str) -> None:
        table = self.tables[pane_name]
        cols = PANE_COLUMNS[pane_name]
        
        # Determine which column to use for color coding based on pane
        color_col = None
        color_type = None
        
        if pane_name in ["Potential Issues", "Client Visits / Audits", "Activities"]:
            # Look for RAG status column
            for idx, name in enumerate(cols):
                if "RAG" in name:
                    color_col = idx
                    color_type = "rag_status"
                    break
        elif pane_name == "Project Details":
            # Look for complexity column
            for idx, name in enumerate(cols):
                if "Complexity" in name:
                    color_col = idx
                    color_type = "complexity"
                    break
        elif pane_name == "Initiatives":
            # Look for priority column
            for idx, name in enumerate(cols):
                if "Priority" in name:
                    color_col = idx
                    color_type = "priority"
                    break
        
        if color_col is None:
            return
            
        for r in range(table.rowCount()):
            item = table.item(r, color_col)
            if item is None:
                continue
            val = item.text().strip().lower()
            base = None
            text_color = QColor("#000000")
            
            if color_type == "rag_status":
                # RAG Status: Red -> Red, Amber -> Yellow, Green -> Green
                if val == "red":
                    base = QColor("#ffebee")  # Light red background
                    text_color = QColor("#c62828")  # Dark red text
                elif val == "amber":
                    base = QColor("#fff8e1")  # Light yellow background
                    text_color = QColor("#f57f17")  # Dark yellow text
                elif val == "green":
                    base = QColor("#e8f5e8")  # Light green background
                    text_color = QColor("#2e7d32")  # Dark green text
                    
            elif color_type == "complexity":
                # Complexity: High -> Red, Medium -> Yellow, Low -> Green
                if val == "high":
                    base = QColor("#ffebee")  # Light red background
                    text_color = QColor("#c62828")  # Dark red text
                elif val == "medium":
                    base = QColor("#fff8e1")  # Light yellow background
                    text_color = QColor("#f57f17")  # Dark yellow text
                elif val == "low":
                    base = QColor("#e8f5e8")  # Light green background
                    text_color = QColor("#2e7d32")  # Dark green text
                    
            elif color_type == "priority":
                # Priority: Critical/High -> Red, Medium -> Yellow, Low -> Green
                if val in ["critical", "high"]:
                    base = QColor("#ffebee")  # Light red background
                    text_color = QColor("#c62828")  # Dark red text
                elif val == "medium":
                    base = QColor("#fff8e1")  # Light yellow background
                    text_color = QColor("#f57f17")  # Dark yellow text
                elif val == "low":
                    base = QColor("#e8f5e8")  # Light green background
                    text_color = QColor("#2e7d32")  # Dark green text
            
            if base:
                for c in range(table.columnCount()):
                    cell = table.item(r, c)
                    if cell:
                        cell.setBackground(base)
                        cell.setForeground(text_color)

    # Filtering
    def filter_table(self, pane_name: str, text: str) -> None:
        table = self.tables[pane_name]
        query = (text or "").lower()
        for r in range(table.rowCount()):
            visible = False
            for c in range(table.columnCount()):
                item = table.item(r, c)
                if item and query in item.text().lower():
                    visible = True
                    break
            table.setRowHidden(r, not visible)

    def filter_pane_rows(self, pane_name: str, status_sel: str, rag_sel: str, text: str) -> None:
        table = self.tables[pane_name]
        cols = PANE_COLUMNS[pane_name]
        sidx = cols.index("Status") if "Status" in cols else None
        ridx = next((i for i, c in enumerate(cols) if "RAG" in c), None)
        text_q = (text or "").lower()
        for r in range(table.rowCount()):
            match_text = False if text_q else True
            if text_q:
                for c in range(table.columnCount()):
                    it = table.item(r, c)
                    if it and text_q in it.text().lower():
                        match_text = True
                        break
            match_status = True
            if sidx is not None and status_sel != "All":
                it = table.item(r, sidx)
                match_status = it and it.text() == status_sel
            match_rag = True
            if ridx is not None and rag_sel != "All":
                it = table.item(r, ridx)
                match_rag = it and it.text() == rag_sel
            # Quick filter: Has new updates (for Potential Issues)
            has_updates_ok = True
            if pane_name == "Potential Issues" and "Updates" in cols and text_q == "has:new-updates":
                uidx = cols.index("Updates")
                it = table.item(r, uidx)
                has_updates_ok = bool(it and it.text().strip())
                match_text = True  # override text matching for this special token
            table.setRowHidden(r, not (match_text and match_status and match_rag and has_updates_ok))

        # Inline validation: highlight invalid dates
        try:
            date_cols = [i for i, c in enumerate(cols) if "Date" in c]
            for r in range(table.rowCount()):
                for dc in date_cols:
                    it = table.item(r, dc)
                    if not it:
                        continue
                    val = it.text().strip()
                    if not val:
                        continue
                    ok = True
                    try:
                        datetime.strptime(val, "%Y-%m-%d")
                    except Exception:
                        ok = False
                    if not ok:
                        it.setBackground(QColor("#fee2e2"))
        except Exception as e:
            ErrorHandler.handle_ui_error("validate dates", e)

        # Persist column order and hidden columns after filtering (cheap safeguards)
        try:
            self._save_column_order(pane_name, table)
            self._persist_hidden_columns(pane_name, table)
        except Exception as e:
            ErrorHandler.handle_ui_error("persist column settings", e)

    # Dashboard
    def update_dashboard(self) -> None:
        # Clear cards
        while self.cards_bar.count():
            w = self.cards_bar.takeAt(0).widget()
            if w:
                w.setParent(None)
        selected_ism = self.ism_filter.currentText() if hasattr(self, 'ism_filter') else "All ISMs"
        def row_matches_ism(pane_name: str, row: list[str]) -> bool:
            if not selected_ism or selected_ism == "All ISMs":
                return True
            cols = PANE_COLUMNS[pane_name]
            if "ISM Name" in cols:
                idx = cols.index("ISM Name")
                return idx < len(row) and row[idx] == selected_ism
            return False
        # Only count rows that have at least one non-empty cell
        def is_non_empty_row(row):
            return any(str(cell).strip() for cell in row if cell is not None)
        
        # Helper function to get RAG status for a row based on pane type
        def get_rag_status(pane_name: str, row: list[str]) -> str:
            cols = PANE_COLUMNS[pane_name]
            if pane_name == "Initiatives":
                # For Initiatives, use Priority field as RAG status
                priority_idx = cols.index("Priority") if "Priority" in cols else None
                if priority_idx is not None and priority_idx < len(row):
                    priority = str(row[priority_idx]).strip()
                    # Map priority values to RAG colors
                    if priority in ["Critical", "High"]:
                        return "Red"
                    elif priority == "Medium":
                        return "Amber"
                    elif priority == "Low":
                        return "Green"
            else:
                # For other panes, look for RAG field
                rag_idx = next((i for i, c in enumerate(cols) if "RAG" in c), None)
                if rag_idx is not None and rag_idx < len(row):
                    return str(row[rag_idx]).strip()
            return ""
        
        # Only count tasks from Potential Issues and Activities panes
        def include_in_totals(pane: str) -> bool:
            return pane in ["Potential Issues", "Activities"]

        total = sum(sum(1 for row in v if include_in_totals(p) and row_matches_ism(p, row) and is_non_empty_row(row)) for p, v in self.data.items())
        red = sum(sum(1 for row in v if include_in_totals(p) and row_matches_ism(p, row) and is_non_empty_row(row) and get_rag_status(p, row) == "Red") for p, v in self.data.items())
        yellow = sum(sum(1 for row in v if include_in_totals(p) and row_matches_ism(p, row) and is_non_empty_row(row) and get_rag_status(p, row) == "Amber") for p, v in self.data.items())
        green = sum(sum(1 for row in v if include_in_totals(p) and row_matches_ism(p, row) and is_non_empty_row(row) and get_rag_status(p, row) == "Green") for p, v in self.data.items())
        for title, value, color, rag_key in [
            ("Total Tasks", total, self.accent, None),
            ("Red Tasks", red, "#e74c3c", "Red"),
            ("Amber Tasks", yellow, "#f1c40f", "Amber"),
            ("Green Tasks", green, "#2ecc71", "Green"),
        ]:
            btn = QPushButton()
            btn.setFlat(True)
            btn.setStyleSheet(
                "QPushButton { border-radius: 12px; background-color: %s; color: white; padding: 14px; text-align:left; }" % color
            )
            shadow = QGraphicsDropShadowEffect()
            shadow.setBlurRadius(16)
            shadow.setOffset(0, 4)
            shadow.setColor(QColor(0, 0, 0, 60))
            btn.setGraphicsEffect(shadow)
            # Compose content
            w = QWidget()
            lay = QVBoxLayout(w)
            lay.setContentsMargins(8, 6, 8, 6)
            t = QLabel(title)
            t.setStyleSheet("font-weight:700; letter-spacing:0.3px; font-size:13px;")
            val_lbl = QLabel(str(value))
            val_font = QFont()
            val_font.setPointSize(26)
            val_font.setWeight(QFont.Weight.Bold)
            val_lbl.setFont(val_font)
            lay.addWidget(t)
            lay.addWidget(val_lbl)
            btn.setLayout(lay)
            btn.setMinimumSize(220, 110)
            if rag_key:
                btn.clicked.connect(lambda _=False, rag=rag_key: self.open_rag_report(rag))
                btn.setToolTip(f"Click to view tasks with {rag_key} status")
            else:
                btn.setDisabled(True)
                btn.setToolTip("Total of all tasks")
            self.cards_bar.addWidget(btn)

        # Add Leads Attention card (clickable)
        # Add Critical Initiatives (In Progress) card
        crit_inprog = 0
        try:
            pane = "Initiatives"
            if pane in self.data:
                cols_i = PANE_COLUMNS[pane]
                pr_idx = cols_i.index("Priority") if "Priority" in cols_i else None
                st_idx = cols_i.index("Status") if "Status" in cols_i else None
                ism_idx = cols_i.index("ISM Name") if "ISM Name" in cols_i else None
                for r in self.data[pane]:
                    if not is_non_empty_row(r):
                        continue
                    if not row_matches_ism(pane, r):
                        continue
                    pr = (r[pr_idx] if pr_idx is not None and pr_idx < len(r) else "").strip()
                    st = (r[st_idx] if st_idx is not None and st_idx < len(r) else "").strip()
                    if pr.lower() == "critical" and st.lower() == "in progress":
                        crit_inprog += 1
        except Exception:
            crit_inprog = 0

        crit_btn = QPushButton()
        crit_btn.setFlat(True)
        crit_btn.setStyleSheet("QPushButton { border-radius: 12px; background-color: #e11d48; color: white; padding: 14px; text-align:left; }")
        crit_container = QWidget(); crit_lay = QVBoxLayout(crit_container); crit_lay.setContentsMargins(8,6,8,6)
        crit_t = QLabel("Critical Initiatives (In Progress)"); crit_t.setStyleSheet("font-weight:700; letter-spacing:0.3px; font-size:13px;")
        crit_v = QLabel(str(crit_inprog)); cf = QFont(); cf.setPointSize(26); cf.setWeight(QFont.Weight.Bold); crit_v.setFont(cf)
        crit_lay.addWidget(crit_t); crit_lay.addWidget(crit_v)
        crit_btn.setLayout(crit_lay); crit_btn.setMinimumSize(220,110)
        try:
            crit_btn.clicked.connect(self.open_critical_initiatives_report)
            crit_btn.setToolTip("Click to view Critical + In Progress initiatives")
        except Exception as e:
            ErrorHandler.handle_ui_error("setup critical initiatives button", e)
        self.cards_bar.addWidget(crit_btn)

        leads_count = 0
        
        # Count from Potential Issues with "Leads Attention Required" = "Yes"
        if "Potential Issues" in self.data:
            cols_pi = PANE_COLUMNS["Potential Issues"]
            if "Leads Attention Required" in cols_pi:
                att_idx = cols_pi.index("Leads Attention Required")
                leads_count += sum(1 for r in self.data["Potential Issues"] if att_idx < len(r) and str(r[att_idx]).strip().lower() == "yes" and (self.ism_filter.currentText() == "All ISMs" or ("ISM Name" in PANE_COLUMNS["Potential Issues"] and r[PANE_COLUMNS["Potential Issues"].index("ISM Name")] == self.ism_filter.currentText())))
        
        # Count from Activities with "Support Required" = "Yes"
        if "Activities" in self.data:
            cols_act = PANE_COLUMNS["Activities"]
            if "Support Required" in cols_act:
                sup_idx = cols_act.index("Support Required")
                leads_count += sum(1 for r in self.data["Activities"] if sup_idx < len(r) and str(r[sup_idx]).strip().lower() == "yes" and (self.ism_filter.currentText() == "All ISMs" or ("ISM Name" in PANE_COLUMNS["Activities"] and r[PANE_COLUMNS["Activities"].index("ISM Name")] == self.ism_filter.currentText())))
        leads_btn = QPushButton()
        leads_btn.setFlat(True)
        leads_btn.setStyleSheet("QPushButton { border-radius: 12px; background-color: #8e44ad; color: white; padding: 14px; text-align:left; }")
        leads_container = QWidget()
        leads_lay = QVBoxLayout(leads_container)
        leads_lay.setContentsMargins(8, 6, 8, 6)
        leads_t = QLabel("Leads Attention")
        leads_t.setStyleSheet("font-weight:700; letter-spacing:0.3px; font-size:13px;")
        leads_v = QLabel(str(leads_count))
        lf = QFont(); lf.setPointSize(26); lf.setWeight(QFont.Weight.Bold)
        leads_v.setFont(lf)
        leads_pb = QProgressBar()
        leads_pb.setMaximum(max(1, total))
        leads_pb.setValue(leads_count)
        leads_pb.setTextVisible(False)
        leads_pb.setStyleSheet("QProgressBar { background-color:#ffffff33; border-radius:6px; height:8px; } QProgressBar::chunk { background-color:#ffffff; border-radius:6px; }")
        leads_lay.addWidget(leads_t)
        leads_lay.addWidget(leads_v)
        leads_lay.addWidget(leads_pb)
        leads_btn.setLayout(leads_lay)
        leads_btn.clicked.connect(self.open_leads_report)
        leads_btn.setToolTip("Click to view items requiring leads attention (Potential Issues + Activities)")
        leads_btn.setMinimumSize(220, 110)
        self.cards_bar.addWidget(leads_btn)

        # Chart
        for i in reversed(range(self.chart_layout.count())):
            w = self.chart_layout.itemAt(i).widget()
            if w:
                w.setParent(None)
        # Replace pie with stacked bar of R/Y/G per pane for better readability
        for i in reversed(range(self.chart_layout.count())):
            w = self.chart_layout.itemAt(i).widget()
            if w:
                w.setParent(None)
        # Exclude Leave Tracker, Project Details, and Accolades from RAG BY pane graph
        panes = [pane for pane in self.data.keys() if pane not in ["Leave Tracker", "Project Details", "Accolades"]]
        rag_by_pane = []
        def row_matches_ism(pane_name: str, row: list[str]) -> bool:
            sel = self.ism_filter.currentText() if hasattr(self, 'ism_filter') else "All ISMs"
            if not sel or sel == "All ISMs":
                return True
            cols = PANE_COLUMNS[pane_name]
            if "ISM Name" in cols:
                idx = cols.index("ISM Name")
                return idx < len(row) and row[idx] == sel
            return False
        # Helper function to get RAG status for chart calculation
        def get_rag_status_chart(pane_name: str, row: list[str]) -> str:
            cols = PANE_COLUMNS[pane_name]
            if pane_name == "Initiatives":
                # For Initiatives, use Priority field as RAG status
                priority_idx = cols.index("Priority") if "Priority" in cols else None
                if priority_idx is not None and priority_idx < len(row):
                    priority = str(row[priority_idx]).strip()
                    # Map priority values to RAG colors
                    if priority in ["Critical", "High"]:
                        return "Red"
                    elif priority == "Medium":
                        return "Amber"
                    elif priority == "Low":
                        return "Green"
            else:
                # For other panes, look for RAG field
                rag_idx = next((i for i, c in enumerate(cols) if "RAG" in c), None)
                if rag_idx is not None and rag_idx < len(row):
                    return str(row[rag_idx]).strip()
            return ""
        
        for pane in panes:
            rc = sum(1 for row in self.data[pane] if row_matches_ism(pane, row) and is_non_empty_row(row) and get_rag_status_chart(pane, row) == "Red")
            yc = sum(1 for row in self.data[pane] if row_matches_ism(pane, row) and is_non_empty_row(row) and get_rag_status_chart(pane, row) == "Amber")
            gc = sum(1 for row in self.data[pane] if row_matches_ism(pane, row) and is_non_empty_row(row) and get_rag_status_chart(pane, row) == "Green")
            rag_by_pane.append((rc, yc, gc))
        # Ensure logs are loaded from backend so Activity/Change logs persist across sessions
        try:
            if hasattr(self, '_load_logs_from_sqlite'):
                self._load_logs_from_sqlite()
        except Exception as e:
            ErrorHandler.handle_ui_error("load logs from backend", e)
        fig, ax = plt.subplots(figsize=(7, 3))
        fig.patch.set_facecolor('white')
        ax.set_facecolor('white')
        x = range(len(panes))
        reds = [v[0] for v in rag_by_pane]
        yellows = [v[1] for v in rag_by_pane]
        greens = [v[2] for v in rag_by_pane]
        bars_g = ax.bar(x, greens, color="#2ecc71", label="Green")
        bars_y = ax.bar(x, yellows, bottom=greens, color="#f1c40f", label="Amber")
        bottom2 = [g + y for g, y in zip(greens, yellows)]
        bars_r = ax.bar(x, reds, bottom=bottom2, color="#e74c3c", label="Red")
        ax.set_xticks(list(x))
        ax.set_xticklabels([p[:12] for p in panes], rotation=0)
        ax.set_ylabel("Tasks (count)")
        ax.set_title("RAG by Pane")
        ax.legend(loc="upper right")
        # Ensure y-axis shows whole numbers only
        ax.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
        # Add value labels
        for bars in (bars_g, bars_y, bars_r):
            for b in bars:
                height = b.get_height()
                if height > 0:
                    ax.annotate(f"{int(height)}",
                                xy=(b.get_x() + b.get_width() / 2, b.get_y() + height/2),
                                ha='center', va='center', color='black', fontsize=8)
        canvas1 = FigureCanvas(fig)
        # Ensure canvas is created in main thread
        self._ensure_main_thread(lambda: self.chart_layout.addWidget(canvas1))
        # Hover tooltip for RAG by Pane
        annot1 = ax.annotate("", xy=(0,0), xytext=(10,10), textcoords="offset points",
                             bbox=dict(boxstyle="round", fc="w"), arrowprops=dict(arrowstyle="->"))
        annot1.set_visible(False)
        all_bars = list(bars_g) + list(bars_y) + list(bars_r)
        def hover1(event):
            vis = annot1.get_visible()
            if event.inaxes == ax:
                for b in all_bars:
                    cont, _ = b.contains(event)
                    if cont:
                        val = int(b.get_height())
                        annot1.xy = (b.get_x() + b.get_width()/2, b.get_y() + b.get_height())
                        annot1.set_text(str(val))
                        annot1.set_visible(True)
                        canvas1.draw_idle()
                        return
            if vis:
                annot1.set_visible(False)
                canvas1.draw_idle()
        canvas1.mpl_connect('motion_notify_event', hover1)
        # Close pyplot figure to avoid accumulation warnings; canvas keeps its own reference
        try:
            plt.close(fig)
        except Exception as e:
            ErrorHandler.handle_ui_error("close matplotlib figure", e)
        # Status bar chart
        for i in reversed(range(self.chart_layout2.count())):
            w = self.chart_layout2.itemAt(i).widget()
            if w:
                w.setParent(None)
        # Exclude "NA" from status overview graph
        status_keys = [k for k in STATUS_OPTIONS if k != "NA"]
        status_counts = {k: 0 for k in status_keys}
        for pane, rows in self.data.items():
            cols = PANE_COLUMNS[pane]
            if "Status" in cols:
                sidx = cols.index("Status")
                for r in rows:
                    if not row_matches_ism(pane, r):
                        continue
                    sval = r[sidx] if sidx < len(r) else "NA"
                    if sval in status_counts:
                        status_counts[sval] += 1
        fig2, ax2 = plt.subplots(figsize=(5, 3))
        fig2.patch.set_facecolor('white')
        ax2.set_facecolor('white')
        bars = ax2.bar(list(status_counts.keys()), list(status_counts.values()), color=[self.accent, "#f1c40f", "#2ecc71", "#95a5a6"])
        ax2.set_ylabel("Tasks (count)")
        ax2.set_title("Status Overview")
        # Ensure y-axis shows whole numbers only
        ax2.yaxis.set_major_locator(plt.MaxNLocator(integer=True))
        for b in bars:
            h = b.get_height()
            ax2.annotate(f"{int(h)}", xy=(b.get_x()+b.get_width()/2, h), xytext=(0,3), textcoords='offset points', ha='center', va='bottom')
        canvas2 = FigureCanvas(fig2)
        # Ensure canvas is created in main thread
        self._ensure_main_thread(lambda: self.chart_layout2.addWidget(canvas2))
        annot2 = ax2.annotate("", xy=(0,0), xytext=(0,10), textcoords="offset points",
                              ha='center', bbox=dict(boxstyle="round", fc="w"))
        annot2.set_visible(False)
        def hover2(event):
            vis = annot2.get_visible()
            if event.inaxes == ax2:
                for b in bars:
                    cont, _ = b.contains(event)
                    if cont:
                        val = int(b.get_height())
                        annot2.xy = (b.get_x() + b.get_width()/2, b.get_height())
                        annot2.set_text(str(val))
                        annot2.set_visible(True)
                        canvas2.draw_idle()
                        return
            if vis:
                annot2.set_visible(False)
                canvas2.draw_idle()
        canvas2.mpl_connect('motion_notify_event', hover2)
        try:
            plt.close(fig2)
        except Exception as e:
            ErrorHandler.handle_ui_error("close matplotlib figure 2", e)
        self.refresh_ism_filter()
        # Overdue table disabled per latest feedback
        # Recent activity
        # self.refresh_activity()  # moved to menu on-demand
        # Client Visits & Audits consolidated
        self.populate_client_visits_sections()
        # Leads attention section
        self.populate_leads_attention()

    def populate_overdue_table(self) -> None:
        # Deprecated per latest feedback
        pass

    def refresh_activity(self) -> None:
        pass

    def show_activity_dialog(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Recent Activity & Audit Log")
        dlg.resize(800, 600)
        lay = QVBoxLayout(dlg)
        
        # Create tab widget for different views
        tab_widget = QTabWidget()
        
        # Simple log tab
        simple_tab = QWidget()
        simple_layout = QVBoxLayout(simple_tab)
        simple_lst = QListWidget()
        for entry in self.change_log[-50:][::-1]:
            simple_lst.addItem(entry)
        simple_layout.addWidget(simple_lst)
        tab_widget.addTab(simple_tab, "Simple Log")
        
        # Detailed audit log tab
        detailed_tab = QWidget()
        detailed_layout = QVBoxLayout(detailed_tab)
        
        # Create table for detailed logs
        table = QTableWidget()
        table.setColumnCount(6)
        table.setHorizontalHeaderLabels(["Timestamp", "User", "Action", "Pane", "Details", "Level"])
        table.horizontalHeader().setStretchLastSection(True)
        table.setAlternatingRowColors(True)
        
        # Populate detailed logs
        recent_logs = self.change_log_data[-100:][::-1]  # Last 100 entries, most recent first
        table.setRowCount(len(recent_logs))
        
        for row, log_entry in enumerate(recent_logs):
            if len(log_entry) >= 6:
                for col, value in enumerate(log_entry[:6]):
                    item = QTableWidgetItem(str(value))
                    # Color code by level
                    if col == 5:  # Level column
                        if value == "ERROR":
                            item.setBackground(QColor("#ffebee"))
                        elif value == "WARNING":
                            item.setBackground(QColor("#fff3e0"))
                        elif value == "SUCCESS":
                            item.setBackground(QColor("#e8f5e8"))
                    table.setItem(row, col, item)
        
        detailed_layout.addWidget(table)
        tab_widget.addTab(detailed_tab, "Detailed Audit Log")
        
        lay.addWidget(tab_widget)
        
        # Close button
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        btn = QPushButton("Close")
        btn.clicked.connect(dlg.accept)
        btn_layout.addWidget(btn)
        lay.addLayout(btn_layout)
        
        dlg.exec()

    def show_help_dialog(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Help & Shortcuts")
        lay = QVBoxLayout(dlg)
        txt = QLabel(
            "<b>Keyboard</b><br>"
            "Enter: Edit cell<br>Esc: Cancel selection<br>Ctrl+Z / Ctrl+Y: Undo/Redo<br>"
            "Alt+1..9: Switch tabs<br>Ctrl+K: Global Search" )
        txt.setTextFormat(Qt.TextFormat.RichText)
        txt.setStyleSheet("color:#0f172a; line-height:1.4;")
        lay.addWidget(txt)
        btn = QPushButton("Got it")
        btn.setObjectName("primary")
        btn.clicked.connect(dlg.accept)
        lay.addWidget(btn, alignment=Qt.AlignmentFlag.AlignRight)
        dlg.resize(420, 220)
        dlg.exec()

    def open_global_search(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Global Search")
        lay = QVBoxLayout(dlg)
        top = QHBoxLayout()
        q = QLineEdit(); q.setPlaceholderText("Search across all panes...")
        pane_filter = QComboBox(); pane_filter.addItem("All Panes"); [pane_filter.addItem(p) for p in self.tables.keys()]
        top.addWidget(QLabel("Pane:")); top.addWidget(pane_filter); top.addWidget(q)
        results = QListWidget()
        lay.addLayout(top); lay.addWidget(results)
        def run_search(text: str):
            text_l = (text or "").lower().strip()
            results.clear()
            if not text_l:
                return
            panes = list(self.tables.keys()) if pane_filter.currentText() == "All Panes" else [pane_filter.currentText()]
            for pane in panes:
                table = self.tables[pane]
                for r in range(table.rowCount()):
                    for c in range(table.columnCount()):
                        it = table.item(r, c)
                        if it and text_l in it.text().lower():
                            results.addItem(f"{pane} — r{r+1} c{c+1}: {it.text()[:80]}")
                            break
        q.textChanged.connect(run_search)
        def open_location():
            it = results.currentItem()
            if not it:
                return
            try:
                txt = it.text()
                pane = txt.split(" — ")[0]
                meta = txt.split(" — ")[1].split(":")[0]
                r = int(meta.split(" ")[0].replace("r", "")) - 1
                c = int(meta.split(" ")[1].replace("c", "")) - 1
                # Focus pane and select cell
                idx = None
                for i in range(self.tabs.count()):
                    if self.tabs.tabText(i) == pane:
                        idx = i; break
                if idx is not None:
                    self.tabs.setCurrentIndex(idx)
                    tbl = self.tables[pane]
                    tbl.setCurrentCell(r, c)
                    tbl.scrollToItem(tbl.item(r, c), QAbstractItemView.ScrollHint.PositionAtCenter)
            except Exception as e:
                ErrorHandler.handle_ui_error("open search result", e)
        results.itemDoubleClicked.connect(lambda _=None: open_location())
        close = QPushButton("Close"); close.clicked.connect(dlg.accept)
        lay.addWidget(close)
        dlg.resize(700, 400)
        dlg.exec()

    def quick_filter_status(self, label: str) -> None:
        # Operate on current tab table
        idx = self.tabs.currentIndex()
        # 0 is Dashboard
        if idx <= 0:
            return
        pane_name = self.tabs.tabText(idx)
        table = self.tables.get(pane_name)
        if not table:
            return
        cols = PANE_COLUMNS[pane_name]
        if "Status" not in cols:
            return
        sidx = cols.index("Status")
        for r in range(table.rowCount()):
            item = table.item(r, sidx)
            val = item.text() if item else ""
            visible = (label == "All") or (val == label)
            table.setRowHidden(r, not visible)
    def _is_change_log_event(self, action: str | None, pane: str | None, details: str | None) -> bool:
        """Classify whether an event should be recorded in Change Log.
        Criteria: row addition/removal/modification and org chart member activities.
        """
        try:
            txt = " ".join([str(action or ""), str(pane or ""), str(details or "")]).lower()
            keywords = [
                "add row", "added row", "insert row", "new row", "create row",
                "delete row", "deleted row", "remove row", "removed row",
                "update cell", "updated cell", "edit cell", "edited cell", "cell change", "edited",
                # org chart activities
                "org", "org chart", "org directory", "add member", "added member", "remove member", "removed member",
                "rename", "edit member", "move member", "reparent"
            ]
            return any(k in txt for k in keywords)
        except Exception:
            return False

    def _log(self, message: str, level: str = "INFO", show_toast: bool = True, pane: str = None, action: str = None) -> None:
        """Enhanced logging with different levels, user info, pane details, and optional toast notifications"""
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        user = getattr(self, 'logged_in_user', 'Unknown')
        
        # Build detailed message with user and pane information
        details = []
        if user != 'Unknown':
            details.append(f"User: {user}")
        if pane:
            details.append(f"Pane: {pane}")
        if action:
            details.append(f"Action: {action}")
        
        if details:
            detailed_message = f"{message} | {' | '.join(details)}"
        else:
            detailed_message = message
            
        formatted_message = f"[{ts}] {LEVEL_EMOJI.get(level, '[i]')} {detailed_message}"
        self.change_log.append(formatted_message)
        
        # Ensure datasets exist
        if not hasattr(self, 'activity_log_data'):
            self.activity_log_data = []
        if not hasattr(self, 'change_log_data'):
            self.change_log_data = []
        # Route to Activity Logs always
        self.activity_log_data.append([ts, user, action or "General", pane or "System", message, level])
        # Route to Change Log only for targeted events (row/org changes)
        if self._is_change_log_event(action, pane, message):
            self.change_log_data.append([ts, user, action or "General", pane or "System", message, level])
        
        # Keep only last 1000 entries to prevent memory issues
        if len(self.change_log_data) > 1000:
            self.change_log_data = self.change_log_data[-1000:]
        
        if show_toast and level in ["SUCCESS", "WARNING", "ERROR"]:
            self._show_toast(message, level=level)
        elif show_toast and level == "INFO":
            self._show_toast(message)

    def _log_change(self, action: str, pane: str, details: str, level: str = "INFO") -> None:
        """Enhanced detailed logging with different levels for tracking and audit purposes"""
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        user = getattr(self, 'logged_in_user', 'Unknown')
        
        # Create comprehensive log entry with all details
        self.change_log_data.append([ts, user, action, pane, details, level])
        
        # Keep only last 1000 entries to prevent memory issues
        if len(self.change_log_data) > 1000:
            self.change_log_data = self.change_log_data[-1000:]
        
        # Save logs to Excel backend
        self._save_logs_to_csv()
        
        # Also add to Activity Logs for comprehensive tracking
        try:
            enhanced_message = f"{action} in {pane}: {details} | User: {user}"
            self._log(enhanced_message, level=level, show_toast=False, pane=pane, action=action)
        except Exception:
            pass

    def get_filtered_logs(self, level: str = None, pane: str = None, action: str = None) -> list:
        """Get filtered logs based on criteria"""
        if not hasattr(self, 'change_log_data'):
            return []
        
        filtered_logs = self.change_log_data.copy()
        
        if level:
            filtered_logs = [log for log in filtered_logs if len(log) > 5 and log[5] == level]
        if pane:
            filtered_logs = [log for log in filtered_logs if log[3] == pane]
        if action:
            filtered_logs = [log for log in filtered_logs if log[2] == action]
        
        return filtered_logs

    def get_log_summary(self) -> dict:
        """Get summary statistics of logs"""
        if not hasattr(self, 'change_log_data'):
            return {}
        
        summary = {
            "total_entries": len(self.change_log_data),
            "by_level": {},
            "by_pane": {},
            "by_action": {},
            "recent_activity": self.change_log_data[-10:] if self.change_log_data else []
        }
        
        for log in self.change_log_data:
            if len(log) > 5:
                level = log[5]
                pane = log[3]
                action = log[2]
                
                summary["by_level"][level] = summary["by_level"].get(level, 0) + 1
                summary["by_pane"][pane] = summary["by_pane"].get(pane, 0) + 1
                summary["by_action"][action] = summary["by_action"].get(action, 0) + 1
        
        return summary

    def _save_logs_to_csv(self) -> None:
        """Save logs to Excel backend file with enhanced error handling and auto-creation"""
        try:
            backend_path = getattr(self, '__backend_path__', None) or self.backend_sqlite_path
            if not backend_path:
                # Create default backend path if none exists
                backend_path = os.path.join(os.getcwd(), "tracker_backend.xlsx")
                self.__backend_path__ = backend_path
                self.backend_sqlite_path = backend_path
            # If SQLite backend, save logs into SQLite tables and return
            if str(backend_path).lower().endswith((".sqlite", ".db")):
                try:
                    with self._sqlite_connect(backend_path) as conn:
                        cur = conn.cursor()
                        # Ensure tables
                        cur.execute("CREATE TABLE IF NOT EXISTS activity_logs (Timestamp TEXT, User TEXT, Action TEXT, Pane TEXT, Details TEXT, Level TEXT)")
                        cur.execute("CREATE TABLE IF NOT EXISTS change_log (Timestamp TEXT, User TEXT, Action TEXT, Pane TEXT, Details TEXT, Level TEXT)")
                        # Replace content
                        cur.execute("DELETE FROM activity_logs")
                        cur.execute("DELETE FROM change_log")
                        cols6 = ["Timestamp","User","Action","Pane","Details","Level"]
                        activity_rows = getattr(self, 'activity_log_data', []) or []
                        change_rows = getattr(self, 'change_log_data', []) or []
                        if activity_rows:
                            cur.executemany(f"INSERT INTO activity_logs({', '.join(cols6)}) VALUES (?,?,?,?,?,?)", activity_rows)
                        if change_rows:
                            cur.executemany(f"INSERT INTO change_log({', '.join(cols6)}) VALUES (?,?,?,?,?,?)", change_rows)
                    return
                except Exception as se:
                    ErrorHandler.handle_ui_error("save logs to sqlite", se)
                    return
            
            # Ensure the backend file exists
            # Backend file is now SQLite, no need to ensure Excel exists
            
            # Try to load the Excel file; if locked, wait/retry instead of recreating
            from openpyxl import load_workbook
            import time, zipfile
            retries = 0
            max_retries = 12  # ~30s at 2.5s per iteration below
            while True:
                try:
                    wb = load_workbook(backend_path)
                    break
                except Exception as load_error:
                    msg = str(load_error).lower()
                    locked = isinstance(load_error, PermissionError) or ("permission denied" in msg) or ("in use" in msg) or ("being used by another process" in msg) or isinstance(load_error, zipfile.BadZipFile) or ("file is not a zip" in msg)
                    if locked and retries < max_retries:
                        time.sleep(2.5)
                        retries += 1
                        continue
                    # If not locked/transient or exceeded retries, only then create new file
                    if not os.path.exists(backend_path):
                        self._create_new_backend_file(backend_path)
                        wb = load_workbook(backend_path)
                        break
                    raise
            
            # Create or get the logging worksheet
            # Ensure both Activity Logs and Change Log sheets exist with headers
            for sheet_name in ('Activity Logs', 'Change Log'):
                if sheet_name not in wb.sheetnames:
                    ws = wb.create_sheet(sheet_name)
                    headers = ["Timestamp", "User", "Action", "Pane", "Details", "Level"]
                    for col, header in enumerate(headers, 1):
                        cell = ws.cell(row=1, column=col, value=header)
                        from openpyxl.styles import Font
                        cell.font = Font(bold=True)
                else:
                    ws = wb[sheet_name]
            
            # Write activity and change logs to respective sheets
            datasets = {
                'Activity Logs': getattr(self, 'activity_log_data', []) or [],
                'Change Log': getattr(self, 'change_log_data', []) or [],
            }
            for sheet_name, rows in datasets.items():
                ws = wb[sheet_name]
                if ws.max_row > 1:
                    ws.delete_rows(2, ws.max_row)
                for row_idx, log_entry in enumerate(rows, 2):
                    if len(log_entry) >= 6:
                        for col_idx, value in enumerate(log_entry[:6], 1):
                            ws.cell(row=row_idx, column=col_idx, value=value)
            
            # Auto-fit columns for both sheets
            for sheet_name in ('Activity Logs', 'Change Log'):
                ws = wb[sheet_name]
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except Exception as e:
                            logger.warning(f"Failed to calculate column width for {sheet_name} {column_letter}: {e}")
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width
            
            wb.save(backend_path)
            
        except Exception as e:
            ErrorHandler.handle_ui_error("Save logs to SQLite", e, {"backend_path": backend_path})
            # Try error recovery
            if self._recover_from_error("Excel Save logs to Excel", e, {"backend_path": backend_path}):
                # Retry the operation after recovery
                try:
                    self._save_logs_to_csv()
                    ErrorHandler.log_operation("Retry save logs after recovery", True)
                except Exception as retry_error:
                    ErrorHandler.handle_ui_error("Retry save logs after recovery", retry_error)
            else:
                ErrorHandler.log_operation("Excel error recovery", False, "Recovery failed")


    def _load_logs_from_sqlite(self) -> None:
        """Load logs from SQLite backend file with enhanced error handling"""
        try:
            backend_path = getattr(self, '__backend_path__', None) or self.backend_sqlite_path
            if not backend_path or not os.path.exists(backend_path):
                return
            # If SQLite backend, load from SQLite instead
            if str(backend_path).lower().endswith((".sqlite", ".db")):
                try:
                    with self._sqlite_connect(backend_path) as conn:
                        cur = conn.cursor()
                        self.activity_log_data = []
                        self.change_log_data = []
                        try:
                            cur.execute("SELECT Timestamp, User, Action, Pane, Details, Level FROM activity_logs")
                            self.activity_log_data = [list(map(lambda x: '' if x is None else str(x), row)) for row in cur.fetchall()]
                        except Exception:
                            pass
                        try:
                            cur.execute("SELECT Timestamp, User, Action, Pane, Details, Level FROM change_log")
                            self.change_log_data = [list(map(lambda x: '' if x is None else str(x), row)) for row in cur.fetchall()]
                        except Exception:
                            pass
                    self._cleanup_old_logs()
                except Exception as _e:
                    print(f"Failed to load logs from SQLite: {_e}")
                return
            
            from openpyxl import load_workbook
            try:
                wb = load_workbook(backend_path)
            except Exception as load_error:
                print(f"Could not load Excel file for logs, skipping: {load_error}")
                return
            
            # Load both sheets independently
            self.activity_log_data = []
            self.change_log_data = []
            if 'Activity Logs' in wb.sheetnames:
                ws = wb['Activity Logs']
                rows = [row for row in ws.iter_rows(min_row=2, values_only=True) if row and row[0]]
                self.activity_log_data = [list(r[:6]) for r in rows]
            if 'Change Log' in wb.sheetnames:
                ws = wb['Change Log']
                rows = [row for row in ws.iter_rows(min_row=2, values_only=True) if row and row[0]]
                self.change_log_data = [list(r[:6]) for r in rows]
                
                # Clean up old logs (older than 30 days)
                self._cleanup_old_logs()
                
        except Exception as e:
            print(f"Error loading logs from Excel: {e}")
            # Initialize empty log data if loading fails
            if not hasattr(self, 'change_log_data'):
                self.change_log_data = []

    def _ensure_backend_sqlite_exists(self, backend_path: str) -> None:
        """Ensure the backend SQLite file exists and is valid; self-heal quietly."""
        try:
            # If using SQLite backend, skip Excel ensure
            if str(backend_path).lower().endswith((".sqlite", ".db")):
                return
            # Normalize to .xlsx
            if not str(backend_path).lower().endswith('.xlsx'):
                backend_path = f"{backend_path}.xlsx"
                self.__backend_path__ = backend_path
                self.backend_sqlite_path = backend_path
            # Create if missing
            if not os.path.exists(backend_path):
                # Offer the user to change path or auto-create if UI available
                try:
                    if QApplication.instance() is not None:
                        self._prompt_missing_backend_file(backend_path)
                    else:
                        # Headless fallback: auto-create
                        self._create_new_backend_file(backend_path)
                except Exception as e:
                    # If prompt fails, fallback to auto-create
                    try:
                        self._create_new_backend_file(backend_path)
                    except Exception:
                        raise e
                return
            # Validate by attempting to open; if locked, wait instead of recreating
            try:
                from openpyxl import load_workbook
                _ = load_workbook(backend_path)
            except Exception as validate_error:
                import time, zipfile
                msg = str(validate_error).lower()
                locked = isinstance(validate_error, PermissionError) or ("permission denied" in msg) or ("in use" in msg) or ("being used by another process" in msg) or isinstance(validate_error, zipfile.BadZipFile) or ("file is not a zip" in msg)
                if locked:
                    # Wait for the file to be released (up to ~60s)
                    attempts = 0
                    while attempts < 24:
                        try:
                            _ = load_workbook(backend_path)
                            return
                        except Exception as e2:
                            m2 = str(e2).lower()
                            still_locked = isinstance(e2, PermissionError) or ("permission denied" in m2) or ("in use" in m2) or ("being used by another process" in m2) or isinstance(e2, zipfile.BadZipFile) or ("file is not a zip" in m2)
                            if not still_locked:
                                break
                            time.sleep(2.5)
                            attempts += 1
                # If we reach here, treat as invalid only after retries; do not recreate on lock
                try:
                    if not getattr(self, '_backend_excel_recovered_once', False):
                        logger.warning("Backend Excel validation failed. Skipping recreation due to possible lock; will retry later.")
                        self._backend_excel_recovered_once = True
                except Exception:
                    self._backend_excel_recovered_once = True
                # Do not recreate here to avoid data loss when file is locked/writing
                return
        except Exception as e:
            print(f"Error ensuring backend Excel exists: {e}")
            # Create new file as fallback
            self._create_new_backend_file(backend_path)

    def _prompt_missing_backend_file(self, backend_path: str) -> None:
        """Prompt the user when the configured backend Excel file is missing.
        Options: Change Path, Auto-Create, or Cancel.
        """
        try:
            # If using SQLite backend, create SQLite backend silently and return
            if str(backend_path).lower().endswith((".sqlite", ".db")):
                self._create_new_backend_sqlite(backend_path)
                self.backend_sqlite_path = backend_path
                self.__backend_path__ = backend_path
                return
            msg = QMessageBox(self)
            msg.setIcon(QMessageBox.Icon.Warning)
            msg.setWindowTitle("Backend File Missing")
            msg.setText("The configured backend Excel file is missing:")
            msg.setInformativeText(str(backend_path))
            change_btn = msg.addButton("Change Path", QMessageBox.ButtonRole.ActionRole)
            create_btn = msg.addButton("Auto-Create", QMessageBox.ButtonRole.AcceptRole)
            cancel_btn = msg.addButton(QMessageBox.StandardButton.Cancel)
            msg.exec()

            clicked = msg.clickedButton()
            if clicked == change_btn:
                # Allow selecting an existing file or specifying a new one
                new_path, _ = QFileDialog.getSaveFileName(self, "Select or Create Backend Excel", os.path.dirname(backend_path) or "", "Excel (*.xlsx)")
                if not new_path:
                    # Treat as cancel
                    raise FileNotFoundError(f"Backend Excel not found and user cancelled path change: {backend_path}")
                if not new_path.lower().endswith('.xlsx'):
                    new_path = f"{new_path}.xlsx"
                # If the chosen file doesn't exist, create the structure
                if not os.path.exists(new_path):
                    self._create_new_backend_file(new_path)
                # Persist runtime path and settings
                self.backend_sqlite_path = new_path
                self.__backend_path__ = new_path
                try:
                    if PANDAS_AVAILABLE:
                        with pd.ExcelWriter(new_path, engine="openpyxl", mode='a', if_sheet_exists='replace') as writer:
                            settings_data = [
                                ["App Version", "1.0.0", "Current application version"],
                                ["Last Updated", datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "Last time the application was updated"],
                                ["Backend Path", new_path, "Path to backend Excel file"],
                            ]
                            pd.DataFrame(settings_data, columns=["Setting","Value","Description"]).to_excel(writer, sheet_name="Settings", index=False)
                except Exception:
                    pass
                return
            elif clicked == create_btn:
                self._create_new_backend_file(backend_path)
                self.backend_sqlite_path = backend_path
                self.__backend_path__ = backend_path
                return
            else:
                # Cancel
                raise FileNotFoundError(f"Backend Excel not found and user cancelled: {backend_path}")
        except Exception as e:
            # Propagate to caller to decide recovery or fallback
            raise e

    def _create_new_backend_file(self, backend_path: str) -> None:
        """Create a new backend Excel file with all required sheets and headers"""
        try:
            # Route to SQLite if chosen
            if str(backend_path).lower().endswith((".sqlite", ".db")):
                self._create_new_backend_sqlite(backend_path)
                return
            from openpyxl import Workbook
            from openpyxl.styles import Font
            
            wb = Workbook()
            
            # Remove default sheet
            wb.remove(wb.active)
            
            # Create all required sheets with headers
            sheets_config = {
                "Activity Logs": ["Timestamp", "User", "Action", "Pane", "Details", "Level"],
                "Change Log": ["Timestamp", "User", "Action", "Pane", "Details", "Level"],
                "Imports": ["Timestamp", "Pane", "Source", "Rows"],
                "Projects": ["Project Name", "Project ID"],
                "ISM Directory": ["Name"],
                "Settings": ["Setting", "Value", "Description"]
            }
            
            # Add all panes from PANE_COLUMNS
            for pane_name, columns in PANE_COLUMNS.items():
                sheet_name = self._sheet_title_for_pane(pane_name)
                sheets_config[sheet_name] = columns
            
            # Create sheets and add headers
            for sheet_name, headers in sheets_config.items():
                ws = wb.create_sheet(sheet_name)
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = Font(bold=True)
            
            # Save the file
            wb.save(backend_path)
            print(f"Created new backend Excel file: {backend_path}")
            
        except Exception as e:
            print(f"Error creating new backend file: {e}")
            raise e

    def _cleanup_old_logs(self) -> None:
        """Remove logs older than 30 days"""
        try:
            from datetime import datetime, timedelta
            cutoff_date = datetime.now() - timedelta(days=30)
            cutoff_str = cutoff_date.strftime("%Y-%m-%d")
            
            # Filter out old logs
            filtered_logs = []
            for log_entry in self.change_log_data:
                if len(log_entry) >= 1:
                    log_date_str = log_entry[0].split(' ')[0]  # Get date part
                    if log_date_str >= cutoff_str:
                        filtered_logs.append(log_entry)
            
            # Update the log data
            old_count = len(self.change_log_data)
            self.change_log_data = filtered_logs
            new_count = len(self.change_log_data)
            
            if old_count != new_count:
                print(f"Cleaned up {old_count - new_count} old log entries (older than 30 days)")
                # Save the cleaned logs back to Excel
                self._save_logs_to_csv()
                
        except Exception as e:
            print(f"Error cleaning up old logs: {e}")

    def reset_logging_data(self) -> None:
        """Reset all logging data with confirmation"""
        reply = QMessageBox.question(
            self, 
            "Reset Logging Data", 
            "Are you sure you want to delete all logging data?\n\nThis action cannot be undone.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            try:
                # Clear in-memory logs
                self.change_log = []
                self.change_log_data = []
                self.activity_log_data = []
                
                # Clear Excel logs
                backend_path = getattr(self, '__backend_path__', None) or self.backend_sqlite_path
                if backend_path and os.path.exists(backend_path):
                    # If SQLite backend, clear tables instead of using openpyxl
                    if str(backend_path).lower().endswith((".sqlite", ".db")):
                        try:
                            with self._sqlite_connect(backend_path) as conn:
                                cur = conn.cursor()
                                # Ensure tables exist
                                cur.execute("CREATE TABLE IF NOT EXISTS activity_logs (Timestamp TEXT, User TEXT, Action TEXT, Pane TEXT, Details TEXT, Level TEXT)")
                                cur.execute("CREATE TABLE IF NOT EXISTS change_log (Timestamp TEXT, User TEXT, Action TEXT, Pane TEXT, Details TEXT, Level TEXT)")
                                # Clear data
                                cur.execute("DELETE FROM activity_logs")
                                cur.execute("DELETE FROM change_log")
                        except Exception as se:
                            QMessageBox.critical(self, "Reset Error", f"Failed to reset SQLite logs:\n{str(se)}")
                            return
                        QMessageBox.information(self, "Reset Complete", "All logging data has been deleted successfully.")
                        self._log_change("Reset", "System", "All logging data cleared", "INFO")
                        return
                    from openpyxl import load_workbook
                    try:
                        wb = load_workbook(backend_path)
                    except Exception as load_error:
                        print(f"Could not load Excel file for log reset, skipping: {load_error}")
                        return
                    
                    if 'Activity Logs' in wb.sheetnames:
                        ws = wb['Activity Logs']
                        # Clear all data except headers
                        if ws.max_row > 1:
                            ws.delete_rows(2, ws.max_row)
                        wb.save(backend_path)
                
                QMessageBox.information(self, "Reset Complete", "All logging data has been deleted successfully.")
                self._log_change("Reset", "System", "All logging data cleared", "INFO")
                
            except Exception as e:
                QMessageBox.critical(self, "Reset Error", f"Failed to reset logging data:\n{str(e)}")

    def _check_log_rotation(self) -> None:
        """Check if log rotation is needed based on size and age"""
        try:
            # Check log size (max 10000 entries)
            max_entries = 10000
            if len(self.change_log_data) > max_entries:
                print(f"Log rotation needed: {len(self.change_log_data)} entries exceeds limit of {max_entries}")
                self._rotate_logs()
                return
            
            # Check log age (older than 90 days)
            from datetime import datetime, timedelta
            cutoff_date = datetime.now() - timedelta(days=90)
            cutoff_str = cutoff_date.strftime("%Y-%m-%d")
            
            old_entries = 0
            for log_entry in self.change_log_data:
                if len(log_entry) >= 1:
                    log_date_str = log_entry[0].split(' ')[0]
                    if log_date_str < cutoff_str:
                        old_entries += 1
            
            if old_entries > 1000:  # If more than 1000 old entries
                print(f"Log rotation needed: {old_entries} old entries found")
                self._rotate_logs()
                
        except Exception as e:
            print(f"Error checking log rotation: {e}")

    def _rotate_logs(self) -> None:
        """Rotate logs by archiving old entries and keeping recent ones"""
        try:
            from datetime import datetime, timedelta
            
            # Keep only last 30 days of logs
            cutoff_date = datetime.now() - timedelta(days=30)
            cutoff_str = cutoff_date.strftime("%Y-%m-%d")
            
            # Separate recent and old logs
            recent_logs = []
            old_logs = []
            
            for log_entry in self.change_log_data:
                if len(log_entry) >= 1:
                    log_date_str = log_entry[0].split(' ')[0]
                    if log_date_str >= cutoff_str:
                        recent_logs.append(log_entry)
                    else:
                        old_logs.append(log_entry)
            
            # Archive old logs if any
            if old_logs:
                self._archive_old_logs(old_logs)
            
            # Update current logs with recent ones only
            old_count = len(self.change_log_data)
            self.change_log_data = recent_logs
            new_count = len(self.change_log_data)
            
            print(f"Log rotation completed: {old_count - new_count} entries archived, {new_count} entries kept")
            
            # Save the rotated logs
            self._save_logs_to_csv()
            
        except Exception as e:
            print(f"Error rotating logs: {e}")

    def _archive_old_logs(self, old_logs: list) -> None:
        """Archive old logs to a separate file"""
        try:
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            archive_path = f"tracker_logs_archive_{timestamp}.xlsx"
            
            # Create archive file
            from openpyxl import Workbook
            from openpyxl.styles import Font
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Archived Logs"
            
            # Add headers
            headers = ["Timestamp", "User", "Action", "Pane", "Details", "Level"]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)
            
            # Add archived data
            for row_idx, log_entry in enumerate(old_logs, 2):
                if len(log_entry) >= 6:
                    for col_idx, value in enumerate(log_entry[:6], 1):
                        ws.cell(row=row_idx, column=col_idx, value=value)
            
            # Auto-fit columns
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except (AttributeError, TypeError, ValueError) as e:
                        # Handle cases where cell.value might be None or non-string
                        ErrorHandler.handle_ui_error("cell_width_calculation", e, {"cell": str(cell)})
                        continue
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            wb.save(archive_path)
            print(f"Archived {len(old_logs)} old log entries to: {archive_path}")
            
        except Exception as e:
            print(f"Error archiving old logs: {e}")

    def _log_change(self, action: str, pane: str, details: str, level: str = "INFO") -> None:
        """Enhanced log change method with automatic rotation check"""
        try:
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            user = getattr(self, 'logged_in_user', 'Unknown')
            
            log_entry = [timestamp, user, action, pane, details, level]
            self.change_log_data.append(log_entry)
            
            # Check if log rotation is needed
            self._check_log_rotation()
            
            # Save to Excel periodically (every 10 entries)
            if len(self.change_log_data) % 10 == 0:
                self._save_logs_to_csv()
                
        except Exception as e:
            print(f"Error logging change: {e}")

    def show_logs_dialog(self) -> None:
        """Show a dialog to view and manage activity logs"""
        try:
            # Ensure log data attributes exist
            if not hasattr(self, 'activity_log_data'):
                self.activity_log_data = []
            if not hasattr(self, 'change_log_data'):
                self.change_log_data = []
            
            dialog = QDialog(self)
            dialog.setWindowTitle("Activity Logs")
            dialog.setGeometry(100, 100, 1200, 800)
            
            layout = QVBoxLayout(dialog)
            
            # Header with controls
            header_layout = QHBoxLayout()
            
            # Filter controls
            filter_label = QLabel("Filter by Level:")
            header_layout.addWidget(filter_label)
            
            level_filter = QComboBox()
            level_filter.addItems(["All", "INFO", "WARNING", "ERROR", "DEBUG"])
            level_filter.currentTextChanged.connect(lambda: self._filter_logs_table(logs_table, level_filter.currentText()))
            header_layout.addWidget(level_filter)
            
            pane_filter_label = QLabel("Filter by Pane:")
            header_layout.addWidget(pane_filter_label)
            
            pane_filter = QComboBox()
            pane_filter.addItems(["All"] + list(PANE_COLUMNS.keys()) + ["System"])
            pane_filter.currentTextChanged.connect(lambda: self._filter_logs_table(logs_table, level_filter.currentText(), pane_filter.currentText()))
            header_layout.addWidget(pane_filter)
            
            header_layout.addStretch()
            
            # View toggle: Activity vs Change logs
            from PyQt6.QtWidgets import QRadioButton, QButtonGroup
            activity_rb = QRadioButton("Activity Logs")
            change_rb = QRadioButton("Change Log")
            
            # Group radios to ensure exclusivity
            rb_group = QButtonGroup(dialog)
            rb_group.setExclusive(True)
            rb_group.addButton(activity_rb, 0)  # ID 0 for activity
            rb_group.addButton(change_rb, 1)    # ID 1 for change
            activity_rb.setChecked(True)
            
            header_layout.addWidget(activity_rb)
            header_layout.addWidget(change_rb)
            
            # Action buttons
            refresh_btn = QPushButton("Refresh")
            header_layout.addWidget(refresh_btn)
            
            export_btn = QPushButton("Export")
            export_btn.clicked.connect(self.export_logs_to_csv)
            header_layout.addWidget(export_btn)
            
            clear_btn = QPushButton("Clear All")
            clear_btn.setStyleSheet("QPushButton { background-color: #e74c3c; color: white; }")
            header_layout.addWidget(clear_btn)
            
            layout.addLayout(header_layout)
            
            # Logs table
            logs_table = QTableWidget()
            logs_table.setColumnCount(6)
            logs_table.setHorizontalHeaderLabels(["Timestamp", "User", "Action", "Pane", "Details", "Level"])
            logs_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            logs_table.setAlternatingRowColors(True)
            logs_table.setSortingEnabled(True)
            
            # Populate table based on selected view
            def _fill_table(rows: list[list[str]]):
                try:
                    logs_table.setRowCount(len(rows))
                    for row_idx, log_entry in enumerate(rows):
                        for col_idx, value in enumerate((log_entry or [])[:6]):
                            item = QTableWidgetItem(str(value))
                            if len(log_entry) > 5:
                                lvl = str(log_entry[5]).lower()
                                if lvl == "error":
                                    item.setBackground(QColor("#ffebee")); item.setForeground(QColor("#c62828"))
                                elif lvl == "warning":
                                    item.setBackground(QColor("#fff8e1")); item.setForeground(QColor("#f57f17"))
                                elif lvl == "debug":
                                    item.setBackground(QColor("#e3f2fd")); item.setForeground(QColor("#1976d2"))
                            logs_table.setItem(row_idx, col_idx, item)
                    logs_table.resizeColumnsToContents()
                except Exception as e:
                    print(f"Error filling logs table: {e}")

            def _refresh_for_mode():
                rows = []
                try:
                    if activity_rb.isChecked():
                        rows = getattr(self, 'activity_log_data', []) or []
                        print(f"Refreshing with Activity Logs: {len(rows)} entries")
                    else:
                        rows = getattr(self, 'change_log_data', []) or []
                        print(f"Refreshing with Change Logs: {len(rows)} entries")
                except Exception as e:
                    print(f"Error refreshing logs: {e}")
                    rows = []
                _fill_table(rows)

            # Connect radio button toggles
            def on_activity_toggled(checked):
                if checked:
                    _refresh_for_mode()
            
            def on_change_toggled(checked):
                if checked:
                    _refresh_for_mode()
            
            activity_rb.toggled.connect(on_activity_toggled)
            change_rb.toggled.connect(on_change_toggled)
            
            # Connect refresh button
            refresh_btn.clicked.connect(_refresh_for_mode)
            
            # Connect clear button with custom handler
            def clear_logs_and_refresh():
                try:
                    # Clear in-memory logs
                    self.activity_log_data = []
                    self.change_log_data = []
                    
                    # Clear SQLite logs
                    backend_path = getattr(self, '__backend_path__', None) or self.backend_sqlite_path
                    if backend_path and str(backend_path).lower().endswith((".sqlite", ".db")):
                        with self._sqlite_connect(backend_path) as conn:
                            cur = conn.cursor()
                            cur.execute("DELETE FROM activity_logs")
                            cur.execute("DELETE FROM change_log")
                    
                    # Refresh the table
                    _refresh_for_mode()
                    QMessageBox.information(dialog, "Clear Complete", "All logging data has been cleared successfully.")
                except Exception as e:
                    QMessageBox.critical(dialog, "Clear Error", f"Failed to clear logs: {str(e)}")
            
            clear_btn.clicked.connect(clear_logs_and_refresh)
            
            # Initial refresh
            _refresh_for_mode()
            
            layout.addWidget(logs_table)
            
            # Close button
            close_btn = QPushButton("Close")
            close_btn.clicked.connect(dialog.accept)
            layout.addWidget(close_btn)
            
            dialog.exec()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to show logs dialog: {str(e)}")

    def _refresh_logs_table(self, table: QTableWidget) -> None:
        """Refresh the logs table with current data"""
        try:
            if not hasattr(self, 'change_log_data'):
                self.change_log_data = []
            
            table.setRowCount(len(self.change_log_data))
            
            for row_idx, log_entry in enumerate(self.change_log_data):
                for col_idx, value in enumerate(log_entry[:6]):
                    item = QTableWidgetItem(str(value))
                    
                    # Color code by level
                    if len(log_entry) > 5:
                        level = log_entry[5].lower()
                        if level == "error":
                            item.setBackground(QColor("#ffebee"))
                            item.setForeground(QColor("#c62828"))
                        elif level == "warning":
                            item.setBackground(QColor("#fff8e1"))
                            item.setForeground(QColor("#f57f17"))
                        elif level == "debug":
                            item.setBackground(QColor("#e3f2fd"))
                            item.setForeground(QColor("#1976d2"))
                    
                    table.setItem(row_idx, col_idx, item)
            
            # Auto-resize columns
            table.resizeColumnsToContents()
            
        except Exception as e:
            print(f"Error refreshing logs table: {e}")

    def _filter_logs_table(self, table: QTableWidget, level_filter: str = "All", pane_filter: str = "All") -> None:
        """Filter the logs table based on level and pane"""
        try:
            for row in range(table.rowCount()):
                should_show = True
                
                # Filter by level
                if level_filter != "All":
                    level_item = table.item(row, 5)
                    if level_item and level_item.text() != level_filter:
                        should_show = False
                
                # Filter by pane
                if pane_filter != "All" and should_show:
                    pane_item = table.item(row, 3)
                    if pane_item and pane_item.text() != pane_filter:
                        should_show = False
                
                table.setRowHidden(row, not should_show)
                
        except Exception as e:
            print(f"Error filtering logs table: {e}")

    def export_logs_to_csv(self) -> None:
        """Export logs to a separate CSV file"""
        try:
            if not hasattr(self, 'change_log_data') or not self.change_log_data:
                QMessageBox.information(self, "Export Logs", "No logs to export.")
                return
            
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            default_filename = f"activity_logs_export_{timestamp}.csv"
            
            path, _ = QFileDialog.getSaveFileName(
                self, 
                "Export Activity Logs", 
                default_filename, 
                "CSV (*.csv)"
            )
            
            if not path:
                return
            
            # Create CSV file with logs
            import csv
            
            with open(path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                
                # Add headers
                headers = ["Timestamp", "User", "Action", "Pane", "Details", "Level"]
                writer.writerow(headers)
                
                # Add log data
                for log_entry in self.change_log_data:
                    writer.writerow(log_entry[:6])
                
            QMessageBox.information(self, "Export Complete", f"Logs exported successfully to:\n{path}")
            
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export logs:\n{str(e)}")

    def _show_toast(self, text: str, ms: int = 2000, level: str = "INFO") -> None:
        """Redirect toast notifications to a color-coded status bar label (toasts hidden)."""
        try:
            self._show_status_bar(text, level=level, ms=ms)
        except Exception as e:
            ErrorHandler.handle_ui_error("show notification banner", e)

    def _show_status_bar(self, text: str, level: str = "INFO", ms: int = 3000) -> None:
        """Show a persistent, color-coded status label in the window's status bar."""
        try:
            # Colors map reused from toast
            color_map = {
                "ERROR": ("#ffffff", "#ef4444"),
                "WARNING": ("#000000", "#f59e0b"),
                "SUCCESS": ("#ffffff", "#10b981"),
                "INFO": ("#ffffff", "#0f172a"),
                "DEBUG": ("#ffffff", "#6366f1"),
            }
            fg, bg = color_map.get(level, ("#ffffff", "#0f172a"))

            # Ensure we have a status bar and a dedicated label
            try:
                sb = self.statusBar() if hasattr(self, 'statusBar') else None
            except Exception:
                sb = None
            if sb is None:
                # Fallback: create a bottom-aligned label overlay if no status bar is available
                if not hasattr(self, 'fallback_status_label') or self.fallback_status_label is None:
                    self.fallback_status_label = QLabel(self)
                    self.fallback_status_label.setObjectName("fallbackStatusLabel")
                    self.fallback_status_label.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
                self.fallback_status_label.setStyleSheet(
                    f"QLabel#fallbackStatusLabel {{ background-color: {bg}; color: {fg}; padding: 6px 10px; border-top-left-radius: 6px; border-top-right-radius: 6px; font-weight: 700; }}"
                )
                self.fallback_status_label.setText(text)
                self.fallback_status_label.adjustSize()
                # Dock to bottom width of window
                try:
                    w = max(260, self.width())
                    h = self.fallback_status_label.height()
                    self.fallback_status_label.setGeometry(0, max(0, self.height() - h), w, h)
                except Exception:
                    pass
                self.fallback_status_label.show()
                try:
                    self._safe_timer_single_shot(ms, lambda: self.fallback_status_label.hide())
                except Exception:
                    self.fallback_status_label.hide()
                return

            # Build or reuse a permanent label widget for the status bar
            if not hasattr(self, '_status_label') or self._status_label is None:
                self._status_label = QLabel()
                self._status_label.setObjectName("updateStatusLabel")
                try:
                    sb.addPermanentWidget(self._status_label, 1)
                except Exception:
                    # Graceful fallback if addPermanentWidget fails
                    sb.showMessage("")
            self._status_label.setStyleSheet(
                f"QLabel#updateStatusLabel {{ background-color: {bg}; color: {fg}; padding: 6px 10px; border-radius: 6px; font-weight: 700; }}"
            )
            self._status_label.setText(text)
            self._status_label.setVisible(True)

            # Auto-clear after ms
            try:
                self._safe_timer_single_shot(ms, lambda: self._status_label.setVisible(False))
            except Exception:
                try:
                    self._status_label.setVisible(False)
                except Exception:
                    pass
        except Exception as e:
            ErrorHandler.handle_ui_error("show status bar", e)

    def _save_column_order(self, pane_name: str, table: QTableWidget) -> None:
        try:
            header = table.horizontalHeader()
            order = [header.logicalIndex(i) for i in range(header.count())]
            self._column_orders[pane_name] = order
            self._save_autosave()
        except Exception as e:
            ErrorHandler.handle_ui_error("save column order", e)

    def _persist_hidden_columns(self, pane_name: str, table: QTableWidget) -> None:
        try:
            hidden = []
            for i in range(table.columnCount()):
                if table.isColumnHidden(i):
                    hidden.append(i)
            self._hidden_columns[pane_name] = hidden
            self._save_autosave()
        except Exception as e:
            ErrorHandler.handle_ui_error("persist hidden columns", e)

    def _undo_last_edit(self) -> None:
        if not self._edit_undo_stack:
            self._show_toast("Nothing to undo")
            return
        pane, row, col, old, new = self._edit_undo_stack.pop()
        try:
            current = self.data[pane][row][col]
        except Exception:
            current = ""
        # Push to redo stack
        self._edit_redo_stack.append((pane, row, col, current, old))
        # Apply old value
        try:
            self.data[pane][row][col] = old
            self.rebuild_table(pane)
            self.update_dashboard()
            self.update_home_stats()  # Update home page stats
            self._save_backend_sqlite()
            self._show_toast("Undone")
            ErrorHandler.handle_success("Undo operation", f"Undid change in {pane}")
        except Exception as e:
            ErrorHandler.handle_ui_error("undo operation", e, {"pane": pane, "row": row, "col": col})

    def _redo_last_edit(self) -> None:
        if not self._edit_redo_stack:
            self._show_toast("Nothing to redo")
            return
        pane, row, col, old, new = self._edit_redo_stack.pop()
        # Push to undo stack
        self._edit_undo_stack.append((pane, row, col, old, new))
        try:
            self.data[pane][row][col] = new
            self.rebuild_table(pane)
            self.update_dashboard()
            self.update_home_stats()  # Update home page stats
            self._save_backend_sqlite()
            self._show_toast("Redone")
            ErrorHandler.handle_success("Redo operation", f"Redid change in {pane}")
        except Exception as e:
            ErrorHandler.handle_ui_error("redo operation", e, {"pane": pane, "row": row, "col": col})

    def refresh_ism_filter(self) -> None:
        if not hasattr(self, 'ism_filter'):
            return
        isms = set()
        for pane, rows in self.data.items():
            cols = PANE_COLUMNS[pane]
            if "ISM Name" in cols:
                idx = cols.index("ISM Name")
                for r in rows:
                    if idx < len(r) and r[idx]:
                        isms.add(r[idx])
        current = self.ism_filter.currentText() if self.ism_filter.count() > 0 else "All ISMs"
        self.ism_filter.blockSignals(True)
        self.ism_filter.clear()
        self.ism_filter.addItem("All ISMs")
        for name in sorted(isms):
            self.ism_filter.addItem(name)
        idx = self.ism_filter.findText(current)
        if idx >= 0:
            self.ism_filter.setCurrentIndex(idx)
        self.ism_filter.blockSignals(False)

    def _collect_all_isms(self) -> set[str]:
        isms: set[str] = set()
        for pane, rows in self.data.items():
            cols = PANE_COLUMNS[pane]
            if "ISM Name" in cols:
                idx = cols.index("ISM Name")
                for r in rows:
                    if idx < len(r) and r[idx]:
                        isms.add(r[idx])
        return isms
    def manage_ism_directory(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Manage ISM Directory")
        lay = QVBoxLayout(dlg)
        info = QLabel("Add or remove ISM names. These appear in dropdowns.")
        lay.addWidget(info)
        input_row = QHBoxLayout(); name_edit = QLineEdit(); name_edit.setPlaceholderText("Enter ISM name"); add_btn = QPushButton("Add")
        input_row.addWidget(name_edit); input_row.addWidget(add_btn)
        lay.addLayout(input_row)
        lst = QListWidget();
        combined = sorted(set((self.ism_directory or []) + list(self._collect_all_isms())))
        for n in combined:
            lst.addItem(n)
        lay.addWidget(lst)
        btns = QHBoxLayout(); remove_btn = QPushButton("Remove Selected"); close = QPushButton("Close")
        btns.addStretch(1); btns.addWidget(remove_btn); btns.addWidget(close)
        lay.addLayout(btns)
        def add_name():
            try:
                n = name_edit.text().strip()
                if not n:
                    ErrorHandler.handle_warning("Add ISM", "ISM name cannot be empty")
                    return
                if n not in self.ism_directory:
                    self.ism_directory.append(n)
                if not any(lst.item(i).text() == n for i in range(lst.count())):
                    lst.addItem(n)
                name_edit.clear()
                self._save_autosave()
                self.refresh_ism_filter()
                ErrorHandler.handle_success("Add ISM", f"Added {n} to directory")
            except Exception as e:
                ErrorHandler.handle_ui_error("add ISM to directory", e, {"ism_name": n})
        def remove_sel():
            items = lst.selectedItems()
            for it in items:
                try:
                    txt = it.text()
                    if txt in self.ism_directory:
                        self.ism_directory.remove(txt)
                    row = lst.row(it)
                    lst.takeItem(row)
                    ErrorHandler.handle_success("Remove ISM", f"Removed {txt} from directory")
                except Exception as e:
                    ErrorHandler.handle_ui_error("remove ISM from directory", e, {"ism_name": txt})
            self._save_autosave(); self.refresh_ism_filter()
        add_btn.clicked.connect(add_name)
        remove_btn.clicked.connect(remove_sel)
        close.clicked.connect(dlg.accept)
        dlg.resize(480, 420)
        dlg.exec()

    def _collect_org_directory_rows(self) -> list[tuple[str, str, str, str, str, str]]:
        rows: list[tuple[str, str, str, str, str, str]] = []
        try:
            # Check if org_tree exists before accessing it
            if not hasattr(self, 'org_tree') or self.org_tree is None:
                return rows
            root = self.org_tree.invisibleRootItem()
            stack = [root]
            while stack:
                cur = stack.pop()
                for i in range(cur.childCount()):
                    ch = cur.child(i)
                    rows.append((ch.text(0), ch.text(1), ch.text(2), ch.text(3), ch.text(4), ch.text(5)))
                    stack.append(ch)
        except Exception as e:
            ErrorHandler.handle_ui_error("collect org directory", e)
        # Deduplicate by Enterprise ID, keep first
        seen = set(); out = []
        for name, desig, ent, email, mgr, loc in rows:
            key = ent or name
            if key in seen:
                continue
            seen.add(key)
            out.append((name, desig, ent, email, mgr, loc))
        return out

    def show_team_directory(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Team Directory")
        lay = QVBoxLayout(dlg)
        top = QHBoxLayout()
        q = QLineEdit(); q.setPlaceholderText("Search name, designation, or ID")
        export_btn = QPushButton("Export CSV")
        top.addWidget(q); top.addStretch(1); top.addWidget(export_btn)
        lay.addLayout(top)
        tbl = QTableWidget(); tbl.setColumnCount(6); tbl.setHorizontalHeaderLabels(["Name", "Designation", "Enterprise ID", "Email ID", "Manager Enterprise ID", "Location"])
        tbl.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        tbl.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        lay.addWidget(tbl)
        btns = QHBoxLayout(); copy_btn = QPushButton("Copy Selected"); close = QPushButton("Close")
        btns.addStretch(1); btns.addWidget(copy_btn); btns.addWidget(close)
        lay.addLayout(btns)

        data = self._collect_org_directory_rows()
        def render(rows: list[tuple[str, str, str, str, str, str]]):
            tbl.setRowCount(len(rows))
            for i, (n, d, e, email, mgr, loc) in enumerate(rows):
                it_name = QTableWidgetItem(n)
                try:
                    it_name.setIcon(QIcon(self._default_avatar(n, 20)))
                except Exception as e:
                    ErrorHandler.handle_ui_error("set org member avatar", e)
                tbl.setItem(i, 0, it_name)
                tbl.setItem(i, 1, QTableWidgetItem(d))
                tbl.setItem(i, 2, QTableWidgetItem(e))
                tbl.setItem(i, 3, QTableWidgetItem(email))
                tbl.setItem(i, 4, QTableWidgetItem(mgr))
                tbl.setItem(i, 5, QTableWidgetItem(loc))
        render(data)

        def on_search(text: str):
            t = (text or "").lower().strip()
            if not t:
                render(data); return
            filtered = []
            for n, d, e, email, mgr, loc in data:
                if t in (n or "").lower() or t in (d or "").lower() or t in (e or "").lower() or t in (email or "").lower() or t in (mgr or "").lower() or t in (loc or "").lower():
                    filtered.append((n, d, e, email, mgr, loc))
            render(filtered)
        q.textChanged.connect(on_search)

        def do_copy():
            sel_rows = sorted({idx.row() for idx in tbl.selectedIndexes()})
            if not sel_rows:
                return
            lines = []
            for r in sel_rows:
                n = tbl.item(r, 0).text() if tbl.item(r, 0) else ""
                d = tbl.item(r, 1).text() if tbl.item(r, 1) else ""
                e = tbl.item(r, 2).text() if tbl.item(r, 2) else ""
                email = tbl.item(r, 3).text() if tbl.item(r, 3) else ""
                mgr = tbl.item(r, 4).text() if tbl.item(r, 4) else ""
                loc = tbl.item(r, 5).text() if tbl.item(r, 5) else ""
                lines.append(f"{n}\t{d}\t{e}\t{email}\t{mgr}\t{loc}")
            QApplication.clipboard().setText("\n".join(lines))
            self._show_toast("Copied to clipboard")
        copy_btn.clicked.connect(do_copy)

        def do_export():
            path, _ = QFileDialog.getSaveFileName(self, "Export Team Directory", "team_directory.csv", "CSV (*.csv)")
            if not path:
                return
            try:
                pd.DataFrame(data, columns=["Name", "Designation", "Enterprise ID", "Email ID", "Manager Enterprise ID", "Location"]).to_csv(path, index=False)
                QMessageBox.information(self, "Export", "Directory exported")
            except Exception as e:
                QMessageBox.critical(self, "Export", str(e))
        export_btn.clicked.connect(do_export)

        close.clicked.connect(dlg.accept)
        dlg.resize(720, 480)
        dlg.exec()

    def _safe_cell(self, cols: list[str], row_tuple: tuple, idx: int) -> str:
        try:
            return str(row_tuple[idx])
        except Exception:
            return ""

    def populate_client_visits_sections(self) -> None:
        pane = "Client Visits / Audits"
        if pane not in self.data:
            self.visits_table.setRowCount(0)
            return
        cols = PANE_COLUMNS[pane]
        # Map required indexes safely
        idx_project = cols.index("Project Name") if "Project Name" in cols else None
        idx_type = cols.index("Audit Type") if "Audit Type" in cols else None
        idx_ism = cols.index("ISM Name") if "ISM Name" in cols else None
        idx_start = cols.index("Audit Start Date") if "Audit Start Date" in cols else None
        idx_end = cols.index("Audit End Date") if "Audit End Date" in cols else None
        idx_status = cols.index("Status") if "Status" in cols else None
        idx_rag = next((i for i, c in enumerate(cols) if "RAG" in c), None)

        merged_rows = []
        today = datetime.today().date()
        cutoff = today + timedelta(days=5)

        for r in self.data[pane]:
            try:
                start_str = r[idx_start] if idx_start is not None and idx_start < len(r) else ""
                start_date = datetime.strptime(str(start_str), "%Y-%m-%d").date()
            except Exception:
                start_date = None
            status_val = (r[idx_status] if idx_status is not None and idx_status < len(r) else "") or ""
            row_tuple = (
                (r[idx_project] if idx_project is not None and idx_project < len(r) else ""),
                (r[idx_type] if idx_type is not None and idx_type < len(r) else ""),
                (r[idx_ism] if idx_ism is not None and idx_ism < len(r) else ""),
                (r[idx_start] if idx_start is not None and idx_start < len(r) else ""),
                (r[idx_end] if idx_end is not None and idx_end < len(r) else ""),
                status_val,
                (r[idx_rag] if idx_rag is not None and idx_rag < len(r) else ""),
            )
            is_upcoming = start_date is not None and today <= start_date <= cutoff and status_val != "Completed"
            is_inprogress = status_val == "In Progress"
            if is_upcoming or is_inprogress:
                merged_rows.append((row_tuple, is_inprogress))

        self.visits_table.setRowCount(0)
        self.visits_table.setRowCount(len(merged_rows))
        for i, (row, highlight) in enumerate(merged_rows):
            tooltip_parts = []
            # Build tooltip: include Audit Scope if exists and Status
            scope_idx = cols.index("Audit Scope") if "Audit Scope" in cols else None
            status_idx = cols.index("Status") if "Status" in cols else None
            if scope_idx is not None:
                tooltip_parts.append(f"Scope: {self._safe_cell(cols, row, scope_idx)}")
            if status_idx is not None:
                tooltip_parts.append(f"Status: {row[5]}")
            tip = " | ".join(tooltip_parts)
            for j, val in enumerate(row):
                item = QTableWidgetItem(str(val))
                if tip:
                    item.setToolTip(tip)
                if highlight:
                    item.setBackground(QColor("#fff3cd"))  # soft amber highlight for In Progress
                self.visits_table.setItem(i, j, item)
        # Apply quick filters
        def apply_visits_filters():
            only_progress = self.visits_only_inprogress.isChecked()
            type_sel = self.visits_audit_type.currentText()
            for r in range(self.visits_table.rowCount()):
                status_val = self.visits_table.item(r, 5).text()
                type_val = self.visits_table.item(r, 1).text()
                show = True
                if only_progress and status_val != "In Progress":
                    show = False
                if type_sel != "All Types" and type_val != type_sel:
                    show = False
                self.visits_table.setRowHidden(r, not show)
        self.visits_only_inprogress.stateChanged.connect(lambda _=None: apply_visits_filters())
        self.visits_audit_type.currentTextChanged.connect(lambda _=None: apply_visits_filters())
        apply_visits_filters()

    def populate_leads_attention(self) -> None:
        # Build or refresh a simple table of potential issues requiring leads attention
        pane = "Potential Issues"
        if pane not in self.data:
            return
        cols = PANE_COLUMNS[pane]
        if "Leads Attention Required" not in cols:
            return
        att_idx = cols.index("Leads Attention Required")
        ism_idx = cols.index("ISM Name") if "ISM Name" in cols else None
        desc_idx = cols.index("Description") if "Description" in cols else None
        proj_idx = cols.index("Project Name") if "Project Name" in cols else None
        filtered = []
        for r in self.data[pane]:
            val = r[att_idx] if att_idx < len(r) else ""
            if str(val).strip().lower() == "yes":
                ism = r[ism_idx] if (ism_idx is not None and ism_idx < len(r)) else ""
                desc = r[desc_idx] if (desc_idx is not None and desc_idx < len(r)) else ""
                proj = r[proj_idx] if (proj_idx is not None and proj_idx < len(r)) else ""
                filtered.append((ism, desc, proj))
        # Render table if any
        if not hasattr(self, 'leads_table'):
            self.dashboard_layout.addWidget(QLabel("Potential Issues Requiring Leads Attention"))
            self.leads_table = QTableWidget()
            self.leads_table.setColumnCount(3)
            self.leads_table.setHorizontalHeaderLabels(["ISM Name", "Project Name", "Description"])
            self.leads_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
            self.dashboard_layout.addWidget(self.leads_table)
        self.leads_table.setRowCount(0)
        self.leads_table.setRowCount(len(filtered))
        for i, (ism, desc, proj) in enumerate(filtered):
            self.leads_table.setItem(i, 0, QTableWidgetItem(str(ism)))
            self.leads_table.setItem(i, 1, QTableWidgetItem(str(proj)))
            self.leads_table.setItem(i, 2, QTableWidgetItem(str(desc)))

    def open_critical_initiatives_report(self) -> None:
        """Show details of Critical Initiatives that are In Progress"""
        dlg = QDialog(self)
        dlg.setWindowTitle("Critical Initiatives (In Progress) - Details")
        lay = QVBoxLayout(dlg)
        tbl = QTableWidget()
        headers = ["ISM Name", "Project Name", "Summary", "Start Date", "Due Date", "Status", "Priority"]
        tbl.setColumnCount(len(headers))
        tbl.setHorizontalHeaderLabels(headers)
        tbl.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        rows_to_add = []
        
        # Use the same helper functions as the dashboard counting logic
        selected_ism = self.ism_filter.currentText() if hasattr(self, 'ism_filter') else "All ISMs"
        
        def row_matches_ism(pane_name: str, row: list[str]) -> bool:
            if not selected_ism or selected_ism == "All ISMs":
                return True
            cols = PANE_COLUMNS[pane_name]
            if "ISM Name" in cols:
                idx = cols.index("ISM Name")
                return idx < len(row) and row[idx] == selected_ism
            return False
        
        def is_non_empty_row(row):
            return any(str(cell).strip() for cell in row if cell is not None)
        
        # Only look at Initiatives pane
        pane = "Initiatives"
        if pane in self.data:
            cols = PANE_COLUMNS[pane]
            
            # Get column indices
            ism_idx = cols.index("ISM Name") if "ISM Name" in cols else None
            proj_idx = cols.index("Project Name") if "Project Name" in cols else None
            summary_idx = 0 if cols else 0  # First column is usually summary
            start_idx = next((cols.index(c) for c in ("Start Date","Created Date") if c in cols), None)
            due_idx = next((cols.index(c) for c in ("Due Date","Target Date","End Date") if c in cols), None)
            status_idx = cols.index("Status") if "Status" in cols else None
            priority_idx = cols.index("Priority") if "Priority" in cols else None
            
            for r in self.data[pane]:
                # Use the same filtering logic as dashboard counting
                if not is_non_empty_row(r):
                    continue
                if not row_matches_ism(pane, r):
                    continue
                
                # Get priority and status
                priority = (r[priority_idx] if priority_idx is not None and priority_idx < len(r) else "").strip()
                status = (r[status_idx] if status_idx is not None and status_idx < len(r) else "").strip()
                
                # Only include Critical + In Progress initiatives
                if priority.lower() == "critical" and status.lower() == "in progress":
                    row_tuple = (
                        r[ism_idx] if (ism_idx is not None and ism_idx < len(r)) else "",
                        r[proj_idx] if (proj_idx is not None and proj_idx < len(r)) else "",
                        r[summary_idx] if summary_idx < len(r) else "",
                        r[start_idx] if (start_idx is not None and start_idx < len(r)) else "",
                        r[due_idx] if (due_idx is not None and due_idx < len(r)) else "",
                        status,
                        priority
                    )
                    rows_to_add.append(row_tuple)
        
        tbl.setRowCount(len(rows_to_add))
        for i, row in enumerate(rows_to_add):
            for j, val in enumerate(row):
                tbl.setItem(i, j, QTableWidgetItem(str(val)))
        
        lay.addWidget(tbl)
        btn = QPushButton("Close")
        btn.clicked.connect(dlg.accept)
        lay.addWidget(btn)
        dlg.resize(900, 500)
        dlg.exec()

    def open_leads_report(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Leads Attention - Details")
        lay = QVBoxLayout(dlg)
        tbl = QTableWidget()
        tbl.setColumnCount(4)
        tbl.setHorizontalHeaderLabels(["Pane", "ISM Name", "Project/Activity", "Description"])
        tbl.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        rows = []
        
        # Add Potential Issues with "Leads Attention Required" = "Yes"
        if "Potential Issues" in self.data:
            pane = "Potential Issues"
            cols = PANE_COLUMNS[pane]
            if "Leads Attention Required" in cols:
                att_idx = cols.index("Leads Attention Required")
                ism_idx = cols.index("ISM Name") if "ISM Name" in cols else None
                desc_idx = cols.index("Description") if "Description" in cols else None
                proj_idx = cols.index("Project Name") if "Project Name" in cols else None
                for r in self.data[pane]:
                    val = r[att_idx] if att_idx < len(r) else ""
                    if str(val).strip().lower() == "yes":
                        ism = r[ism_idx] if (ism_idx is not None and ism_idx < len(r)) else ""
                        desc = r[desc_idx] if (desc_idx is not None and desc_idx < len(r)) else ""
                        proj = r[proj_idx] if (proj_idx is not None and proj_idx < len(r)) else ""
                        rows.append(("Potential Issues", ism, proj, desc))
        
        # Add Activities with "Support Required" = "Yes"
        if "Activities" in self.data:
            pane = "Activities"
            cols = PANE_COLUMNS[pane]
            if "Support Required" in cols:
                sup_idx = cols.index("Support Required")
                ism_idx = cols.index("ISM Name") if "ISM Name" in cols else None
                desc_idx = cols.index("Activity/Issue") if "Activity/Issue" in cols else None
                proj_idx = cols.index("Project Name") if "Project Name" in cols else None
                for r in self.data[pane]:
                    val = r[sup_idx] if sup_idx < len(r) else ""
                    if str(val).strip().lower() == "yes":
                        ism = r[ism_idx] if (ism_idx is not None and ism_idx < len(r)) else ""
                        desc = r[desc_idx] if (desc_idx is not None and desc_idx < len(r)) else ""
                        proj = r[proj_idx] if (proj_idx is not None and proj_idx < len(r)) else ""
                        rows.append(("Activities", ism, proj, desc))
        
        tbl.setRowCount(len(rows))
        for i, (pane, ism, proj, desc) in enumerate(rows):
            tbl.setItem(i, 0, QTableWidgetItem(str(pane)))
            tbl.setItem(i, 1, QTableWidgetItem(str(ism)))
            tbl.setItem(i, 2, QTableWidgetItem(str(proj)))
            tbl.setItem(i, 3, QTableWidgetItem(str(desc)))
        lay.addWidget(tbl)
        btn = QPushButton("Close")
        btn.clicked.connect(dlg.accept)
        lay.addWidget(btn)
        dlg.resize(1000, 500)
        dlg.exec()

    def open_critical_initiatives_report(self) -> None:
        pane = "Initiatives"
        if pane not in self.data:
            return
        cols = PANE_COLUMNS[pane]
        pr_idx = cols.index("Priority") if "Priority" in cols else None
        st_idx = cols.index("Status") if "Status" in cols else None
        act_idx = cols.index("Action") if "Action" in cols else None
        desc_idx = cols.index("Description") if "Description" in cols else None
        trk_idx = cols.index("Tracker") if "Tracker" in cols else None
        own_idx = cols.index("Ownership") if "Ownership" in cols else None
        dlg = QDialog(self); dlg.setWindowTitle("Critical Initiatives — In Progress"); lay = QVBoxLayout(dlg)
        headers = ["Action","Description","Tracker","Ownership"]
        tbl = QTableWidget(); tbl.setColumnCount(len(headers)); tbl.setHorizontalHeaderLabels(headers); tbl.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        rows = []
        for r in self.data[pane]:
            pr = r[pr_idx] if pr_idx is not None and pr_idx < len(r) else ""
            st = r[st_idx] if st_idx is not None and st_idx < len(r) else ""
            if str(pr).strip().lower() == "critical" and str(st).strip().lower() == "in progress":
                action = r[act_idx] if (act_idx is not None and act_idx < len(r)) else ""
                desc = r[desc_idx] if (desc_idx is not None and desc_idx < len(r)) else ""
                trk = r[trk_idx] if (trk_idx is not None and trk_idx < len(r)) else ""
                own = r[own_idx] if (own_idx is not None and own_idx < len(r)) else ""
                rows.append((action, desc, trk, own))
        tbl.setRowCount(len(rows))
        for i,(a,d,tr,ow) in enumerate(rows):
            tbl.setItem(i,0,QTableWidgetItem(str(a)))
            tbl.setItem(i,1,QTableWidgetItem(str(d)))
            tbl.setItem(i,2,QTableWidgetItem(str(tr)))
            tbl.setItem(i,3,QTableWidgetItem(str(ow)))
        lay.addWidget(tbl)
        btn = QPushButton("Close"); btn.clicked.connect(dlg.accept); lay.addWidget(btn)
        dlg.resize(900, 420); dlg.exec()

    def apply_light_palette(self) -> None:
        # Reset any custom style and apply a light palette
        app = QApplication.instance()
        if app:
            app.setStyleSheet("")
            pal = app.palette()
            pal.setColor(QPalette.ColorRole.Window, QColor("#ffffff"))
            pal.setColor(QPalette.ColorRole.Base, QColor("#ffffff"))
            pal.setColor(QPalette.ColorRole.Text, QColor("#000000"))
            pal.setColor(QPalette.ColorRole.WindowText, QColor("#000000"))
            pal.setColor(QPalette.ColorRole.Button, QColor("#f5f5f5"))
            pal.setColor(QPalette.ColorRole.ButtonText, QColor("#000000"))
            app.setPalette(pal)
            # Improve visibility of menubar/menus/headers in light mode
            light_styles = """
                QMenuBar { background-color: #ffffff; color: #000000; font-weight: 600; font-size: 13px; }
                QMenuBar::item { background: transparent; padding: 6px 10px; font-weight: 600; font-size: 13px; }
                QMenuBar::item:selected { background: #e6f2ff; color: #000000; font-weight: 700; }
                QMenu { background-color: #ffffff; color: #000000; border: 1px solid #cccccc; }
                QMenu::item { padding: 6px 18px; }
                QMenu::item:selected { background-color: #e6f2ff; color: #000000; }
                QHeaderView::section { background-color: #f5f5f5; color: #000000; border: 1px solid #dddddd; padding: 4px; }
                QTableWidget::item:selected { background-color: #d6ebff; color: #000000; }
                QTableWidget::item { padding: 4px; }
                QPushButton { background-color: #f8f9fb; color: #000000; border: 1px solid #d6d6d6; padding: 6px 10px; border-radius: 6px; }
                QPushButton:hover { background-color: #eaeaea; }
                QPushButton#primary { background-color: #2a7de1; color: #ffffff; border: none; }
                QPushButton#primary:hover { background-color: #246fc7; }
                QPushButton#secondary { background-color: #f1f4f8; color: #0f172a; border: 1px solid #d6d6d6; }
                QLineEdit { background-color: #ffffff; color: #000000; border: 1px solid #cccccc; padding: 6px 8px; border-radius: 6px; }
                QComboBox { background-color: #ffffff; color: #000000; border: 1px solid #cccccc; padding: 6px 8px; border-radius: 6px; }
                QComboBox QListView::item { padding: 6px 10px; }
                QTabBar::tab { 
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #f8fafc, stop:1 #f1f5f9);
                    color: #64748b; 
                    padding: 12px 20px; 
                    border: 1px solid #e2e8f0;
                    border-bottom: none;
                    margin: 0px 2px 0px 0px;
                    font-weight: 500;
                    font-size: 13px;
                    min-width: 100px;
                    text-align: center;
                    border-radius: 8px 8px 0px 0px;
                }
                QTabBar::tab:hover { 
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #f1f5f9, stop:1 #e2e8f0);
                    color: #374151;
                    border-color: #cbd5e1;
                }
                QTabBar::tab:selected { 
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ffffff, stop:1 #f8fafc);
                    color: #1e40af;
                    border-color: #3b82f6;
                    border-bottom: 4px solid #3b82f6;
                    font-weight: 600;
                }
                QTabBar::tab:first {
                    border-top-left-radius: 8px;
                    margin-left: 0px;
                }
                QTabBar::tab:last {
                    border-top-right-radius: 8px;
                    margin-right: 0px;
                }
                QTabWidget::pane { 
                    border: 1px solid #e2e8f0; 
                    background: #ffffff;
                    border-top: none;
                    border-radius: 0px 0px 8px 8px;
                }
                QTabWidget::tab-bar {
                    alignment: center;
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #f8fafc, stop:1 #f1f5f9);
                    border-radius: 8px 8px 0px 0px;
                    padding: 4px 4px 0px 4px;
                }
                QTabWidget {
                    background: #f8fafc;
                    border-radius: 8px;
                }
                QTabBar {
                    background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #f8fafc, stop:1 #f1f5f9);
                    border: none;
                    spacing: 0px;
                    border-radius: 8px 8px 0px 0px;
                    padding: 4px 4px 0px 4px;
                }
                QWidget#headerBar { background: #ffffff; border-bottom: 1px solid #e6e6e6; }
                QLabel#appTitle { color: #0f172a; }
            """
            app.setStyleSheet(light_styles)

    def _load_autosave(self) -> None:
        # Load autosave snapshot from SQLite if present
        if os.path.exists(self.autosave_path):
            try:
                self._load_backend_sqlite(self.autosave_path)
            except Exception as e:
                ErrorHandler.handle_ui_error("load autosave sqlite", e)

    def _save_autosave(self) -> None:
        try:
            if not hasattr(self, 'autosave_path'):
                return
            # Snapshot current in-memory state into autosave SQLite
            self._save_backend_sqlite(self.autosave_path)
        except Exception as e:
            ErrorHandler.handle_ui_error("save autosave sqlite", e)

    def _serialize_org_tree(self) -> list[dict]:
        if not hasattr(self, 'org_tree'):
            return []
        items: list[dict] = []
        def walk(node, manager_ent: str | None):
            for i in range(node.childCount()):
                ch = node.child(i)
                obj = {
                    "name": ch.text(0),
                    "designation": ch.text(1),
                    "ent_id": ch.text(2),
                    "email_id": ch.text(3) if ch.columnCount() > 3 else "",
                    "manager": manager_ent or "",
                    "location": ch.text(5) if ch.columnCount() > 5 else "",
                }
                items.append(obj)
                walk(ch, ch.text(2))
        root = self.org_tree.invisibleRootItem()
        walk(root, None)
        return items

    def _save_org_to_csv(self) -> None:
        try:
            directory = self._serialize_org_tree()
            backend_path = getattr(self, '__backend_path__', None) or self.backend_sqlite_path
            # If SQLite backend, persist Org Directory into org_directory table
            if str(backend_path).lower().endswith((".sqlite", ".db")):
                try:
                    with self._sqlite_connect(backend_path) as conn:
                        cur = conn.cursor()
                        cur.execute("CREATE TABLE IF NOT EXISTS org_directory (\"Name\" TEXT, \"Designation\" TEXT, \"Enterprise ID\" TEXT, \"Email ID\" TEXT, \"Manager Enterprise ID\" TEXT, \"Location\" TEXT)")
                        cur.execute("DELETE FROM org_directory")
                        if directory:
                            cur.executemany(
                                "INSERT INTO org_directory(\"Name\", \"Designation\", \"Enterprise ID\", \"Email ID\", \"Manager Enterprise ID\", \"Location\") VALUES (?,?,?,?,?,?)",
                                [
                                    (
                                        item.get('name',''),
                                        item.get('designation',''),
                                        item.get('ent_id',''),
                                        item.get('email_id',''),
                                        item.get('manager',''),
                                        item.get('location','')
                                    ) for item in directory
                                ]
                            )
                except Exception as e:
                    ErrorHandler.handle_ui_error("save org to sqlite", e)
                return
            # This should not happen since we're using SQLite
            ErrorHandler.handle_ui_error("save org to csv", Exception("Non-SQLite backend path provided"))
        except Exception as e:
            ErrorHandler.handle_ui_error("save org to csv", e)
    
    def _export_org_chart(self) -> None:
        """Export org chart data to file"""
        try:
            # Get file path from user
            file_path, file_type = QFileDialog.getSaveFileName(
                self, 
                "Export Org Chart Data", 
                "org_chart_export.csv",
                "CSV Files (*.csv);;Excel Files (*.xlsx);;JSON Files (*.json)"
            )
            
            if not file_path:
                return
            
            # Get org chart data
            directory = self._serialize_org_tree()
            if not directory:
                QMessageBox.information(self, "Export", "No org chart data to export.")
                return
            
            # Convert to list of lists for export
            org_data = []
            for item in directory:
                org_data.append([
                    item.get('name', ''),
                    item.get('designation', ''),
                    item.get('ent_id', ''),
                    item.get('email_id', ''),
                    item.get('manager', ''),
                    item.get('location', '')
                ])
            
            # Export based on file type
            if file_path.lower().endswith('.csv'):
                self._export_org_to_csv(file_path, org_data)
            elif file_path.lower().endswith('.xlsx'):
                self._export_org_to_excel(file_path, org_data)
            elif file_path.lower().endswith('.json'):
                self._export_org_to_json(file_path, org_data)
            
            QMessageBox.information(self, "Export", f"Org chart data exported successfully to {file_path}")
            
        except Exception as e:
            ErrorHandler.handle_ui_error("export org chart", e)
    
    def _import_org_chart(self) -> None:
        """Import org chart data from file with import guide"""
        try:
            # Show import guide dialog first
            result = self._show_org_import_guide()
            
            # Only proceed with import if user clicked "Continue to Import"
            if not result:
                return
            
            # Get file path from user
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Import Org Chart Data",
                "",
                "CSV Files (*.csv);;Excel Files (*.xlsx);;JSON Files (*.json)"
            )
            
            if not file_path:
                return
            
            # Import data based on file type
            org_data = []
            if file_path.lower().endswith('.csv'):
                org_data = self._import_org_from_csv(file_path)
            elif file_path.lower().endswith('.xlsx'):
                org_data = self._import_org_from_excel(file_path)
            elif file_path.lower().endswith('.json'):
                org_data = self._import_org_from_json(file_path)
            
            if not org_data:
                QMessageBox.warning(self, "Import", "No valid data found in the selected file.")
                return
            
            # Clear existing org chart
            self.org_tree.clear()
            
            # Add imported data to org tree
            self._populate_org_tree_from_data(org_data)
            
            # Render the org chart
            self.render_org_chart()
            self._sync_ism_directory_with_org()
            
            # Save to backend
            self._save_org_to_csv()
            self._save_autosave()
            
            QMessageBox.information(self, "Import", f"Successfully imported {len(org_data)} org chart entries.")
            
        except Exception as e:
            ErrorHandler.handle_ui_error("import org chart", e)
    
    def _show_org_import_guide(self) -> bool:
        """Show import guide dialog for org chart. Returns True if user wants to continue with import."""
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle("Org Chart Import Guide")
            dialog.setModal(True)
            dialog.resize(500, 400)
            dialog.setStyleSheet("""
                QDialog {
                    background-color: #f8f9fa;
                }
            """)
            
            layout = QVBoxLayout(dialog)
            layout.setSpacing(15)
            layout.setContentsMargins(20, 20, 20, 20)
            
            # Title
            title = QLabel("📋 Org Chart Import Guide")
            title.setStyleSheet("""
                font-size: 18px;
                font-weight: bold;
                color: #2c3e50;
                margin-bottom: 10px;
            """)
            layout.addWidget(title)
            
            # Import steps
            steps_text = QLabel("""
<b>How to Import Org Chart Data:</b><br><br>
1. <b>Download Sample:</b> Click "📋 Sample" to get the template with proper headers<br><br>
2. <b>Fill Data:</b> Add your team members with required fields:<br>
   • <b>Name:</b> Full name of the person<br>
   • <b>Designation:</b> Job title or role<br>
   • <b>Enterprise ID:</b> Unique identifier (e.g., john.smith)<br>
   • <b>Email ID:</b> Email address<br>
   • <b>Manager Enterprise ID:</b> Enterprise ID of their manager<br>
   • <b>Location:</b> Office location<br><br>
3. <b>Save File:</b> Save as .csv or .xlsx format<br><br>
4. <b>Import:</b> Click "📥 Import" and select your file<br><br>
5. <b>Verify:</b> Check the org chart displays correctly<br><br>
<b>Important Notes:</b><br>
• Manager Enterprise ID should match an existing Enterprise ID in the data for proper hierarchy<br>
• Leave Manager Enterprise ID empty for top-level employees<br>
• All fields are required except Manager Enterprise ID
            """)
            steps_text.setWordWrap(True)
            steps_text.setStyleSheet("""
                font-size: 12px;
                color: #495057;
                line-height: 1.4;
                padding: 15px;
                background-color: #ffffff;
                border: 1px solid #dee2e6;
                border-radius: 8px;
            """)
            layout.addWidget(steps_text)
            
            # Buttons
            button_layout = QHBoxLayout()
            
            # Download sample button
            sample_btn = QPushButton("📋 Download Sample")
            sample_btn.setStyleSheet("""
                QPushButton {
                    background-color: #9b59b6;
                    color: white;
                    padding: 10px 20px;
                    border: none;
                    border-radius: 6px;
                    font-size: 14px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #8e44ad;
                }
            """)
            sample_btn.clicked.connect(lambda: (dialog.accept(), self._download_org_chart_sample()))
            button_layout.addWidget(sample_btn)
            
            button_layout.addStretch()
            
            # Continue button
            continue_btn = QPushButton("Continue to Import")
            continue_btn.setStyleSheet("""
                QPushButton {
                    background-color: #2a7de1;
                    color: white;
                    padding: 10px 20px;
                    border: none;
                    border-radius: 6px;
                    font-size: 14px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #1e6bb8;
                }
            """)
            continue_btn.clicked.connect(dialog.accept)
            button_layout.addWidget(continue_btn)
            
            # Cancel button
            cancel_btn = QPushButton("Cancel")
            cancel_btn.setStyleSheet("""
                QPushButton {
                    background-color: #6c757d;
                    color: white;
                    padding: 10px 20px;
                    border: none;
                    border-radius: 6px;
                    font-size: 14px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #5a6268;
                }
            """)
            cancel_btn.clicked.connect(dialog.reject)
            button_layout.addWidget(cancel_btn)
            
            layout.addLayout(button_layout)
            
            # Return True if user clicked "Continue to Import", False if "Cancel" or "Download Sample"
            result = dialog.exec()
            return result == QDialog.DialogCode.Accepted
            
        except Exception as e:
            ErrorHandler.handle_ui_error("show org import guide", e)
            return False
    
    def _download_org_chart_sample(self) -> None:
        """Download sample org chart template"""
        try:
            # Get file path from user
            file_path, file_type = QFileDialog.getSaveFileName(
                self,
                "Save Org Chart Sample Template",
                "org_chart_sample.csv",
                "CSV Files (*.csv);;Excel Files (*.xlsx)"
            )
            
            if not file_path:
                return
            
            # Create sample data
            sample_data = [
                ["John Smith", "Senior Manager", "john.smith", "john.smith@accenture.com", "", "New York"],
                ["Sarah Johnson", "Manager", "sarah.johnson", "sarah.johnson@accenture.com", "john.smith", "San Francisco"],
                ["Mike Wilson", "Senior Analyst", "mike.wilson", "mike.wilson@accenture.com", "sarah.johnson", "Chicago"],
                ["Lisa Brown", "Analyst", "lisa.brown", "lisa.brown@accenture.com", "sarah.johnson", "Chicago"]
            ]
            
            # Export sample based on file type
            if file_path.lower().endswith('.csv'):
                self._export_org_to_csv(file_path, sample_data, include_headers=True)
            elif file_path.lower().endswith('.xlsx'):
                self._export_org_to_excel(file_path, sample_data, include_headers=True)
            
            QMessageBox.information(self, "Sample Download", f"Sample template saved to {file_path}")
            
        except Exception as e:
            ErrorHandler.handle_ui_error("download org chart sample", e)
    
    def _export_org_to_csv(self, file_path: str, org_data: list, include_headers: bool = True) -> None:
        """Export org chart data to CSV file"""
        try:
            import csv
            with open(file_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                if include_headers:
                    writer.writerow(["Name", "Designation", "Enterprise ID", "Email ID", "Manager Enterprise ID", "Location"])
                writer.writerows(org_data)
        except Exception as e:
            raise Exception(f"Failed to export org chart to CSV: {e}")
    
    def _export_org_to_excel(self, file_path: str, org_data: list, include_headers: bool = True) -> None:
        """Export org chart data to Excel file"""
        try:
            if not PANDAS_AVAILABLE:
                raise Exception("Pandas not available for Excel export")
            import pandas as pd
            
            df = pd.DataFrame(org_data, columns=["Name", "Designation", "Enterprise ID", "Email ID", "Manager Enterprise ID", "Location"] if include_headers else None)
            df.to_excel(file_path, index=False)
        except Exception as e:
            raise Exception(f"Failed to export org chart to Excel: {e}")
    
    def _export_org_to_json(self, file_path: str, org_data: list) -> None:
        """Export org chart data to JSON file"""
        try:
            import json
            # Convert to list of dictionaries
            json_data = []
            for row in org_data:
                json_data.append({
                    "name": row[0],
                    "designation": row[1],
                    "enterprise_id": row[2],
                    "email_id": row[3],
                    "manager_enterprise_id": row[4],
                    "location": row[5]
                })
            
            with open(file_path, 'w', encoding='utf-8') as jsonfile:
                json.dump(json_data, jsonfile, indent=2, ensure_ascii=False)
        except Exception as e:
            raise Exception(f"Failed to export org chart to JSON: {e}")
    
    def _import_org_from_csv(self, file_path: str) -> list:
        """Import org chart data from CSV file"""
        try:
            import csv
            org_data = []
            with open(file_path, 'r', encoding='utf-8') as csvfile:
                reader = csv.reader(csvfile)
                # Skip header row
                next(reader, None)
                for row in reader:
                    if len(row) >= 6:  # Ensure we have all required columns
                        org_data.append(row[:6])  # Take only first 6 columns
            return org_data
        except Exception as e:
            raise Exception(f"Failed to import org chart from CSV: {e}")
    
    def _import_org_from_excel(self, file_path: str) -> list:
        """Import org chart data from Excel file"""
        try:
            if not PANDAS_AVAILABLE:
                raise Exception("Pandas not available for Excel import")
            import pandas as pd
            
            df = pd.read_excel(file_path)
            # Ensure we have the required columns
            required_columns = ["Name", "Designation", "Enterprise ID", "Email ID", "Manager Enterprise ID", "Location"]
            if not all(col in df.columns for col in required_columns):
                raise Exception(f"Excel file must contain columns: {', '.join(required_columns)}")
            
            org_data = []
            for _, row in df.iterrows():
                org_data.append([
                    str(row.get("Name", "")),
                    str(row.get("Designation", "")),
                    str(row.get("Enterprise ID", "")),
                    str(row.get("Email ID", "")),
                    str(row.get("Manager Enterprise ID", "")),
                    str(row.get("Location", ""))
                ])
            return org_data
        except Exception as e:
            raise Exception(f"Failed to import org chart from Excel: {e}")
    
    def _import_org_from_json(self, file_path: str) -> list:
        """Import org chart data from JSON file"""
        try:
            import json
            with open(file_path, 'r', encoding='utf-8') as jsonfile:
                json_data = json.load(jsonfile)
            
            org_data = []
            for item in json_data:
                if isinstance(item, dict):
                    org_data.append([
                        str(item.get("name", "")),
                        str(item.get("designation", "")),
                        str(item.get("enterprise_id", "")),
                        str(item.get("email_id", "")),
                        str(item.get("manager_enterprise_id", "")),
                        str(item.get("location", ""))
                    ])
            return org_data
        except Exception as e:
            raise Exception(f"Failed to import org chart from JSON: {e}")
    
    def _populate_org_tree_from_data(self, org_data: list) -> None:
        """Populate org tree from imported data"""
        try:
            # Create a mapping of enterprise_id to tree item for hierarchy building
            item_map = {}
            
            # First pass: create all items
            for row in org_data:
                name, designation, ent_id, email_id, manager_id, location = row
                item = QTreeWidgetItem()
                item.setText(0, name)
                item.setText(1, designation)
                item.setText(2, ent_id)
                item.setText(3, email_id)
                item.setText(4, manager_id)
                item.setText(5, location)
                
                if ent_id:
                    item_map[ent_id] = item
            
            # Second pass: build hierarchy
            root = self.org_tree.invisibleRootItem()
            for row in org_data:
                name, designation, ent_id, email_id, manager_id, location = row
                if ent_id in item_map:
                    item = item_map[ent_id]
                    if manager_id and manager_id in item_map:
                        # Add as child of manager
                        parent_item = item_map[manager_id]
                        parent_item.addChild(item)
                    else:
                        # Add as root item
                        root.addChild(item)
            
        except Exception as e:
            ErrorHandler.handle_ui_error("populate org tree from data", e)

    def _save_org_to_csv_with_writer(self, writer) -> None:
        """Save org chart data using the provided CSV writer"""
        try:
            directory = self._serialize_org_tree()
            if directory:
                org_data = []
                for item in directory:
                    org_data.append([
                        item.get('name', ''),
                        item.get('designation', ''),
                        item.get('ent_id', ''),
                        item.get('email_id', ''),
                        item.get('manager', ''),
                        item.get('location', '')
                    ])
                df = pd.DataFrame(org_data, columns=["Name", "Designation", "Enterprise ID", "Email ID", "Manager Enterprise ID", "Location"])
                df.to_csv(writer, sheet_name="Org Directory", index=False)
        except Exception as e:
            ErrorHandler.handle_ui_error("save org to excel with writer", e)

    def _rebuild_org_tree_from_saved(self, saved_items: list[dict]) -> None:
        # Clear existing
        try:
            while self.org_tree.topLevelItemCount() > 0:
                self.org_tree.takeTopLevelItem(0)
        except Exception as e:
            ErrorHandler.handle_ui_error("clear org tree", e)
        ent_to_item: dict[str, QTreeWidgetItem] = {}
        # First pass create all nodes
        for entry in saved_items:
            name = str(entry.get("name", ""))
            desig = str(entry.get("designation", ""))
            ent = str(entry.get("ent_id", ""))
            email = str(entry.get("email_id", ""))
            manager = str(entry.get("manager", ""))
            location = str(entry.get("location", ""))
            node = QTreeWidgetItem([name, desig, ent, email, manager, location])
            node.setIcon(0, QIcon(self._default_avatar(name, 32)))
            ent_to_item[ent] = node
        # Second pass attach to managers or top level
        for entry in saved_items:
            ent = str(entry.get("ent_id", ""))
            mgr = str(entry.get("manager", "") or "")
            node = ent_to_item.get(ent)
            if not node:
                continue
            if mgr and mgr in ent_to_item:
                ent_to_item[mgr].addChild(node)
            else:
                self.org_tree.addTopLevelItem(node)
        # Render
        self.render_org_chart()
        # Sync ISM directory with org chart
        self._sync_ism_directory_with_org()

    def _create_new_backend_file(self, backend_path: str = None) -> None:
        """Create a new backend Excel file with required sheets and headers, even without pandas."""
        if backend_path is None:
            backend_path = self.backend_sqlite_path
        # Normalize extension
        if not str(backend_path).lower().endswith('.xlsx'):
            backend_path = f"{backend_path}.xlsx"
            self.backend_sqlite_path = backend_path
            self.__backend_path__ = backend_path
        try:
            if PANDAS_AVAILABLE:
                with pd.ExcelWriter(backend_path, engine="openpyxl") as writer:
                    for pane_name, columns in PANE_COLUMNS.items():
                        pd.DataFrame(columns=columns).to_excel(writer, sheet_name=self._sheet_title_for_pane(pane_name), index=False)
                    pd.DataFrame(columns=["Name", "Designation", "Enterprise ID", "Email ID", "Manager Enterprise ID", "Location"]).to_excel(writer, sheet_name="Org Directory", index=False)
                    pd.DataFrame(columns=["Project Name","Project ID"]).to_excel(writer, sheet_name="Projects", index=False)
                    pd.DataFrame(columns=["Timestamp","Pane","Source","Rows"]).to_excel(writer, sheet_name="Imports", index=False)
                    pd.DataFrame(columns=["Name"]).to_excel(writer, sheet_name="ISM Directory", index=False)
                    pd.DataFrame(columns=["Timestamp", "User", "Action", "Pane", "Details"]).to_excel(writer, sheet_name="Change Log", index=False)
                    settings_data = [
                        ["App Version", "1.0.0", "Current application version"],
                        ["Last Updated", "", "Last time the application was updated"],
                        ["Default ISM Filter", "All ISMs", "Default ISM filter setting"],
                        ["Auto Save", "True", "Enable automatic saving"],
                        ["Backend Path", backend_path, "Path to backend Excel file"],
                    ]
                    pd.DataFrame(settings_data, columns=["Setting","Value","Description"]).to_excel(writer, sheet_name="Settings", index=False)
            else:
                # Fallback: build a valid workbook with openpyxl
                from openpyxl import Workbook
                from openpyxl.styles import Font
                wb = Workbook()
                ws = wb.active
                ws.title = self._sheet_title_for_pane(next(iter(PANE_COLUMNS))) if PANE_COLUMNS else 'Data'
                # Replace with proper panes
                for pane_name, columns in PANE_COLUMNS.items():
                    ws = wb.create_sheet(self._sheet_title_for_pane(pane_name))
                    for col, header in enumerate(columns, 1):
                        cell = ws.cell(row=1, column=col, value=header)
                        cell.font = Font(bold=True)
                ws = wb.create_sheet('Org Directory'); ws.append(["Name","Designation","Enterprise ID","Email ID","Manager Enterprise ID","Location"])
                ws = wb.create_sheet('Projects'); ws.append(["Project Name","Project ID"])
                ws = wb.create_sheet('Imports'); ws.append(["Timestamp","Pane","Source","Rows"])
                ws = wb.create_sheet('ISM Directory'); ws.append(["Name"])
                ws = wb.create_sheet('Change Log'); ws.append(["Timestamp","User","Action","Pane","Details"])
                ws = wb.create_sheet('Settings'); ws.append(["Setting","Value","Description"]); ws.append(["Backend Path", backend_path, "Path to backend Excel file"])
                # Remove default if still present and empty
                try:
                    default = wb['Sheet']
                    wb.remove(default)
                except Exception:
                    pass
                wb.save(backend_path)
            # Initialize the backend path
            self.__backend_path__ = backend_path
        except Exception as e:
            ErrorHandler.handle_ui_error("create backend file", e, {"backend_path": backend_path})
            print(f"Failed to create backend file: {e}")
            raise e

    def animated_load_data(self) -> None:
        """Animated load data function that opens file dialog"""
        try:
            # Animate the button click
            if hasattr(self, 'animated_load_btn'):
                self.animated_load_btn.animate_click()
                self.animated_load_btn.setText("⏳ Loading...")
                self.animated_load_btn.stop_pulse()
            
            # Open file dialog
            path, _ = QFileDialog.getOpenFileName(self, "Load Data File", "", "Excel (*.xlsx)")
            if not path:
                # Reset button if cancelled
                if hasattr(self, 'animated_load_btn'):
                    self.animated_load_btn.setText("📁 Load Data")
                    self.animated_load_btn.start_pulse()
                return
            
            # Show loading overlay
            self.loading_overlay.show_loading("Loading Excel data...", show_progress=True)
            
            # Set backend path and load data
            self.__backend_path__ = path
            self._load_backend_sqlite(path)
            self._save_autosave()  # Save the backend path
            
            # Hide loading overlay and show success
            self.loading_overlay.hide_loading()
            self.notifications.show_success("Data loaded successfully!")
            ErrorHandler.handle_success("animated load data", f"Loaded data from {path}")
            
            # Reset button
            if hasattr(self, 'animated_load_btn'):
                self.animated_load_btn.setText("📁 Load Data")
                self.animated_load_btn.start_pulse()
                
        except Exception as e:
            self.loading_overlay.hide_loading()
            ErrorHandler.handle_ui_error("animated load data", e)
            self.notifications.show_error(f"Failed to load data: {str(e)}")
            if hasattr(self, 'animated_load_btn'):
                self.animated_load_btn.setText("📁 Load Data")
                self.animated_load_btn.start_pulse()
    def _load_backend_sqlite(self, db_path: str) -> None:
        try:
            if not os.path.exists(db_path):
                self._create_new_backend_sqlite(db_path)
            with self._sqlite_connect(db_path) as conn:
                cur = conn.cursor()
                loaded_panes = []
                # Load pane tables
                for pane_name, columns in PANE_COLUMNS.items():
                    table_name = pane_name.lower().replace(' ', '_').replace('/', '_').replace('\\', '_').replace(':', '_').replace('?', '_').replace('*', '_').replace('[', '(').replace(']', ')')
                    cur.execute(f"SELECT * FROM {table_name}")
                    rows = cur.fetchall()
                    if rows:
                        # Convert to list of lists
                        self.data[pane_name] = [list(row) for row in rows]
                    loaded_panes.append(pane_name)
            
            # Load logs from SQLite after loading all data
            self._load_logs_from_sqlite()
            # Org Directory and Projects are now loaded from SQLite in the main load function

            # Change Log is now loaded from SQLite in the main load function

            # Rebuild tables for all loaded panes
            for pane in loaded_panes:
                try:
                    self.rebuild_table(pane)
                    print(f"Successfully rebuilt table for {pane}")
                except Exception as e:
                    ErrorHandler.handle_ui_error("rebuild table", e, {"pane": pane})
                    print(f"Failed to rebuild table for {pane}: {e}")
            
            # Also rebuild tables for any panes that weren't loaded from SQLite but exist in data
            for pane in self.data.keys():
                if pane not in loaded_panes:
                    try:
                        self.rebuild_table(pane)
                        print(f"Successfully rebuilt existing table for {pane}")
                    except Exception as e:
                        print(f"Failed to rebuild existing table for {pane}: {e}")
            
            # Update dashboard and refresh
            self.update_dashboard()
            self.refresh_ism_filter()
            # Ensure Project Details custom UI reflects newly loaded data
            try:
                if hasattr(self, 'projects_table') and self.projects_table is not None:
                    self._load_projects_data()
                    if hasattr(self, '_update_summary_metrics'):
                        self._update_summary_metrics()
            except Exception:
                pass
            
            # Reconstruct org chart if org data was loaded
            if hasattr(self, 'org_tree'):
                self.render_org_chart()
            
            # Rebuild/refresh Leave Tracker UI
            try:
                    # The Leave Tracker uses custom UI; ensure decorations and day list refresh
                    if hasattr(self, '_refresh_calendar_decorations'):
                        self._refresh_calendar_decorations()
                    # Trigger the current date refresh via selectionChanged slot by setting the same date
                    if hasattr(self, 'leave_calendar') and self.leave_calendar:
                        d = self.leave_calendar.selectedDate()
                        self.leave_calendar.setSelectedDate(d)
                    print("Refreshed Leave Tracker UI")
            except Exception as e:
                self.notifications.show_warning(f"Failed to refresh Leave Tracker UI: {str(e)}")
                self._log_change("Error", "Leave Tracker", f"Failed to refresh UI: {str(e)}")
            
            # Update home page stats after loading data
            self.update_home_stats()
            
            # Sync projects from Project Details data
            self.sync_projects_from_details()
            
            # Auto-update ageing calculations after loading data
            self.auto_update_ageing()
            
            # Show success animation on the load button instead of message box
            if hasattr(self, 'animated_load_btn'):
                self.animated_load_btn.setText("✅ Loaded!")
                self.animated_load_btn.stop_pulse()
                # Reset button text after 2 seconds (ensure main thread)
                try:
                    self._safe_timer_single_shot(2000, lambda: self._safe_reset_load_button())
                except Exception as e:
                    print(f"Load button timer error: {e}")
                    # Fallback: reset immediately if timer fails
                    try:
                        self.animated_load_btn.setText("📁 Load Data")
                        self.animated_load_btn.start_pulse()
                    except Exception as e:
                        ErrorHandler.handle_ui_error("save badge timer", e)
            
        except Exception as e:
            ErrorHandler.handle_ui_error("load data from SQLite", e, {"file_path": db_path})
            QMessageBox.critical(self, "Load Data Failed", f"Failed to load data from {db_path}:\n{str(e)}")
            # Reset button on error
            if hasattr(self, 'animated_load_btn'):
                self.animated_load_btn.setText("📁 Load Data")
                self.animated_load_btn.start_pulse()


    def export_all_sheets(self) -> None:
        """Export all panes to a single Excel file without blocking the UI."""
        path, _ = QFileDialog.getSaveFileName(self, "Export All", "dashboard_export.xlsx", "Excel (*.xlsx)")
        if not path:
            return
        if not PANDAS_AVAILABLE:
            QMessageBox.warning(self, "Export All", "Pandas is not available. Excel export is disabled.")
            return

        # Lightweight modal progress dialog (busy indicator)
        prog = QDialog(self)
        prog.setWindowTitle("Exporting…")
        v = QVBoxLayout(prog)
        lbl = QLabel("Exporting data to Excel… This may take a moment.")
        bar = QProgressBar(); bar.setRange(0, 0)
        btn = QPushButton("Cancel")
        h = QHBoxLayout(); h.addStretch(1); h.addWidget(btn)
        v.addWidget(lbl); v.addWidget(bar); v.addLayout(h)
        canceled = {"flag": False}
        btn.clicked.connect(lambda: (prog.reject(), canceled.__setitem__("flag", True)))

        def do_export():
            error: Optional[Exception] = None
            try:
                with pd.ExcelWriter(path, engine="openpyxl") as writer:
                    for pane, rows in self.data.items():
                        if canceled["flag"]:
                            break
                        if pane == "Leave Tracker":
                            continue
                        df = pd.DataFrame(rows, columns=PANE_COLUMNS[pane])
                        df.to_csv(writer, sheet_name=self._sheet_title_for_pane(pane), index=False)
            except Exception as e:  # capture error to report on UI thread
                error = e
            finally:
                # marshal UI updates back to main thread
                self._safe_timer_single_shot(0, lambda: self._on_export_all_done(error, prog, canceled["flag"]))

        t = threading.Thread(target=do_export, daemon=True)
        t.start()
        prog.exec()

    def _on_export_all_done(self, error: Optional[Exception], dialog: QDialog, was_canceled: bool) -> None:
        try:
            if dialog.isVisible():
                dialog.accept()
        except Exception:
            pass
        if was_canceled:
            self._show_toast("Export canceled", ms=2000)
            return
        if error is not None:
            try:
                QMessageBox.critical(self, "Export All Failed", str(error))
            except Exception:
                self.notifications.show_warning(f"Export failed: {error}")
            return
        try:
            QMessageBox.information(self, "Export All", "Export completed")
        except Exception:
            self._show_toast("Export completed", ms=2000)

    # Save on window close
    def closeEvent(self, event) -> None:
        # Persist both JSON snapshot and backend SQLite when app closes
        self._save_autosave()
        try:
            self._save_org_to_csv()  # This now uses CSV for export
        except Exception as e:
            self.notifications.show_warning(f"Failed to save Org Chart: {str(e)}")
        self._save_backend_sqlite()  # This now uses SQLite internally
        
        # Cleanup resources to prevent memory leaks
        self._cleanup_resources()
        
        super().closeEvent(event)
    
    def _get_backend_path_from_settings(self) -> str:
        """Get backend path - SQLite is now the primary backend, Excel only for export."""
        try:
            # Prefer already loaded backend path
            if hasattr(self, '__backend_path__') and self.__backend_path__:
                return self.__backend_path__
            
            # Always default to SQLite - Excel is only for export
            script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
            sqlite_default = os.path.join(script_dir, "backend_data.sqlite")
            
            # If SQLite file doesn't exist, create it
            if not os.path.exists(sqlite_default):
                self._create_new_backend_sqlite(sqlite_default)
            
            return sqlite_default
        except Exception:
            script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
            sqlite_default = os.path.join(script_dir, "backend_data.sqlite")
            return sqlite_default

    def _open_backend_path_settings_dialog(self) -> None:
        """Show a dialog to view/update the Backend Excel file path stored in Settings sheet."""
        try:
            current_path = getattr(self, '__backend_path__', None) or getattr(self, 'backend_sqlite_path', "")
            dlg = QDialog(self)
            dlg.setWindowTitle("Backend File")
            lay = QVBoxLayout(dlg)
            lbl = QLabel("Current Backend File:")
            path_edit = QLineEdit(current_path)
            path_edit.setReadOnly(True)
            choose_btn = QPushButton("Choose New File…")
            save_btn = QPushButton("Save to Settings")
            close_btn = QPushButton("Close")
            lay.addWidget(lbl)
            lay.addWidget(path_edit)
            btns = QHBoxLayout(); btns.addWidget(choose_btn); btns.addStretch(1); btns.addWidget(save_btn); btns.addWidget(close_btn)
            lay.addLayout(btns)
            def choose():
                new_path, _ = QFileDialog.getOpenFileName(self, "Select Backend File", current_path or "", "SQLite (*.sqlite);;Excel (*.xlsx)")
                if new_path:
                    path_edit.setText(new_path)
            choose_btn.clicked.connect(choose)
            
            def do_save():
                new_path = path_edit.text().strip()
                if not new_path:
                    QMessageBox.warning(self, "Backend", "Please select a valid backend file path.")
                    return
                # Persist to Settings sheet
                try:
                    # Ensure file exists or create new backend structure
                    if not os.path.exists(new_path):
                        self.backend_sqlite_path = new_path
                        if new_path.lower().endswith((".sqlite", ".db")):
                            self._create_new_backend_sqlite(new_path)
                        else:
                            self._create_new_backend_file(new_path)
                    # Write settings path
                    if new_path.lower().endswith((".sqlite", ".db")):
                        with self._sqlite_connect(new_path) as conn:
                            conn.execute("CREATE TABLE IF NOT EXISTS settings (setting TEXT PRIMARY KEY, value TEXT, description TEXT)")
                            conn.execute("INSERT OR REPLACE INTO settings(setting, value, description) VALUES (?,?,?)", ("Backend Path", new_path, "Path to backend SQLite file"))
                    else:
                        if not PANDAS_AVAILABLE:
                            QMessageBox.warning(self, "Backend", "Pandas not available. Cannot write Settings sheet.")
                            return
                        with pd.ExcelWriter(new_path, engine="openpyxl", mode='a', if_sheet_exists='replace') as writer:
                            settings_data = [
                                ["App Version", "1.0.0", "Current application version"],
                                ["Last Updated", datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "Last time the application was updated"],
                                ["Default ISM Filter", getattr(self, 'ism_filter', QComboBox()).currentText() if hasattr(self, 'ism_filter') else "All ISMs", "Default ISM filter setting"],
                                ["Auto Save", "True", "Enable automatic saving"],
                                ["Backend Path", new_path, "Path to backend Excel file"],
                            ]
                            pd.DataFrame(settings_data, columns=["Setting","Value","Description"]).to_excel(writer, sheet_name="Settings", index=False)
                    # Update runtime paths
                    self.backend_sqlite_path = new_path
                    self.__backend_path__ = new_path
                    QMessageBox.information(self, "Backend", "Backend file path updated in Settings.")
                except Exception as e:
                    QMessageBox.critical(self, "Backend", f"Failed to save backend path: {str(e)}")
            save_btn.clicked.connect(do_save)
            close_btn.clicked.connect(dlg.accept)
            dlg.resize(640, 160)
            dlg.exec()
        except Exception as e:
            ErrorHandler.handle_ui_error("backend path settings dialog", e)

    # =========================
    # SQLite backend utilities
    # =========================
    def _sqlite_connect(self, db_path: str) -> sqlite3.Connection:
        conn = sqlite3.connect(db_path, timeout=30, isolation_level=None)
        # Enable WAL for better concurrency
        try:
            conn.execute("PRAGMA journal_mode=WAL;")
            conn.execute("PRAGMA synchronous=NORMAL;")
        except Exception:
            pass
        return conn
    
    def _sheet_title_for_pane(self, pane_name: str) -> str:
        """Convert pane name to safe table/sheet name"""
        safe = pane_name.replace("/", "_").replace("\\", "_").replace(":", "_").replace("?", "_").replace("*", "_").replace("[", "(").replace("]", ")")
        return safe[:31]  # Limit to 31 characters for compatibility

    def _create_new_backend_sqlite(self, db_path: str) -> None:
        try:
            os.makedirs(os.path.dirname(db_path) or '.', exist_ok=True)
            with self._sqlite_connect(db_path) as conn:
                cur = conn.cursor()
                # Settings table
                cur.execute("CREATE TABLE IF NOT EXISTS settings (setting TEXT PRIMARY KEY, value TEXT, description TEXT)")
                # Logs tables
                cur.execute("CREATE TABLE IF NOT EXISTS activity_logs (Timestamp TEXT, User TEXT, Action TEXT, Pane TEXT, Details TEXT, Level TEXT)")
                cur.execute("CREATE TABLE IF NOT EXISTS change_log (Timestamp TEXT, User TEXT, Action TEXT, Pane TEXT, Details TEXT, Level TEXT)")
                # Imports, Projects, ISM Directory
                cur.execute("CREATE TABLE IF NOT EXISTS imports (Timestamp TEXT, Pane TEXT, Source TEXT, Rows TEXT)")
                cur.execute("CREATE TABLE IF NOT EXISTS projects (\"Project Name\" TEXT, \"Project ID\" TEXT)")
                cur.execute("CREATE TABLE IF NOT EXISTS ism_directory (Name TEXT)")
                # Org Directory table mirrors Excel sheet
                cur.execute("CREATE TABLE IF NOT EXISTS org_directory (\"Name\" TEXT, \"Designation\" TEXT, \"Enterprise ID\" TEXT, \"Email ID\" TEXT, \"Manager Enterprise ID\" TEXT, \"Location\" TEXT)")
                # Enhanced Wiki / Knowledge Base tables
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_articles (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        title TEXT NOT NULL,
                        slug TEXT UNIQUE,
                        content_md TEXT,
                        content_html TEXT,
                        tags TEXT,
                        author TEXT,
                        status TEXT DEFAULT 'Draft',
                        created_at TEXT,
                        updated_at TEXT,
                        review_at TEXT
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_versions (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        article_id INTEGER,
                        version_no INTEGER,
                        title TEXT,
                        content_md TEXT,
                        content_html TEXT,
                        tags TEXT,
                        author TEXT,
                        created_at TEXT,
                        FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_attachments (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        article_id INTEGER,
                        filename TEXT,
                        path TEXT,
                        uploaded_at TEXT,
                        FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_links (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        from_article_id INTEGER,
                        to_article_id INTEGER,
                        FOREIGN KEY(from_article_id) REFERENCES kb_articles(id),
                        FOREIGN KEY(to_article_id) REFERENCES kb_articles(id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_comments (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        article_id INTEGER,
                        author TEXT,
                        content TEXT,
                        created_at TEXT,
                        FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_approvals (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        article_id INTEGER,
                        state TEXT,
                        reviewer TEXT,
                        decided_at TEXT,
                        comment TEXT,
                        FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                    )
                """)
                # FTS5 for full-text search if available
                try:
                    cur.execute("""
                        CREATE VIRTUAL TABLE IF NOT EXISTS kb_articles_fts USING fts5(
                            title, content, tags, content='kb_articles', content_rowid='id'
                        )
                    """)
                except Exception:
                    pass
                # Dynamic pane tables
                for pane_name, columns in PANE_COLUMNS.items():
                    cols_sql = ", ".join([f'"{c}" TEXT' for c in columns])
                    table_name = self._sheet_title_for_pane(pane_name).lower().replace(' ', '_')
                    cur.execute(f"CREATE TABLE IF NOT EXISTS {table_name} ({cols_sql})")
                # Seed settings
                cur.execute("INSERT OR REPLACE INTO settings(setting, value, description) VALUES (?,?,?)", ("App Version", "1.0.0", "Current application version"))
                cur.execute("INSERT OR REPLACE INTO settings(setting, value, description) VALUES (?,?,?)", ("Backend Path", db_path, "Path to backend SQLite file"))
                conn.commit()
        except Exception as e:
            self.notifications.show_warning(f"Failed to create SQLite backend: {str(e)}")

    def _ensure_wiki_tables_exist(self, conn) -> None:
        """Ensure Wiki/Knowledge Base tables exist in the database."""
        try:
            cur = conn.cursor()
            # Check if kb_articles table exists
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='kb_articles'")
            if not cur.fetchone():
                # Create Wiki tables if they don't exist
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_articles (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        title TEXT NOT NULL,
                        slug TEXT UNIQUE,
                        content_md TEXT,
                        content_html TEXT,
                        tags TEXT,
                        author TEXT,
                        status TEXT DEFAULT 'Draft',
                        created_at TEXT,
                        updated_at TEXT,
                        review_at TEXT
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_versions (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        article_id INTEGER,
                        version_no INTEGER,
                        title TEXT,
                        content_md TEXT,
                        content_html TEXT,
                        tags TEXT,
                        author TEXT,
                        created_at TEXT,
                        FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_attachments (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        article_id INTEGER,
                        filename TEXT,
                        path TEXT,
                        uploaded_at TEXT,
                        FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_links (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        from_article_id INTEGER,
                        to_article_id INTEGER,
                        FOREIGN KEY(from_article_id) REFERENCES kb_articles(id),
                        FOREIGN KEY(to_article_id) REFERENCES kb_articles(id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_comments (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        article_id INTEGER,
                        author TEXT,
                        content TEXT,
                        created_at TEXT,
                        FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS kb_approvals (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        article_id INTEGER,
                        state TEXT,
                        reviewer TEXT,
                        decided_at TEXT,
                        comment TEXT,
                        FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                    )
                """)
                # FTS5 for full-text search if available
                try:
                    cur.execute("""
                        CREATE VIRTUAL TABLE IF NOT EXISTS kb_articles_fts USING fts5(
                            title, content, tags, content='kb_articles', content_rowid='id'
                        )
                    """)
                except Exception:
                    pass
                conn.commit()
        except Exception as e:
            logger.debug(f"Failed to ensure Wiki tables exist: {e}")

    def _load_backend_sqlite(self, db_path: str) -> None:
        try:
            if not os.path.exists(db_path):
                self._create_new_backend_sqlite(db_path)
            with self._sqlite_connect(db_path) as conn:
                cur = conn.cursor()
                loaded_panes = []
                
                # Ensure Wiki/Knowledge Base tables exist (for existing databases)
                self._ensure_wiki_tables_exist(conn)
                
                # Load pane tables
                for pane_name, columns in PANE_COLUMNS.items():
                    # Try both table naming conventions
                    table_name_lower = self._sheet_title_for_pane(pane_name).lower().replace(' ', '_')
                    table_name_original = self._sheet_title_for_pane(pane_name)
                    
                    # Check which table exists
                    cur.execute(f"SELECT name FROM sqlite_master WHERE type='table' AND name IN ('{table_name_lower}', '{table_name_original}')")
                    existing_table = cur.fetchone()
                    
                    if existing_table:
                        table_name = existing_table[0]
                        try:
                            # Quote table name if it contains spaces
                            quoted_table_name = f'"{table_name}"' if ' ' in table_name else table_name
                            cur.execute(f"SELECT {', '.join([f'"{c}"' for c in columns])} FROM {quoted_table_name}")
                            rows = cur.fetchall()
                            self.data[pane_name] = [[str(v) if v is not None else '' for v in row] for row in rows]
                            loaded_panes.append(pane_name)
                            print(f"Loaded {pane_name} from table {table_name} with {len(rows)} rows")
                        except Exception as e:
                            print(f"Error loading {pane_name} from table {table_name}: {e}")
                            self.data[pane_name] = []
                    else:
                        print(f"No table found for {pane_name} (tried {table_name_lower} and {table_name_original})")
                        self.data[pane_name] = []
                # Load Projects
                try:
                    cur.execute("SELECT \"Project Name\", \"Project ID\" FROM projects")
                    self.projects = [(str(n) if n else '', str(pid) if pid else '') for (n, pid) in cur.fetchall()]
                except Exception:
                    self.projects = []
                # Load ISM Directory
                try:
                    cur.execute("SELECT Name FROM ism_directory")
                    self.ism_directory = [str(n) for (n,) in cur.fetchall() if n]
                except Exception:
                    self.ism_directory = []
                # Load Org Directory and rebuild org tree
                try:
                    cur.execute("SELECT \"Name\", \"Designation\", \"Enterprise ID\", \"Email ID\", \"Manager Enterprise ID\", \"Location\" FROM org_directory")
                    rows = cur.fetchall()
                    saved_items = []
                    for r in rows:
                        saved_items.append({
                            'name': str(r[0] or ''),
                            'designation': str(r[1] or ''),
                            'ent_id': str(r[2] or ''),
                            'email_id': str(r[3] or ''),
                            'manager': str(r[4] or ''),
                            'location': str(r[5] or ''),
                        })
                    if saved_items:
                        self._rebuild_org_tree_from_saved(saved_items)
                except Exception:
                    pass
                # Load logs
                try:
                    cur.execute("SELECT Timestamp, User, Action, Pane, Details, Level FROM change_log")
                    self.change_log_data = [list(map(lambda x: '' if x is None else str(x), row)) for row in cur.fetchall()]
                except Exception:
                    self.change_log_data = []
                try:
                    cur.execute("SELECT Timestamp, User, Action, Pane, Details, Level FROM activity_logs")
                    self.activity_log_data = [list(map(lambda x: '' if x is None else str(x), row)) for row in cur.fetchall()]
                except Exception:
                    self.activity_log_data = []
            # Rebuild UI
            for pane in loaded_panes:
                try:
                    self.rebuild_table(pane)
                except Exception:
                    pass
            self.update_dashboard()
            self.refresh_ism_filter()
            self.update_home_stats()
            self.sync_projects_from_details()
            self.auto_update_ageing()
            # Render Org chart after loading data
            try:
                if hasattr(self, 'render_org_chart'):
                    self.render_org_chart()
            except Exception as e:
                ErrorHandler.handle_ui_error("render org chart after load", e)
            # Ensure Project Details custom UI reflects newly loaded data
            try:
                if hasattr(self, 'projects_table') and self.projects_table is not None:
                    self._load_projects_data()
                    if hasattr(self, '_update_summary_metrics'):
                        self._update_summary_metrics()
            except Exception:
                pass
            # Ensure Leave Tracker UI reflects newly loaded data
            try:
                if hasattr(self, 'leave_tab') and self.leave_tab is not None:
                    self.refresh_leave_tracker()
            except Exception:
                pass
        except Exception as e:
            QMessageBox.critical(self, "Load Data Failed", f"Failed to load data from {db_path}:\n{str(e)}")

    def _save_backend_sqlite(self, db_path: str = None) -> None:
        try:
            # Use default backend path if no specific path is set
            if db_path is None:
                db_path = getattr(self, '__backend_path__', None) or self._get_backend_path_from_settings()
            
            # Show UI feedback
            if hasattr(self, '_save_badge') and self._save_badge is not None:
                self._save_badge.setText("Saving…")
            
            # Show saving status near Load Data button
            if hasattr(self, 'save_status_label'):
                self.save_status_label.setText("Saving...")
                self.save_status_label.setStyleSheet("""
                    QLabel {
                        color: #f39c12;
                        font-size: 12px;
                        font-weight: bold;
                        padding: 4px 8px;
                        background-color: rgba(243, 156, 18, 0.1);
                        border: 1px solid #f39c12;
                        border-radius: 4px;
                        min-width: 80px;
                        text-align: center;
                    }
                """)
                self.save_status_label.setVisible(True)
            
            # Show loading overlay for save operation
            if hasattr(self, 'loading_overlay'):
                self.loading_overlay.show_loading("Saving data...", show_progress=True)
            
            with self._sqlite_connect(db_path) as conn:
                cur = conn.cursor()
                # Ensure schema exists (create tables if they don't exist)
                for pane_name, columns in PANE_COLUMNS.items():
                    table_name = self._sheet_title_for_pane(pane_name).lower().replace(' ', '_')
                    cols_sql = ", ".join([f'"{c}" TEXT' for c in columns])
                    cur.execute(f"CREATE TABLE IF NOT EXISTS {table_name} ({cols_sql})")
                # Save panes
                for pane_name, rows in self.data.items():
                    columns = PANE_COLUMNS.get(pane_name, [])
                    # Try both table naming conventions
                    table_name_lower = self._sheet_title_for_pane(pane_name).lower().replace(' ', '_')
                    table_name_original = self._sheet_title_for_pane(pane_name)
                    
                    # Check which table exists
                    cur.execute(f"SELECT name FROM sqlite_master WHERE type='table' AND name IN ('{table_name_lower}', '{table_name_original}')")
                    existing_table = cur.fetchone()
                    
                    if existing_table:
                        table_name = existing_table[0]
                    else:
                        # Use lowercase convention for new tables
                        table_name = table_name_lower
                    
                    # Recreate table to keep schema and remove deleted rows
                    cols_sql = ", ".join([f'"{c}" TEXT' for c in columns])
                    try:
                        # Quote table name if it contains spaces
                        quoted_table_name = f'"{table_name}"' if ' ' in table_name else table_name
                        cur.execute(f"CREATE TABLE IF NOT EXISTS {quoted_table_name} ({cols_sql})")
                        cur.execute(f"DELETE FROM {quoted_table_name}")
                        if columns and rows:
                            placeholders = ",".join(["?"] * len(columns))
                            insert_sql = f"INSERT INTO {quoted_table_name} ({', '.join([f'"{c}"' for c in columns])}) VALUES ({placeholders})"
                            # Normalize rows to column count
                            norm_rows = []
                            for r in rows:
                                rlist = list(r)
                                if len(rlist) < len(columns):
                                    rlist += [""] * (len(columns) - len(rlist))
                                norm_rows.append(tuple(None if (v == '' or v is None) else str(v) for v in rlist[:len(columns)]))
                            cur.executemany(insert_sql, norm_rows)
                        print(f"Successfully saved {pane_name} to table {table_name} with {len(rows)} rows")
                    except Exception as e:
                        print(f"Error saving {pane_name} to table {table_name}: {e}")
                        ErrorHandler.handle_ui_error(f"save pane {pane_name}", e)
                # Projects
                cur.execute("DELETE FROM projects")
                if getattr(self, 'projects', None):
                    cur.executemany("INSERT INTO projects(\"Project Name\", \"Project ID\") VALUES (?, ?)", [(n, pid) for (n, pid) in self.projects])
                # ISM Directory
                cur.execute("DELETE FROM ism_directory")
                if getattr(self, 'ism_directory', None):
                    cur.executemany("INSERT INTO ism_directory(Name) VALUES (?)", [(n,) for n in self.ism_directory])
                # Org Directory
                try:
                    directory = self._serialize_org_tree()
                except Exception:
                    directory = []
                cur.execute("DELETE FROM org_directory")
                if directory:
                    cur.executemany(
                        "INSERT INTO org_directory(\"Name\", \"Designation\", \"Enterprise ID\", \"Email ID\", \"Manager Enterprise ID\", \"Location\") VALUES (?,?,?,?,?,?)",
                        [
                            (
                                item.get('name',''),
                                item.get('designation',''),
                                item.get('ent_id',''),
                                item.get('email_id',''),
                                item.get('manager',''),
                                item.get('location','')
                            ) for item in directory
                        ]
                    )
                # Imports
                try:
                    imports = getattr(self, '_imports_log', []) or []
                    cur.execute("DELETE FROM imports")
                    if imports:
                        cur.executemany("INSERT INTO imports(Timestamp, Pane, Source, Rows) VALUES (?,?,?,?)", imports)
                except Exception:
                    pass
                # Logs
                cur.execute("DELETE FROM activity_logs")
                cur.execute("DELETE FROM change_log")
                cols6 = ["Timestamp","User","Action","Pane","Details","Level"]
                activity_rows = getattr(self, 'activity_log_data', []) or []
                change_rows = getattr(self, 'change_log_data', []) or []
                if activity_rows:
                    cur.executemany(f"INSERT INTO activity_logs({', '.join(cols6)}) VALUES (?,?,?,?,?,?)", activity_rows)
                if change_rows:
                    cur.executemany(f"INSERT INTO change_log({', '.join(cols6)}) VALUES (?,?,?,?,?,?)", change_rows)
                # Settings
                cur.execute("INSERT OR REPLACE INTO settings(setting, value, description) VALUES (?,?,?)", ("Backend Path", db_path, "Path to backend SQLite file"))
                conn.commit()
            
            # Update the backend path if it wasn't set
            if not hasattr(self, '__backend_path__') or not self.__backend_path__:
                self.__backend_path__ = db_path
            
            # Show success feedback
            if hasattr(self, '_save_badge') and self._save_badge is not None:
                self._save_badge.setText("Saved")
                try:
                    self._safe_timer_single_shot(1500, lambda: self._safe_set_badge_text(""))
                except Exception:
                    pass
            
            if hasattr(self, 'loading_overlay'):
                self.loading_overlay.hide_loading()
            
            if hasattr(self, 'show_save_status'):
                self.show_save_status("Data Saved", 3000)
            
            ErrorHandler.handle_success("save backend SQLite", f"Saved {len(self.data)} panes to {db_path}")
            
            # Refresh calendar after successful save
            try:
                if hasattr(self, 'calendar_tab') and self.calendar_tab:
                    self.calendar_tab.refresh_calendar_if_visible()
            except Exception:
                pass
            
        except Exception as e:
            if hasattr(self, 'loading_overlay'):
                self.loading_overlay.hide_loading()
            ErrorHandler.handle_ui_error("save backend SQLite", e)
            if hasattr(self, 'notifications'):
                self.notifications.show_error(f"Failed to save data: {str(e)}")
            if hasattr(self, '_save_badge') and self._save_badge is not None:
                self._save_badge.setText("Error")
                try:
                    self._safe_timer_single_shot(2000, lambda: self._safe_set_badge_text(""))
                except Exception:
                    pass
            raise

    def _backup_sqlite_database(self, backup_dir: str = None) -> bool:
        """Create a backup of the current SQLite database"""
        try:
            backend_path = getattr(self, '__backend_path__', None) or self.backend_sqlite_path
            if not backend_path or not str(backend_path).lower().endswith(('.sqlite', '.db')):
                QMessageBox.warning(self, "Backup Failed", "Backup is only available for SQLite databases.")
                return False
            
            if not os.path.exists(backend_path):
                QMessageBox.warning(self, "Backup Failed", "Backend database file not found.")
                return False
            
            # Get backup directory
            if not backup_dir:
                backup_dir = self.preferences.get("backup_directory", "")
                if not backup_dir or not os.path.exists(backup_dir):
                    backup_dir = QFileDialog.getExistingDirectory(
                        self, 
                        "Select Backup Directory",
                        os.path.dirname(backend_path) if backend_path else "",
                        QFileDialog.Option.ShowDirsOnly
                    )
                    if not backup_dir:
                        return False
                    # Save the selected directory as preference
                    self.preferences.set("backup_directory", backup_dir)
            
            # Generate backup filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            db_name = os.path.splitext(os.path.basename(backend_path))[0]
            backup_filename = f"{db_name}_backup_{timestamp}.sqlite"
            backup_path = os.path.join(backup_dir, backup_filename)
            
            # Create backup using SQLite's backup API
            with self._sqlite_connect(backend_path) as source_conn:
                with sqlite3.connect(backup_path) as backup_conn:
                    source_conn.backup(backup_conn)
            
            # Clean up old backups if backup_count is set
            backup_count = self.preferences.get("backup_count", 5)
            if backup_count > 0:
                self._cleanup_old_backups(backup_dir, db_name, backup_count)
            
            QMessageBox.information(
                self, 
                "Backup Successful", 
                f"Database backed up successfully to:\n{backup_path}"
            )
            return True
            
        except Exception as e:
            QMessageBox.critical(
                self, 
                "Backup Failed", 
                f"Failed to create backup:\n{str(e)}"
            )
            return False
    
    def _cleanup_old_backups(self, backup_dir: str, db_name: str, max_backups: int):
        """Remove old backup files to maintain backup count limit"""
        try:
            # Find all backup files for this database
            backup_pattern = f"{db_name}_backup_*.sqlite"
            backup_files = []
            
            for filename in os.listdir(backup_dir):
                if filename.startswith(f"{db_name}_backup_") and filename.endswith(".sqlite"):
                    file_path = os.path.join(backup_dir, filename)
                    if os.path.isfile(file_path):
                        backup_files.append((file_path, os.path.getmtime(file_path)))
            
            # Sort by modification time (newest first)
            backup_files.sort(key=lambda x: x[1], reverse=True)
            
            # Remove excess backups
            if len(backup_files) > max_backups:
                for file_path, _ in backup_files[max_backups:]:
                    try:
                        os.remove(file_path)
                        logger.info(f"Removed old backup: {file_path}")
                    except Exception as e:
                        logger.warning(f"Failed to remove old backup {file_path}: {e}")
                        
        except Exception as e:
            logger.warning(f"Failed to cleanup old backups: {e}")

    def _register_for_cleanup(self, obj, obj_type="widget"):
        """Register an object for cleanup when the application closes"""
        try:
            if obj_type == "timer":
                self._cleanup_timers.append(obj)
            elif obj_type == "widget":
                self._cleanup_widgets.append(obj)
            else:
                self._active_timers.append(obj)
        except Exception as e:
            # Use the consolidated ErrorHandler
            pass  # Will be fixed when we update all references

    def _handle_error_recovery(self, error_context: str, error: Exception) -> bool:
        """Handle error recovery and prevent application crashes"""
        try:
            current_time = datetime.now()
            
            # Track error frequency
            if self._last_error_time and (current_time - self._last_error_time).seconds < 5:
                self._error_count += 1
            else:
                self._error_count = 1
                
            self._last_error_time = current_time
            
            # If too many errors in short time, disable recovery temporarily
            if self._error_count > 10:
                self._error_recovery_enabled = False
                self.notifications.show_warning("Too many errors detected. Error recovery disabled temporarily.")
                return False
                
            # Basic recovery actions
            if "table" in error_context.lower():
                # Try to rebuild the problematic table
                try:
                    current_tab = self.tabs.currentWidget()
                    if hasattr(current_tab, 'objectName'):
                        pane_name = current_tab.objectName()
                        if pane_name in PANE_COLUMNS:
                            self.rebuild_table(pane_name)
                            self.notifications.show_info(f"Table {pane_name} rebuilt after error")
                            return True
                except Exception:
                    pass
                    
            elif "data" in error_context.lower():
                # Try to reload data
                try:
                    self._load_backend_sqlite(self.backend_sqlite_path)
                    self.notifications.show_info("Data reloaded after error")
                    return True
                except Exception:
                    pass
                    
            return False
            
        except Exception:
            return False

    def _cleanup_resources(self):
        """Clean up resources to prevent memory leaks"""
        try:
            # Clear all active timers
            for timer in self._active_timers:
                if timer and timer.isActive():
                    timer.stop()
                    timer.deleteLater()
            self._active_timers.clear()
            
            # Clear timers
            for timer in self._cleanup_timers:
                if timer and timer.isActive():
                    timer.stop()
                    timer.deleteLater()
            self._cleanup_timers.clear()
            
            # Clear widgets
            for widget in self._cleanup_widgets:
                if widget:
                    widget.deleteLater()
            self._cleanup_widgets.clear()
            
            # Clear specific timers
            if hasattr(self, 'log_cleanup_timer') and self.log_cleanup_timer:
                self.log_cleanup_timer.stop()
                self.log_cleanup_timer.deleteLater()
                
            if hasattr(self, '_save_status_timer') and self._save_status_timer:
                self._save_status_timer.stop()
                self._save_status_timer.deleteLater()
                
            if hasattr(self, '_main_thread_timer') and self._main_thread_timer:
                self._main_thread_timer.stop()
                self._main_thread_timer.deleteLater()
            
            # Clear large data structures
            if hasattr(self, 'data'):
                for pane in self.data:
                    if isinstance(self.data[pane], list):
                        self.data[pane].clear()
            
            # Clear other collections
            if hasattr(self, 'projects'):
                self.projects.clear()
            if hasattr(self, 'ism_directory'):
                self.ism_directory.clear()
            if hasattr(self, 'change_log_data'):
                self.change_log_data.clear()
            if hasattr(self, '_imports_log'):
                self._imports_log.clear()
                
        except Exception as e:
            ErrorHandler.handle_ui_error("cleanup", e)
            print(f"Error during cleanup: {e}")
    
    def _validate_project_details_data(self):
        """Validate Project Details data integrity"""
        try:
            if "Project Details" not in self.data:
                return True, "No Project Details data found"
            
            project_data = self.data["Project Details"]
            if not isinstance(project_data, list):
                return False, "Project Details data is not a list"
            
            expected_columns = len(PANE_COLUMNS["Project Details"])
            issues = []
            
            for i, row in enumerate(project_data):
                if not isinstance(row, list):
                    issues.append(f"Row {i}: Not a list")
                    continue
                
                if len(row) != expected_columns:
                    issues.append(f"Row {i}: Expected {expected_columns} columns, got {len(row)}")
                    continue
                
                # Validate required fields
                if not row[0] or not str(row[0]).strip():  # Project Name
                    issues.append(f"Row {i}: Project Name is required")
                if not row[1] or not str(row[1]).strip():  # Project ID
                    issues.append(f"Row {i}: Project ID is required")
            
            if issues:
                return False, f"Data validation issues: {'; '.join(issues)}"
            
            return True, f"Validated {len(project_data)} rows successfully"
            
        except Exception as e:
            ErrorHandler.handle_data_error("validate project details data", e)
            return False, f"Validation error: {str(e)}"

    # Projects management
    def load_projects_csv(self) -> None:
        path, _ = QFileDialog.getOpenFileName(self, "Load Projects CSV", "", "CSV (*.csv)")
        if not path:
            return
        try:
            df = pd.read_csv(path)
            # Expect columns Project Name, Project ID (case-insensitive)
            cols = {c.lower(): c for c in df.columns}
            name_col = cols.get("project name")
            id_col = cols.get("project id")
            if not name_col or not id_col:
                raise ValueError("CSV must include 'Project Name' and 'Project ID' columns")
            self.projects = list(zip(df[name_col].astype(str).tolist(), df[id_col].astype(str).tolist()))
            QMessageBox.information(self, "Projects", f"Loaded {len(self.projects)} projects")
            ErrorHandler.handle_success("load projects CSV", f"Loaded {len(self.projects)} projects from {path}")
        except Exception as e:
            ErrorHandler.handle_ui_error("load projects CSV", e, {"file_path": path})
            QMessageBox.critical(self, "Load Projects Failed", str(e))

    def add_project_manual(self) -> None:
        try:
            name, ok1 = QInputDialog.getText(self, "Add Project", "Project Name:")
            if not ok1 or not name:
                return
            pid, ok2 = QInputDialog.getText(self, "Add Project", "Project ID:")
            if not ok2 or not pid:
                return
            
            # Validate project name and ID
            is_valid_name, name_error = InputValidator.validate_required_field(name, "Project Name")
            if not is_valid_name:
                ErrorHandler.handle_validation_error("Project Name", name, name_error)
                return
                
            is_valid_id, id_error = InputValidator.validate_project_id(pid)
            if not is_valid_id:
                ErrorHandler.handle_validation_error("Project ID", pid, id_error)
                return
            
            # Check for duplicates
            if any(p[0] == name for p in self.projects):
                ErrorHandler.handle_warning("Add Project", f"Project with name '{name}' already exists")
                return
            if any(p[1] == pid for p in self.projects):
                ErrorHandler.handle_warning("Add Project", f"Project with ID '{pid}' already exists")
                return
            
            self.projects.append((name, pid))
            QMessageBox.information(self, "Projects", "Project added")
            ErrorHandler.handle_success("add project", f"Added project '{name}' with ID '{pid}'")
        except Exception as e:
            ErrorHandler.handle_ui_error("add project manually", e)

    def view_projects_dialog(self) -> None:
        dlg = QDialog(self)
        dlg.setWindowTitle("Projects Repository")
        lay = QVBoxLayout(dlg)
        tbl = QTableWidget(); tbl.setColumnCount(2); tbl.setHorizontalHeaderLabels(["Project Name","Project ID"])
        tbl.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        tbl.setRowCount(len(self.projects))
        for i,(n,pid) in enumerate(self.projects):
            tbl.setItem(i,0,QTableWidgetItem(str(n)))
            tbl.setItem(i,1,QTableWidgetItem(str(pid)))
        lay.addWidget(tbl)
        btns = QHBoxLayout(); close = QPushButton("Close"); close.clicked.connect(dlg.accept); btns.addStretch(1); btns.addWidget(close); lay.addLayout(btns)
        dlg.resize(600,400)
        dlg.exec()

    # Org Chart helpers
    def _find_org_item_by_ent_id(self, ent_id: str):
        root = self.org_tree.invisibleRootItem()
        stack = [root]
        while stack:
            cur = stack.pop()
            for i in range(cur.childCount()):
                ch = cur.child(i)
                if ch.text(2) == ent_id:
                    return ch
                stack.append(ch)
        return None

    def _is_descendant(self, potential_parent, node) -> bool:
        stack = [potential_parent]
        while stack:
            cur = stack.pop()
            if cur is node:
                return True
            for i in range(cur.childCount()):
                stack.append(cur.child(i))
        return False

    def _reparent_member(self, node, new_mgr_id: str) -> bool:
        if not new_mgr_id:
            return False
        target = self._find_org_item_by_ent_id(new_mgr_id)
        if target is None or target is node or self._is_descendant(node, target):
            QMessageBox.warning(self, "Move Member", "Invalid manager ID or cycle detected.")
            return False
        # Detach
        parent = node.parent()
        if parent:
            parent.removeChild(node)
        else:
            idx = self.org_tree.indexOfTopLevelItem(node)
            self.org_tree.takeTopLevelItem(idx)
        # Attach
        target.addChild(node)
        # Update manager field in the node
        if node.columnCount() > 4:
            node.setText(4, new_mgr_id)
        # Sync after move
        self._sync_ism_directory_with_org(); self._save_autosave()
        try:
            self._save_org_to_csv()
        except Exception as e:
            ErrorHandler.handle_ui_error("save org to csv", e)
        # Save to backend Excel
        self._save_backend_sqlite()
        return True

    def add_org_member(self, manager_ent_id: str | None = None) -> None:
        name, ok1 = QInputDialog.getText(self, "Add Member", "Name:")
        if not ok1 or not name:
            return
        desig, ok2 = QInputDialog.getText(self, "Add Member", "Designation:")
        if not ok2:
            return
        ent_id, ok3 = QInputDialog.getText(self, "Add Member", "Enterprise ID:")
        if not ok3 or not ent_id:
            return
        # Auto-fill email address using Enterprise ID + @accenture.com
        auto_email = f"{ent_id}@accenture.com"
        email_id, ok4 = QInputDialog.getText(self, "Add Member", "Email ID:", text=auto_email)
        if not ok4:
            return
        location, ok5 = QInputDialog.getText(self, "Add Member", "Location:")
        if not ok5:
            return
        prompt = "Manager Enterprise ID (optional):"
        default_mgr = manager_ent_id or ""
        mgr_id, _ = QInputDialog.getText(self, "Add Member", prompt, text=default_mgr)
        # Create tree item with default avatar
        cols = [name, desig, ent_id, email_id, mgr_id, location]
        node = QTreeWidgetItem(cols)
        node.setIcon(0, QIcon(self._default_avatar(name, 32)))
        # Attach under manager if found
        parent = self._find_org_item_by_ent_id(mgr_id) if mgr_id else None
        if parent is None:
            self.org_tree.addTopLevelItem(node)
        else:
            parent.addChild(node)
        self._log(f"Added org member {name}")
        self.render_org_chart()
        # Persist org changes immediately
        self._sync_ism_directory_with_org()
        self._save_autosave()
        try:
            self._save_org_to_csv()
        except Exception as e:
            ErrorHandler.handle_ui_error("save org to csv", e)
        # Save to backend Excel
        self._save_backend_sqlite()
        
    def add_sample_org_data(self) -> None:
        """Add sample organizational data for testing"""
        sample_data = [
            {"name": "John Smith", "designation": "Senior Manager", "ent_id": "john.smith", "email_id": "john.smith@accenture.com", "manager": "", "location": "New York"},
            {"name": "Sarah Johnson", "designation": "Manager", "ent_id": "sarah.johnson", "email_id": "sarah.johnson@accenture.com", "manager": "john.smith", "location": "San Francisco"},
            {"name": "Mike Wilson", "designation": "Associate Manager", "ent_id": "mike.wilson", "email_id": "mike.wilson@accenture.com", "manager": "sarah.johnson", "location": "Seattle"},
            {"name": "Lisa Brown", "designation": "Specialist", "ent_id": "lisa.brown", "email_id": "lisa.brown@accenture.com", "manager": "mike.wilson", "location": "Austin"},
            {"name": "David Lee", "designation": "Senior Analyst", "ent_id": "david.lee", "email_id": "david.lee@accenture.com", "manager": "lisa.brown", "location": "Chicago"},
            {"name": "Emma Davis", "designation": "Analyst", "ent_id": "emma.davis", "email_id": "emma.davis@accenture.com", "manager": "david.lee", "location": "Boston"},
            {"name": "Alex Chen", "designation": "Senior Manager", "ent_id": "alex.chen", "email_id": "alex.chen@accenture.com", "manager": "john.smith", "location": "London"},
            {"name": "Maria Garcia", "designation": "Manager", "ent_id": "maria.garcia", "email_id": "maria.garcia@accenture.com", "manager": "alex.chen", "location": "Madrid"},
            {"name": "James Wilson", "designation": "Specialist", "ent_id": "james.wilson", "email_id": "james.wilson@accenture.com", "manager": "maria.garcia", "location": "Dublin"},
            {"name": "Anna Kumar", "designation": "Senior Analyst", "ent_id": "anna.kumar", "email_id": "anna.kumar@accenture.com", "manager": "james.wilson", "location": "Bangalore"},
        ]
        
        # Clear existing data
        self.org_tree.clear()
        
        # Add sample data
        for entry in sample_data:
            name = entry["name"]
            desig = entry["designation"]
            ent_id = entry["ent_id"]
            email_id = entry["email_id"]
            manager_id = entry["manager"]
            location = entry["location"]
            
            cols = [name, desig, ent_id, email_id, manager_id, location]
            node = QTreeWidgetItem(cols)
            node.setIcon(0, QIcon(self._default_avatar(name, 32)))
            
            # Find manager if specified
            parent = self._find_org_item_by_ent_id(manager_id) if manager_id else None
            if parent is None:
                self.org_tree.addTopLevelItem(node)
            else:
                parent.addChild(node)
        
        self.render_org_chart()
        self._sync_ism_directory_with_org()
        self._save_autosave()
        try:
            self._save_org_to_csv()
        except Exception as e:
            ErrorHandler.handle_ui_error("save org to csv", e)
        # Save to backend Excel
        self._save_backend_sqlite()
        QMessageBox.information(self, "Sample Data", "Sample organizational data has been added!")
    def _open_org_context(self, pos) -> None:
        item = self.org_tree.itemAt(pos)
        if item is None:
            return
        menu = QMenu(self)
        # Context menu with non-destructive and removal options
        view_details = QAction("View Details", self)
        rename_act = QAction("Rename/Edit", self)
        def show_details():
            name = item.text(0)
            desig = item.text(1)
            ent = item.text(2)
            email = item.text(3)
            mgr = item.text(4)
            loc = item.text(5)
            msg = QMessageBox(self)
            msg.setWindowTitle("Member Details")
            msg.setText(f"Name: {name}\nDesignation: {desig}\nEnterprise ID: {ent}\nEmail ID: {email}\nManager Enterprise ID: {mgr}\nLocation: {loc}")
            # Show default avatar
            msg.setIconPixmap(self._default_avatar(name, 64))
            msg.addButton("Close", QMessageBox.ButtonRole.AcceptRole)
            msg.exec()
        view_details.triggered.connect(show_details)
        def do_rename():
            name, ok1 = QInputDialog.getText(self, "Edit Member", "Name:", text=item.text(0))
            if not ok1:
                return
            desig, ok2 = QInputDialog.getText(self, "Edit Member", "Designation:", text=item.text(1))
            if not ok2:
                return
            ent, ok3 = QInputDialog.getText(self, "Edit Member", "Enterprise ID:", text=item.text(2))
            if not ok3:
                return
            # Auto-fill email address using Enterprise ID + @accenture.com
            current_email = item.text(3) if item.columnCount() > 3 else ""
            auto_email = f"{ent}@accenture.com"
            email, ok4 = QInputDialog.getText(self, "Edit Member", "Email ID:", text=auto_email)
            if not ok4:
                return
            mgr, ok5 = QInputDialog.getText(self, "Edit Member", "Manager Enterprise ID:", text=item.text(4) if item.columnCount() > 4 else "")
            if not ok5:
                return
            loc, ok6 = QInputDialog.getText(self, "Edit Member", "Location:", text=item.text(5) if item.columnCount() > 5 else "")
            if not ok6:
                return
            item.setText(0, name); item.setText(1, desig); item.setText(2, ent); item.setText(3, email); item.setText(4, mgr); item.setText(5, loc)
            self.render_org_chart(); self._sync_ism_directory_with_org(); self._save_autosave();
            try:
                self._save_org_to_csv()
            except Exception as e:
                ErrorHandler.handle_ui_error("save org to csv", e)
            # Save to backend Excel
            self._save_backend_sqlite()
        rename_act.triggered.connect(do_rename)
        expand = QAction("Expand All Under", self, triggered=lambda: item.setExpanded(True))
        collapse = QAction("Collapse", self, triggered=lambda: item.setExpanded(False))
        remove_act = QAction("Remove Member", self)
        def do_remove():
            # Confirm deletion (will remove any children as well)
            total_children = item.childCount()
            msg = "Remove this member?"
            if total_children > 0:
                msg = f"Remove this member and {total_children} direct report(s)?"
            res = QMessageBox.question(self, "Confirm Removal", msg, QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if res != QMessageBox.StandardButton.Yes:
                return
            parent = item.parent()
            if parent:
                parent.removeChild(item)
            else:
                idx = self.org_tree.indexOfTopLevelItem(item)
                self.org_tree.takeTopLevelItem(idx)
            self.render_org_chart()
            self._sync_ism_directory_with_org(); self._save_autosave()
            # Save to backend Excel
            self._save_backend_sqlite()
        remove_act.triggered.connect(do_remove)
        menu.addAction(view_details)
        menu.addAction(rename_act)
        menu.addSeparator()
        menu.addAction(expand)
        menu.addAction(collapse)
        menu.addSeparator()
        menu.addAction(remove_act)
        menu.exec(self.org_tree.viewport().mapToGlobal(pos))

    # --- Org Graph Rendering ---
    def _collect_org_roots(self) -> list:
        roots = []
        for i in range(self.org_tree.topLevelItemCount()):
            roots.append(self.org_tree.topLevelItem(i))
        return roots

    def _default_avatar(self, name: str, size: int = 56) -> QPixmap:
        # Generate a modern circular avatar with prominent initials
        initials = "".join([part[0].upper() for part in (name or "?").split() if part][:2]) or "?"
        pix = QPixmap(size, size)
        pix.fill(Qt.GlobalColor.transparent)
        painter = QPainter(pix)
        
        # Enhanced color palette with better contrast
        palette = ["#2563eb", "#059669", "#7c3aed", "#d97706", "#dc2626", "#0891b2", "#be185d", "#65a30d"]
        color = QColor(palette[hash(name) % len(palette)])
        
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        painter.setRenderHint(QPainter.RenderHint.TextAntialiasing)
        
        # Draw background circle with subtle gradient effect
        painter.setBrush(QBrush(color))
        painter.setPen(Qt.PenStyle.NoPen)
        painter.drawEllipse(2, 2, size-4, size-4)
        
        # Add subtle inner highlight
        highlight_color = QColor(255, 255, 255, 30)
        painter.setBrush(QBrush(highlight_color))
        painter.drawEllipse(4, 4, size-8, size-8)
        
        # Draw initials with compact styling - subtle but readable
        painter.setPen(QColor("#ffffff"))
        font = painter.font()
        font.setBold(True)
        font.setPointSize(int(size * 0.38))  # Compact size for subtle appearance
        font.setWeight(QFont.Weight.Bold)  # Bold for readability
        painter.setFont(font)
        
        # Draw text with shadow for better readability
        text_rect = QRectF(0, 0, size, size)
        painter.drawText(text_rect, Qt.AlignmentFlag.AlignCenter, initials)
        
        painter.end()
        return pix

    def _measure_subtree(self, node) -> int:
        # Returns required width in units (card widths)
        if node.childCount() == 0:
            return 1
        return max(1, sum(self._measure_subtree(node.child(i)) for i in range(node.childCount())))

    def _org_card_colors(self, level: int, designation: str = "", color_scheme: str = "level") -> tuple[str, str]:
        # return (header_bg, body_bg)
        
        # Designation-based color mapping
        designation_colors = {
            "Senior Manager": ("#e74c3c", "#fdf2f2"),  # Red
            "Manager": ("#f39c12", "#fef7ed"),        # Orange
            "Associate Manager": ("#f1c40f", "#fefce8"), # Amber
            "Specialist": ("#9b59b6", "#faf5ff"),      # Purple
            "Senior Analyst": ("#3498db", "#eff6ff"),  # Blue
            "Analyst": ("#2ecc71", "#f0fdf4"),         # Green
        }
        
        # Level-based color mapping (original)
        level_palettes = [
            ("#1fb6ff", "#e6f7ff"),  # top teal
            ("#2bd576", "#ecfbf3"),  # green
            ("#f6a23c", "#fff3e6"),  # orange
            ("#8e44ad", "#f3e8ff"),  # purple
            ("#f55f77", "#ffe9ee"),  # rose
        ]
        
        # Combined color mapping (designation + level)
        combined_colors = {
            "Senior Manager": [
                ("#c0392b", "#fdf2f2"),  # Dark red
                ("#e74c3c", "#fdf2f2"),  # Red
                ("#ec7063", "#fdf2f2"),  # Light red
            ],
            "Manager": [
                ("#d68910", "#fef7ed"),  # Dark orange
                ("#f39c12", "#fef7ed"),  # Orange
                ("#f7dc6f", "#fef7ed"),  # Light orange
            ],
            "Associate Manager": [
                ("#d4ac0d", "#fefce8"),  # Dark yellow
                ("#f1c40f", "#fefce8"),  # Amber
                ("#f4d03f", "#fefce8"),  # Light yellow
            ],
            "Specialist": [
                ("#8e44ad", "#faf5ff"),  # Dark purple
                ("#9b59b6", "#faf5ff"),  # Purple
                ("#bb8fce", "#faf5ff"),  # Light purple
            ],
            "Senior Analyst": [
                ("#2980b9", "#eff6ff"),  # Dark blue
                ("#3498db", "#eff6ff"),  # Blue
                ("#85c1e9", "#eff6ff"),  # Light blue
            ],
            "Analyst": [
                ("#27ae60", "#f0fdf4"),  # Dark green
                ("#2ecc71", "#f0fdf4"),  # Green
                ("#82e0aa", "#f0fdf4"),  # Light green
            ],
        }
        
        if color_scheme == "designation" and designation in designation_colors:
            return designation_colors[designation]
        elif color_scheme == "combined" and designation in combined_colors:
            return combined_colors[designation][level % len(combined_colors[designation])]
        else:
            # Default to level-based coloring
            return level_palettes[level % len(level_palettes)]

    def render_org_chart(self) -> None:
        if not hasattr(self, 'org_scene'):
            return
        self.org_scene.clear()
        self._org_item_map.clear()

        card_w = 240
        card_h = 100
        h_gap = 30
        v_gap = 60

        def layout(node, x, y, level) -> tuple[float, float]:
            # returns center x of node
            subtree_units = self._measure_subtree(node)
            total_width = subtree_units * (card_w + h_gap) - h_gap
            left = x - total_width / 2
            # Place children first to compute center
            child_centers = []
            cur = left
            for i in range(node.childCount()):
                child = node.child(i)
                units = self._measure_subtree(child)
                width = units * (card_w + h_gap) - h_gap
                cx, cy = layout(child, cur + width / 2, y + card_h + v_gap, level + 1)
                child_centers.append((cx, cy))
                cur += width + h_gap

            # Draw this node
            center_x = x
            top_left_x = center_x - card_w / 2
            designation = node.text(1)  # Get designation from the node
            header_bg, body_bg = self._org_card_colors(level, designation, self._org_color_scheme)
            # Simple rounded card background
            path = QPainterPath()
            path.addRoundedRect(QRectF(top_left_x, y, card_w, card_h), 8, 8)
            rect = self.org_scene.addPath(path, QPen(QColor("#d1d5db"), 1), QBrush(QColor(body_bg)))
            
            # Simple header
            header = self.org_scene.addRect(top_left_x, y, card_w, 30, QPen(Qt.PenStyle.NoPen), QBrush(QColor(header_bg)))
            # Title
            name = node.text(0)
            desig = node.text(1)
            ent = node.text(2)
            email = node.text(3) if node.columnCount() > 3 else ""
            mgr = node.text(4) if node.columnCount() > 4 else ""
            loc = node.text(5) if node.columnCount() > 5 else ""
            txt_name = self.org_scene.addText(name)
            f = txt_name.font(); f.setPointSize(10); f.setBold(True); txt_name.setFont(f)
            txt_name.setDefaultTextColor(QColor("#ffffff"))
            txt_name.setPos(top_left_x + 8, y + 6)
            # Body - Only show Designation and Location on card
            def add_line(t, i):
                item = self.org_scene.addText(t)
                ff = item.font(); ff.setPointSize(8); item.setFont(ff)
                item.setDefaultTextColor(QColor("#374151"))
                item.setPos(top_left_x + 70, y + 35 + i * 15)
                return item
            add_line(f"Designation: {desig}", 0)
            if loc:
                add_line(f"Location: {loc}", 1)

            # Avatar (always use default)
            pix = self._default_avatar(name, 40)
            pm = self.org_scene.addPixmap(pix)
            pm.setPos(top_left_x + 12, y + 35)

            # Map selection & tooltip
            rect.setData(0, node)
            header.setData(0, node)
            self._org_item_map[rect] = node
            self._org_item_map[header] = node
            tooltip = f"<b>{name}</b><br/>{desig}<br/>Enterprise ID: {ent}"
            if email:
                tooltip += f"<br/>Email ID: {email}"
            if mgr:
                tooltip += f"<br/>Manager: {mgr}"
            if loc:
                tooltip += f"<br/>Location: {loc}"
            rect.setToolTip(tooltip)
            header.setToolTip(tooltip)

            # Connectors
            dot_color = QColor(header_bg)
            for cx, cy in child_centers:
                # vertical from this bottom center to mid gap
                parent_bottom = y + card_h
                child_top = cy
                mid_y = parent_bottom + v_gap / 2
                pen = QPen(QColor("#94a3b8"), 2)
                self.org_scene.addLine(center_x, parent_bottom, center_x, mid_y, pen)
                self.org_scene.addEllipse(center_x - 3, mid_y - 3, 6, 6, QPen(Qt.PenStyle.NoPen), QBrush(dot_color))
                self.org_scene.addLine(cx, mid_y, cx, child_top, pen)
                self.org_scene.addLine(center_x, mid_y, cx, mid_y, pen)
            return center_x, y

        roots = self._collect_org_roots()
        if not roots:
            # Empty state
            msg = self.org_scene.addText("No members yet. Use 'Add Team Member' to get started.")
            f = msg.font(); f.setPointSize(12); msg.setFont(f)
            msg.setDefaultTextColor(QColor("#64748b"))
            msg.setPos(40, 40)
            return

        x_cursor = 0
        total_units = sum(self._measure_subtree(r) for r in roots)
        total_width = total_units * (card_w + h_gap) - h_gap
        x_origin = total_width / 2
        cur = 0
        for r in roots:
            units = self._measure_subtree(r)
            width = units * (card_w + h_gap) - h_gap
            layout(r, -x_origin + cur + width / 2 + 40, 20, 0)
            cur += width + h_gap
        # Center the view on the content instead of fitting to view
        if self.org_scene.itemsBoundingRect().isValid():
            self.org_view.centerOn(self.org_scene.itemsBoundingRect().center())

    def _open_org_graph_context(self, pos) -> None:
        if not hasattr(self, 'org_view'):
            return
        scene_pos = self.org_view.mapToScene(pos)
        items = self.org_scene.items(scene_pos)
        node_item = None
        linked = None
        for it in items:
            linked = it.data(0)
            if linked:
                node_item = it
                break
        menu = QMenu(self)
        if linked is None:
            act = QAction("Add Member (Top)", self, triggered=lambda: self.add_org_member())
            menu.addAction(act)
        else:
            view_details = QAction("View Details", self)
            def show_details():
                name = linked.text(0); desig = linked.text(1); ent = linked.text(2); email = linked.text(3); mgr = linked.text(4); loc = linked.text(5)
                msg = QMessageBox(self)
                msg.setWindowTitle("Member Details")
                msg.setText(f"Name: {name}\nDesignation: {desig}\nEnterprise ID: {ent}\nEmail ID: {email}\nManager Enterprise ID: {mgr}\nLocation: {loc}")
                msg.setIconPixmap(self._default_avatar(name, 64))
                msg.exec()
            view_details.triggered.connect(show_details)
            add_report = QAction("Add Direct Report", self, triggered=lambda: self.add_org_member(linked.text(2)))
            move_under = QAction("Change Manager...", self)
            def do_move():
                mgr_id, ok = QInputDialog.getText(self, "Change Manager", "New Manager Enterprise ID:")
                if not ok or not mgr_id:
                    return
                if self._reparent_member(linked, mgr_id):
                    self.render_org_chart()
            move_under.triggered.connect(do_move)
            remove_act = QAction("Remove Member", self)
            def do_remove():
                res = QMessageBox.question(self, "Confirm Removal", "Remove this member?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                if res != QMessageBox.StandardButton.Yes:
                    return
                parent = linked.parent()
                if parent:
                    parent.removeChild(linked)
                else:
                    idx = self.org_tree.indexOfTopLevelItem(linked)
                    self.org_tree.takeTopLevelItem(idx)
                self.render_org_chart()
                # Save to backend Excel
                self._save_backend_sqlite()
            remove_act.triggered.connect(do_remove)
            rename_act = QAction("Rename/Edit", self)
            def do_rename_graph():
                name, ok1 = QInputDialog.getText(self, "Edit Member", "Name:", text=linked.text(0))
                if not ok1:
                    return
                desig, ok2 = QInputDialog.getText(self, "Edit Member", "Designation:", text=linked.text(1))
                if not ok2:
                    return
                ent, ok3 = QInputDialog.getText(self, "Edit Member", "Enterprise ID:", text=linked.text(2))
                if not ok3:
                    return
                # Auto-fill email address using Enterprise ID + @accenture.com
                current_email = linked.text(3) if linked.columnCount() > 3 else ""
                auto_email = f"{ent}@accenture.com"
                email, ok4 = QInputDialog.getText(self, "Edit Member", "Email ID:", text=auto_email)
                if not ok4:
                    return
                mgr, ok5 = QInputDialog.getText(self, "Edit Member", "Manager Enterprise ID:", text=linked.text(4) if linked.columnCount() > 4 else "")
                if not ok5:
                    return
                loc, ok6 = QInputDialog.getText(self, "Edit Member", "Location:", text=linked.text(5) if linked.columnCount() > 5 else "")
                if not ok6:
                    return
                linked.setText(0, name); linked.setText(1, desig); linked.setText(2, ent); linked.setText(3, email); linked.setText(4, mgr); linked.setText(5, loc)
                self.render_org_chart(); self._sync_ism_directory_with_org(); self._save_autosave()
                # Save to backend Excel
                self._save_backend_sqlite()
            rename_act.triggered.connect(do_rename_graph)
            menu.addAction(view_details)
            menu.addSeparator()
            menu.addAction(add_report)
            menu.addAction(move_under)
            menu.addAction(rename_act)
            menu.addSeparator()
            menu.addAction(remove_act)
        menu.exec(self.org_view.viewport().mapToGlobal(pos))

    def _on_color_scheme_changed(self, scheme_text: str):
        """Handle color scheme change"""
        scheme_map = {
            "Level-based": "level",
            "Designation-based": "designation", 
            "Combined": "combined"
        }
        self._org_color_scheme = scheme_map.get(scheme_text, "level")
        self.render_org_chart()
        self._log(f"Changed org chart color scheme to: {scheme_text}")

    def _show_color_scheme_info(self):
        """Show color scheme information dialog"""
        dlg = QDialog(self)
        dlg.setWindowTitle("Org Chart Color Schemes")
        dlg.resize(600, 500)
        
        layout = QVBoxLayout(dlg)
        
        # Title
        title = QLabel("🎨 Org Chart Color Schemes")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        layout.addWidget(title)
        
        # Level-based scheme
        level_group = QGroupBox("Level-based Colors")
        level_layout = QVBoxLayout(level_group)
        level_layout.addWidget(QLabel("Colors are determined by hierarchy level (top to bottom):"))
        
        level_colors = [
            ("Level 0 (Top)", "#1fb6ff", "#e6f7ff"),
            ("Level 1", "#2bd576", "#ecfbf3"),
            ("Level 2", "#f6a23c", "#fff3e6"),
            ("Level 3", "#8e44ad", "#f3e8ff"),
            ("Level 4+", "#f55f77", "#ffe9ee")
        ]
        
        for level, header, body in level_colors:
            color_row = QHBoxLayout()
            color_row.addWidget(QLabel(f"{level}:"))
            color_row.addWidget(self._create_color_preview(header, body))
            color_row.addStretch()
            level_layout.addLayout(color_row)
        
        layout.addWidget(level_group)
        
        # Designation-based scheme
        desig_group = QGroupBox("Designation-based Colors")
        desig_layout = QVBoxLayout(desig_group)
        desig_layout.addWidget(QLabel("Colors are determined by employee designation:"))
        
        desig_colors = [
            ("Senior Manager", "#e74c3c", "#fdf2f2"),
            ("Manager", "#f39c12", "#fef7ed"),
            ("Associate Manager", "#f1c40f", "#fefce8"),
            ("Specialist", "#9b59b6", "#faf5ff"),
            ("Senior Analyst", "#3498db", "#eff6ff"),
            ("Analyst", "#2ecc71", "#f0fdf4")
        ]
        
        for desig, header, body in desig_colors:
            color_row = QHBoxLayout()
            color_row.addWidget(QLabel(f"{desig}:"))
            color_row.addWidget(self._create_color_preview(header, body))
            color_row.addStretch()
            desig_layout.addLayout(color_row)
        
        layout.addWidget(desig_group)
        
        # Combined scheme
        combined_group = QGroupBox("Combined Colors")
        combined_layout = QVBoxLayout(combined_group)
        combined_layout.addWidget(QLabel("Colors combine designation with hierarchy level:"))
        combined_layout.addWidget(QLabel("• Each designation has 3 color variations (dark, medium, light)"))
        combined_layout.addWidget(QLabel("• Level determines which variation to use"))
        combined_layout.addWidget(QLabel("• Provides both role and hierarchy information"))
        
        layout.addWidget(combined_group)
        
        # Close button
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dlg.accept)
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        layout.addWidget(close_btn)
        
        dlg.exec()
    
    def _create_color_preview(self, header_color: str, body_color: str) -> QFrame:
        """Create a color preview widget"""
        preview = QFrame()
        preview.setFixedSize(60, 30)
        preview.setStyleSheet(f"""
            QFrame {{
                background-color: {body_color};
                border: 2px solid {header_color};
                border-radius: 4px;
            }}
        """)
        return preview
class ProjectDashboardWindow(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Project Summary Dashboard")
        self.setGeometry(100, 100, 1500, 1000)
        
        # Set window style
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f5f5f5;
            }
            QGroupBox {
                font-weight: bold;
                font-size: 12px;
                color: #2c3e50;
                border: 1px solid #bdc3c7;
                border-radius: 4px;
                margin-top: 8px;
                padding-top: 8px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 8px;
                padding: 0 4px 0 4px;
            }
        """)
        
        # Create central widget and layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)
        layout.setSpacing(20)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Create header section
        header_layout = QHBoxLayout()
        title_label = QLabel("Project Summary Dashboard")
        title_label.setStyleSheet("font-size: 18px; font-weight: bold; color: #2c3e50; margin: 8px 0;")
        header_layout.addWidget(title_label)
        header_layout.addStretch()
        
        # Refresh button in header
        self.refresh_button = QPushButton("Refresh Data")
        self.refresh_button.setStyleSheet("""
            QPushButton {
                background-color: #34495e;
                color: white;
                padding: 8px 16px;
                border: none;
                border-radius: 4px;
                font-size: 11px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2c3e50;
            }
        """)
        self.refresh_button.clicked.connect(self.refresh_data)
        header_layout.addWidget(self.refresh_button)
        
        layout.addLayout(header_layout)
        
        # Create summary section with enterprise look
        summary_group = QGroupBox("Project Overview")
        summary_layout = QVBoxLayout(summary_group)
        summary_layout.setSpacing(15)
        
        # First row - Main metrics
        first_row = QHBoxLayout()
        first_row.setSpacing(15)
        
        # Total projects
        self.total_projects_card = self.create_enterprise_card("Total Projects", "0", "#2c3e50")
        first_row.addWidget(self.total_projects_card)
        
        # Active projects
        self.active_projects_card = self.create_enterprise_card("Active Projects", "0", "#27ae60")
        first_row.addWidget(self.active_projects_card)
        
        # ISM hours
        self.ism_hours_card = self.create_enterprise_card("ISM Hours", "0", "#3498db")
        first_row.addWidget(self.ism_hours_card)
        
        # Voice solutions
        self.voice_solutions_card = self.create_enterprise_card("Voice Solutions", "0", "#8e44ad")
        first_row.addWidget(self.voice_solutions_card)
        
        # Audit required
        self.audit_required_card = self.create_enterprise_card("Audit Required", "0", "#e74c3c")
        first_row.addWidget(self.audit_required_card)
        
        first_row.addStretch()
        summary_layout.addLayout(first_row)
        
        # Second row - Additional metrics
        second_row = QHBoxLayout()
        second_row.setSpacing(15)
        
        # Completed projects
        self.completed_projects_card = self.create_enterprise_card("Completed", "0", "#16a085")
        second_row.addWidget(self.completed_projects_card)
        
        # In progress
        self.in_progress_card = self.create_enterprise_card("In Progress", "0", "#f39c12")
        second_row.addWidget(self.in_progress_card)
        
        # Pending
        self.pending_card = self.create_enterprise_card("Pending", "0", "#95a5a6")
        second_row.addWidget(self.pending_card)
        
        # Overdue
        self.overdue_card = self.create_enterprise_card("Overdue", "0", "#e67e22")
        second_row.addWidget(self.overdue_card)
        
        # High priority
        self.high_priority_card = self.create_enterprise_card("High Priority", "0", "#c0392b")
        second_row.addWidget(self.high_priority_card)
        
        second_row.addStretch()
        summary_layout.addLayout(second_row)
        layout.addWidget(summary_group)
        
        # ISM Summary section
        ism_group = QGroupBox("ISM Summary")
        ism_layout = QVBoxLayout(ism_group)
        ism_layout.setSpacing(10)
        
        # ISM table
        self.ism_table = QTableWidget()
        self.ism_table.setColumnCount(4)
        self.ism_table.setHorizontalHeaderLabels(["ISM Name", "Primary Projects", "Secondary Projects", "Total Projects"])
        self.ism_table.horizontalHeader().setStretchLastSection(True)
        self.ism_table.setMinimumHeight(300)
        self.ism_table.setMaximumHeight(400)
        self.ism_table.setAlternatingRowColors(True)
        self.ism_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.ism_table.setStyleSheet("""
            QTableWidget {
                gridline-color: #bdc3c7;
                background-color: white;
                alternate-background-color: #f8f9fa;
                border: 1px solid #dee2e6;
            }
            QHeaderView::section {
                background-color: #34495e;
                color: white;
                padding: 8px;
                border: none;
                font-weight: bold;
            }
            QTableWidget::item {
                padding: 8px;
                border-bottom: 1px solid #ecf0f1;
            }
            QTableWidget::item:selected {
                background-color: #3498db;
                color: white;
            }
        """)
        
        ism_layout.addWidget(self.ism_table)
        layout.addWidget(ism_group)
        
        # Project Filters section
        filters_group = QGroupBox("Project Filters & Actions")
        filters_layout = QVBoxLayout(filters_group)
        filters_layout.setSpacing(15)
        
        # Filter buttons row
        filters_row = QHBoxLayout()
        filters_row.setSpacing(15)
        
        self.audit_yes_card = self.create_enterprise_card("Audit Required", "0", "#e74c3c")
        self.audit_yes_card.clicked.connect(self.show_audit_projects)
        filters_row.addWidget(self.audit_yes_card)
        
        self.high_card = self.create_enterprise_card("High Priority", "0", "#c0392b")
        self.high_card.clicked.connect(lambda: self.show_priority_projects("High"))
        filters_row.addWidget(self.high_card)
        
        self.medium_card = self.create_enterprise_card("Medium Priority", "0", "#f39c12")
        self.medium_card.clicked.connect(lambda: self.show_priority_projects("Medium"))
        filters_row.addWidget(self.medium_card)
        
        self.low_card = self.create_enterprise_card("Low Priority", "0", "#27ae60")
        self.low_card.clicked.connect(lambda: self.show_priority_projects("Low"))
        filters_row.addWidget(self.low_card)
        
        # Additional filter cards
        self.voice_solutions_filter_card = self.create_enterprise_card("Voice Solutions", "0", "#8e44ad")
        filters_row.addWidget(self.voice_solutions_filter_card)
        
        self.ism_hours_filter_card = self.create_enterprise_card("ISM Hours", "0", "#3498db")
        filters_row.addWidget(self.ism_hours_filter_card)
        
        filters_row.addStretch()
        filters_layout.addLayout(filters_row)
        layout.addWidget(filters_group)
        
        # Footer with close button
        footer_layout = QHBoxLayout()
        footer_layout.addStretch()
        
        self.close_button = QPushButton("✖️ Close Dashboard")
        self.close_button.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                padding: 12px 24px;
                border: none;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        self.close_button.clicked.connect(self.close)
        footer_layout.addWidget(self.close_button)
        
        layout.addLayout(footer_layout)
        
        # Load initial data
        self.refresh_data()
    
    def create_summary_card(self, title, value, color, icon):
        """Create a summary card widget"""
        card = QFrame()
        card.setFixedSize(200, 120)
        card.setStyleSheet(f"""
            QFrame {{
                background-color: white;
                border: 2px solid {color};
                border-radius: 10px;
                margin: 5px;
            }}
            QFrame:hover {{
                border-width: 3px;
                background-color: #f8f9fa;
            }}
        """)
        
        layout = QVBoxLayout(card)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(8)
        
        # Icon and title row
        top_row = QHBoxLayout()
        icon_label = QLabel(icon)
        icon_label.setStyleSheet("font-size: 24px;")
        title_label = QLabel(title)
        title_label.setStyleSheet("font-size: 12px; font-weight: bold; color: #2c3e50;")
        title_label.setWordWrap(True)
        
        top_row.addWidget(icon_label)
        top_row.addWidget(title_label)
        top_row.addStretch()
        
        # Value
        value_label = QLabel(value)
        value_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        value_label.setStyleSheet(f"font-size: 28px; font-weight: bold; color: {color};")
        
        layout.addLayout(top_row)
        layout.addWidget(value_label)
        layout.addStretch()
        
        return card
    
    def create_clickable_card(self, title, count, color, icon):
        """Create a clickable card widget"""
        card = QPushButton()
        card.setFixedSize(180, 100)
        card.setStyleSheet(f"""
            QPushButton {{
                background-color: {color};
                color: white;
                border: none;
                border-radius: 10px;
                font-size: 14px;
                font-weight: bold;
                margin: 5px;
            }}
            QPushButton:hover {{
                background-color: {self.darken_color(color)};
            }}
        """)
        
        layout = QVBoxLayout(card)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(5)
        
        # Icon and title
        top_row = QHBoxLayout()
        icon_label = QLabel(icon)
        icon_label.setStyleSheet("font-size: 20px;")
        title_label = QLabel(title)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setStyleSheet("color: white; font-size: 11px; font-weight: bold;")
        title_label.setWordWrap(True)
        
        top_row.addWidget(icon_label)
        top_row.addWidget(title_label)
        top_row.addStretch()
        
        # Count
        count_label = QLabel(count)
        count_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        count_label.setStyleSheet("color: white; font-size: 24px; font-weight: bold;")
        
        layout.addLayout(top_row)
        layout.addWidget(count_label)
        
        return card
    
    def darken_color(self, color):
        """Darken a hex color for hover effect"""
        color_map = {
            "#e74c3c": "#c0392b",
            "#f39c12": "#d68910", 
            "#27ae60": "#229954",
            "#95a5a6": "#7f8c8d"
        }
        return color_map.get(color, color)
    
    def refresh_data(self):
        """Refresh the dashboard data"""
        # In real implementation, this would load data from your database
        # For now, we'll show empty/zero values to indicate no data loaded
        
        # Update main overview cards
        self.update_enterprise_card(self.total_projects_card, "0")
        self.update_enterprise_card(self.active_projects_card, "0")
        self.update_enterprise_card(self.ism_hours_card, "0")
        self.update_enterprise_card(self.voice_solutions_card, "0")
        self.update_enterprise_card(self.audit_required_card, "0")
        
        # Update second row cards
        self.update_enterprise_card(self.completed_projects_card, "0")
        self.update_enterprise_card(self.in_progress_card, "0")
        self.update_enterprise_card(self.pending_card, "0")
        self.update_enterprise_card(self.overdue_card, "0")
        self.update_enterprise_card(self.high_priority_card, "0")
        
        # Update filter cards
        self.update_enterprise_card(self.audit_yes_card, "0")
        self.update_enterprise_card(self.high_card, "0")
        self.update_enterprise_card(self.medium_card, "0")
        self.update_enterprise_card(self.low_card, "0")
        self.update_enterprise_card(self.voice_solutions_filter_card, "0")
        self.update_enterprise_card(self.ism_hours_filter_card, "0")
        
        # Clear ISM table
        self.ism_table.setRowCount(0)
    
    def update_enterprise_card(self, card, value):
        """Update the value in an enterprise card"""
        if hasattr(card, 'value_label'):
            card.value_label.setText(value)
    
    def update_summary_card(self, card, value):
        """Update the value in a summary card"""
        # Find the value label (last QLabel in the card)
        value_label = None
        for child in card.findChildren(QLabel):
            if child.parent() == card and child.styleSheet().find("font-size: 28px") != -1:
                value_label = child
                break
        if value_label:
            value_label.setText(value)
    
    def update_clickable_card(self, card, count):
        """Update the count in a clickable card"""
        # Find the count label (last QLabel in the card)
        count_label = None
        for child in card.findChildren(QLabel):
            if child.parent() == card and child.styleSheet().find("font-size: 24px") != -1:
                count_label = child
                break
        if count_label:
            count_label.setText(count)
    
    def update_metric_card(self, card, value):
        """Update the value in a metric card"""
        # Find the value label (last QLabel in the card)
        value_label = None
        for child in card.findChildren(QLabel):
            if child.parent() == card and child.styleSheet().find("font-size: 28px") != -1:
                value_label = child
                break
        if value_label:
            value_label.setText(value)
    
    def create_metric_card(self, title: str, value: str, color: str, bg_color: str) -> QFrame:
        """Create a styled metric card with enhanced visual design"""
        card = QFrame()
        card.setStyleSheet(f"""
            QFrame {{
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 {bg_color}, stop:1 {color}15);
                border: 2px solid {color};
                border-radius: 12px;
                padding: 16px;
                margin: 4px;
            }}
            QFrame:hover {{
                border: 2px solid {color};
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 {color}20, stop:1 {color}30);
            }}
        """)
        card.setFixedHeight(100)
        card.setMinimumWidth(140)
        
        layout = QVBoxLayout(card)
        layout.setContentsMargins(12, 8, 12, 8)
        layout.setSpacing(6)
        
        # Title with icon
        title_label = QLabel(title)
        title_label.setStyleSheet(f"""
            color: {color}; 
            font-size: 11px; 
            font-weight: 600; 
            background: transparent;
            border: none;
        """)
        title_label.setWordWrap(True)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)
        
        # Value with enhanced styling
        value_label = QLabel(value)
        value_label.setStyleSheet(f"""
            color: {color}; 
            font-size: 24px; 
            font-weight: 800; 
            background: transparent;
            border: none;
        """)
        value_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(value_label)
        
        # Add a subtle shadow effect
        shadow = QGraphicsDropShadowEffect()
        shadow.setBlurRadius(8)
        shadow.setXOffset(0)
        shadow.setYOffset(2)
        shadow.setColor(QColor(0, 0, 0, 25))
        card.setGraphicsEffect(shadow)
        
        # Store reference for updating
        card.value_label = value_label
        
        return card
    
    def create_enterprise_card(self, title: str, value: str, color: str) -> QPushButton:
        """Create a simple enterprise-style card"""
        card = QPushButton()
        card.setFixedHeight(100)
        card.setMinimumWidth(180)
        card.setMaximumWidth(200)
        card.setStyleSheet(f"""
            QPushButton {{
                background-color: white;
                border: 2px solid {color};
                border-radius: 6px;
                color: {color};
                font-size: 12px;
                font-weight: bold;
                text-align: center;
            }}
            QPushButton:hover {{
                background-color: {color};
                color: white;
            }}
        """)
        
        # Create layout for the card content
        layout = QVBoxLayout(card)
        layout.setContentsMargins(12, 12, 12, 12)
        layout.setSpacing(6)
        
        # Title
        title_label = QLabel(title)
        title_label.setStyleSheet(f"color: {color}; font-size: 11px; font-weight: bold;")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setWordWrap(True)
        layout.addWidget(title_label)
        
        # Value
        value_label = QLabel(value)
        value_label.setStyleSheet(f"color: {color}; font-size: 24px; font-weight: bold;")
        value_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(value_label)
        
        # Store reference for updating
        card.value_label = value_label
        
        return card
    
    def show_audit_projects(self):
        """Show projects with Audit = Yes"""
        # In real implementation, this would query the database for audit projects
        audit_projects = []  # Empty list - no sample data
        
        self.show_project_details("Audit Required Projects", audit_projects)
    
    def show_priority_projects(self, priority):
        """Show projects filtered by priority"""
        # In real implementation, this would query the database for priority projects
        priority_projects = []  # Empty list - no sample data
        
        self.show_project_details(f"{priority} Priority Projects", priority_projects)
    
    def show_project_details(self, title, projects):
        """Show project details in a dialog"""
        dialog = QDialog(self)
        dialog.setWindowTitle(title)
        dialog.setModal(True)
        dialog.resize(700, 500)
        dialog.setStyleSheet("""
            QDialog {
                background-color: #f8f9fa;
            }
        """)
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        if not projects:
            # Show empty state
            empty_label = QLabel("📋 No projects found")
            empty_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            empty_label.setStyleSheet("""
                font-size: 18px;
                color: #7f8c8d;
                padding: 40px;
            """)
            layout.addWidget(empty_label)
            
            info_label = QLabel("No data available. Please ensure your database is connected and contains project information.")
            info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            info_label.setWordWrap(True)
            info_label.setStyleSheet("""
                font-size: 12px;
                color: #95a5a6;
                padding: 20px;
            """)
            layout.addWidget(info_label)
        else:
            # Create table with styling
            table = QTableWidget()
            table.setColumnCount(3)
            table.setHorizontalHeaderLabels(["Project Name", "Project ID", "Primary ISM Name"])
            table.horizontalHeader().setStretchLastSection(True)
            table.setAlternatingRowColors(True)
            table.setStyleSheet("""
                QTableWidget {
                    gridline-color: #bdc3c7;
                    background-color: white;
                    alternate-background-color: #f8f9fa;
                }
                QHeaderView::section {
                    background-color: #34495e;
                    color: white;
                    padding: 10px;
                    border: none;
                    font-weight: bold;
                }
                QTableWidget::item {
                    padding: 10px;
                    border-bottom: 1px solid #ecf0f1;
                }
            """)
            
            table.setRowCount(len(projects))
            for row, project in enumerate(projects):
                for col, data in enumerate(project):
                    table.setItem(row, col, QTableWidgetItem(data))
            
            layout.addWidget(table)
        
        # Close button
        close_btn = QPushButton("✖️ Close")
        close_btn.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                padding: 10px 20px;
                border: none;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        close_btn.clicked.connect(dialog.accept)
        
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        button_layout.addWidget(close_btn)
        layout.addLayout(button_layout)
        
        dialog.exec()


# Global instances - created after all classes are defined (respect feature flags)
search_engine = SearchEngine() if FEATURE_FLAGS.get('search_engine', True) else None
collaboration_manager = CollaborationManager() if FEATURE_FLAGS.get('collaboration', False) else None
report_builder = ReportBuilder() if FEATURE_FLAGS.get('report_builder', False) else None
scheduled_report_manager = ScheduledReportManager() if FEATURE_FLAGS.get('scheduled_reports', False) else None
data_manager = DataManager() if FEATURE_FLAGS.get('data_manager', True) else None
help_system = HelpSystem() if FEATURE_FLAGS.get('help_system', False) else None
documentation_generator = DocumentationGenerator() if FEATURE_FLAGS.get('documentation_generator', False) else None

class ContentManager:
    """SQLite-backed content manager for the Wiki/Knowledge Base with versioning and attachments."""
    def __init__(self, db_path: str):
        self.db_path = db_path
        # Schema is now created by the main application in _create_new_backend_sqlite
        # No need to call _ensure_schema() here

    def _conn(self):
        """Create SQLite connection with WAL mode enabled (same as main app)"""
        conn = sqlite3.connect(self.db_path, timeout=30, isolation_level=None)
        # Enable WAL for better concurrency
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA synchronous=NORMAL")
        conn.execute("PRAGMA cache_size=10000")
        conn.execute("PRAGMA temp_store=MEMORY")
        return conn

    def _ensure_tables_exist(self) -> None:
        """Ensure Wiki tables exist in the database."""
        try:
            with self._conn() as conn:
                cur = conn.cursor()
                # Check if kb_articles table exists
                cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='kb_articles'")
                if not cur.fetchone():
                    # Create Wiki tables if they don't exist
                    cur.execute("""
                        CREATE TABLE IF NOT EXISTS kb_articles (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            title TEXT NOT NULL,
                            slug TEXT UNIQUE,
                            content_md TEXT,
                            content_html TEXT,
                            tags TEXT,
                            author TEXT,
                            status TEXT DEFAULT 'Draft',
                            created_at TEXT,
                            updated_at TEXT,
                            review_at TEXT
                        )
                    """)
                    cur.execute("""
                        CREATE TABLE IF NOT EXISTS kb_versions (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            article_id INTEGER,
                            version_no INTEGER,
                            title TEXT,
                            content_md TEXT,
                            content_html TEXT,
                            tags TEXT,
                            author TEXT,
                            created_at TEXT,
                            FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                        )
                    """)
                    cur.execute("""
                        CREATE TABLE IF NOT EXISTS kb_attachments (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            article_id INTEGER,
                            filename TEXT,
                            path TEXT,
                            uploaded_at TEXT,
                            FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                        )
                    """)
                    cur.execute("""
                        CREATE TABLE IF NOT EXISTS kb_links (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            from_article_id INTEGER,
                            to_article_id INTEGER,
                            FOREIGN KEY(from_article_id) REFERENCES kb_articles(id),
                            FOREIGN KEY(to_article_id) REFERENCES kb_articles(id)
                        )
                    """)
                    cur.execute("""
                        CREATE TABLE IF NOT EXISTS kb_comments (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            article_id INTEGER,
                            author TEXT,
                            content TEXT,
                            created_at TEXT,
                            FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                        )
                    """)
                    cur.execute("""
                        CREATE TABLE IF NOT EXISTS kb_approvals (
                            id INTEGER PRIMARY KEY AUTOINCREMENT,
                            article_id INTEGER,
                            state TEXT,
                            reviewer TEXT,
                            decided_at TEXT,
                            comment TEXT,
                            FOREIGN KEY(article_id) REFERENCES kb_articles(id)
                        )
                    """)
                    # FTS5 for full-text search if available
                    try:
                        cur.execute("""
                            CREATE VIRTUAL TABLE IF NOT EXISTS kb_articles_fts USING fts5(
                                title, content, tags, content='kb_articles', content_rowid='id'
                            )
                        """)
                    except Exception:
                        pass
                    conn.commit()
        except Exception as e:
            logger.debug(f"Failed to ensure Wiki tables exist: {e}")

    def _now(self) -> str:
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def create_article(self, title: str, content_md: str, tags: list[str], author: str, status: str = 'Draft', review_at: str | None = None) -> int:
        self._ensure_tables_exist()  # Ensure tables exist before creating article
        content_html = self._render_html(content_md)
        slug = self._generate_unique_slug(title)
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute(
                "INSERT INTO kb_articles(title, slug, content_md, content_html, tags, author, status, created_at, updated_at, review_at) VALUES (?,?,?,?,?,?,?,?,?,?)",
                (title, slug, content_md, content_html, ",".join(tags), author, status, self._now(), self._now(), review_at or "")
            )
            article_id = cur.lastrowid
            # Initial version
            cur.execute(
                "INSERT INTO kb_versions(article_id, version_no, title, content_md, content_html, tags, author, created_at) VALUES (?,?,?,?,?,?,?,?)",
                (article_id, 1, title, content_md, content_html, ",".join(tags), author, self._now())
            )
            # FTS index
            try:
                cur.execute("INSERT INTO kb_articles_fts(rowid, title, content, tags) VALUES (?,?,?,?)", (article_id, title, content_md, ",".join(tags)))
            except Exception:
                pass
            conn.commit()
            return article_id

    def update_article(self, article_id: int, title: str, content_md: str, tags: list[str], author: str, status: str | None = None) -> None:
        self._ensure_tables_exist()  # Ensure tables exist before updating article
        content_html = self._render_html(content_md)
        slug = self._generate_unique_slug_for_update(article_id, title)
        with self._conn() as conn:
            cur = conn.cursor()
            if status is None:
                cur.execute(
                    "UPDATE kb_articles SET title=?, slug=?, content_md=?, content_html=?, tags=?, author=?, updated_at=? WHERE id=?",
                    (title, slug, content_md, content_html, ",".join(tags), author, self._now(), article_id)
                )
            else:
                cur.execute(
                    "UPDATE kb_articles SET title=?, slug=?, content_md=?, content_html=?, tags=?, author=?, status=?, updated_at=? WHERE id=?",
                    (title, slug, content_md, content_html, ",".join(tags), author, status, self._now(), article_id)
                )
            # Versioning (increment)
            cur.execute("SELECT COALESCE(MAX(version_no),0) FROM kb_versions WHERE article_id=?", (article_id,))
            v = (cur.fetchone() or [0])[0] + 1
            cur.execute(
                "INSERT INTO kb_versions(article_id, version_no, title, content_md, content_html, tags, author, created_at) VALUES (?,?,?,?,?,?,?,?)",
                (article_id, v, title, content_md, content_html, ",".join(tags), author, self._now())
            )
            # FTS update
            try:
                cur.execute("DELETE FROM kb_articles_fts WHERE rowid=?", (article_id,))
                cur.execute("INSERT INTO kb_articles_fts(rowid, title, content, tags) VALUES (?,?,?,?)", (article_id, title, content_md, ",".join(tags)))
            except Exception:
                pass
            conn.commit()

    def get_article(self, article_id: int) -> dict | None:
        self._ensure_tables_exist()  # Ensure tables exist before getting article
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT id, title, slug, content_md, content_html, tags, author, status, created_at, updated_at, review_at FROM kb_articles WHERE id=?", (article_id,))
            row = cur.fetchone()
            if not row:
                return None
            keys = ["id","title","slug","content_md","content_html","tags","author","status","created_at","updated_at","review_at"]
            rec = {k: row[i] for i, k in enumerate(keys)}
            rec["tags"] = (rec["tags"] or "").split(",") if rec.get("tags") else []
            return rec

    def list_articles(self, query: str = "", tag_filter: list[str] | None = None, author: str | None = None) -> list[dict]:
        self._ensure_tables_exist()  # Ensure tables exist before listing articles
        with self._conn() as conn:
            cur = conn.cursor()
            rows = []
            if query:
                try:
                    cur.execute("SELECT rowid FROM kb_articles_fts WHERE kb_articles_fts MATCH ? LIMIT 200", (query,))
                    ids = [r[0] for r in cur.fetchall()]
                    if not ids:
                        return []
                    placeholders = ",".join(["?"] * len(ids))
                    cur.execute(f"SELECT id,title,tags,author,updated_at,status FROM kb_articles WHERE id IN ({placeholders}) ORDER BY updated_at DESC", ids)
                    rows = cur.fetchall()
                except Exception:
                    cur.execute("SELECT id,title,tags,author,updated_at,status FROM kb_articles ORDER BY updated_at DESC LIMIT 200")
                    rows = cur.fetchall()
            else:
                cur.execute("SELECT id,title,tags,author,updated_at,status FROM kb_articles ORDER BY updated_at DESC LIMIT 200")
                rows = cur.fetchall()
            out = []
            for r in rows:
                _id, _title, _tags, _author, _upd, _status = r
                if author and str(_author or "") != author:
                    continue
                if tag_filter:
                    have = set(((_tags or "").split(",")))
                    if not set(tag_filter).issubset(have):
                        continue
                out.append({"id": _id, "title": _title, "tags": ((_tags or "").split(",") if _tags else []), "author": _author, "updated_at": _upd, "status": _status})
            return out

    def save_comment(self, article_id: int, author: str, content: str) -> None:
        with self._conn() as conn:
            conn.execute("INSERT INTO kb_comments(article_id, author, content, created_at) VALUES (?,?,?,?)", (article_id, author, content, self._now()))

    def list_comments(self, article_id: int) -> list[tuple]:
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT author, content, created_at FROM kb_comments WHERE article_id=? ORDER BY created_at DESC", (article_id,))
            return cur.fetchall()

    def set_approval_state(self, article_id: int, state: str, reviewer: str, comment: str = "") -> None:
        with self._conn() as conn:
            conn.execute("INSERT INTO kb_approvals(article_id, state, reviewer, decided_at, comment) VALUES (?,?,?,?,?)", (article_id, state, reviewer, self._now(), comment))
            conn.execute("UPDATE kb_articles SET status=?, updated_at=? WHERE id=?", (state, self._now(), article_id))

    def list_versions(self, article_id: int) -> list[tuple]:
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT version_no, title, author, created_at FROM kb_versions WHERE article_id=? ORDER BY version_no DESC", (article_id,))
            return cur.fetchall()

    def get_version(self, article_id: int, version_no: int) -> dict | None:
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT title, content_md, content_html, tags, author, created_at FROM kb_versions WHERE article_id=? AND version_no=?", (article_id, version_no))
            r = cur.fetchone()
            if not r:
                return None
            return {"title": r[0], "content_md": r[1], "content_html": r[2], "tags": (r[3] or "").split(","), "author": r[4], "created_at": r[5]}

    def rollback(self, article_id: int, version_no: int, author: str) -> None:
        ver = self.get_version(article_id, version_no)
        if not ver:
            return
        self.update_article(article_id, ver["title"], ver["content_md"], ver["tags"], author)

    def list_backlinks(self, article_id: int) -> list[tuple[int, str]]:
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT from_article_id FROM kb_links WHERE to_article_id=?", (article_id,))
            ids = [r[0] for r in cur.fetchall()]
            if not ids:
                return []
            placeholders = ",".join(["?"] * len(ids))
            cur.execute(f"SELECT id, title FROM kb_articles WHERE id IN ({placeholders})", ids)
            return cur.fetchall()

    def related_by_tags(self, article_id: int, limit: int = 10) -> list[tuple[int, str]]:
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT tags FROM kb_articles WHERE id=?", (article_id,))
            r = cur.fetchone()
            if not r:
                return []
            tags = [t.strip() for t in (r[0] or '').split(',') if t.strip()]
            if not tags:
                return []
            like = [f"%{t}%" for t in tags]
            cond = " OR ".join(["tags LIKE ?"] * len(like))
            cur.execute(f"SELECT id,title FROM kb_articles WHERE id<>? AND ({cond}) ORDER BY updated_at DESC LIMIT {limit}", [article_id, *like])
            return cur.fetchall()

    def aging_articles(self) -> list[tuple[int, str, str]]:
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT id,title,review_at FROM kb_articles WHERE review_at IS NOT NULL AND LENGTH(review_at)>0 ORDER BY review_at ASC")
            return cur.fetchall()

    def update_links_from_markdown(self, article_id: int, md: str) -> None:
        # Auto-linking by naive title mention [[Title]] pattern
        import re
        titles = re.findall(r"\[\[([^\]]+)\]\]", md or "")
        if not titles:
            return
        with self._conn() as conn:
            cur = conn.cursor()
            # Resolve titles to ids
            ids: list[int] = []
            for t in titles:
                cur.execute("SELECT id FROM kb_articles WHERE title=?", (t.strip(),))
                r = cur.fetchone()
                if r:
                    ids.append(r[0])
            # refresh links
            cur.execute("DELETE FROM kb_links WHERE from_article_id=?", (article_id,))
            for to_id in set(ids):
                if to_id != article_id:
                    cur.execute("INSERT INTO kb_links(from_article_id, to_article_id) VALUES (?,?)", (article_id, to_id))

    def export_all_zip(self, zip_path: str) -> None:
        import zipfile, io
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT id,title,content_md,content_html FROM kb_articles")
            rows = cur.fetchall()
        with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as z:
            for _id, title, md, html in rows:
                base = f"{_id}_{self._slugify(title)}"
                z.writestr(f"{base}.md", md or "")
                z.writestr(f"{base}.html", html or "")

    def export_docx(self, article_id: int, path: str) -> None:
        try:
            from docx import Document  # type: ignore[import-not-found]
        except Exception as e:
            raise RuntimeError("python-docx not installed. pip install python-docx") from e
        art = self.get_article(article_id)
        if not art:
            raise ValueError("Article not found")
        doc = Document()
        doc.add_heading(art['title'] or '', level=1)
        for line in (art['content_md'] or '').splitlines():
            doc.add_paragraph(line)
        doc.save(path)

    def all_tags(self) -> list[str]:
        with self._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT tags FROM kb_articles WHERE tags IS NOT NULL AND LENGTH(tags)>0")
            tag_set: set[str] = set()
            for (tstr,) in cur.fetchall():
                for t in (tstr or '').split(','):
                    tt = t.strip()
                    if tt:
                        tag_set.add(tt)
            return sorted(tag_set)

    def delete_article(self, article_id: int) -> None:
        with self._conn() as conn:
            cur = conn.cursor()
            # Remove dependencies first
            cur.execute("DELETE FROM kb_versions WHERE article_id=?", (article_id,))
            cur.execute("DELETE FROM kb_attachments WHERE article_id=?", (article_id,))
            cur.execute("DELETE FROM kb_comments WHERE article_id=?", (article_id,))
            cur.execute("DELETE FROM kb_approvals WHERE article_id=?", (article_id,))
            cur.execute("DELETE FROM kb_links WHERE from_article_id=? OR to_article_id=?", (article_id, article_id))
            try:
                cur.execute("DELETE FROM kb_articles_fts WHERE rowid=?", (article_id,))
            except Exception:
                pass
            cur.execute("DELETE FROM kb_articles WHERE id=?", (article_id,))
            conn.commit()

    def export_article(self, article_id: int, fmt: str, path: str) -> None:
        art = self.get_article(article_id)
        if not art:
            raise ValueError("Article not found")
        if fmt == 'md':
            with open(path, 'w', encoding='utf-8') as f:
                f.write(f"# {art['title']}\n\n")
                f.write(art['content_md'] or "")
        elif fmt == 'html':
            with open(path, 'w', encoding='utf-8') as f:
                f.write(art['content_html'] or "")
        elif fmt == 'pdf':
            try:
                from reportlab.lib.pagesizes import letter  # type: ignore[import-not-found]
                from reportlab.pdfgen import canvas  # type: ignore[import-not-found]
                c = canvas.Canvas(path, pagesize=letter)
                textobject = c.beginText(40, 750)
                for line in (art['content_md'] or "").splitlines():
                    textobject.textLine(line)
                c.drawText(textobject)
                c.showPage()
                c.save()
            except Exception as e:
                raise e
        else:
            raise ValueError("Unsupported export format")

    def import_markdown(self, md_path: str, title: str, tags: list[str], author: str) -> int:
        with open(md_path, 'r', encoding='utf-8') as f:
            content_md = f.read()
        return self.create_article(title, content_md, tags, author)

    def _render_html(self, md: str) -> str:
        """Render markdown to HTML with enhanced styling and features."""
        try:
            if MARKDOWN2_AVAILABLE:
                # Use markdown2 with extra features for better rendering
                extras = ['fenced-code-blocks', 'tables', 'header-ids', 'toc', 'footnotes', 'strike', 'task-lists']
                html = markdown2.markdown(md or "", extras=extras)
                return self._wrap_html_with_styles(html)
        except Exception as e:
            logger.debug(f"Markdown2 rendering failed: {e}")
        
        try:
            if MISTUNE_AVAILABLE:
                # Use mistune with enhanced features
                renderer = mistune.HTMLRenderer()
                markdown = mistune.Markdown(renderer)
                html = markdown(md or "")
                return self._wrap_html_with_styles(html)
        except Exception as e:
            logger.debug(f"Mistune rendering failed: {e}")
        
        # Fallback: basic HTML with line breaks
        html = (md or "").replace('\n', '<br>')
        return self._wrap_html_with_styles(html)
    
    def _wrap_html_with_styles(self, html: str) -> str:
        """Wrap HTML content with modern CSS styling for better preview."""
        css_styles = """
        <style>
            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Oxygen', 'Ubuntu', 'Cantarell', sans-serif;
                line-height: 1.6;
                color: #333;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background: #fff;
            }
            h1, h2, h3, h4, h5, h6 {
                color: #2c3e50;
                margin-top: 1.5em;
                margin-bottom: 0.5em;
                font-weight: 600;
            }
            h1 { font-size: 2em; border-bottom: 2px solid #3498db; padding-bottom: 0.3em; }
            h2 { font-size: 1.5em; border-bottom: 1px solid #bdc3c7; padding-bottom: 0.2em; }
            h3 { font-size: 1.25em; }
            h4 { font-size: 1.1em; }
            h5 { font-size: 1em; }
            h6 { font-size: 0.9em; color: #7f8c8d; }
            
            p { margin-bottom: 1em; }
            
            code {
                background: #f8f9fa;
                border: 1px solid #e9ecef;
                border-radius: 3px;
                padding: 2px 4px;
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                font-size: 0.9em;
                color: #e74c3c;
            }
            
            pre {
                background: #f8f9fa;
                border: 1px solid #e9ecef;
                border-radius: 5px;
                padding: 15px;
                overflow-x: auto;
                margin: 1em 0;
            }
            
            pre code {
                background: none;
                border: none;
                padding: 0;
                color: #333;
            }
            
            blockquote {
                border-left: 4px solid #3498db;
                margin: 1em 0;
                padding: 0.5em 1em;
                background: #f8f9fa;
                color: #555;
            }
            
            table {
                border-collapse: collapse;
                width: 100%;
                margin: 1em 0;
            }
            
            th, td {
                border: 1px solid #ddd;
                padding: 8px 12px;
                text-align: left;
            }
            
            th {
                background: #f8f9fa;
                font-weight: 600;
                color: #2c3e50;
            }
            
            tr:nth-child(even) {
                background: #f8f9fa;
            }
            
            ul, ol {
                margin: 1em 0;
                padding-left: 2em;
            }
            
            li {
                margin-bottom: 0.5em;
            }
            
            a {
                color: #3498db;
                text-decoration: none;
            }
            
            a:hover {
                text-decoration: underline;
            }
            
            img {
                max-width: 100%;
                height: auto;
                border-radius: 5px;
                box-shadow: 0 2px 8px rgba(0,0,0,0.1);
                margin: 1em 0;
            }
            
            hr {
                border: none;
                height: 2px;
                background: linear-gradient(to right, transparent, #bdc3c7, transparent);
                margin: 2em 0;
            }
            
            .task-list-item {
                list-style: none;
                margin-left: -1.5em;
            }
            
            .task-list-item input[type="checkbox"] {
                margin-right: 0.5em;
            }
            
            .toc {
                background: #f8f9fa;
                border: 1px solid #e9ecef;
                border-radius: 5px;
                padding: 15px;
                margin: 1em 0;
            }
            
            .toc ul {
                margin: 0;
                padding-left: 1.5em;
            }
            
            .footnote {
                font-size: 0.9em;
                color: #7f8c8d;
            }
        </style>
        """
        return f"<html><head>{css_styles}</head><body>{html}</body></html>"

    def _slugify(self, text: str) -> str:
        import re
        s = re.sub(r"[^a-zA-Z0-9\-\s]", "", text or "").strip().lower()
        s = re.sub(r"\s+", "-", s)
        return s[:80]

    def _generate_unique_slug(self, title: str) -> str:
        """Generate a unique slug for the given title, appending numbers if necessary."""
        base_slug = self._slugify(title)
        slug = base_slug
        
        with self._conn() as conn:
            cur = conn.cursor()
            counter = 1
            
            while True:
                # Check if slug already exists
                cur.execute("SELECT COUNT(*) FROM kb_articles WHERE slug = ?", (slug,))
                if cur.fetchone()[0] == 0:
                    break
                
                # Generate new slug with counter
                if counter == 1:
                    # For first duplicate, append -2
                    slug = f"{base_slug}-2"
                else:
                    # For subsequent duplicates, increment counter
                    slug = f"{base_slug}-{counter + 1}"
                counter += 1
                
                # Safety check to prevent infinite loop
                if counter > 1000:
                    # Fallback to timestamp-based slug
                    import time
                    timestamp = int(time.time())
                    slug = f"{base_slug}-{timestamp}"
                    break
        
        return slug

    def _generate_unique_slug_for_update(self, article_id: int, title: str) -> str:
        """Generate a unique slug for updating an article, only changing if necessary."""
        base_slug = self._slugify(title)
        
        with self._conn() as conn:
            cur = conn.cursor()
            
            # Check if the current article already has this slug
            cur.execute("SELECT slug FROM kb_articles WHERE id = ?", (article_id,))
            current_slug = cur.fetchone()
            if current_slug and current_slug[0] == base_slug:
                # Title hasn't changed or generates the same slug, keep current slug
                return base_slug
            
            # Check if the new slug conflicts with any other article
            cur.execute("SELECT COUNT(*) FROM kb_articles WHERE slug = ? AND id != ?", (base_slug, article_id))
            if cur.fetchone()[0] == 0:
                # No conflict, use the base slug
                return base_slug
            
            # Conflict exists, generate unique slug
            slug = base_slug
            counter = 1
            
            while True:
                # Check if slug already exists (excluding current article)
                cur.execute("SELECT COUNT(*) FROM kb_articles WHERE slug = ? AND id != ?", (slug, article_id))
                if cur.fetchone()[0] == 0:
                    break
                
                # Generate new slug with counter
                if counter == 1:
                    slug = f"{base_slug}-2"
                else:
                    slug = f"{base_slug}-{counter + 1}"
                counter += 1
                
                # Safety check to prevent infinite loop
                if counter > 1000:
                    import time
                    timestamp = int(time.time())
                    slug = f"{base_slug}-{timestamp}"
                    break
        
        return slug

class SearchManager:
    """Wrapper around ContentManager for FTS operations and filters."""
    def __init__(self, content: ContentManager):
        self.content = content

    def search(self, query: str, tags: list[str] | None = None, author: str | None = None) -> list[dict]:
        return self.content.list_articles(query=query, tag_filter=tags, author=author)

class UserManager:
    def __init__(self, current_user: str):
        self.current_user = current_user

class ThemeManager:
    def apply_enterprise_theme(self, widget: QWidget) -> None:
        try:
            widget.setStyleSheet("""
                QWidget { font-size: 12px; }
                QFrame#kbCard { background: #ffffff; border: 1px solid #e2e8f0; border-radius: 10px; }
                QLineEdit, QTextEdit { border: 1px solid #d1d5db; border-radius: 6px; padding: 6px; }
                QPushButton { background-color: #2a7de1; color: white; border: none; border-radius: 6px; padding: 8px 12px; }
                QPushButton:hover { background-color: #1e60b8; }
                QTreeWidget { background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 8px; }
            """)
        except Exception:
            pass

class GraphManager:
    """Maintains backlinks graph based on kb_links table."""
    def __init__(self, content: ContentManager):
        self.content = content

    def get_backlinks(self, article_id: int) -> list[int]:
        with self.content._conn() as conn:
            cur = conn.cursor()
            cur.execute("SELECT from_article_id FROM kb_links WHERE to_article_id=?", (article_id,))
            return [r[0] for r in cur.fetchall()]

# ModernKnowledgeBase class removed
# ModernKnowledgeBase class content removed

# ====== Kanban Board Pane ======
class KanbanListWidget(QListWidget):
    """Custom QListWidget for Kanban columns with drag and drop support"""
    
    def __init__(self, parent, column_type: str):
        super().__init__(parent)
        self.column_type = column_type
        self.setDragDropMode(QListWidget.DragDropMode.DragDrop)
        self.setDefaultDropAction(Qt.DropAction.MoveAction)
    
    def dragEnterEvent(self, event):
        if event.mimeData().hasFormat("application/x-qlistwidgetitem"):
            event.accept()
        else:
            event.ignore()
    
    def dropEvent(self, event):
        if event.mimeData().hasFormat("application/x-qlistwidgetitem"):
            # Get the parent KanbanPane to handle the drop
            kanban_pane = self.parent()
            while kanban_pane and not isinstance(kanban_pane, KanbanPane):
                kanban_pane = kanban_pane.parent()
            
            if kanban_pane:
                kanban_pane.handle_task_drop(event, self.column_type)
            event.accept()
        else:
            event.ignore()

class KanbanPane(QWidget):
    """User-specific Kanban board view with drag-and-drop functionality"""
    
    def __init__(self, main_window, current_user: str):
        super().__init__()
        self.main_window = main_window
        self.current_user = current_user
        self.setup_ui()
        self.load_tasks()
    
    def setup_ui(self):
        """Setup the Kanban board UI with modern design"""
        # Main layout with modern styling
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        
        # Set background color for the main widget
        self.setStyleSheet("""
            QWidget {
                background-color: #f8fafc;
                font-family: 'Segoe UI', Arial, sans-serif;
            }
        """)
        
        # Modern header with enhanced styling
        header_widget = QWidget()
        header_widget.setStyleSheet("""
            QWidget {
                background-color: white;
                border-radius: 12px;
                padding: 16px;
                box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
            }
        """)
        header_layout = QHBoxLayout(header_widget)
        header_layout.setContentsMargins(20, 16, 20, 16)
        header_layout.setSpacing(20)
        
        # Title section
        title_section = QHBoxLayout()
        title_label = QLabel("📋 Kanban Board")
        title_label.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: 700;
                color: #1e293b;
                margin-right: 20px;
            }
        """)
        title_section.addWidget(title_label)
        title_section.addStretch()
        header_layout.addLayout(title_section)
        
        # Filters section with modern styling
        filters_widget = QWidget()
        filters_layout = QHBoxLayout(filters_widget)
        filters_layout.setSpacing(16)
        
        # Project filter with modern styling
        project_label = QLabel("Project")
        project_label.setStyleSheet("""
            QLabel {
                font-weight: 600;
                color: #64748b;
                font-size: 14px;
            }
        """)
        filters_layout.addWidget(project_label)
        
        self.project_filter = QComboBox()
        self.project_filter.addItem("All Projects")
        self.project_filter.setStyleSheet("""
            QComboBox {
                background-color: #f1f5f9;
                border: 2px solid #e2e8f0;
                border-radius: 8px;
                padding: 8px 12px;
                font-size: 14px;
                min-width: 120px;
            }
            QComboBox:hover {
                border-color: #3b82f6;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #64748b;
                margin-right: 8px;
            }
        """)
        self.project_filter.currentTextChanged.connect(self.apply_filters)
        filters_layout.addWidget(self.project_filter)
        
        # Priority filter with modern styling
        priority_label = QLabel("Priority")
        priority_label.setStyleSheet("""
            QLabel {
                font-weight: 600;
                color: #64748b;
                font-size: 14px;
            }
        """)
        filters_layout.addWidget(priority_label)
        
        self.priority_filter = QComboBox()
        self.priority_filter.addItem("All Priorities")
        self.priority_filter.addItems(PRIORITY_OPTIONS)
        self.priority_filter.setStyleSheet("""
            QComboBox {
                background-color: #f1f5f9;
                border: 2px solid #e2e8f0;
                border-radius: 8px;
                padding: 8px 12px;
                font-size: 14px;
                min-width: 120px;
            }
            QComboBox:hover {
                border-color: #3b82f6;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #64748b;
                margin-right: 8px;
            }
        """)
        self.priority_filter.currentTextChanged.connect(self.apply_filters)
        filters_layout.addWidget(self.priority_filter)
        
        header_layout.addWidget(filters_widget)
        
        # New Task button with modern styling
        self.new_task_btn = QPushButton("+ New Task")
        self.new_task_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #3b82f6, stop:1 #2563eb);
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 10px;
                font-weight: 600;
                font-size: 14px;
                min-width: 120px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #2563eb, stop:1 #1d4ed8);
                transform: translateY(-1px);
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #1d4ed8, stop:1 #1e40af);
            }
        """)
        self.new_task_btn.clicked.connect(self.create_new_task)
        header_layout.addWidget(self.new_task_btn)
        
        layout.addWidget(header_widget)
        
        # Kanban columns with modern layout
        columns_widget = QWidget()
        columns_widget.setStyleSheet("""
            QWidget {
                background-color: transparent;
            }
        """)
        columns_layout = QHBoxLayout(columns_widget)
        columns_layout.setContentsMargins(0, 0, 0, 0)
        columns_layout.setSpacing(20)
        
        # To Do column with modern styling
        self.todo_column = self.create_column("To Do", "#ef4444", "🔴")
        columns_layout.addWidget(self.todo_column)
        
        # In Progress column with modern styling
        self.in_progress_column = self.create_column("In Progress", "#f59e0b", "🟡")
        columns_layout.addWidget(self.in_progress_column)
        
        # Done column with modern styling
        self.done_column = self.create_column("Done", "#10b981", "🟢")
        columns_layout.addWidget(self.done_column)
        
        layout.addWidget(columns_widget)
    
    def create_column(self, title: str, color: str, emoji: str = "") -> QWidget:
        """Create a modern Kanban column with drag-and-drop support"""
        column = QWidget()
        column.setFixedWidth(320)
        column.setStyleSheet(f"""
            QWidget {{
                background-color: #ffffff;
                border: 1px solid #e2e8f0;
                border-radius: 16px;
                margin: 8px;
            }}
        """)
        
        layout = QVBoxLayout(column)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)
        
        # Modern column header with emoji and count
        header_widget = QWidget()
        header_widget.setStyleSheet(f"""
            QWidget {{
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
                    stop:0 {color}, stop:1 {self._lighten_color(color)});
                border-radius: 12px;
                padding: 12px;
            }}
        """)
        header_layout = QHBoxLayout(header_widget)
        header_layout.setContentsMargins(16, 12, 16, 12)
        
        # Title with emoji
        title_label = QLabel(f"{emoji} {title}")
        title_label.setStyleSheet("""
            QLabel {
                color: white;
                font-size: 16px;
                font-weight: 700;
                background: transparent;
            }
        """)
        header_layout.addWidget(title_label)
        
        # Task count badge
        count_label = QLabel("0")
        count_label.setObjectName("task_count")
        count_label.setStyleSheet("""
            QLabel {
                background-color: rgba(255, 255, 255, 0.2);
                color: white;
                border-radius: 12px;
                padding: 4px 8px;
                font-size: 12px;
                font-weight: 600;
                min-width: 20px;
                text-align: center;
            }
        """)
        header_layout.addWidget(count_label)
        
        layout.addWidget(header_widget)
        
        # Modern task list with enhanced styling
        task_list = KanbanListWidget(column, title)
        task_list.setStyleSheet("""
            QListWidget {
                background-color: transparent;
                border: none;
                padding: 8px;
                outline: none;
            }
            QListWidget::item {
                background-color: #ffffff;
                border: 1px solid #e2e8f0;
                border-radius: 12px;
                padding: 16px;
                margin: 6px 0;
                min-height: 60px;
            }
            QListWidget::item:hover {
                border-color: #3b82f6;
                box-shadow: 0 4px 12px rgba(59, 130, 246, 0.15);
                transform: translateY(-2px);
            }
            QListWidget::item:selected {
                border-color: #3b82f6;
                background-color: #eff6ff;
            }
        """)
        task_list.setDragDropMode(QListWidget.DragDropMode.DragDrop)
        task_list.setDefaultDropAction(Qt.DropAction.MoveAction)
        layout.addWidget(task_list)
        
        # Store references for later use
        column.task_list = task_list
        column.count_label = count_label
        
        return column
    
    def _lighten_color(self, color: str) -> str:
        """Lighten a hex color for gradient effect"""
        color_map = {
            "#ef4444": "#f87171",  # Red
            "#f59e0b": "#fbbf24",  # Amber
            "#10b981": "#34d399",  # Green
        }
        return color_map.get(color, color)
    
    def load_tasks(self):
        """Load tasks from the main window data"""
        # Clear existing tasks
        self.todo_column.task_list.clear()
        self.in_progress_column.task_list.clear()
        self.done_column.task_list.clear()
        
        # Load projects for filter
        self.project_filter.blockSignals(True)
        self.project_filter.clear()
        self.project_filter.addItem("All Projects")
        for project_name, _ in self.main_window.projects:
            self.project_filter.addItem(project_name)
        self.project_filter.blockSignals(False)
        
        # Helper function for flexible user matching
        def user_matches(user_in_data, current_user):
            if not user_in_data or not current_user:
                return False
            user_in_data = str(user_in_data).strip().lower()
            current_user = str(current_user).strip().lower()
            
            # Direct match
            if user_in_data == current_user:
                return True
            
            # Handle cases like "s.h.prabhakaran" vs "santhosh prabhakaran"
            # Extract last name and first initial
            def extract_name_parts(name):
                parts = name.replace('.', ' ').split()
                if len(parts) >= 2:
                    first_part = parts[0]
                    last_part = parts[-1]
                    return first_part, last_part
                return name, ""
            
            user_first, user_last = extract_name_parts(user_in_data)
            current_first, current_last = extract_name_parts(current_user)
            
            # Match if last names are the same and first parts match (initial or full)
            if user_last and current_last and user_last == current_last:
                if (user_first == current_first or 
                    user_first.startswith(current_first) or 
                    current_first.startswith(user_first) or
                    (len(user_first) == 1 and current_first.startswith(user_first)) or
                    (len(current_first) == 1 and user_first.startswith(current_first))):
                    return True
            
            # Original matching logic
            return (user_in_data.endswith(current_user) or 
                    current_user.endswith(user_in_data) or
                    user_in_data.replace('.', '') == current_user.replace('.', '') or
                    current_user.replace('.', '') == user_in_data.replace('.', ''))
        
        # Get tasks from Potential Issues and Activities panes
        tasks = []
        
        # From Potential Issues
        if "Potential Issues" in self.main_window.data:
            pi_data = self.main_window.data["Potential Issues"]
            pi_cols = PANE_COLUMNS["Potential Issues"]
            ism_idx = pi_cols.index("ISM Name") if "ISM Name" in pi_cols else None
            
            for row in pi_data:
                if ism_idx is not None and len(row) > ism_idx and user_matches(row[ism_idx], self.current_user):
                    # Get title from Description, Updates, or use Project Name as fallback
                    title = ""
                    if "Description" in pi_cols and pi_cols.index("Description") < len(row) and row[pi_cols.index("Description")]:
                        title = row[pi_cols.index("Description")]
                    elif "Updates" in pi_cols and pi_cols.index("Updates") < len(row) and row[pi_cols.index("Updates")]:
                        title = row[pi_cols.index("Updates")]
                    elif "Project Name" in pi_cols and pi_cols.index("Project Name") < len(row) and row[pi_cols.index("Project Name")]:
                        title = f"Potential Issue - {row[pi_cols.index('Project Name')]}"
                    else:
                        title = "Potential Issue"
                    
                    task = {
                        'title': title,
                        'due_date': row[pi_cols.index("Due Date")] if "Due Date" in pi_cols else "",
                        'priority': row[pi_cols.index("RAG Status")] if "RAG Status" in pi_cols else "GREEN",
                        'project': row[pi_cols.index("Project Name")] if "Project Name" in pi_cols else "",
                        'status': row[pi_cols.index("Status")] if "Status" in pi_cols else "To Do",
                        'pane': "Potential Issues",
                        'row_data': row
                    }
                    tasks.append(task)
        
        # From Activities
        if "Activities" in self.main_window.data:
            act_data = self.main_window.data["Activities"]
            act_cols = PANE_COLUMNS["Activities"]
            ism_idx = act_cols.index("ISM Name") if "ISM Name" in act_cols else None
            
            for row in act_data:
                if ism_idx is not None and len(row) > ism_idx and user_matches(row[ism_idx], self.current_user):
                    # Get title from Activity/Issue, Brief Update, or use Project Name as fallback
                    title = ""
                    if "Activity/Issue" in act_cols and act_cols.index("Activity/Issue") < len(row) and row[act_cols.index("Activity/Issue")]:
                        title = row[act_cols.index("Activity/Issue")]
                    elif "Brief Update" in act_cols and act_cols.index("Brief Update") < len(row) and row[act_cols.index("Brief Update")]:
                        title = row[act_cols.index("Brief Update")]
                    elif "Project Name" in act_cols and act_cols.index("Project Name") < len(row) and row[act_cols.index("Project Name")]:
                        title = f"Activity - {row[act_cols.index('Project Name')]}"
                    else:
                        title = "Activity"
                    
                    task = {
                        'title': title,
                        'due_date': row[act_cols.index("Target Date")] if "Target Date" in act_cols else "",
                        'priority': row[act_cols.index("RAG")] if "RAG" in act_cols else "GREEN",
                        'project': row[act_cols.index("Project Name")] if "Project Name" in act_cols else "",
                        'status': row[act_cols.index("Status")] if "Status" in act_cols else "To Do",
                        'pane': "Activities",
                        'row_data': row
                    }
                    tasks.append(task)
        
        # Add tasks to appropriate columns
        for task in tasks:
            self.add_task_to_column(task)
    
    def add_task_to_column(self, task: dict):
        """Add a modern task card to the appropriate column"""
        # Determine which column based on status
        if task['status'] in ['Done', 'Completed', 'Closed']:
            column = self.done_column.task_list
            column_widget = self.done_column
        elif task['status'] in ['In Progress', 'Active', 'Working']:
            column = self.in_progress_column.task_list
            column_widget = self.in_progress_column
        else:
            column = self.todo_column.task_list
            column_widget = self.todo_column
        
        # Create modern task card
        card = QListWidgetItem()
        card.setData(Qt.ItemDataRole.UserRole, task)
        
        # Format task display text with modern styling
        title = task.get('title', 'Untitled Task')
        due_date = task.get('due_date', '')
        priority = task.get('priority', 'GREEN')
        project = task.get('project', '')
        
        # Create rich text for the card with modern design
        card_text = f"""
        <div style="font-family: 'Segoe UI', Arial, sans-serif;">
            <div style="font-size: 16px; font-weight: 700; color: #1e293b; margin-bottom: 8px; line-height: 1.3;">
                {title}
            </div>
        """
        
        if project:
            card_text += f"""
            <div style="display: flex; align-items: center; margin-bottom: 6px;">
                <span style="color: #64748b; font-size: 12px; margin-right: 6px;">📁</span>
                <span style="color: #64748b; font-size: 12px; font-weight: 500;">{project}</span>
            </div>
            """
        
        if due_date:
            card_text += f"""
            <div style="display: flex; align-items: center; margin-bottom: 6px;">
                <span style="color: #64748b; font-size: 12px; margin-right: 6px;">📅</span>
                <span style="color: #64748b; font-size: 12px; font-weight: 500;">{due_date}</span>
            </div>
            """
        
        if priority:
            priority_colors = {
                'RED': '#ef4444',
                'AMBER': '#f59e0b', 
                'GREEN': '#10b981',
                'YELLOW': '#f59e0b',
                'ORANGE': '#f59e0b'
            }
            priority_color = priority_colors.get(priority.upper(), '#10b981')
            card_text += f"""
            <div style="display: flex; align-items: center; justify-content: space-between; margin-top: 8px;">
                <div style="display: flex; align-items: center;">
                    <div style="width: 8px; height: 8px; background-color: {priority_color}; border-radius: 50%; margin-right: 6px;"></div>
                    <span style="color: {priority_color}; font-size: 11px; font-weight: 600; text-transform: uppercase;">{priority}</span>
                </div>
                <div style="color: #94a3b8; font-size: 10px; font-weight: 500;">
                    {task.get('pane', '').replace(' ', '')}
                </div>
            </div>
            """
        
        card_text += "</div>"
        
        card.setText(card_text)
        column.addItem(card)
        
        # Update task count in column header
        if hasattr(column_widget, 'count_label'):
            count = column.count()
            column_widget.count_label.setText(str(count))
    
    def handle_task_drop(self, event, target_column_type: str):
        """Handle drag and drop of task cards"""
        # Get the source item from the event
        source_item = None
        source_list = None
        
        # Try to get the source item from the event
        if hasattr(event, 'source') and event.source():
            source_list = event.source()
            source_item = source_list.currentItem()
        
        if not source_item:
            return
        
        # Get task data
        task = source_item.data(Qt.ItemDataRole.UserRole)
        if not task:
            return
        
        # Determine new status based on target column type
        if target_column_type == "To Do":
            new_status = "To Do"
        elif target_column_type == "In Progress":
            new_status = "In Progress"
        elif target_column_type == "Done":
            new_status = "Done"
        else:
            return
        
        # Update task status in main data
        self.update_task_status(task, new_status)
        
        # Move item to new column
        if source_list:
            source_list.takeItem(source_list.row(source_item))
        
        # Add to target column
        target_list = None
        if target_column_type == "To Do":
            target_list = self.todo_column.task_list
        elif target_column_type == "In Progress":
            target_list = self.in_progress_column.task_list
        elif target_column_type == "Done":
            target_list = self.done_column.task_list
        
        if target_list:
            source_item.setData(Qt.ItemDataRole.UserRole, {**task, 'status': new_status})
            target_list.addItem(source_item)
    
    def update_task_status(self, task: dict, new_status: str):
        """Update task status in the main window data"""
        try:
            pane_name = task['pane']
            if pane_name in self.main_window.data:
                rows = self.main_window.data[pane_name]
                columns = PANE_COLUMNS[pane_name]
                
                # Find the row and update status
                for i, row in enumerate(rows):
                    if row == task['row_data']:
                        status_idx = columns.index("Status") if "Status" in columns else None
                        if status_idx is not None and status_idx < len(row):
                            rows[i][status_idx] = new_status
                            break
        except Exception as e:
            ErrorHandler.handle_ui_error("update task status", e)
    
    def apply_filters(self):
        """Apply project and priority filters"""
        self.load_tasks()  # Reload with filters applied
    
    def create_new_task(self):
        """Open modern dialog to create a new task"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Create New Task")
        dialog.setModal(True)
        dialog.resize(480, 420)
        dialog.setStyleSheet("""
            QDialog {
                background-color: #ffffff;
                border-radius: 16px;
            }
        """)
        
        # Main layout with modern styling
        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(32, 32, 32, 32)
        layout.setSpacing(20)
        
        # Header
        header_label = QLabel("✨ Create New Task")
        header_label.setStyleSheet("""
            QLabel {
                font-size: 24px;
                font-weight: 700;
                color: #1e293b;
                margin-bottom: 8px;
            }
        """)
        layout.addWidget(header_label)
        
        # Form fields with modern styling
        fields_layout = QVBoxLayout()
        fields_layout.setSpacing(16)
        
        # Task title
        title_label = QLabel("Task Title")
        title_label.setStyleSheet("""
            QLabel {
                font-weight: 600;
                color: #374151;
                font-size: 14px;
                margin-bottom: 4px;
            }
        """)
        fields_layout.addWidget(title_label)
        
        title_edit = QLineEdit()
        title_edit.setStyleSheet("""
            QLineEdit {
                background-color: #f9fafb;
                border: 2px solid #e5e7eb;
                border-radius: 8px;
                padding: 12px 16px;
                font-size: 14px;
                color: #1f2937;
            }
            QLineEdit:focus {
                border-color: #3b82f6;
                background-color: #ffffff;
            }
        """)
        title_edit.setPlaceholderText("Enter task title...")
        fields_layout.addWidget(title_edit)
        
        # Project selection
        project_label = QLabel("Project")
        project_label.setStyleSheet("""
            QLabel {
                font-weight: 600;
                color: #374151;
                font-size: 14px;
                margin-bottom: 4px;
            }
        """)
        fields_layout.addWidget(project_label)
        
        project_combo = QComboBox()
        project_combo.addItem("")
        for project_name, _ in self.main_window.projects:
            project_combo.addItem(project_name)
        project_combo.setStyleSheet("""
            QComboBox {
                background-color: #f9fafb;
                border: 2px solid #e5e7eb;
                border-radius: 8px;
                padding: 12px 16px;
                font-size: 14px;
                color: #1f2937;
            }
            QComboBox:focus {
                border-color: #3b82f6;
                background-color: #ffffff;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #6b7280;
                margin-right: 8px;
            }
        """)
        fields_layout.addWidget(project_combo)
        
        # Priority and Due Date row
        priority_due_layout = QHBoxLayout()
        priority_due_layout.setSpacing(16)
        
        # Priority selection
        priority_widget = QWidget()
        priority_layout = QVBoxLayout(priority_widget)
        priority_layout.setContentsMargins(0, 0, 0, 0)
        
        priority_label = QLabel("Priority")
        priority_label.setStyleSheet("""
            QLabel {
                font-weight: 600;
                color: #374151;
                font-size: 14px;
                margin-bottom: 4px;
            }
        """)
        priority_layout.addWidget(priority_label)
        
        priority_combo = QComboBox()
        priority_combo.addItems(PRIORITY_OPTIONS)
        priority_combo.setStyleSheet("""
            QComboBox {
                background-color: #f9fafb;
                border: 2px solid #e5e7eb;
                border-radius: 8px;
                padding: 12px 16px;
                font-size: 14px;
                color: #1f2937;
            }
            QComboBox:focus {
                border-color: #3b82f6;
                background-color: #ffffff;
            }
            QComboBox::drop-down {
                border: none;
                width: 20px;
            }
            QComboBox::down-arrow {
                image: none;
                border-left: 5px solid transparent;
                border-right: 5px solid transparent;
                border-top: 5px solid #6b7280;
                margin-right: 8px;
            }
        """)
        priority_layout.addWidget(priority_combo)
        priority_due_layout.addWidget(priority_widget)
        
        # Due date
        due_widget = QWidget()
        due_layout = QVBoxLayout(due_widget)
        due_layout.setContentsMargins(0, 0, 0, 0)
        
        due_label = QLabel("Due Date")
        due_label.setStyleSheet("""
            QLabel {
                font-weight: 600;
                color: #374151;
                font-size: 14px;
                margin-bottom: 4px;
            }
        """)
        due_layout.addWidget(due_label)
        
        due_date_edit = QLineEdit()
        due_date_edit.setPlaceholderText("YYYY-MM-DD")
        due_date_edit.setStyleSheet("""
            QLineEdit {
                background-color: #f9fafb;
                border: 2px solid #e5e7eb;
                border-radius: 8px;
                padding: 12px 16px;
                font-size: 14px;
                color: #1f2937;
            }
            QLineEdit:focus {
                border-color: #3b82f6;
                background-color: #ffffff;
            }
        """)
        due_layout.addWidget(due_date_edit)
        priority_due_layout.addWidget(due_widget)
        
        fields_layout.addLayout(priority_due_layout)
        
        # Description
        desc_label = QLabel("Description")
        desc_label.setStyleSheet("""
            QLabel {
                font-weight: 600;
                color: #374151;
                font-size: 14px;
                margin-bottom: 4px;
            }
        """)
        fields_layout.addWidget(desc_label)
        
        desc_edit = QTextEdit()
        desc_edit.setMaximumHeight(100)
        desc_edit.setStyleSheet("""
            QTextEdit {
                background-color: #f9fafb;
                border: 2px solid #e5e7eb;
                border-radius: 8px;
                padding: 12px 16px;
                font-size: 14px;
                color: #1f2937;
                font-family: 'Segoe UI', Arial, sans-serif;
            }
            QTextEdit:focus {
                border-color: #3b82f6;
                background-color: #ffffff;
            }
        """)
        desc_edit.setPlaceholderText("Enter task description...")
        fields_layout.addWidget(desc_edit)
        
        layout.addLayout(fields_layout)
        
        # Modern buttons
        button_layout = QHBoxLayout()
        button_layout.setSpacing(12)
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #f3f4f6;
                color: #374151;
                border: 2px solid #e5e7eb;
                border-radius: 8px;
                padding: 12px 24px;
                font-weight: 600;
                font-size: 14px;
                min-width: 100px;
            }
            QPushButton:hover {
                background-color: #e5e7eb;
                border-color: #d1d5db;
            }
        """)
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)
        
        create_btn = QPushButton("Create Task")
        create_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #3b82f6, stop:1 #2563eb);
                color: white;
                border: none;
                border-radius: 8px;
                padding: 12px 24px;
                font-weight: 600;
                font-size: 14px;
                min-width: 120px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #2563eb, stop:1 #1d4ed8);
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, 
                    stop:0 #1d4ed8, stop:1 #1e40af);
            }
        """)
        create_btn.clicked.connect(lambda: self.save_new_task(
            title_edit.text(), project_combo.currentText(), 
            priority_combo.currentText(), due_date_edit.text(), 
            desc_edit.toPlainText(), dialog
        ))
        button_layout.addWidget(create_btn)
        
        layout.addLayout(button_layout)
        
        dialog.exec()
    
    def save_new_task(self, title: str, project: str, priority: str, due_date: str, description: str, dialog: QDialog):
        """Save the new task to the main window data"""
        if not title.strip():
            QMessageBox.warning(self, "Invalid Input", "Task title is required.")
            return
        
        try:
            # Add to Potential Issues pane
            new_row = [""] * len(PANE_COLUMNS["Potential Issues"])
            new_row[PANE_COLUMNS["Potential Issues"].index("Task Type")] = "Task"
            new_row[PANE_COLUMNS["Potential Issues"].index("Status")] = "To Do"
            new_row[PANE_COLUMNS["Potential Issues"].index("Created Date")] = datetime.now().strftime("%Y-%m-%d")
            new_row[PANE_COLUMNS["Potential Issues"].index("Due Date")] = due_date
            new_row[PANE_COLUMNS["Potential Issues"].index("Project Name")] = project
            new_row[PANE_COLUMNS["Potential Issues"].index("RAG Status")] = priority
            new_row[PANE_COLUMNS["Potential Issues"].index("ISM Name")] = self.current_user
            new_row[PANE_COLUMNS["Potential Issues"].index("Description")] = title
            new_row[PANE_COLUMNS["Potential Issues"].index("Updates")] = description
            
            self.main_window.data["Potential Issues"].append(new_row)
            self.load_tasks()
            dialog.accept()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("save new task", e)


# ====== Calendar View Pane ======
class CalendarPane(QWidget):
    """User-specific calendar view for deadlines and leaves"""
    
    def __init__(self, main_window, current_user: str):
        super().__init__()
        self.main_window = main_window
        self.current_user = current_user
        self.setup_ui()
        self.load_calendar_data()
    
    def setup_ui(self):
        """Setup the calendar UI"""
        layout = QVBoxLayout(self)
        
        # Header with view controls
        header_layout = QHBoxLayout()
        
        # View toggle buttons
        self.month_view_btn = QPushButton("Month View")
        self.week_view_btn = QPushButton("Week View")
        self.month_view_btn.setCheckable(True)
        self.week_view_btn.setCheckable(True)
        self.month_view_btn.setChecked(True)
        
        self.month_view_btn.clicked.connect(lambda: self.set_view_mode("month"))
        self.week_view_btn.clicked.connect(lambda: self.set_view_mode("week"))
        
        header_layout.addWidget(self.month_view_btn)
        header_layout.addWidget(self.week_view_btn)
        header_layout.addStretch()
        
        # Refresh button for real-time updates
        refresh_btn = QPushButton("🔄 Refresh")
        refresh_btn.clicked.connect(self.refresh_calendar_data)
        refresh_btn.setToolTip("Refresh calendar data from all panes")
        header_layout.addWidget(refresh_btn)
        
        # Add simplified legend
        legend_label = QLabel("Legend:")
        legend_label.setStyleSheet("font-weight: bold; margin-right: 10px;")
        header_layout.addWidget(legend_label)
        
        # Today indicator
        today_legend = QLabel("📅 Today")
        today_legend.setStyleSheet("background-color: #ffff00; color: #000000; font-weight: bold; padding: 2px 6px; border-radius: 3px;")
        header_layout.addWidget(today_legend)
        
        # Task indicator
        task_legend = QLabel("🟡 Tasks")
        task_legend.setStyleSheet("background-color: #fff8dc; color: #ff8c00; font-weight: bold; padding: 2px 6px; border-radius: 3px;")
        header_layout.addWidget(task_legend)
        
        # Overdue indicator
        overdue_legend = QLabel("⚠️ Overdue")
        overdue_legend.setStyleSheet("background-color: #ffc8c8; color: #dc2626; font-weight: bold; padding: 2px 6px; border-radius: 3px;")
        header_layout.addWidget(overdue_legend)
        
        # Leave indicator
        leave_legend = QLabel("🔵 Leaves")
        leave_legend.setStyleSheet("background-color: #dcf8ff; color: #007bff; font-weight: bold; padding: 2px 6px; border-radius: 3px;")
        header_layout.addWidget(leave_legend)
        
        # Sync button (stub)
        sync_btn = QPushButton("Sync with System Calendar")
        sync_btn.clicked.connect(self.sync_with_system_calendar)
        header_layout.addWidget(sync_btn)
        
        layout.addLayout(header_layout)
        
        # Main content area
        content_layout = QHBoxLayout()
        
        # Calendar widget
        self.calendar = QCalendarWidget()
        self.calendar.setGridVisible(True)
        self.calendar.clicked.connect(self.on_date_clicked)
        self.calendar.activated.connect(self.on_date_double_clicked)
        
        # Enable context menu for right-click
        self.calendar.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.calendar.customContextMenuRequested.connect(self.show_date_context_menu)
        
        content_layout.addWidget(self.calendar, 2)
        
        # Right panel for selected date details
        self.details_panel = QWidget()
        self.details_panel.setFixedWidth(300)
        self.details_panel.setStyleSheet("""
            QWidget {
                background-color: #f8fafc;
                border: 1px solid #e2e8f0;
                border-radius: 8px;
                margin: 4px;
            }
        """)
        
        details_layout = QVBoxLayout(self.details_panel)
        
        # Selected date label
        self.selected_date_label = QLabel("Select a date")
        self.selected_date_label.setStyleSheet("""
            QLabel {
                font-weight: bold;
                font-size: 14px;
                padding: 8px;
                background-color: #3b82f6;
                color: white;
                border-radius: 4px;
            }
        """)
        self.selected_date_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        details_layout.addWidget(self.selected_date_label)
        
        # Tasks section
        tasks_label = QLabel("Tasks & Deadlines")
        tasks_label.setStyleSheet("font-weight: bold; margin-top: 8px;")
        details_layout.addWidget(tasks_label)
        
        self.tasks_list = QListWidget()
        self.tasks_list.setDragDropMode(QListWidget.DragDropMode.InternalMove)
        self.tasks_list.setDefaultDropAction(Qt.DropAction.MoveAction)
        self.tasks_list.setStyleSheet("""
            QListWidget {
                background-color: white;
                border: 1px solid #e2e8f0;
                border-radius: 4px;
                padding: 4px;
            }
            QListWidget::item {
                padding: 6px;
                border-bottom: 1px solid #f1f5f9;
            }
            QListWidget::item:hover {
                background-color: #f8fafc;
            }
        """)
        
        # Enable context menu for tasks
        self.tasks_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.tasks_list.customContextMenuRequested.connect(self.show_task_context_menu)
        
        # Enable double-click editing
        self.tasks_list.itemDoubleClicked.connect(self.edit_task_inline)
        
        details_layout.addWidget(self.tasks_list)
        
        # Leaves section
        leaves_label = QLabel("Leaves")
        leaves_label.setStyleSheet("font-weight: bold; margin-top: 8px;")
        details_layout.addWidget(leaves_label)
        
        self.leaves_list = QListWidget()
        self.leaves_list.setStyleSheet("""
            QListWidget {
                background-color: white;
                border: 1px solid #e2e8f0;
                border-radius: 4px;
                padding: 4px;
            }
            QListWidget::item {
                padding: 6px;
                border-bottom: 1px solid #f1f5f9;
            }
        """)
        details_layout.addWidget(self.leaves_list)
        
        content_layout.addWidget(self.details_panel)
        layout.addLayout(content_layout)
    
    def load_calendar_data(self):
        """Load tasks and leaves data for calendar display"""
        # Clear existing data
        self.tasks_data = []
        self.leaves_data = []
        
        # Helper function to check if user matches (handle different name formats)
        def user_matches(data_user, current_user):
            if not data_user or not current_user:
                return False
            
            # Normalize names for comparison
            data_user_norm = data_user.lower().strip()
            current_user_norm = current_user.lower().strip()
            
            # Direct match
            if data_user_norm == current_user_norm:
                return True
            
            # Check if current_user is contained in data_user (e.g., "s.h.prabhakaran" in "Santhosh Prabhakaran")
            if current_user_norm in data_user_norm:
                return True
            
            # Check if data_user is contained in current_user
            if data_user_norm in current_user_norm:
                return True
            
            # Check for partial matches (e.g., "santhosh" in "Santhosh Prabhakaran")
            # Split names and check for partial matches
            data_parts = data_user_norm.split()
            current_parts = current_user_norm.split()
            
            for current_part in current_parts:
                for data_part in data_parts:
                    if len(current_part) > 2 and len(data_part) > 2:  # Only check meaningful parts
                        if current_part in data_part or data_part in current_part:
                            return True
            
            # Check for initials match (e.g., "s.h.prabhakaran" matches "Santhosh Prabhakaran")
            if '.' in current_user_norm:
                # Extract initials from current_user
                initials = ''.join([part[0] for part in current_user_norm.split('.') if part])
                if initials and len(initials) >= 2:
                    # Check if initials match the start of data_user
                    if data_user_norm.startswith(initials.lower()):
                        return True
            
            return False
        
        # Load tasks from Potential Issues
        if "Potential Issues" in self.main_window.data:
            pi_data = self.main_window.data["Potential Issues"]
            pi_cols = PANE_COLUMNS["Potential Issues"]
            ism_idx = pi_cols.index("ISM Name") if "ISM Name" in pi_cols else None
            
            for row in pi_data:
                if ism_idx is not None and len(row) > ism_idx and user_matches(row[ism_idx], self.current_user):
                    due_date = row[pi_cols.index("Due Date")] if "Due Date" in pi_cols and pi_cols.index("Due Date") < len(row) else ""
                    if due_date:
                        # Get title from Description or Updates field
                        title = ""
                        if "Description" in pi_cols and pi_cols.index("Description") < len(row) and row[pi_cols.index("Description")]:
                            title = row[pi_cols.index("Description")]
                        elif "Updates" in pi_cols and pi_cols.index("Updates") < len(row) and row[pi_cols.index("Updates")]:
                            title = row[pi_cols.index("Updates")]
                        else:
                            title = "Potential Issue"
                        
                        task = {
                            'title': title,
                            'due_date': due_date,
                            'priority': row[pi_cols.index("RAG Status")] if "RAG Status" in pi_cols and pi_cols.index("RAG Status") < len(row) else "GREEN",
                            'project': row[pi_cols.index("Project Name")] if "Project Name" in pi_cols and pi_cols.index("Project Name") < len(row) else "",
                            'status': row[pi_cols.index("Status")] if "Status" in pi_cols and pi_cols.index("Status") < len(row) else "To Do",
                            'type': 'potential_issue',
                            'source': 'Potential Issues'
                        }
                        self.tasks_data.append(task)
        
        # Load tasks from Activities
        if "Activities" in self.main_window.data:
            act_data = self.main_window.data["Activities"]
            act_cols = PANE_COLUMNS["Activities"]
            ism_idx = act_cols.index("ISM Name") if "ISM Name" in act_cols else None
            
            for row in act_data:
                if ism_idx is not None and len(row) > ism_idx and user_matches(row[ism_idx], self.current_user):
                    target_date = row[act_cols.index("Target Date")] if "Target Date" in act_cols and act_cols.index("Target Date") < len(row) else ""
                    if target_date:
                        # Get title from Activity/Issue or Brief Update field
                        title = ""
                        if "Activity/Issue" in act_cols and act_cols.index("Activity/Issue") < len(row) and row[act_cols.index("Activity/Issue")]:
                            title = row[act_cols.index("Activity/Issue")]
                        elif "Brief Update" in act_cols and act_cols.index("Brief Update") < len(row) and row[act_cols.index("Brief Update")]:
                            title = row[act_cols.index("Brief Update")]
                        else:
                            title = "Activity"
                        
                        task = {
                            'title': title,
                            'due_date': target_date,
                            'priority': row[act_cols.index("RAG")] if "RAG" in act_cols and act_cols.index("RAG") < len(row) else "GREEN",
                            'project': row[act_cols.index("Project Name")] if "Project Name" in act_cols and act_cols.index("Project Name") < len(row) else "",
                            'status': row[act_cols.index("Status")] if "Status" in act_cols and act_cols.index("Status") < len(row) else "To Do",
                            'type': 'activity',
                            'source': 'Activities'
                        }
                        self.tasks_data.append(task)
        
        # Load leaves
        if "Leave Tracker" in self.main_window.data:
            leave_data = self.main_window.data["Leave Tracker"]
            leave_cols = PANE_COLUMNS["Leave Tracker"]
            ism_idx = leave_cols.index("ISM Name") if "ISM Name" in leave_cols else None
            
            for row in leave_data:
                if ism_idx is not None and len(row) > ism_idx and user_matches(row[ism_idx], self.current_user):
                    leave_date = row[leave_cols.index("Date")] if "Date" in leave_cols and leave_cols.index("Date") < len(row) else ""
                    if leave_date:
                        # Get title from Description or Type field
                        title = ""
                        if "Description" in leave_cols and leave_cols.index("Description") < len(row) and row[leave_cols.index("Description")]:
                            title = row[leave_cols.index("Description")]
                        elif "Type" in leave_cols and leave_cols.index("Type") < len(row) and row[leave_cols.index("Type")]:
                            title = row[leave_cols.index("Type")]
                        else:
                            title = "Leave"
                        
                        leave = {
                            'title': title,
                            'date': leave_date,
                            'type': row[leave_cols.index("Type")] if "Type" in leave_cols and leave_cols.index("Type") < len(row) else "Leave",
                            'duration': row[leave_cols.index("Duration")] if "Duration" in leave_cols and leave_cols.index("Duration") < len(row) else "Full Day",
                            'approval_status': row[leave_cols.index("Approval Status")] if "Approval Status" in leave_cols and leave_cols.index("Approval Status") < len(row) else "Pending"
                        }
                        self.leaves_data.append(leave)
        
        # Update calendar display
        self.update_calendar_display()
    
    def update_calendar_display(self):
        """Update calendar with task and leave highlights"""
        # Group tasks and leaves by date for quick lookup
        self.tasks_by_date = {}
        self.leaves_by_date = {}
        
        # Group tasks by date
        for task in self.tasks_data:
            date = task['due_date']
            if date not in self.tasks_by_date:
                self.tasks_by_date[date] = []
            self.tasks_by_date[date].append(task)
        
        # Group leaves by date
        for leave in self.leaves_data:
            date = leave['date']
            if date not in self.leaves_by_date:
                self.leaves_by_date[date] = []
            self.leaves_by_date[date].append(leave)
        
        # Add visual indicators to calendar dates
        from PyQt6.QtGui import QTextCharFormat, QColor
        from datetime import datetime, date
        
        # Get today's date for comparison
        today = date.today()
        
        # Format for today's date
        today_format = QTextCharFormat()
        today_format.setBackground(QColor(255, 255, 0))  # Bright yellow background
        today_format.setForeground(QColor(0, 0, 0))     # Black text
        today_format.setFontWeight(900)  # Extra bold
        
        # Format for dates with tasks (general)
        task_format = QTextCharFormat()
        task_format.setBackground(QColor(255, 248, 220))  # Light yellow background
        task_format.setForeground(QColor(255, 140, 0))   # Orange text
        task_format.setFontWeight(700)  # Bold
        
        # Format for dates with leaves
        leave_format = QTextCharFormat()
        leave_format.setBackground(QColor(220, 248, 255))  # Light blue background
        leave_format.setForeground(QColor(0, 123, 255))   # Blue text
        leave_format.setFontWeight(700)  # Bold
        
        # Format for overdue tasks (red background)
        overdue_format = QTextCharFormat()
        overdue_format.setBackground(QColor(255, 200, 200))  # Light red background
        overdue_format.setForeground(QColor(220, 38, 38))   # Dark red text
        overdue_format.setFontWeight(700)  # Bold
        
        # Apply formatting to dates
        for date_str in self.tasks_by_date:
            try:
                date = QDate.fromString(date_str, "yyyy-MM-dd")
                if date.isValid():
                    tasks = self.tasks_by_date[date_str]
                    leaves = self.leaves_by_date.get(date_str, [])
                    
                    # Check if this is today's date
                    task_date = datetime.strptime(date_str, "%Y-%m-%d").date()
                    is_today = task_date == today
                    
                    # Check for overdue tasks
                    has_overdue = False
                    for task in tasks:
                        if task_date < today:
                            has_overdue = True
                            break
                    
                    # Apply formatting based on status
                    if is_today:
                        self.calendar.setDateTextFormat(date, today_format)
                    elif has_overdue:
                        self.calendar.setDateTextFormat(date, overdue_format)
                    elif tasks:
                        self.calendar.setDateTextFormat(date, task_format)
                    elif leaves and not tasks:
                        self.calendar.setDateTextFormat(date, leave_format)
            except Exception:
                pass
        
        # Apply formatting to dates with only leaves (not covered above)
        for date_str in self.leaves_by_date:
            try:
                date = QDate.fromString(date_str, "yyyy-MM-dd")
                if date.isValid():
                    if date_str not in self.tasks_by_date:  # Only leaves, no tasks
                        self.calendar.setDateTextFormat(date, leave_format)
            except Exception:
                pass
    
    def set_view_mode(self, mode: str):
        """Set calendar view mode"""
        if mode == "month":
            self.month_view_btn.setChecked(True)
            self.week_view_btn.setChecked(False)
            # Show month view
            self.calendar.show()
            if hasattr(self, 'week_view_widget'):
                self.week_view_widget.hide()
        else:  # week view
            self.week_view_btn.setChecked(True)
            self.month_view_btn.setChecked(False)
            # Show week view
            self.calendar.hide()
            if not hasattr(self, 'week_view_widget'):
                self.create_week_view()
            self.week_view_widget.show()
        
        # Update the calendar display with visual indicators
        self.update_calendar_display()
    
    def create_week_view(self):
        """Create a custom week view widget"""
        from PyQt6.QtWidgets import QWidget, QVBoxLayout, QHBoxLayout, QLabel, QListWidget
        from PyQt6.QtCore import QDate, Qt
        
        self.week_view_widget = QWidget()
        week_layout = QVBoxLayout(self.week_view_widget)
        
        # Week header
        week_header = QHBoxLayout()
        week_title = QLabel("Week View")
        week_title.setStyleSheet("font-size: 16px; font-weight: bold; margin: 10px;")
        week_header.addWidget(week_title)
        week_header.addStretch()
        
        # Navigation buttons
        prev_week_btn = QPushButton("◀ Previous Week")
        next_week_btn = QPushButton("Next Week ▶")
        prev_week_btn.clicked.connect(self.previous_week)
        next_week_btn.clicked.connect(self.next_week)
        week_header.addWidget(prev_week_btn)
        week_header.addWidget(next_week_btn)
        week_layout.addLayout(week_header)
        
        # Week days container
        self.week_days_container = QHBoxLayout()
        week_layout.addLayout(self.week_days_container)
        
        # Initialize with current week
        self.current_week_start = QDate.currentDate()
        self.update_week_view()
        
        # Add to main layout
        layout = self.layout()
        layout.addWidget(self.week_view_widget)
    
    def update_week_view(self):
        """Update the week view display"""
        if not hasattr(self, 'week_view_widget'):
            return
        
        # Clear existing week days
        for i in reversed(range(self.week_days_container.count())):
            child = self.week_days_container.itemAt(i).widget()
            if child:
                child.setParent(None)
        
        # Calculate week start (Monday)
        week_start = self.current_week_start
        while week_start.dayOfWeek() != 1:  # Monday
            week_start = week_start.addDays(-1)
        
        # Create day widgets for the week
        for i in range(7):
            day = week_start.addDays(i)
            day_widget = self.create_day_widget(day)
            self.week_days_container.addWidget(day_widget)
    
    def create_day_widget(self, date: QDate):
        """Create a widget for a single day in week view"""
        from PyQt6.QtWidgets import QWidget, QVBoxLayout, QLabel, QListWidget
        from PyQt6.QtCore import Qt
        
        day_widget = QWidget()
        day_widget.setStyleSheet("border: 1px solid #ddd; margin: 2px; padding: 5px;")
        day_layout = QVBoxLayout(day_widget)
        
        # Day header
        day_label = QLabel(date.toString("ddd, MMM d"))
        day_label.setStyleSheet("font-weight: bold; text-align: center; padding: 5px;")
        day_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        day_layout.addWidget(day_label)
        
        # Day content list
        day_list = QListWidget()
        day_list.setMaximumHeight(200)
        day_list.setStyleSheet("font-size: 11px;")
        day_layout.addWidget(day_list)
        
        # Load content for this day
        date_str = date.toString("yyyy-MM-dd")
        tasks = self.tasks_by_date.get(date_str, [])
        leaves = self.leaves_by_date.get(date_str, [])
        
        # Add tasks
        for task in tasks:
            source_icon = "⚠️" if task.get('source') == 'Potential Issues' else "📋"
            item_text = f"{source_icon} {task['title']}"
            day_list.addItem(item_text)
        
        # Add leaves
        for leave in leaves:
            item_text = f"🗓️ {leave['title']}"
            day_list.addItem(item_text)
        
        # If no items, show message
        if not tasks and not leaves:
            day_list.addItem("No items")
        
        return day_widget
    
    def previous_week(self):
        """Go to previous week"""
        self.current_week_start = self.current_week_start.addDays(-7)
        self.update_week_view()
    
    def next_week(self):
        """Go to next week"""
        self.current_week_start = self.current_week_start.addDays(7)
        self.update_week_view()
    
    def show_date_context_menu(self, position):
        """Show context menu for calendar dates"""
        from PyQt6.QtWidgets import QMenu, QMessageBox
        from PyQt6.QtCore import QDate
        
        # Use the currently selected date for the context menu
        date = self.calendar.selectedDate()
        if not date.isValid():
            return
        
        menu = QMenu(self)
        
        # Add Potential Issue
        add_potential_issue = menu.addAction("➕ Add Potential Issue")
        add_potential_issue.triggered.connect(lambda: self.quick_add_task(date, "Potential Issues"))
        
        # Add Activity
        add_activity = menu.addAction("➕ Add Activity")
        add_activity.triggered.connect(lambda: self.quick_add_task(date, "Activities"))
        
        # Add Leave
        add_leave = menu.addAction("➕ Add Leave")
        add_leave.triggered.connect(lambda: self.quick_add_leave(date))
        
        menu.addSeparator()
        
        # View date details
        view_details = menu.addAction("👁️ View Details")
        view_details.triggered.connect(lambda: self.on_date_clicked(date))
        
        # Show menu at cursor position
        from PyQt6.QtCore import QPoint
        if position:
            menu.exec(self.calendar.mapToGlobal(position))
        else:
            # Fallback to showing at center of calendar
            menu.exec(self.calendar.mapToGlobal(QPoint(self.calendar.width()//2, self.calendar.height()//2)))
    
    def show_task_context_menu(self, position):
        """Show context menu for tasks"""
        from PyQt6.QtWidgets import QMenu, QMessageBox
        
        item = self.tasks_list.itemAt(position)
        if not item:
            return
        
        menu = QMenu(self)
        
        # Edit task
        edit_task = menu.addAction("✏️ Edit Task")
        edit_task.triggered.connect(lambda: self.edit_task_inline(item))
        
        # Delete task
        delete_task = menu.addAction("🗑️ Delete Task")
        delete_task.triggered.connect(lambda: self.delete_task(item))
        
        # Move to different date
        move_task = menu.addAction("📅 Move to Different Date")
        move_task.triggered.connect(lambda: self.move_task_to_date(item))
        
        # Show menu at cursor position
        from PyQt6.QtCore import QPoint
        if position:
            menu.exec(self.tasks_list.mapToGlobal(position))
        else:
            # Fallback to showing at center of task list
            menu.exec(self.tasks_list.mapToGlobal(QPoint(self.tasks_list.width()//2, self.tasks_list.height()//2)))
    
    def quick_add_task(self, date: QDate, task_type: str):
        """Quickly add a task to a specific date using the same logic as the main app"""
        from PyQt6.QtWidgets import QInputDialog, QMessageBox
        
        title, ok = QInputDialog.getText(self, f"Add {task_type}", f"Enter {task_type.lower()} title:")
        if not ok or not title.strip():
            return
        
        try:
            date_str = date.toString("yyyy-MM-dd")
            
            # Use the same logic as the main app's add_row method
            if task_type == "Potential Issues":
                # Create new row using the app's default logic
                cols = PANE_COLUMNS["Potential Issues"]
                new_row = default_row_for_columns(cols)
                
                # Set specific values for calendar-created items
                if "ISM Name" in cols:
                    idx = cols.index("ISM Name")
                    mapped_name = self.main_window._resolve_ism_name_from_eid(self.current_user)
                    new_row[idx] = mapped_name if mapped_name else self.current_user
                
                if "Due Date" in cols:
                    new_row[cols.index("Due Date")] = date_str
                
                if "Description" in cols:
                    new_row[cols.index("Description")] = title
                
                if "Task Type" in cols:
                    new_row[cols.index("Task Type")] = "To Do"
                
                if "Status" in cols:
                    new_row[cols.index("Status")] = "In Progress"
                
                if "Updates" in cols:
                    new_row[cols.index("Updates")] = f"Created via Calendar on {datetime.now().strftime('%Y-%m-%d')}"
                
                # Add to data and refresh using main app's methods
                self.main_window.data["Potential Issues"].append(new_row)
                self.main_window.rebuild_table("Potential Issues")
                self.main_window._save_backend_sqlite()
                
            elif task_type == "Activities":
                # Create new row using the app's default logic
                cols = PANE_COLUMNS["Activities"]
                new_row = default_row_for_columns(cols)
                
                # Set specific values for calendar-created items
                if "ISM Name" in cols:
                    idx = cols.index("ISM Name")
                    mapped_name = self.main_window._resolve_ism_name_from_eid(self.current_user)
                    new_row[idx] = mapped_name if mapped_name else self.current_user
                
                if "Target Date" in cols:
                    new_row[cols.index("Target Date")] = date_str
                
                if "Activity/Issue" in cols:
                    new_row[cols.index("Activity/Issue")] = title
                
                if "Status" in cols:
                    new_row[cols.index("Status")] = "Yet to Start"
                
                if "Brief Update" in cols:
                    new_row[cols.index("Brief Update")] = f"Created via Calendar on {datetime.now().strftime('%Y-%m-%d')}"
                
                # Add to data and refresh using main app's methods
                self.main_window.data["Activities"].append(new_row)
                self.main_window.rebuild_table("Activities")
                self.main_window._save_backend_sqlite()
            
            # Refresh calendar display
            self.load_calendar_data()
            self.on_date_clicked(date)
            
            QMessageBox.information(self, "Success", f"{task_type} added successfully!\nEntry created in {task_type} pane.")
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to add {task_type}: {str(e)}")
    
    def quick_add_leave(self, date: QDate):
        """Quickly add a leave to a specific date using the same logic as the main app"""
        from PyQt6.QtWidgets import QInputDialog, QMessageBox, QDialog, QVBoxLayout, QHBoxLayout, QLabel, QComboBox, QPushButton, QLineEdit
        
        # Create a custom dialog for leave input
        dialog = QDialog(self)
        dialog.setWindowTitle("Add Leave")
        dialog.setModal(True)
        dialog.resize(400, 200)
        
        layout = QVBoxLayout(dialog)
        
        # Title input
        title_label = QLabel("Leave Title:")
        layout.addWidget(title_label)
        title_input = QLineEdit()
        title_input.setPlaceholderText("Enter leave title...")
        layout.addWidget(title_input)
        
        # Leave type selection
        type_label = QLabel("Leave Type:")
        layout.addWidget(type_label)
        type_combo = QComboBox()
        type_combo.addItems(["WFH", "Planned Leave", "Public Holiday", "Earned Leave", "Casual Leave"])
        type_combo.setCurrentText("Earned Leave")  # Default selection
        layout.addWidget(type_combo)
        
        # Buttons
        button_layout = QHBoxLayout()
        ok_button = QPushButton("Add Leave")
        cancel_button = QPushButton("Cancel")
        button_layout.addWidget(ok_button)
        button_layout.addWidget(cancel_button)
        layout.addLayout(button_layout)
        
        # Connect buttons
        ok_button.clicked.connect(dialog.accept)
        cancel_button.clicked.connect(dialog.reject)
        
        # Show dialog
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
            
        title = title_input.text().strip()
        if not title:
            QMessageBox.warning(self, "Error", "Please enter a leave title.")
            return
            
        leave_type = type_combo.currentText()
        
        try:
            date_str = date.toString("yyyy-MM-dd")
            
            # Use the same logic as the main app's add_row method for Leave Tracker
            cols = PANE_COLUMNS["Leave Tracker"]
            new_row = default_row_for_columns(cols)
            
            # Apply the same logic as the main app's add_row method for Leave Tracker
            if "ISM Name" in cols:
                idx = cols.index("ISM Name")
                mapped_name = self.main_window._resolve_ism_name_from_eid(self.current_user)
                new_row[idx] = mapped_name if mapped_name else self.current_user
            
            # Set specific values for calendar-created leave
            if "Description" in cols:
                new_row[cols.index("Description")] = title
            
            if "Type" in cols:
                new_row[cols.index("Type")] = leave_type
            
            if "Leave Date" in cols:
                new_row[cols.index("Leave Date")] = date_str
            
            if "Date" in cols:
                new_row[cols.index("Date")] = date_str
            
            # Apply the same Leave Tracker defaults as the main app
            if "Approval Status" in cols:
                new_row[cols.index("Approval Status")] = "Pending"
            
            # Set up manager and approver information like the main app
            req_eid = self.current_user
            req_name = self.main_window._resolve_name_from_eid(req_eid) or self.current_user
            mgr_eid = self.main_window._resolve_manager_eid_for_user_eid(req_eid) or ""
            mgr_name = self.main_window._resolve_name_from_eid(mgr_eid) if mgr_eid else ""
            
            if "Requested By Enterprise ID" in cols:
                new_row[cols.index("Requested By Enterprise ID")] = req_eid
            if "Requested By Name" in cols:
                new_row[cols.index("Requested By Name")] = req_name
            if "Approver Enterprise ID" in cols:
                new_row[cols.index("Approver Enterprise ID")] = mgr_eid
            if "Approver Name" in cols:
                new_row[cols.index("Approver Name")] = mgr_name or ""
            
            # Add to data and refresh using main app's methods
            self.main_window.data["Leave Tracker"].append(new_row)
            self.main_window.rebuild_table("Leave Tracker")
            self.main_window._save_backend_sqlite()
            
            # Notify approver if current user is approver (same as main app)
            try:
                self.main_window._notify_approver_new_leave(mgr_eid, date_str, req_name)
            except Exception:
                pass
            
            # Refresh calendar display
            self.load_calendar_data()
            self.on_date_clicked(date)
            
            QMessageBox.information(self, "Success", "Leave added successfully!\nEntry created in Leave Tracker pane.")
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to add leave: {str(e)}")
    
    def edit_task_inline(self, item):
        """Edit task inline"""
        from PyQt6.QtWidgets import QInputDialog, QMessageBox
        
        if not item:
            return
        
        current_text = item.text()
        new_text, ok = QInputDialog.getText(self, "Edit Task", "Edit task title:", text=current_text)
        if ok and new_text.strip():
            item.setText(new_text)
            # Update the underlying data
            self.update_task_in_data(item, new_text)
    
    def delete_task(self, item):
        """Delete a task"""
        from PyQt6.QtWidgets import QMessageBox
        
        reply = QMessageBox.question(self, "Delete Task", "Are you sure you want to delete this task?",
                                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.tasks_list.takeItem(self.tasks_list.row(item))
            # Remove from underlying data
            self.remove_task_from_data(item)
    
    def move_task_to_date(self, item):
        """Move task to a different date"""
        from PyQt6.QtWidgets import QInputDialog, QMessageBox
        from PyQt6.QtCore import QDate
        
        current_text = item.text()
        new_date_str, ok = QInputDialog.getText(self, "Move Task", "Enter new date (YYYY-MM-DD):")
        if ok and new_date_str.strip():
            try:
                new_date = QDate.fromString(new_date_str, "yyyy-MM-dd")
                if new_date.isValid():
                    # Update the task date in data
                    self.update_task_date_in_data(item, new_date_str)
                    QMessageBox.information(self, "Success", "Task moved successfully!")
                    # Refresh display
                    self.load_calendar_data()
                else:
                    QMessageBox.warning(self, "Error", "Invalid date format. Use YYYY-MM-DD")
            except Exception as e:
                QMessageBox.warning(self, "Error", f"Failed to move task: {str(e)}")
    
    def update_task_in_data(self, item, new_text):
        """Update task in underlying data"""
        # This is a simplified implementation
        # In a full implementation, you would update the actual data source
        pass
    
    def remove_task_from_data(self, item):
        """Remove task from underlying data"""
        # This is a simplified implementation
        # In a full implementation, you would remove from the actual data source
        pass
    
    def update_task_date_in_data(self, item, new_date_str):
        """Update task date in underlying data"""
        # This is a simplified implementation
        # In a full implementation, you would update the actual data source
        pass
    
    def on_date_clicked(self, date: QDate):
        """Handle date click to show tasks and leaves for that date"""
        date_str = date.toString("yyyy-MM-dd")
        self.selected_date_label.setText(f"Selected: {date_str}")
        
        # Reload data to ensure it's up-to-date
        self.load_calendar_data()
        
        # Clear lists
        self.tasks_list.clear()
        self.leaves_list.clear()
        
        # Show tasks for this date using the grouped data
        tasks_for_date = self.tasks_by_date.get(date_str, [])
        for task in tasks_for_date:
            priority_colors = {
                'RED': '#ef4444',
                'YELLOW': '#f59e0b',
                'GREEN': '#10b981',
                'AMBER': '#f59e0b'
            }
            priority_color = priority_colors.get(task['priority'], '#6b7280')
            
            # Status color
            status_colors = {
                'To Do': '#6b7280',
                'In Progress': '#3b82f6',
                'Done': '#10b981',
                'Yet to Start': '#f59e0b',
                'NA': '#6b7280'
            }
            status_color = status_colors.get(task.get('status', 'To Do'), '#6b7280')
            
            # Source indicator
            source_icon = "⚠️" if task.get('source') == 'Potential Issues' else "📋"
            
            item_text = f"{source_icon} {task['title']}\n● {task['priority']} • Status: {task.get('status', 'To Do')}\nProject: {task.get('project', 'N/A')}"
            
            item = QListWidgetItem()
            item.setText(item_text)
            self.tasks_list.addItem(item)
        
        # Show leaves for this date using the grouped data
        leaves_for_date = self.leaves_by_date.get(date_str, [])
        for leave in leaves_for_date:
            status_color = "#10b981" if leave['approval_status'] == "Approved" else "#f59e0b"
            
            item_text = f"🗓️ {leave['title']}\n● {leave['approval_status']}\n{leave['type']} - {leave['duration']}"
            
            item = QListWidgetItem()
            item.setText(item_text)
            self.leaves_list.addItem(item)
        
        # Update the display to show counts
        tasks_count = self.tasks_list.count()
        leaves_count = self.leaves_list.count()
        
        if tasks_count == 0 and leaves_count == 0:
            # Show a message when no items for this date
            no_items = QListWidgetItem()
            no_items.setText("No tasks or leaves scheduled for this date")
            # Set a custom role for styling
            no_items.setData(Qt.ItemDataRole.UserRole, "no_items")
            self.tasks_list.addItem(no_items)
    
    def on_date_double_clicked(self, date: QDate):
        """Handle double-click to add new task or leave"""
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Add Item for {date.toString('yyyy-MM-dd')}")
        dialog.setModal(True)
        dialog.resize(400, 200)
        
        layout = QVBoxLayout(dialog)
        
        # Item type selection
        layout.addWidget(QLabel("Item Type:"))
        type_combo = QComboBox()
        type_combo.addItems(["Task", "Leave"])
        layout.addWidget(type_combo)
        
        # Title
        layout.addWidget(QLabel("Title:"))
        title_edit = QLineEdit()
        layout.addWidget(title_edit)
        
        # Additional fields based on type
        self.desc_edit = QTextEdit()
        self.desc_edit.setMaximumHeight(60)
        self.desc_edit.setPlaceholderText("Description")
        layout.addWidget(self.desc_edit)
        
        # Buttons
        button_layout = QHBoxLayout()
        add_btn = QPushButton("Add Item")
        cancel_btn = QPushButton("Cancel")
        button_layout.addWidget(add_btn)
        button_layout.addWidget(cancel_btn)
        layout.addLayout(button_layout)
        
        # Connect buttons
        add_btn.clicked.connect(lambda: self.add_calendar_item(
            date.toString("yyyy-MM-dd"), type_combo.currentText(),
            title_edit.text(), self.desc_edit.toPlainText(), dialog
        ))
        cancel_btn.clicked.connect(dialog.reject)
        
        dialog.exec()
    
    def add_calendar_item(self, date: str, item_type: str, title: str, description: str, dialog: QDialog):
        """Add new task or leave item"""
        if not title.strip():
            QMessageBox.warning(self, "Invalid Input", "Title is required.")
            return
        
        try:
            if item_type == "Task":
                # Add to Potential Issues
                new_row = [""] * len(PANE_COLUMNS["Potential Issues"])
                new_row[PANE_COLUMNS["Potential Issues"].index("Task Type")] = "Task"
                new_row[PANE_COLUMNS["Potential Issues"].index("Status")] = "To Do"
                new_row[PANE_COLUMNS["Potential Issues"].index("Created Date")] = datetime.now().strftime("%Y-%m-%d")
                new_row[PANE_COLUMNS["Potential Issues"].index("Due Date")] = date
                new_row[PANE_COLUMNS["Potential Issues"].index("RAG Status")] = "GREEN"
                new_row[PANE_COLUMNS["Potential Issues"].index("ISM Name")] = self.current_user
                new_row[PANE_COLUMNS["Potential Issues"].index("Description")] = title
                new_row[PANE_COLUMNS["Potential Issues"].index("Updates")] = description
                
                self.main_window.data["Potential Issues"].append(new_row)
            else:
                # Add to Leave Tracker
                new_row = [""] * len(PANE_COLUMNS["Leave Tracker"])
                new_row[PANE_COLUMNS["Leave Tracker"].index("Date")] = date
                new_row[PANE_COLUMNS["Leave Tracker"].index("Type")] = "Personal Leave"
                new_row[PANE_COLUMNS["Leave Tracker"].index("Duration")] = "Full Day"
                new_row[PANE_COLUMNS["Leave Tracker"].index("Description")] = title
                new_row[PANE_COLUMNS["Leave Tracker"].index("ISM Name")] = self.current_user
                new_row[PANE_COLUMNS["Leave Tracker"].index("Approval Status")] = "Pending"
                new_row[PANE_COLUMNS["Leave Tracker"].index("Requested By Enterprise ID")] = self.current_user
                new_row[PANE_COLUMNS["Leave Tracker"].index("Requested By Name")] = self.current_user
            
            self.load_calendar_data()
            dialog.accept()
            
        except Exception as e:
            ErrorHandler.handle_ui_error("add calendar item", e)
    
    def refresh_calendar_data(self):
        """Refresh calendar data from all panes"""
        try:
            self.load_calendar_data()
            QMessageBox.information(self, "Calendar Refreshed", f"Calendar data refreshed successfully!\n\nTasks: {len(self.tasks_data)}\nLeaves: {len(self.leaves_data)}")
        except Exception as e:
            QMessageBox.warning(self, "Refresh Error", f"Failed to refresh calendar data: {str(e)}")
    
    def refresh_calendar_if_visible(self):
        """Refresh calendar data if calendar tab is currently visible"""
        try:
            if hasattr(self.main_window, 'calendar_tab') and self.main_window.calendar_tab:
                if hasattr(self.main_window, 'tabs') and self.main_window.tabs.currentWidget() == self.main_window.calendar_tab:
                    self.load_calendar_data()
                    # Also refresh the calendar display to show updated visual indicators
                    self.update_calendar_display()
        except Exception as e:
            # Log the error for debugging
            print(f"Calendar refresh error: {e}")
            pass
    
    def showEvent(self, event):
        """Called when the calendar tab becomes visible - refresh data"""
        super().showEvent(event)
        # Refresh data when tab becomes visible
        try:
            self.load_calendar_data()
            self.update_calendar_display()
        except Exception as e:
            print(f"Calendar showEvent refresh error: {e}")
            pass
    
    def sync_with_system_calendar(self):
        """Sync with system calendar (stub method)"""
        QMessageBox.information(self, "Sync Calendar", 
                               "System calendar sync is not yet implemented.\nThis feature will be available in a future update.")


def main():
    app = QApplication(sys.argv)
    
    # Initialize global timers after QApplication is created
    initialize_global_timers()
    
    win = MainWindow()
    try:
        install_feature_toggle_ui(win)
    except Exception as e:
        ErrorHandler.handle_ui_error("install feature toggle ui (main)", e)
    win.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
